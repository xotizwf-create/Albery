"""docconvert.py — конвертация между PDF и Word для сценариев, где нужна структурная правка.

Два разных инструмента для двух разных задач, и путать их дорого:

* Точечная правка («поменяй сумму», «убери персональные данные») — это `pdfedit`: он меняет
  фрагмент прямо в исходном PDF, и файл остаётся тем же самым. Конвертация здесь только
  навредит.
* Структурная переработка («перепиши раздел», «добавь таблицу», «собери договор заново»)
  требует редактируемого документа. Тогда: `pdf_to_docx` → правка в Word/через `docedit` →
  `docx_to_pdf`. Документ при этом ВЁРСТСЯ ЗАНОВО: сохраняются текст, порядок, таблицы,
  начертания и выравнивание, но переносы строк и положение блоков на странице будут своими.
  Обещать «тот же самый PDF» после круга конвертации нельзя, и здесь этого не делается.

PDF → DOCX собирается на PyMuPDF + python-docx: без тяжёлых зависимостей, которые слабому
серверу пришлось бы держать ради одного сценария. DOCX → PDF выполняет LibreOffice в
headless-режиме — единственный способ получить вёрстку Word без самого Word.
"""
from __future__ import annotations

import io
import logging
import os
import shutil
import subprocess
import tempfile

log = logging.getLogger(__name__)

MAX_INPUT_BYTES = int(os.getenv("DOCCONVERT_MAX_BYTES", str(40 * 1024 * 1024)) or str(40 * 1024 * 1024))
CONVERT_TIMEOUT_SECONDS = int(os.getenv("DOCCONVERT_TIMEOUT_SECONDS", "180") or "180")
_SOFFICE_CANDIDATES = ("soffice", "libreoffice", "/usr/bin/soffice", "/usr/lib/libreoffice/program/soffice")


class ConvertError(Exception):
    """Конвертацию выполнить нельзя — причина понятна человеку из сообщения."""


# --------------------------------------------------------------------------- PDF → DOCX


def _alignment(line_boxes: list, body_left: float, body_right: float):
    """Выравнивание абзаца по его геометрии внутри ТЕКСТОВОЙ области страницы.

    Считать от края листа нельзя: поля документа сдвигают весь текст, и тогда каждый
    абзац выглядит «отступленным справа и слева», то есть ложно центрированным.
    """

    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if not line_boxes:
        return WD_ALIGN_PARAGRAPH.LEFT
    left = min(box[0] for box in line_boxes)
    right = max(box[2] for box in line_boxes)
    left_gap = left - body_left
    right_gap = body_right - right
    if left_gap > 12 and right_gap > 12 and abs(left_gap - right_gap) < 24:
        return WD_ALIGN_PARAGRAPH.CENTER
    # Абзац из нескольких строк, дотянутых до правого края, — выключка по ширине;
    # в договорах это основной режим, и терять его при переносе в Word заметно.
    if len(line_boxes) > 1 and right_gap < 6:
        return WD_ALIGN_PARAGRAPH.JUSTIFY
    return WD_ALIGN_PARAGRAPH.LEFT


def _font_name(span_font: str) -> str:
    name = (span_font or "").lower()
    if "arial" in name or "helvetica" in name or "sans" in name:
        return "Arial"
    if "courier" in name or "mono" in name:
        return "Courier New"
    return "Times New Roman"


def pdf_to_docx(data: bytes, *, file_name: str = "document.pdf") -> tuple[bytes, list[str]]:
    """PDF → редактируемый DOCX. Возвращает (байты, предупреждения)."""

    try:
        import fitz
    except ImportError as exc:  # pragma: no cover
        raise ConvertError("Для конвертации PDF нужен пакет pymupdf — он не установлен.") from exc
    from docx import Document
    from docx.shared import Pt

    if not data:
        raise ConvertError("Файл пустой.")
    if len(data) > MAX_INPUT_BYTES:
        raise ConvertError(f"PDF больше допустимых {MAX_INPUT_BYTES // (1024 * 1024)} МБ.")

    warnings: list[str] = []
    document = Document()
    document.styles["Normal"].font.name = "Times New Roman"
    document.styles["Normal"].font.size = Pt(12)

    try:
        opened = fitz.open(stream=data, filetype="pdf")
    except Exception as exc:  # noqa: BLE001 - битый файл, а не сбой сервера
        raise ConvertError("Файл не открывается как PDF — он повреждён или это другой формат.") from exc

    with opened as doc:
        if doc.is_encrypted and not doc.authenticate(""):
            raise ConvertError("PDF защищён паролем — снимите защиту и пришлите файл заново.")
        images_total = 0
        for page_index, page in enumerate(doc):
            if page_index:
                document.add_page_break()
            images_total += len(page.get_images())

            try:
                tables = page.find_tables()
                table_list = list(tables.tables)
            except Exception:  # noqa: BLE001 - поиск таблиц не должен ронять конвертацию
                table_list = []
            table_zones = [fitz.Rect(table.bbox) for table in table_list]

            text_blocks = [
                block for block in page.get_text("dict").get("blocks", []) if block.get("type") == 0
            ]
            line_boxes_all = [
                line["bbox"] for block in text_blocks for line in block.get("lines", [])
            ]
            body_left = min((box[0] for box in line_boxes_all), default=page.rect.x0)
            body_right = max((box[2] for box in line_boxes_all), default=page.rect.x1)

            for block in page.get_text("dict").get("blocks", []):
                if block.get("type") != 0:
                    continue
                block_rect = fitz.Rect(block["bbox"])
                if any(zone.intersects(block_rect) for zone in table_zones):
                    continue  # текст таблицы выйдет вместе с самой таблицей
                lines = block.get("lines", [])
                if not lines:
                    continue
                paragraph = document.add_paragraph()
                line_boxes = [line["bbox"] for line in lines]
                paragraph.alignment = _alignment(line_boxes, body_left, body_right)
                for line_index, line in enumerate(lines):
                    for span in line.get("spans", []):
                        text = span.get("text") or ""
                        if not text.strip() and not paragraph.runs:
                            continue
                        run = paragraph.add_run(text.replace(" ", " "))
                        flags = int(span.get("flags") or 0)
                        run.bold = bool(flags & 2 ** 4) or "bold" in str(span.get("font", "")).lower()
                        run.italic = bool(flags & 2 ** 1) or "italic" in str(span.get("font", "")).lower()
                        run.font.size = Pt(round(float(span.get("size") or 12), 1))
                        run.font.name = _font_name(span.get("font"))
                    if line_index < len(lines) - 1:
                        # Внутри абзаца PDF переносит строки сам; в Word перенос сделает Word,
                        # поэтому строки склеиваются пробелом, а не жёстким переводом строки.
                        if paragraph.runs and not paragraph.runs[-1].text.endswith(" "):
                            paragraph.runs[-1].text += " "

            for table in table_list:
                extracted = table.extract()
                if not extracted:
                    continue
                columns = max(len(row) for row in extracted)
                word_table = document.add_table(rows=0, cols=columns)
                word_table.style = "Table Grid"
                for row in extracted:
                    cells = word_table.add_row().cells
                    for column_index in range(columns):
                        value = row[column_index] if column_index < len(row) else ""
                        cells[column_index].text = (value or "").replace(" ", " ").strip()

    if images_total:
        warnings.append(
            f"В PDF {images_total} изображений (печати, подписи, логотипы) — в Word они не переносятся."
        )

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue(), warnings


# --------------------------------------------------------------------------- DOCX → PDF


def soffice_path() -> str | None:
    for candidate in _SOFFICE_CANDIDATES:
        found = shutil.which(candidate) if not candidate.startswith("/") else (
            candidate if os.path.isfile(candidate) else None
        )
        if found:
            return found
    return None


def docx_to_pdf(data: bytes, *, file_name: str = "document.docx") -> bytes:
    """DOCX → PDF через LibreOffice. Вёрстка — как в Word при экспорте в PDF."""

    if not data:
        raise ConvertError("Файл пустой.")
    if len(data) > MAX_INPUT_BYTES:
        raise ConvertError(f"Документ больше допустимых {MAX_INPUT_BYTES // (1024 * 1024)} МБ.")
    binary = soffice_path()
    if binary is None:
        raise ConvertError(
            "На сервере нет LibreOffice — конвертация Word → PDF недоступна "
            "(ставится пакетом libreoffice-writer)."
        )

    with tempfile.TemporaryDirectory(prefix="docconvert-") as workdir:
        source = os.path.join(workdir, "input.docx")
        with open(source, "wb") as handle:
            handle.write(data)
        # Свой профиль на каждый запуск: общий профиль LibreOffice не переживает
        # параллельных вызовов и молча отдаёт пустой результат.
        profile = os.path.join(workdir, "profile")
        try:
            completed = subprocess.run(
                [
                    binary,
                    f"-env:UserInstallation=file://{profile}",
                    "--headless",
                    "--norestore",
                    "--convert-to",
                    "pdf:writer_pdf_Export",
                    "--outdir",
                    workdir,
                    source,
                ],
                capture_output=True,
                timeout=CONVERT_TIMEOUT_SECONDS,
                check=False,
            )
        except subprocess.TimeoutExpired as exc:
            raise ConvertError(
                f"LibreOffice не уложился в {CONVERT_TIMEOUT_SECONDS} с — документ слишком тяжёлый."
            ) from exc
        result = os.path.join(workdir, "input.pdf")
        if not os.path.isfile(result):
            detail = (completed.stderr or completed.stdout or b"").decode("utf-8", "replace")[:300]
            raise ConvertError(f"LibreOffice не смог преобразовать «{file_name}»: {detail or 'причина не сообщена'}")
        with open(result, "rb") as handle:
            return handle.read()
