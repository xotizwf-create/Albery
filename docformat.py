"""Dependency-free HTML → DOCX converter for agent-authored documents.

Powers the `export_document` MCP tool: the AGENT fully controls the layout by writing
HTML (page breaks, alignment, font sizes, bordered/borderless tables), and this module
mechanically renders it with python-docx — no app-side opinions about how a contract
should look. That is the whole point: when a user asks «приложения с новой страницы» or
«уберите лишнее», the agent fixes its own HTML on the next turn, no code deploy.

Supported subset (documented in the tool description; unknown tags are ignored, their
text flows through):
  blocks:  h1..h4, p, div, li (ul/ol), table/tr/th/td, br, hr
  inline:  b/strong, i/em, u, s
  styles:  text-align: left|center|right|justify (or align="...")
           page-break-before: always  (or class="page-break" on any block/hr)
           font-size: NNpt · text-indent: N.NNcm · line-height: N.N
           border: none / border="0" on <table> → borderless (реквизиты сторон)
Emoji and other pictographs are stripped defensively — Times New Roman has no glyphs
for them, they render as boxes («иероглифы») in Word.
"""
from __future__ import annotations

import io
import re

from html.parser import HTMLParser

_EMOJI_RE = re.compile(
    "[\U0001F000-\U0001FAFF\U00002600-\U000027BF\U0001F1E6-\U0001F1FF"
    "\U0000FE0E\U0000FE0F\U0000200D\U00002B00-\U00002BFF\U00002190-\U000021FF]+"
)

_ALIGN = {"left": 0, "center": 1, "right": 2, "justify": 3}  # WD_ALIGN_PARAGRAPH values


def _style_dict(attrs: dict[str, str]) -> dict[str, str]:
    out: dict[str, str] = {}
    for part in (attrs.get("style") or "").split(";"):
        if ":" in part:
            k, v = part.split(":", 1)
            out[k.strip().lower()] = v.strip().lower()
    return out


class _DocBuilder(HTMLParser):
    _BLOCKS = ("h1", "h2", "h3", "h4", "p", "div", "li")
    _HEAD_SIZES = {"h1": 16, "h2": 14, "h3": 14, "h4": 13}

    def __init__(self, doc, base_size: float, base_line: float):
        super().__init__(convert_charrefs=True)
        self.doc = doc
        self.base_size = base_size
        self.base_line = base_line
        self.bold = 0
        self.italic = 0
        self.under = 0
        self.strike = 0
        self.para = None          # open docx paragraph
        self.block_tag = ""
        self.block_style: dict[str, str] = {}
        self.pending_page_break = False
        self.list_stack: list[str] = []
        self.skip_depth = 0       # inside <script>/<style>
        # table collection state
        self.table_rows: list[list[dict]] | None = None
        self.table_borders = True
        self.row: list[dict] | None = None
        self.cell: dict | None = None

    # --- helpers ---------------------------------------------------------------------------

    def _close_para(self) -> None:
        self.para = None
        self.block_tag = ""
        self.block_style = {}

    def _open_para(self, tag: str = "p", attrs: dict[str, str] | None = None):
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm, Pt

        attrs = attrs or {}
        style = _style_dict(attrs)
        if self.list_stack and tag == "li":
            kind = self.list_stack[-1]
            p = self.doc.add_paragraph(style="List Number" if kind == "ol" else "List Bullet")
        else:
            p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(4)
        try:
            pf.line_spacing = float(style.get("line-height") or self.base_line)
        except ValueError:
            pf.line_spacing = self.base_line
        indent = style.get("text-indent") or ""
        if indent.endswith("cm"):
            try:
                pf.first_line_indent = Cm(float(indent[:-2]))
            except ValueError:
                pass
        align = (style.get("text-align") or attrs.get("align") or "").strip().lower()
        if tag in self._HEAD_SIZES and not align:
            align = "center" if tag == "h1" else "left"
        # GOST default for official documents: plain paragraphs are justified unless the
        # HTML says otherwise. The model no longer has to remember this on every <p> —
        # explicit text-align always wins, lists/headers/tables keep their own defaults.
        if not align and tag in ("p", "div") and not self.list_stack:
            align = "justify"
        if align in _ALIGN:
            p.alignment = {0: WD_ALIGN_PARAGRAPH.LEFT, 1: WD_ALIGN_PARAGRAPH.CENTER,
                           2: WD_ALIGN_PARAGRAPH.RIGHT, 3: WD_ALIGN_PARAGRAPH.JUSTIFY}[_ALIGN[align]]
        if self.pending_page_break or "always" in style.get("page-break-before", "") \
                or "page-break" in (attrs.get("class") or ""):
            pf.page_break_before = True
            self.pending_page_break = False
        self.para = p
        self.block_tag = tag
        self.block_style = style
        return p

    def _run_size(self) -> float:
        # Headings keep the BASE size (bold is what distinguishes them) — official RU documents
        # are one type size throughout («в договоре всё должно быть 12 шрифтом», owner 2026-07-14).
        # An explicit font-size style still overrides below.
        base = self.base_size
        override = self.block_style.get("font-size") or ""
        if override.endswith("pt"):
            try:
                base = float(override[:-2])
            except ValueError:
                pass
        return base

    def _add_text(self, text: str) -> None:
        from docx.shared import Pt, RGBColor
        from docx.oxml.ns import qn

        if self.cell is not None:
            self.cell["runs"].append((text, self.bold > 0))
            return
        if self.para is None:
            if not text.strip():
                return
            self._open_para()
        run = self.para.add_run(text)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(self._run_size())
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.bold = self.bold > 0 or self.block_tag in self._HEAD_SIZES
        run.italic = self.italic > 0
        run.underline = self.under > 0
        run.font.strike = self.strike > 0

    def _flush_table(self) -> None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt, RGBColor

        rows = [r for r in (self.table_rows or []) if r]
        self.table_rows, self.row, self.cell = None, None, None
        if not rows:
            return
        width = max(len(r) for r in rows)
        table = self.doc.add_table(rows=len(rows), cols=width)
        if self.table_borders:
            tblPr = table._tbl.tblPr
            borders = OxmlElement("w:tblBorders")
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                el = OxmlElement(f"w:{edge}")
                el.set(qn("w:val"), "single")
                el.set(qn("w:sz"), "6")
                el.set(qn("w:color"), "000000")
                borders.append(el)
            tblPr.append(borders)
        for ri, cells in enumerate(rows):
            for ci in range(width):
                spec = cells[ci] if ci < len(cells) else {"runs": [], "header": False, "align": ""}
                cell = table.cell(ri, ci)
                cell.text = ""
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                align = spec.get("align") or ("center" if spec["header"] else "")
                if align == "center":
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif align == "right":
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for text, bold in spec["runs"] or [("", False)]:
                    run = p.add_run(text)
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    run.font.size = Pt(max(10.0, self.base_size - 1))
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run.bold = bold or spec["header"]
        self.doc.add_paragraph().paragraph_format.space_after = Pt(0)

    # --- HTMLParser hooks ------------------------------------------------------------------

    def handle_starttag(self, tag, attrs_list):
        attrs = {k.lower(): (v or "") for k, v in attrs_list}
        if tag in ("script", "style"):
            self.skip_depth += 1
            return
        if tag in ("b", "strong"):
            self.bold += 1
        elif tag in ("i", "em"):
            self.italic += 1
        elif tag == "u":
            self.under += 1
        elif tag in ("s", "del", "strike"):
            self.strike += 1
        elif tag == "br":
            if self.cell is not None:
                self.cell["runs"].append(("\n", False))
            elif self.para is not None:
                self.para.add_run().add_break()
        elif tag == "hr":
            self._close_para()
            if "page-break" in attrs.get("class", "") or "always" in _style_dict(attrs).get("page-break-before", ""):
                self.pending_page_break = True
        elif tag in ("ul", "ol"):
            self._close_para()
            self.list_stack.append(tag)
        elif tag == "table":
            self._close_para()
            style = _style_dict(attrs)
            self.table_borders = not (attrs.get("border") == "0" or "none" in style.get("border", ""))
            self.table_rows = []
        elif tag == "tr" and self.table_rows is not None:
            self.row = []
            self.table_rows.append(self.row)
        elif tag in ("td", "th") and self.row is not None:
            style = _style_dict(attrs)
            self.cell = {"runs": [], "header": tag == "th",
                         "align": (style.get("text-align") or attrs.get("align") or "").lower()}
            self.row.append(self.cell)
        elif tag in self._BLOCKS:
            if self.table_rows is not None:
                return  # block tags inside a table cell: text flows into the cell
            self._close_para()
            style = _style_dict(attrs)
            if "always" in style.get("page-break-before", "") or "page-break" in attrs.get("class", ""):
                self.pending_page_break = True
            self._open_para(tag, attrs)

    def handle_endtag(self, tag):
        if tag in ("script", "style"):
            self.skip_depth = max(0, self.skip_depth - 1)
            return
        if tag in ("b", "strong"):
            self.bold = max(0, self.bold - 1)
        elif tag in ("i", "em"):
            self.italic = max(0, self.italic - 1)
        elif tag == "u":
            self.under = max(0, self.under - 1)
        elif tag in ("s", "del", "strike"):
            self.strike = max(0, self.strike - 1)
        elif tag in ("ul", "ol"):
            if self.list_stack:
                self.list_stack.pop()
            self._close_para()
        elif tag in ("td", "th"):
            self.cell = None
        elif tag == "tr":
            self.row = None
        elif tag == "table":
            self._flush_table()
        elif tag in self._BLOCKS:
            self._close_para()

    def handle_data(self, data):
        if self.skip_depth:
            return
        text = _EMOJI_RE.sub("", data)
        text = re.sub(r"\s+", " ", text)
        if not text:
            return
        if self.cell is None and self.para is None and not text.strip():
            return
        self._add_text(text)


def html_to_docx(html: str, *, font_size_pt: float = 12.0, line_spacing: float = 1.5) -> bytes:
    """Render the supported HTML subset into a .docx (A4, GOST-ish margins: left 3 cm,
    right 1.5 cm, top/bottom 2 cm; Times New Roman throughout)."""
    from docx import Document
    from docx.shared import Cm

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    # the default empty first paragraph would push content down — reuse it as a no-op
    builder = _DocBuilder(doc, base_size=float(font_size_pt), base_line=float(line_spacing))
    builder.feed(str(html or ""))
    builder.close()
    if builder.table_rows is not None:  # unclosed <table> at EOF
        builder._flush_table()
    # drop the implicit empty leading paragraph python-docx creates
    first = doc.paragraphs[0] if doc.paragraphs else None
    if first is not None and not first.runs and len(doc.paragraphs) > 1:
        first._element.getparent().remove(first._element)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
