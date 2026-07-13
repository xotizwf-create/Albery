"""docedit.py — targeted, structure-preserving edits to user-sent documents.

Why (2026-07-13): employees ask the agent to «внеси правку в файл, который я прислал».
Regenerating the file from the model's own retelling loses everything it didn't retell —
the Natalya incident: a rebuilt «Регламент встреч и контроля» xlsx dropped the Zoom links,
agendas and half the columns. This module edits the ORIGINAL bytes instead:

- xlsx/xlsm — openpyxl, full (non read_only) load: only string cells are touched, formulas
  are skipped, styles/merges/hyperlinks survive the round-trip. Files with charts/images get
  an explicit warning (openpyxl drops those parts on save).
- docx — python-docx: text is replaced inside runs, so formatting, tables and images survive.
- plain text (txt/csv/md/…) — str.replace with encoding detection (utf-8 / cp1251).

apply_edits() returns (new_bytes, per-edit match counts, warnings). It never mutates the
input; the caller stores the result as a NEW attachment.
"""
from __future__ import annotations

import io
import logging
import os
import zipfile

log = logging.getLogger(__name__)


class UnsupportedFormat(Exception):
    """Raised when the file type cannot be edited in place."""


_TEXT_EXTS = {"txt", "csv", "tsv", "md", "markdown", "log", "json", "yaml", "yml", "htm", "html"}
_XLSX_EDIT_MAX_BYTES = int(os.getenv("DOCEDIT_XLSX_MAX_BYTES", str(6 * 1024 * 1024)) or str(6 * 1024 * 1024))
_DOCX_EDIT_MAX_BYTES = int(os.getenv("DOCEDIT_DOCX_MAX_BYTES", str(20 * 1024 * 1024)) or str(20 * 1024 * 1024))


def apply_edits(data: bytes, file_name: str, edits: list[tuple[str, str]],
                sheet: str | None = None) -> tuple[bytes, list[int], list[str]]:
    """Apply exact-substring replacements to a document. Returns (bytes, counts, warnings);
    counts[i] = how many occurrences of edits[i].find were replaced."""
    ext = file_name.rsplit(".", 1)[-1].lower() if "." in (file_name or "") else ""
    if ext in ("xlsx", "xlsm"):
        return _edit_xlsx(data, edits, sheet)
    if ext == "docx":
        return _edit_docx(data, edits)
    if ext in _TEXT_EXTS:
        return _edit_text(data, edits)
    raise UnsupportedFormat(
        f"Файл .{ext or '?'} не поддерживает точечную правку (умею: xlsx/xlsm, docx, txt/csv/md/html). "
        "Для PDF или скана попроси пользователя прислать исходный docx/xlsx."
    )


# --- xlsx ------------------------------------------------------------------------------------

def _edit_xlsx(data: bytes, edits: list[tuple[str, str]], sheet: str | None) -> tuple[bytes, list[int], list[str]]:
    if len(data) > _XLSX_EDIT_MAX_BYTES:
        raise UnsupportedFormat(
            f"Excel-файл слишком большой для правки на месте ({len(data) // 1024} КБ, лимит "
            f"{_XLSX_EDIT_MAX_BYTES // 1024} КБ на этом сервере).")
    warnings: list[str] = []
    try:
        zf = zipfile.ZipFile(io.BytesIO(data))
        names = set(zf.namelist())
        has_visuals = any(n.startswith(("xl/media/", "xl/charts/")) for n in names)
        if not has_visuals:
            # a drawing part counts only when it actually anchors objects (empty ones are noise)
            has_visuals = any(
                n.startswith("xl/drawings/") and n.endswith(".xml")
                and b"CellAnchor" in zf.read(n)
                for n in names
            )
        if has_visuals:
            warnings.append("В исходном Excel есть изображения/диаграммы — при пересохранении они "
                            "могут потеряться. Предупреди пользователя и предложи проверить файл.")
    except Exception:  # noqa: BLE001
        pass
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(data))  # keep formulas (no data_only)
    counts = [0] * len(edits)
    sheet_norm = (sheet or "").strip().lower()
    matched_sheet = False
    for ws in wb.worksheets:
        if sheet_norm and str(ws.title).strip().lower() != sheet_norm:
            continue
        matched_sheet = True
        for row in ws.iter_rows():
            for cell in row:
                value = cell.value
                if not isinstance(value, str) or value.startswith("="):
                    continue  # only literal text cells; formulas stay untouched
                new_value = value
                for i, (find, replace) in enumerate(edits):
                    hit = new_value.count(find)
                    if hit:
                        counts[i] += hit
                        new_value = new_value.replace(find, replace)
                if new_value != value:
                    cell.value = new_value
    if sheet_norm and not matched_sheet:
        raise UnsupportedFormat(f"Листа «{sheet}» в файле нет. Доступные листы: "
                                + ", ".join(str(w.title) for w in wb.worksheets))
    out = io.BytesIO()
    wb.save(out)
    return out.getvalue(), counts, warnings


# --- docx ------------------------------------------------------------------------------------

def _replace_across_runs(runs, find: str, replace: str) -> int:
    """Replace every occurrence of ``find`` in the concatenated run text, editing runs in place
    so each occurrence keeps the formatting of the run it starts in."""
    count = 0
    search_from = 0
    while True:
        texts = [r.text or "" for r in runs]
        full = "".join(texts)
        start = full.find(find, search_from)
        if start < 0:
            return count
        end = start + len(find)
        pos = 0
        first = True
        for run, text in zip(runs, texts):
            begin, stop = pos, pos + len(text)
            pos = stop
            if stop <= start or begin >= end:
                continue
            cut_from, cut_to = max(start, begin) - begin, min(end, stop) - begin
            if first:
                run.text = text[:cut_from] + replace + text[cut_to:]
                first = False
            else:
                run.text = text[:cut_from] + text[cut_to:]
        count += 1
        search_from = start + len(replace)


def _iter_docx_paragraphs(doc):
    def walk_tables(tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from cell.paragraphs
                    yield from walk_tables(cell.tables)
    yield from doc.paragraphs
    yield from walk_tables(doc.tables)


def _edit_docx(data: bytes, edits: list[tuple[str, str]]) -> tuple[bytes, list[int], list[str]]:
    if len(data) > _DOCX_EDIT_MAX_BYTES:
        raise UnsupportedFormat(f"Word-файл слишком большой для правки ({len(data) // 1024} КБ).")
    from docx import Document
    doc = Document(io.BytesIO(data))
    counts = [0] * len(edits)
    for paragraph in _iter_docx_paragraphs(doc):
        for i, (find, replace) in enumerate(edits):
            if find in paragraph.text:
                counts[i] += _replace_across_runs(paragraph.runs, find, replace)
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue(), counts, []


# --- plain text ------------------------------------------------------------------------------

def _edit_text(data: bytes, edits: list[tuple[str, str]]) -> tuple[bytes, list[int], list[str]]:
    warnings: list[str] = []
    encoding = "utf-8"
    try:
        text = data.decode("utf-8")
    except UnicodeDecodeError:
        encoding = "cp1251"
        text = data.decode("cp1251", "replace")
        warnings.append("Файл был в кодировке Windows-1251 — сохранён в ней же.")
    counts = []
    for find, replace in edits:
        counts.append(text.count(find))
        text = text.replace(find, replace)
    return text.encode(encoding, "replace"), counts, warnings
