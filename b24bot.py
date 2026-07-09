"""Bitrix24 chat-bot: the agentic assistant living inside Bitrix24 chats.

Moved verbatim out of app.py (2026-07-02 refactor, step Ш2.1 — move-only, no logic
changes): the agent-access management API (/api/agent-access, the "Настройки Агента"
tab) and the whole chat-bot block (LLM tool-calling sandbox bot, the production
imbot local-application event flow, document/vision ingestion, file exports,
error-report flow, Telegram/Bitrix notification bridges, /bitrix/imbot/<secret>).

This module registers its routes directly on the Flask `app` at import time; app.py
imports it at the bottom, after every helper both sides need is defined.
"""
from __future__ import annotations

import base64
import hmac
import json
import logging
import os
import re
import signal
import subprocess
import threading
import time
import urllib.request
import uuid

from datetime import datetime
from datetime import timezone
from flask import jsonify
from flask import request
from typing import Any
import requests

from app import (  # noqa: E501 — single home for shared helpers until services/ split
    BitrixClient,
    MSK_TZ,
    Path,
    ZOOM_EXPORT_DIR,
    _refresh_bitrix_download_url,
    _zoom_export_public_url,
    absolute_bitrix_url,
    app,
    cleanup_zoom_exports,
    first_non_empty,
    flatten_request_payload,
    llm_api_key,
    llm_api_url,
    llm_auth_headers,
    llm_post_with_retry,
    msk_now,
    pg_connect,
    reportlab_font_paths,
    to_int,
)

# --- Agent access management API (/api/agent-access) --------------------------------------
# Backs the "Настройки Агента" tab in the SPA. Lets an admin grant/revoke the chat-bot tier
# (admin/ops/faq) per Bitrix user, live (the bot reads agent_access each turn via a short
# cache). Behind the site's admin session login + /api origin check (require_admin_auth).


@app.get("/api/agent-access")
def agent_access_list():
    rows = []
    try:
        with pg_connect() as conn:
            with conn.cursor() as cur:
                cur.execute(
                    "SELECT bitrix_user_id, tier, display_name, updated_at FROM agent_access "
                    "ORDER BY CASE tier WHEN 'admin' THEN 0 WHEN 'ops' THEN 1 ELSE 2 END, bitrix_user_id"
                )
                for r in cur.fetchall():
                    rows.append({
                        "bitrix_user_id": int(r["bitrix_user_id"]),
                        "tier": r["tier"],
                        "display_name": r["display_name"],
                        "updated_at": r["updated_at"].astimezone(MSK_TZ).strftime("%d.%m.%Y %H:%M") if r["updated_at"] else "",
                    })
    except Exception:  # noqa: BLE001
        logging.exception("agent_access list failed")
        return jsonify({"error": "Не удалось загрузить доступы."}), 500
    return jsonify({"rows": rows, "bootstrap_admin_ids": []})


_PORTAL_USERS_CACHE: dict[str, Any] = {"at": 0.0, "map": {}}


def _b24_portal_user_directory(force: bool = False) -> dict[int, dict[str, str]]:
    """Cached {bitrix_user_id: {name, email, position}} for active accounts on the live bot portal.
    Most portal accounts have empty NAME, so a human name is resolved by email from the synced org
    directory (users table). Cached ~10 min — used by the access UI and by requester-name lookups."""
    now = time.monotonic()
    if not force and _PORTAL_USERS_CACHE["map"] and (now - _PORTAL_USERS_CACHE["at"]) < 600:
        return _PORTAL_USERS_CACHE["map"]
    out: dict[int, dict[str, str]] = {}
    try:
        client = b24_testbot_client()
        data = _b24_testbot_call(client, "user.get", {"ACTIVE": True})
        email_dir: dict[str, tuple[str | None, str | None]] = {}
        with pg_connect() as conn:
            with conn.cursor() as cur:
                cur.execute("SELECT lower(email) AS email, full_name, work_position FROM users WHERE email <> ''")
                email_dir = {r["email"]: (r["full_name"], r["work_position"]) for r in cur.fetchall()}
        for u in data.get("result") or []:
            if not isinstance(u, dict):
                continue
            uid = to_int(u.get("ID"))
            if not uid:
                continue
            email = (u.get("EMAIL") or "").strip()
            portal_name = " ".join(p for p in (u.get("NAME"), u.get("LAST_NAME")) if p).strip()
            dir_name, dir_pos = email_dir.get(email.lower(), (None, None))
            out[uid] = {
                "name": portal_name or (dir_name or "").strip() or email or f"#{uid}",
                "email": email,
                "position": (u.get("WORK_POSITION") or dir_pos or "").strip(),
            }
    except Exception:  # noqa: BLE001
        logging.exception("portal user directory load failed")
        return _PORTAL_USERS_CACHE["map"]
    _PORTAL_USERS_CACHE.update(at=now, map=out)
    return out


def _b24_requester_name(from_user_id: Any) -> str:
    info = _b24_portal_user_directory().get(to_int(from_user_id))
    return (info or {}).get("name") or "Сотрудник"


@app.get("/api/agent-access/bitrix-users")
def agent_access_bitrix_users():
    """Active Bitrix users from the live bot portal — ids map 1:1 to access rows; names resolved
    by email from the synced org directory (see _b24_portal_user_directory)."""
    directory = _b24_portal_user_directory(force=True)
    if not directory:
        return jsonify({"error": "Не удалось получить пользователей Bitrix."}), 502
    users = [{"id": uid, **info} for uid, info in directory.items()]
    users.sort(key=lambda x: x["name"].lower())
    return jsonify({"users": users})


@app.post("/api/agent-access")
def agent_access_upsert():
    body = request.get_json(silent=True) or {}
    uid = to_int(body.get("bitrix_user_id"))
    tier = str(body.get("tier") or "").strip().lower()
    name = (str(body.get("display_name")).strip() or None) if body.get("display_name") else None
    if not uid or uid < 1:
        return jsonify({"error": "Укажите корректный Bitrix user ID."}), 400
    if tier not in ("none", "admin", "ops", "faq"):
        return jsonify({"error": "Уровень должен быть none, faq, ops или admin."}), 400
    try:
        _agent_access_set(uid, tier, name)
    except Exception:  # noqa: BLE001
        logging.exception("agent_access upsert failed")
        return jsonify({"error": "Не удалось сохранить."}), 500
    return jsonify({"ok": True})


@app.delete("/api/agent-access/<int:uid>")
def agent_access_delete(uid: int):
    try:
        _agent_access_remove(uid)
    except Exception:  # noqa: BLE001
        logging.exception("agent_access delete failed")
        return jsonify({"error": "Не удалось убрать."}), 500
    return jsonify({"ok": True})
# ---------------------------------------------------------------------------
# Bitrix24 chat-bot (sandbox): an agentic assistant living on a SEPARATE inbound
# webhook (B24_TESTBOT_WEBHOOK_BASE) so it never touches the production
# BITRIX_WEBHOOK_BASE integration. It reuses the existing OpenAI-compatible LLM
# layer (llm_* helpers) for tool-calling, and the existing BitrixClient for REST.
# ---------------------------------------------------------------------------

B24_TESTBOT_SYSTEM_PROMPT = (
    "Ты — ассистент внутри Битрикс24 (тестовый портал-песочница). Отвечай кратко, по-русски. "
    "Ты умеешь работать с задачами и сотрудниками портала через инструменты. "
    "Прежде чем создать или закрыть задачу, убедись, что знаешь ответственного: если назван человек по "
    "имени — сначала вызови list_users и найди его ID; при неоднозначности переспроси. "
    "Не выдумывай ID, задачи и имена — опирайся только на данные инструментов. "
    "После действия коротко подтверди результат (номер задачи, кому, дедлайн)."
)

B24_TESTBOT_TOOLS = [
    {
        "type": "function",
        "function": {
            "name": "list_users",
            "description": "Список активных сотрудников портала (id, имя, должность). Используй для резолва имени в ID.",
            "parameters": {"type": "object", "properties": {}},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "list_tasks",
            "description": "Список задач портала. Можно отфильтровать по ответственному.",
            "parameters": {
                "type": "object",
                "properties": {
                    "responsible_id": {"type": "integer", "description": "ID ответственного (опц.)"},
                },
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "create_task",
            "description": "Создать задачу в Битриксе.",
            "parameters": {
                "type": "object",
                "properties": {
                    "title": {"type": "string"},
                    "responsible_id": {"type": "integer"},
                    "description": {"type": "string"},
                    "deadline": {"type": "string", "description": "Дедлайн ISO8601, напр. 2026-06-20T18:00:00, опц."},
                },
                "required": ["title", "responsible_id"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "complete_task",
            "description": "Завершить задачу по её ID.",
            "parameters": {
                "type": "object",
                "properties": {"task_id": {"type": "integer"}},
                "required": ["task_id"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "send_message",
            "description": "Отправить личное сообщение сотруднику портала по его ID.",
            "parameters": {
                "type": "object",
                "properties": {
                    "user_id": {"type": "integer"},
                    "text": {"type": "string"},
                },
                "required": ["user_id", "text"],
            },
        },
    },
]


def b24_testbot_secret_valid(secret: str) -> bool:
    expected = os.getenv("B24_TESTBOT_SECRET", "").strip()
    if not expected:
        return False
    return hmac.compare_digest(secret, expected)


def b24_testbot_client() -> BitrixClient:
    base = os.getenv("B24_TESTBOT_WEBHOOK_BASE", "").strip()
    if not base:
        raise ValueError("B24_TESTBOT_WEBHOOK_BASE не задан в .env")
    return BitrixClient(base)


def _imbot_event_param(payload: dict[str, Any], name: str) -> Any:
    """Read a Bitrix imbot form field flattened as data[PARAMS][<name>]."""
    for key in (f"data[PARAMS][{name}]", f"data[params][{name}]"):
        if key in payload:
            value = payload[key]
            return value[0] if isinstance(value, list) and value else value
    return None


def _imbot_scan(payload: dict[str, Any], field: str) -> str:
    """Robustly pull a field whose flattened key ENDS with `field`, regardless of layout
    (command events put COMMAND/DIALOG_ID at data[..] or data[PARAMS][..], unlike messages).
    `field` is matched on the alpha/underscore-normalized key tail, so 'COMMAND' won't grab
    'COMMAND_ID'/'COMMAND_PARAMS'."""
    suf = field.upper()
    for key, value in payload.items():
        norm = re.sub(r"[^A-Z_]", "", str(key).upper())
        if norm.endswith(suf):
            val = value[0] if isinstance(value, list) and value else value
            if val not in (None, ""):
                return str(val)
    return ""


def _b24_testbot_call(client: BitrixClient, method: str, payload: dict[str, Any]) -> dict[str, Any]:
    # Plain webhook path first (standard /rest/<id>/<token>/), API path as fallback.
    return client.call_with_fallback(method, payload, prefer_api=False)


def _b24_testbot_exec_tool(client: BitrixClient, name: str, args: dict[str, Any]) -> str:
    try:
        if name == "list_users":
            data = _b24_testbot_call(client, "user.get", {"ACTIVE": True})
            rows = data.get("result") or []
            users = [
                {
                    "id": to_int(u.get("ID")),
                    "name": " ".join(p for p in (u.get("NAME"), u.get("LAST_NAME")) if p).strip() or u.get("EMAIL"),
                    "position": u.get("WORK_POSITION") or "",
                }
                for u in rows
                if isinstance(u, dict)
            ]
            return json.dumps({"users": users}, ensure_ascii=False)

        if name == "list_tasks":
            filt: dict[str, Any] = {}
            resp_id = to_int(args.get("responsible_id"))
            if resp_id:
                filt["RESPONSIBLE_ID"] = resp_id
            data = _b24_testbot_call(
                client,
                "tasks.task.list",
                {"filter": filt, "select": ["ID", "TITLE", "STATUS", "RESPONSIBLE_ID", "DEADLINE"]},
            )
            result = data.get("result") or {}
            tasks = result.get("tasks", result) if isinstance(result, dict) else result
            return json.dumps({"tasks": tasks}, ensure_ascii=False)[:6000]

        if name == "create_task":
            title = str(args.get("title") or "").strip()
            resp_id = to_int(args.get("responsible_id"))
            if not title or not resp_id:
                return json.dumps({"error": "title и responsible_id обязательны"}, ensure_ascii=False)
            fields: dict[str, Any] = {"TITLE": title, "RESPONSIBLE_ID": resp_id}
            if args.get("description"):
                fields["DESCRIPTION"] = str(args["description"])
            if args.get("deadline"):
                fields["DEADLINE"] = str(args["deadline"])
            fields["SE_PARAMETER"] = [{"CODE": 3, "VALUE": "Y"}]
            data = _b24_testbot_call(client, "tasks.task.add", {"fields": fields})
            task = (data.get("result") or {}).get("task") or {}
            return json.dumps({"created_task_id": to_int(task.get("id")), "title": task.get("title")}, ensure_ascii=False)

        if name == "complete_task":
            task_id = to_int(args.get("task_id"))
            if not task_id:
                return json.dumps({"error": "task_id обязателен"}, ensure_ascii=False)
            _b24_testbot_call(client, "tasks.task.complete", {"taskId": task_id})
            return json.dumps({"completed_task_id": task_id}, ensure_ascii=False)

        if name == "send_message":
            user_id = to_int(args.get("user_id"))
            text = str(args.get("text") or "").strip()
            if not user_id or not text:
                return json.dumps({"error": "user_id и text обязательны"}, ensure_ascii=False)
            _b24_testbot_call(client, "im.message.add", {"DIALOG_ID": str(user_id), "MESSAGE": text})
            return json.dumps({"sent_to": user_id}, ensure_ascii=False)

        return json.dumps({"error": f"неизвестный инструмент {name}"}, ensure_ascii=False)
    except Exception as exc:  # noqa: BLE001
        return json.dumps({"error": str(exc)[:300]}, ensure_ascii=False)


def b24_testbot_run_agent(user_text: str) -> str:
    api_key = llm_api_key()
    if not api_key:
        return "LLM не настроен (нет OPENAI_API_KEY)."
    model = os.getenv("B24_TESTBOT_MODEL", "").strip() or os.getenv("OPENAI_MODEL", "gpt-4.1-mini").strip()
    client = b24_testbot_client()
    messages: list[dict[str, Any]] = [
        {"role": "system", "content": B24_TESTBOT_SYSTEM_PROMPT},
        {"role": "user", "content": user_text},
    ]
    max_steps = max(1, int(os.getenv("B24_TESTBOT_MAX_STEPS", "6")))
    for _ in range(max_steps):
        response = llm_post_with_retry(
            llm_api_url("/chat/completions"),
            llm_auth_headers(api_key),
            {
                "model": model,
                "messages": messages,
                "tools": B24_TESTBOT_TOOLS,
                "tool_choice": "auto",
                "temperature": 0.2,
            },
            timeout=max(60, int(os.getenv("B24_TESTBOT_TIMEOUT_SECONDS", "90"))),
        )
        if not response.ok:
            return f"Ошибка LLM: HTTP {response.status_code} {response.text[:200]}"
        message = (response.json().get("choices") or [{}])[0].get("message") or {}
        messages.append(message)
        tool_calls = message.get("tool_calls") or []
        if not tool_calls:
            return (message.get("content") or "").strip() or "(пустой ответ)"
        for call in tool_calls:
            fn = call.get("function") or {}
            try:
                args = json.loads(fn.get("arguments") or "{}")
            except json.JSONDecodeError:
                args = {}
            result = _b24_testbot_exec_tool(client, fn.get("name") or "", args)
            messages.append({"role": "tool", "tool_call_id": call.get("id"), "content": result})
    return "Не уложился в лимит шагов — переформулируй запрос."


def _b24_testbot_reply(dialog_id: str, text: str) -> None:
    bot_id = os.getenv("B24_TESTBOT_BOT_ID", "").strip()
    if not bot_id or not dialog_id:
        return
    try:
        b24_testbot_client().call_with_fallback(
            "imbot.message.add",
            {"BOT_ID": bot_id, "DIALOG_ID": dialog_id, "MESSAGE": text},
            prefer_api=False,
        )
    except Exception:  # noqa: BLE001
        logging.exception("b24 testbot: failed to send reply")


def _b24_testbot_process(dialog_id: str, user_text: str) -> None:
    try:
        answer = b24_testbot_run_agent(user_text)
    except Exception as exc:  # noqa: BLE001
        logging.exception("b24 testbot: agent failed")
        answer = f"Ошибка: {str(exc)[:200]}"
    _b24_testbot_reply(dialog_id, answer)


# --- Local-application (app-context) flow ---------------------------------
# Bitrix forbids registering a chat-bot from an inbound webhook ("Client ID not
# specified"); a local application is required. The app delivers events with an
# `auth` block (access_token + application_token + client_endpoint); we register
# the bot on ONAPPINSTALL and reply with the per-event access token. Persistent
# state (application_token + bot_id + endpoint) lives in a small JSON file so the
# handler needs no restart after install.

B24_APP_HANDLER_URL = "https://mcp.m4s.ru/bitrix/imbot/app"


def _b24_state_path() -> str:
    return os.getenv("B24_TESTBOT_STATE", "").strip() or "/var/www/albery/.b24_testbot_state.json"


def _b24_load_state() -> dict[str, Any]:
    try:
        with open(_b24_state_path(), encoding="utf-8") as fh:
            return json.load(fh) or {}
    except (OSError, json.JSONDecodeError):
        return {}


def _b24_save_state(state: dict[str, Any]) -> None:
    try:
        with open(_b24_state_path(), "w", encoding="utf-8") as fh:
            json.dump(state, fh)
    except OSError:
        logging.exception("b24 testbot: failed to persist state")


def _imbot_auth(payload: dict[str, Any], name: str) -> str:
    for key in (f"auth[{name}]", f"AUTH[{name}]"):
        if key in payload:
            value = payload[key]
            return str((value[0] if isinstance(value, list) and value else value) or "")
    return ""


def _b24_app_call(client_endpoint: str, access_token: str, method: str, payload: dict[str, Any] | None = None) -> dict[str, Any]:
    base = (client_endpoint or "").strip().rstrip("/")
    if not base or not access_token:
        raise ValueError("client_endpoint/access_token пусты")
    resp = requests.post(f"{base}/{method}.json?auth={access_token}", json=payload or {}, timeout=60)
    if not resp.ok:
        raise RuntimeError(f"{method}: HTTP {resp.status_code} {resp.text[:300]}")
    data = resp.json()
    if isinstance(data, dict) and "error" in data:
        raise RuntimeError(f"{method}: {data.get('error')} ({data.get('error_description')})")
    return data if isinstance(data, dict) else {"result": data}


def _b24_capture_tokens(payload: dict[str, Any], state: dict[str, Any]) -> None:
    """Persist the app OAuth tokens carried in every event's `auth` block, so cron jobs (which have
    no live event) can later post AS THE BOT via a refresh_token grant. Best-effort; saves state.
    state['app_tokens'] = {access_token, refresh_token, expires (unix), client_endpoint}."""
    access_token = _imbot_auth(payload, "access_token")
    refresh_token = _imbot_auth(payload, "refresh_token")
    client_endpoint = _imbot_auth(payload, "client_endpoint")
    if not access_token and not refresh_token:
        return
    try:
        exp = int(_imbot_auth(payload, "expires") or 0)
    except ValueError:
        exp = 0
    if not exp:
        try:
            exp = int(time.time()) + int(_imbot_auth(payload, "expires_in") or 3600)
        except ValueError:
            exp = int(time.time()) + 3600
    tok = dict(state.get("app_tokens") or {})
    if access_token:
        tok["access_token"], tok["expires"] = access_token, exp
    if refresh_token:
        tok["refresh_token"] = refresh_token
    if client_endpoint:
        tok["client_endpoint"] = client_endpoint
    state["app_tokens"] = tok
    _b24_save_state(state)


def _b24_app_access_token() -> tuple[str, str]:
    """Return (client_endpoint, access_token) usable for app REST calls OUTSIDE a live event.
    Reuses the cached token while valid, otherwise refreshes via oauth.bitrix.info using the stored
    refresh_token and persists the rotated pair — so the weekly cron keeps the chain alive forever.
    Returns ('', '') if not bootstrapped yet (the bot must receive at least one event first)."""
    state = _b24_load_state()
    tok = state.get("app_tokens") or {}
    endpoint = (tok.get("client_endpoint") or state.get("client_endpoint") or "").strip()
    access_token = (tok.get("access_token") or "").strip()
    try:
        expires = int(tok.get("expires") or 0)
    except (ValueError, TypeError):
        expires = 0
    if access_token and expires - 120 > int(time.time()):
        return endpoint, access_token
    refresh_token = (tok.get("refresh_token") or "").strip()
    client_id = os.getenv("B24_TESTBOT_CLIENT_ID", "").strip()
    client_secret = os.getenv("B24_TESTBOT_CLIENT_SECRET", "").strip()
    if not (refresh_token and client_id and client_secret):
        return endpoint, access_token
    resp = requests.get(
        "https://oauth.bitrix.info/oauth/token/",
        params={"grant_type": "refresh_token", "client_id": client_id,
                "client_secret": client_secret, "refresh_token": refresh_token},
        timeout=30,
    )
    data = resp.json() if resp.content else {}
    new_access = (data.get("access_token") or "").strip() if isinstance(data, dict) else ""
    if not new_access:
        # A sibling process (the other weekly digest) may have just rotated the token; reuse it.
        fresh = (_b24_load_state().get("app_tokens") or {})
        try:
            fresh_exp = int(fresh.get("expires") or 0)
        except (ValueError, TypeError):
            fresh_exp = 0
        if fresh.get("access_token") and fresh_exp - 120 > int(time.time()):
            return (fresh.get("client_endpoint") or endpoint).strip(), fresh["access_token"].strip()
        raise RuntimeError(f"b24 token refresh failed: {str(data)[:300]}")
    state = _b24_load_state()  # re-read to merge a possibly newer event-stored pair
    tok = dict(state.get("app_tokens") or {})
    tok["access_token"] = new_access
    tok["refresh_token"] = (data.get("refresh_token") or refresh_token).strip()
    try:
        tok["expires"] = int(data.get("expires") or (int(time.time()) + int(data.get("expires_in") or 3600)))
    except (ValueError, TypeError):
        tok["expires"] = int(time.time()) + 3600
    new_endpoint = (data.get("client_endpoint") or data.get("server_endpoint") or endpoint).strip()
    if new_endpoint:
        tok["client_endpoint"] = new_endpoint
    state["app_tokens"] = tok
    _b24_save_state(state)
    return tok.get("client_endpoint") or endpoint, new_access


def b24_app_method_call(method: str, payload: dict[str, Any] | None = None) -> dict[str, Any]:
    """Bitrix REST call with the local-app OAuth token (auto-refreshing). The app token carries the
    `crm` scope the incoming webhooks lack, so CRM tools (funnels/deals) must go through here."""
    client_endpoint, access_token = _b24_app_access_token()
    if not access_token:
        raise RuntimeError(
            "Bitrix app OAuth token is unavailable (the bot has not stored app_tokens yet — "
            "it needs to receive at least one portal event first).")
    return _b24_app_call(client_endpoint, access_token, method, payload)


def _b24_app_register_bot(client_endpoint: str, access_token: str) -> Any:
    payload = {
        "CODE": "hermes_agent",
        "TYPE": "B",
        "OPENLINE": "N",
        "EVENT_MESSAGE_ADD": B24_APP_HANDLER_URL,
        "EVENT_WELCOME_MESSAGE": B24_APP_HANDLER_URL,
        "EVENT_BOT_DELETE": B24_APP_HANDLER_URL,
        "PROPERTIES": {"NAME": "Гермес-ассистент (тест)", "COLOR": "AQUA", "WORK_POSITION": "ИИ-ассистент"},
    }
    return _b24_app_call(client_endpoint, access_token, "imbot.register", payload).get("result")


def _b24_disclaimer() -> str:
    """The grey footnote shown under every bot message (configurable / disable via B24_DISCLAIMER="")."""
    return os.getenv(
        "B24_DISCLAIMER",
        "Ответы Албери AI могут быть неточными. Проверяйте важную информацию.",
    ).strip()


def _b24_keyboard() -> list[dict[str, Any]]:
    """Buttons under the bot's latest reply: '🆕 Новая сессия' (resets via `new`), '⚠️ Сообщить об
    ошибке' (error-report flow via `report_error`) and '❓ Как пользоваться' (re-shows the onboarding
    guide via `help`). Only the most recent bot message carries these — _b24_app_reply strips the
    keyboard off the previous one (see _b24_strip_keyboard). Typed '/new' / 'новая сессия' still
    resets via the _b24_is_reset_command keyword detector."""
    return [
        {
            "TEXT": "🆕 Новая сессия",
            "COMMAND": "new",
            "COMMAND_PARAMS": "/new",
            "DISPLAY": "LINE",
            "BG_COLOR": "#29619b",
            "TEXT_COLOR": "#FFFFFF",
            "BLOCK": "N",
        },
        {
            "TEXT": "⚠️ Сообщить об ошибке",
            "COMMAND": "report_error",
            "COMMAND_PARAMS": "/report_error",
            "DISPLAY": "LINE",
            "BG_COLOR": "#9b3029",
            "TEXT_COLOR": "#FFFFFF",
            "BLOCK": "N",
        },
        {
            "TEXT": "❓ Как пользоваться",
            "COMMAND": "help",
            "COMMAND_PARAMS": "/help",
            "DISPLAY": "LINE",
            "BG_COLOR": "#3a7a3a",
            "TEXT_COLOR": "#FFFFFF",
            "BLOCK": "N",
        },
    ]


def _b24_onb_next_keyboard() -> list[dict[str, Any]]:
    """A single 'Далее ▶️' button to advance the onboarding (command `onb_next`)."""
    return [{
        "TEXT": "Далее ▶️",
        "COMMAND": "onb_next",
        "COMMAND_PARAMS": "/onb_next",
        "DISPLAY": "LINE",
        "BG_COLOR": "#29619b",
        "TEXT_COLOR": "#FFFFFF",
        "BLOCK": "N",
    }]


def _b24_welcome_keyboard() -> list[dict[str, Any]]:
    """First-open greeting: a prominent '🚀 Пройти обучение' (command `help`) — the user can also
    just start asking."""
    return [{
        "TEXT": "🚀 Пройти обучение",
        "COMMAND": "help",
        "COMMAND_PARAMS": "/help",
        "DISPLAY": "LINE",
        "BG_COLOR": "#5440F6",
        "TEXT_COLOR": "#FFFFFF",
        "BLOCK": "N",
    }]


# --- Onboarding: a short 3-step "how to use me" walkthrough driven by buttons -------------
_B24_ONB_LAST_STEP = 3


def _b24_onboarding_text(step: int, tier: str) -> str:  # noqa: ARG001 — tier kept for caller compat
    """Content for an onboarding step. Step 1 lists all capabilities with a note that the actual
    level depends on the user's access (they can ask «Что ты умеешь?» to see their own set)."""
    if step == 1:
        return "\n".join([
            "[b]Шаг 1 из 3 — что я умею[/b]",
            "Я многое умею, вот основное:",
            "- 🔎 Глубокий поиск информации в интернете — могу найти информацию из глубоких архивов",
            "- 📚 Поиск по базе знаний и базе данных компании (регламенты, процессы, оргструктура, документы)",
            "- 📄 Работа с документами: Word, Excel, PDF, Markdown — пришлите файл прямо в чат, я прочитаю и разберу его; могу вернуть ответ файлом (например, PDF)",
            "- 📊 Работа с Google-таблицами — могу сам поставить нужные формулы, сделать красивое оформление или полноценную автоматизацию таблицы",
            "- 🎧 Разбор Zoom-созвонов: итоги, задачи, участники",
            "- ✅ Задачи в Bitrix24: поиск, постановка, закрытие",
            "- 📈 Отчёты по компании, задачам и чатам",
            "- 💬 Сообщения сотрудникам",
            "",
            "━━━━━━━━━━━━━━",
            "⭐ [b]И ГЛАВНОЕ — СОЗДАЮ ПРИЛОЖЕНИЯ[/b]",
            "Для сложных автоматизаций обычной таблицы порой мало — нужно отдельное приложение. "
            "Я соберу его сам на базе ваших Google-таблиц: от вас только идея и логика словами, "
            "а я напишу, опубликую и пришлю готовую рабочую ссылку (доступ по ссылке для всех, "
            "настраивать ничего не нужно). Примеры: калькуляторы, формы, дашборды, мини-сервисы. 🚀",
            "━━━━━━━━━━━━━━",
            "",
            "Важно: уровень моих возможностей зависит от вашего доступа. Чтобы узнать, что доступно "
            "именно вам — задайте вопрос «Что ты умеешь?».",
        ])
    if step == 2:
        return "\n".join([
            "[b]Шаг 2 из 3 — как формулировать запрос[/b]",
            "- Одна задача — одно сообщение.",
            "- Указывайте конкретику: кто / что / период / документ / в каком формате хотите получить результат.",
            "- Можно задать формат ответа (списком, кратко, по пунктам, таблицей, PDF).",
            "",
            "[b]Примеры:[/b]",
            "Вместо «Сделай отчёт по задачам» →",
            "«Сделай отчёт по задачам за период 15.06–19.06, оформи таблицей: столбцы — Ответственный, "
            "кол-во задач за период, кол-во просроченных задач, вывод по ответственному. Под таблицей — "
            "общий управленческий вывод. Пришли мне ответ PDF-файлом».",
            "",
            "Вместо «Кто согласует отпуск для менеджера маркетплейса» →",
            "«Кто согласует отпуск для менеджера маркетплейса исходя из регламентов компании — дай "
            "ссылку на документ и точную цитату из него».",
        ])
    return "\n".join([
        "[b]Шаг 3 из 3 — попробуйте сами[/b]",
        "Примеры запросов:",
        "- «Что в регламенте про пятничную планёрку?»",
        "- «Можешь ли ты работать в Google-таблицах?»",
        "- «Сделай краткую сводку по Zoom-созвонам за вчера»",
        "",
        "Напишите свой первый запрос — я помогу! 🚀",
    ])


def _b24_bkey(bot_id: Any, dialog_id: Any) -> str:
    """Per-bot key for ephemeral per-dialog UI state (awaiting-error / onboarding step / last
    keyboard). A Bitrix private chat is keyed by the USER id, so the SAME dialog_id is reused by
    every bot that user talks to — keying this state by dialog_id alone leaked it ACROSS bots (e.g.
    'Report error' pressed on one bot, then a message to ANOTHER bot got captured as the report).
    bot_id is available at every call site and uniquely identifies the bot."""
    return f"{bot_id}:{dialog_id}"


def _b24_send_onboarding(client_endpoint: str, access_token: str, bot_id: Any, dialog_id: str,
                         step: int, tier: str, message_id: Any = "") -> None:
    """Send one onboarding step. Steps 1-2 carry a 'Далее ▶️' button; the final step restores the
    normal keyboard. The current step is remembered per (bot, dialog) so 'Далее' knows what's next."""
    step = max(1, min(step, _B24_ONB_LAST_STEP))
    if message_id:
        _b24_app_react(client_endpoint, access_token, message_id, "eyes", add=False)
        _b24_app_react(client_endpoint, access_token, message_id, "like", add=True)
    keyboard = _b24_onb_next_keyboard() if step < _B24_ONB_LAST_STEP else _b24_keyboard()
    _b24_app_reply(client_endpoint, access_token, bot_id, dialog_id,
                   _b24_onboarding_text(step, tier), keyboard=keyboard)
    st = _b24_load_state()
    onb = st.get("onboarding") or {}
    if step < _B24_ONB_LAST_STEP:
        onb[_b24_bkey(bot_id, dialog_id)] = step
    else:
        onb.pop(_b24_bkey(bot_id, dialog_id), None)
    st["onboarding"] = onb
    _b24_save_state(st)
    _b24_ensure_command_registered(client_endpoint, access_token, bot_id)


def _b24_onboarding_step(bot_id: Any, dialog_id: str) -> int:
    return int((_b24_load_state().get("onboarding") or {}).get(_b24_bkey(bot_id, dialog_id), 0))


# --- "keyboard only under the last message" bookkeeping ----------------------
# Bitrix attaches a keyboard per-message, so replying with one each turn stacks a fresh pair of
# buttons under every bot message (visual spam). We remember the id+text of the message that
# currently holds the keyboard (per dialog, in the small JSON state) and, when sending a new
# keyboard'd reply, edit the previous one to drop its keyboard — leaving exactly one live set.

def _b24_get_last_kb(bot_id: Any, dialog_id: str) -> dict[str, Any] | None:
    return (_b24_load_state().get("last_kb") or {}).get(_b24_bkey(bot_id, dialog_id))


def _b24_set_last_kb(bot_id: Any, dialog_id: str, message_id: Any, text: str) -> None:
    st = _b24_load_state()
    last_kb = st.get("last_kb") or {}
    last_kb[_b24_bkey(bot_id, dialog_id)] = {"id": message_id, "text": text}
    if len(last_kb) > 300:  # keep the map bounded on a long-lived box
        last_kb = dict(list(last_kb.items())[-300:])
    st["last_kb"] = last_kb
    _b24_save_state(st)


def _b24_set_awaiting_error(bot_id: Any, dialog_id: str) -> None:
    st = _b24_load_state()
    awaiting = st.get("awaiting_error") or {}
    awaiting[_b24_bkey(bot_id, dialog_id)] = int(time.time())
    st["awaiting_error"] = awaiting
    _b24_save_state(st)


def _b24_pop_awaiting_error(bot_id: Any, dialog_id: str, ttl_seconds: int = 3600) -> bool:
    """True (and clears the flag) if THIS bot's dialog is awaiting an error description set within
    ttl. Scoped per bot so 'Report error' on one bot never captures a message sent to another."""
    st = _b24_load_state()
    awaiting = st.get("awaiting_error") or {}
    ts = awaiting.pop(_b24_bkey(bot_id, dialog_id), None)
    if ts is None:
        return False
    st["awaiting_error"] = awaiting
    _b24_save_state(st)
    return (int(time.time()) - int(ts)) < ttl_seconds


def _b24_strip_keyboard(client_endpoint: str, access_token: str, bot_id: Any,
                        message_id: Any, text: str) -> None:
    """Best-effort: edit a previous bot message to drop its keyboard (KEYBOARD='N' clears it).
    imbot.message.update needs the message body, so we re-send the stored text (and re-attach
    the same grey disclaimer plashka). Never raises — a lingering button is only cosmetic."""
    if not (client_endpoint and access_token and bot_id and message_id):
        return
    params: dict[str, Any] = {
        "BOT_ID": bot_id, "MESSAGE_ID": message_id,
        "MESSAGE": text or " ", "KEYBOARD": "N",
    }
    disclaimer = _b24_disclaimer()
    if disclaimer:
        with_attach = dict(params)
        with_attach["ATTACH"] = [
            {"DELIMITER": {"SIZE": 200, "COLOR": "#C8C8C8"}},
            {"MESSAGE": disclaimer},
        ]
        try:
            _b24_app_call(client_endpoint, access_token, "imbot.message.update", with_attach)
            return
        except Exception as exc:  # noqa: BLE001
            logging.debug("b24 testbot: strip-keyboard ATTACH update failed (%s) — retry plain", exc)
    try:
        _b24_app_call(client_endpoint, access_token, "imbot.message.update", params)
    except Exception:  # noqa: BLE001
        logging.debug("b24 testbot: strip-keyboard update failed", exc_info=True)


def _b24_app_reply(client_endpoint: str, access_token: str, bot_id: Any, dialog_id: str,
                   text: str, keyboard: list[dict[str, Any]] | None = None) -> Any:
    """Send a bot message; returns the new message id (or None). When a keyboard is attached, the
    buttons are moved off the previous bot message so only the latest reply shows them.

    The disclaimer footnote under every bot message is a separated grey ATTACH block (closest a
    third-party imbot gets to Bitrix CoPilot's native footer). If the portal rejects the ATTACH
    format, we retry with an INLINE italic footnote so the answer is always delivered."""
    if not (client_endpoint and access_token and bot_id and dialog_id):
        return None
    disclaimer = _b24_disclaimer()
    base: dict[str, Any] = {"BOT_ID": bot_id, "DIALOG_ID": dialog_id, "MESSAGE": text}
    if keyboard:
        base["KEYBOARD"] = keyboard

    new_message_id: Any = None
    sent_text = text
    if disclaimer:
        attach_params = dict(base)
        # Bitrix REST attach = flat array of block objects (NOT the JS-UI COLOR/BLOCKS shape).
        attach_params["ATTACH"] = [
            {"DELIMITER": {"SIZE": 200, "COLOR": "#C8C8C8"}},
            {"MESSAGE": disclaimer},
        ]
        try:
            res = _b24_app_call(client_endpoint, access_token, "imbot.message.add", attach_params)
            new_message_id = res.get("result")
        except Exception as exc:  # noqa: BLE001
            logging.warning("b24 testbot: ATTACH disclaimer rejected (%s) — using inline footnote", exc)
            sent_text = f"{text}\n\n[i]{disclaimer}[/i]"
            base["MESSAGE"] = sent_text
    if new_message_id is None:
        try:
            res = _b24_app_call(client_endpoint, access_token, "imbot.message.add", base)
            new_message_id = res.get("result")
        except Exception:  # noqa: BLE001
            logging.exception("b24 testbot: app reply failed")
            return None

    # Keep the buttons only under the latest message: drop the previous keyboard, remember this one.
    # Scoped per bot so stripping never touches ANOTHER bot's message in the same shared dialog.
    if keyboard and new_message_id:
        prev = _b24_get_last_kb(bot_id, dialog_id)
        if prev and str(prev.get("id")) != str(new_message_id):
            _b24_strip_keyboard(client_endpoint, access_token, bot_id, prev.get("id"), prev.get("text") or "")
        _b24_set_last_kb(bot_id, dialog_id, new_message_id, sent_text)
    return new_message_id


_B24_COMMANDS = [
    {"COMMAND": "new", "ru": "Новая сессия", "en": "New session"},
    {"COMMAND": "report_error", "ru": "Сообщить об ошибке", "en": "Report an error"},
    {"COMMAND": "help", "ru": "Как пользоваться", "en": "How to use"},
    {"COMMAND": "onb_next", "ru": "Далее", "en": "Next"},
]


def _b24_ensure_command_registered(client_endpoint: str, access_token: str, bot_id: Any) -> None:
    """Best-effort: register the bot commands (`/new`, `/report_error`) once so their keyboard
    buttons resolve and they show in the chat '/' menu. Runs in a background thread so it NEVER
    blocks the event response; guarded by a state flag so it registers only once per command set.
    Must use the app access_token (the webhook lacks the client id → ACCESS_DENIED). Already-
    registered commands return an API error (not a transport one) — treated as success."""
    if not (client_endpoint and access_token and bot_id):
        return
    # Commands are registered PER BOT (imbot.command.register takes BOT_ID), so the guard must be
    # keyed by bot_id — a single global flag let the main bot register and then permanently blocked
    # every subagent (e.g. the lawyer bot 70), leaving their keyboard buttons dead.
    reg_key = f"cmds_registered_v3_{bot_id}"
    if _b24_load_state().get(reg_key):
        return

    def _do() -> None:
        transport_failed = False
        for spec in _B24_COMMANDS:
            try:
                _b24_app_call(client_endpoint, access_token, "imbot.command.register", {
                    "BOT_ID": bot_id,
                    "COMMAND": spec["COMMAND"],
                    "COMMON": "N",
                    "HIDDEN": "N",
                    "EXTRANET_SUPPORT": "N",
                    "LANG": [
                        {"LANGUAGE_ID": "ru", "TITLE": spec["ru"], "PARAMS": ""},
                        {"LANGUAGE_ID": "en", "TITLE": spec["en"], "PARAMS": ""},
                    ],
                    "EVENT_COMMAND_ADD": B24_APP_HANDLER_URL,
                })
                logging.info("b24 testbot: /%s command registered", spec["COMMAND"])
            except requests.RequestException as exc:
                transport_failed = True
                logging.warning("b24 testbot: command register transport error for %s: %s", spec["COMMAND"], exc)
            except Exception as exc:  # noqa: BLE001 — API error (e.g. already registered) is fine
                logging.debug("b24 testbot: command register note for %s: %s", spec["COMMAND"], exc)
        if not transport_failed:  # only retry on a real connectivity failure
            st = _b24_load_state()
            st[reg_key] = True
            _b24_save_state(st)

    threading.Thread(target=_do, daemon=True).start()


def _b24_subagent_bot_ids() -> list[str]:
    """bitrix_bot_id of every registered subagent, so we can register their keyboard commands too."""
    try:
        with pg_connect() as conn:
            with conn.cursor() as cur:
                cur.execute("SELECT bitrix_bot_id FROM agents WHERE bitrix_bot_id IS NOT NULL")
                return [str(r["bitrix_bot_id"]) for r in cur.fetchall() if r["bitrix_bot_id"]]
    except Exception:  # noqa: BLE001
        logging.exception("b24 testbot: subagent bot ids query failed")
        return []


_b24_bootstrap_lock = threading.Lock()
_b24_bootstrapped = False


def _b24_bootstrap_all_commands(client_endpoint: str, access_token: str, main_bot_id: Any) -> None:
    """Proactively register keyboard commands for EVERY bot (main + all subagents), so their
    buttons fire ONIMCOMMANDADD without that bot first needing a text message. Without this a
    subagent's buttons just hang (a click can't itself trigger registration → chicken-and-egg).
    Runs ONCE per process, in the background, off the request path; per-bot flags keep it idempotent.
    Uses the live event's app token, which owns all bots this application registered."""
    global _b24_bootstrapped
    if not (client_endpoint and access_token):
        return
    with _b24_bootstrap_lock:
        if _b24_bootstrapped:
            return
        _b24_bootstrapped = True

    def _do() -> None:
        bot_ids: list[str] = []
        if main_bot_id:
            bot_ids.append(str(main_bot_id))
        bot_ids.extend(b for b in _b24_subagent_bot_ids() if b not in bot_ids)
        for bid in bot_ids:
            _b24_ensure_command_registered(client_endpoint, access_token, bid)
        logging.info("b24 testbot: bootstrapped keyboard commands for bots %s", bot_ids)

    threading.Thread(target=_do, daemon=True).start()


def _b24_inflight_register(bot_id: Any, dialog_id: str, agent_slug: str | None, from_user_id: Any,
                           message_id: Any, status_message_id: Any, user_preview: str) -> str | None:
    """Record that a brain turn is starting. Returns the row id (or None on failure — the turn
    still runs; this is only a safety net). Cleared by _b24_inflight_clear when the turn finishes.
    A row that survives = a turn killed mid-flight (restart/OOM/crash), recovered at next boot."""
    turn_id = str(uuid.uuid4())
    try:
        with pg_connect() as conn:
            with conn.transaction():
                with conn.cursor() as cur:
                    cur.execute(
                        "INSERT INTO bitrix_inflight_turns"
                        " (id, bot_id, dialog_id, agent_slug, from_user_id, message_id, status_message_id, user_preview)"
                        " VALUES (%s, %s, %s, %s, %s, %s, %s, %s)",
                        (turn_id, str(bot_id or ""), str(dialog_id), agent_slug, str(from_user_id or ""),
                         str(message_id or ""), str(status_message_id or ""), (user_preview or "")[:200]),
                    )
        return turn_id
    except Exception:  # noqa: BLE001
        logging.warning("b24 testbot: inflight register failed", exc_info=True)
        return None


def _b24_inflight_clear(turn_id: str | None) -> None:
    """Turn finished (ok or handled error) — drop its in-flight row so boot recovery ignores it."""
    if not turn_id:
        return
    try:
        with pg_connect() as conn:
            with conn.transaction():
                with conn.cursor() as cur:
                    cur.execute("DELETE FROM bitrix_inflight_turns WHERE id = %s", (turn_id,))
    except Exception:  # noqa: BLE001
        logging.warning("b24 testbot: inflight clear failed", exc_info=True)


def _b24_recover_inflight_turns(endpoint: str = "", token: str = "") -> int:
    """At boot, every row left in bitrix_inflight_turns belongs to a turn that was killed
    mid-flight (deploy restart, OOM, crash) — its worker thread died with the process, so the
    user got no answer and a stuck 'typing…'. For each, delete the stale progress message and
    tell that user (as the exact bot they wrote to) to resend, then remove the row. This is the
    hard guarantee that a killed turn NEVER looks like an eternal hang. Best-effort, never raises."""
    try:
        with pg_connect() as conn:
            with conn.cursor() as cur:
                cur.execute(
                    "SELECT id::text AS id, bot_id, dialog_id, status_message_id"
                    " FROM bitrix_inflight_turns ORDER BY started_at"
                )
                rows = [dict(r) for r in cur.fetchall()]
    except Exception:  # noqa: BLE001
        logging.warning("b24 testbot: inflight recovery read failed", exc_info=True)
        return 0
    if not rows:
        return 0
    if not (endpoint and token):
        try:
            endpoint, token = _b24_app_access_token()
        except Exception:  # noqa: BLE001
            endpoint, token = "", ""
    for r in rows:
        dlg, bot, smid = r["dialog_id"], r["bot_id"], r["status_message_id"]
        if endpoint and token and smid:
            try:
                _b24_app_call(endpoint, token, "imbot.message.delete",
                              {"BOT_ID": bot, "MESSAGE_ID": smid, "COMPLETE": "Y"})
            except Exception:  # noqa: BLE001
                logging.debug("b24 testbot: inflight recovery status-delete failed", exc_info=True)
        try:
            _albery_bitrix_notify(
                "🙏 Извините — я перезапустился и не успел ответить на ваше прошлое сообщение. "
                "Пожалуйста, отправьте его ещё раз, и я сразу отвечу.",
                dialog_id=dlg, client_endpoint=endpoint, access_token=token, bot_id=bot)
        except Exception:  # noqa: BLE001
            logging.warning("b24 testbot: inflight recovery notify failed dlg=%s", dlg, exc_info=True)
    try:
        ids = [r["id"] for r in rows]
        with pg_connect() as conn:
            with conn.transaction():
                with conn.cursor() as cur:
                    cur.execute("DELETE FROM bitrix_inflight_turns WHERE id::text = ANY(%s)", (ids,))
    except Exception:  # noqa: BLE001
        logging.warning("b24 testbot: inflight recovery cleanup failed", exc_info=True)
    logging.info("b24 testbot: inflight recovery notified %d interrupted turn(s)", len(rows))
    return len(rows)


def _b24_startup_register_commands() -> None:
    """On process start, register keyboard commands for ALL bots without waiting for an event, so
    buttons are live immediately after a deploy/restart (not only after the first message). Uses the
    stored app refresh_token to mint a token off-event; best-effort and idempotent (per-bot flags).
    Also runs the in-flight turn recovery net so any turn cut off by the restart is answered."""
    def _do() -> None:
        time.sleep(15)  # let the app + DB pool settle after startup
        try:
            endpoint, token = _b24_app_access_token()
        except Exception:  # noqa: BLE001
            logging.debug("b24 testbot: startup command registration — token not ready", exc_info=True)
            endpoint, token = "", ""
        if endpoint and token:
            _b24_bootstrap_all_commands(endpoint, token, _b24_load_state().get("bot_id"))
            try:
                _b24_ensure_task_comment_event_bound(endpoint, token)
            except Exception:  # noqa: BLE001
                logging.warning("b24 testbot: task-comment event bind on startup failed", exc_info=True)
        try:
            _b24_recover_inflight_turns(endpoint, token)
        except Exception:  # noqa: BLE001
            logging.warning("b24 testbot: inflight recovery failed", exc_info=True)

    threading.Thread(target=_do, daemon=True).start()


# Kick off proactive command registration + interrupted-turn recovery once, right after import.
_b24_startup_register_commands()


def _b24_app_typing(client_endpoint: str, access_token: str, bot_id: Any, dialog_id: str) -> None:
    """Show the bot as 'typing…' — the read/processing signal (no eye reaction on this portal)."""
    if not (client_endpoint and access_token and bot_id and dialog_id):
        return
    try:
        _b24_app_call(client_endpoint, access_token, "imbot.chat.sendTyping",
                      {"BOT_ID": bot_id, "DIALOG_ID": dialog_id})
    except Exception:  # noqa: BLE001
        logging.debug("b24 testbot: typing failed", exc_info=True)


def _b24_app_react(client_endpoint: str, access_token: str, message_id: Any, reaction: str, add: bool = True) -> None:
    """Add/remove a chat reaction (👀 eyes on read, 👍 like on done) via the IM v2 API."""
    if not (client_endpoint and access_token and message_id):
        return
    method = "im.v2.Chat.Message.Reaction.add" if add else "im.v2.Chat.Message.Reaction.delete"
    try:
        _b24_app_call(client_endpoint, access_token, method, {"messageId": message_id, "reaction": reaction})
    except Exception:  # noqa: BLE001
        logging.debug("b24 testbot: reaction %s/%s failed", reaction, add, exc_info=True)


def _b24_full_user_ids() -> set[int]:
    return {int(x) for x in re.findall(r"\d+", os.getenv("B24_TESTBOT_FULL_USER_IDS", "14,16"))}


def _b24_owner_user_ids() -> set[int]:
    """Bitrix user ids that get the ADMIN connector — can edit AI instructions/capabilities and
    run destructive deletes (OWNER_ONLY_TOOL_NAMES). Default: 16 = Александр on the b24-0xrp3s
    portal. Everyone else (including other 'full' users) is capped at the operational connector."""
    return {int(x) for x in re.findall(r"\d+", os.getenv("B24_TESTBOT_OWNER_USER_IDS", "16"))}


# Access tiers are stored in the agent_access table and managed live from the "Настройки Агента"
# tab (/api/agent-access). We cache the map briefly so tier resolution skips the DB per message.
_AGENT_ACCESS_CACHE: dict[str, Any] = {"at": 0.0, "map": {}}
_AGENT_ACCESS_TTL = 20.0


def _agent_access_map(force: bool = False) -> dict[int, str]:
    now = time.monotonic()
    if not force and (now - _AGENT_ACCESS_CACHE["at"]) < _AGENT_ACCESS_TTL:
        return _AGENT_ACCESS_CACHE["map"]
    try:
        with pg_connect() as conn:
            with conn.cursor() as cur:
                cur.execute("SELECT bitrix_user_id, tier FROM agent_access")
                mapping = {int(r["bitrix_user_id"]): r["tier"] for r in cur.fetchall()}
    except Exception:  # noqa: BLE001 — keep serving with the last good map on a DB hiccup
        logging.exception("agent_access map load failed")
        return _AGENT_ACCESS_CACHE["map"]
    _AGENT_ACCESS_CACHE.update(at=now, map=mapping)
    return mapping


def _agent_access_set(bitrix_user_id: int, tier: str, display_name: str | None = None,
                      note: str | None = None) -> None:
    with pg_connect() as conn:
        with conn.transaction():
            with conn.cursor() as cur:
                cur.execute(
                    """
                    INSERT INTO agent_access (bitrix_user_id, tier, display_name, note, updated_at)
                    VALUES (%s, %s, %s, %s, now())
                    ON CONFLICT (bitrix_user_id) DO UPDATE SET
                        tier = EXCLUDED.tier,
                        display_name = COALESCE(EXCLUDED.display_name, agent_access.display_name),
                        note = COALESCE(EXCLUDED.note, agent_access.note),
                        updated_at = now()
                    """,
                    (int(bitrix_user_id), tier, display_name, note),
                )
    _agent_access_map(force=True)


def _agent_access_remove(bitrix_user_id: int) -> None:
    with pg_connect() as conn:
        with conn.transaction():
            with conn.cursor() as cur:
                cur.execute("DELETE FROM agent_access WHERE bitrix_user_id = %s", (int(bitrix_user_id),))
    _agent_access_map(force=True)


def _b24_tier_for(from_user_id: Any) -> str:
    """Access tier from the Bitrix-trusted sender id (cannot be spoofed in chat):
    'admin' = full incl. instruction/settings edits + deletes; 'ops' = full operational access
    minus the admin tools; 'faq' = knowledge base; 'none' = no access (the bot does not respond).
    The agent_access table (managed from the "Настройки Агента" tab) decides; a user with no row
    defaults to 'faq' (knowledge base). 'none' is an explicit stored deny. No id is hard-pinned —
    every level, including the owner's, is editable from the UI."""
    tier = _agent_access_map().get(to_int(from_user_id))
    return tier if tier in ("admin", "ops", "faq", "none") else "faq"


def _b24_main_allows(from_user_id: Any) -> bool:
    """Access to the UNIVERSAL (main) agent is a strict allowlist = its «Команда и доступы»
    (agent_access rows with a non-'none' tier). Anyone NOT in that list gets no answer — the
    team list IS the access list (owner rule 2026-07-03). The bootstrap owner id(s) are always
    allowed so nobody can lock themselves out by removing their own row."""
    uid = to_int(from_user_id)
    if uid in _b24_owner_user_ids():
        return True
    return _agent_access_map().get(uid) not in (None, "none")


# --- Session lifecycle: 30-min idle reset + turn-cap rotation with carried summary ----
# Hermes auto-compression is disabled on this box (it failed on codex), so we bound the
# context ourselves: each Bitrix dialog maps to an epoch'd session. Idle >30 min starts a
# genuinely fresh epoch (history floor raised + no carried summary, like the manual «Новая
# сессия» button); after a turn cap we rotate to a new epoch seeded with a short summary of
# the previous one (conversation-summary-buffer) so long threads never blow the window.
B24_IDLE_RESET_SECONDS = int(os.getenv("B24_TESTBOT_IDLE_RESET_SECONDS", "1800"))
B24_TURN_CAP = int(os.getenv("B24_TESTBOT_TURN_CAP", "16"))


def _b24_scope(dialog_id: Any, agent_slug: str | None) -> str:
    """Per-agent session key. In Bitrix a private chat is keyed by the USER id, so the SAME
    dialog_id is reused by every bot that person talks to — keying the hermes session and the
    lifecycle row by dialog_id alone made the main bot and every subagent SHARE one conversation.
    Scope the key by agent so each agent has its own sessions (main agent = bare dialog_id, keeps
    existing rows working). The real dialog_id is still used for Bitrix replies and for filtering
    bitrix_bot_interactions (which carries its own agent_slug column)."""
    slug = (agent_slug or "").strip()
    return f"{slug}:{dialog_id}" if slug else str(dialog_id)


def _b24_summarize_segment(dialog_id: str, agent_slug: str | None = None) -> str | None:
    try:
        with pg_connect() as conn:
            with conn.cursor() as cur:
                cur.execute(
                    "SELECT question, answer FROM bitrix_bot_interactions "
                    "WHERE dialog_id=%s AND agent_slug IS NOT DISTINCT FROM %s::text ORDER BY id DESC LIMIT 20",
                    (str(dialog_id), agent_slug),
                )
                rows = cur.fetchall()
        if not rows:
            return None
        convo = "\n".join(f"- {r['question']} → {r['answer']}" for r in reversed(rows))[:6000]
        proc = subprocess.run(
            ["hermes", "-z", "Сожми диалог в 3-4 предложениях как контекст для продолжения, по-русски:\n\n" + convo,
             "-t", "albery-faq", "--yolo"],
            capture_output=True, text=True, timeout=90, cwd="/root", env={**os.environ, "HOME": "/root"},
        )
        return (proc.stdout or "").strip() or None
    except Exception:  # noqa: BLE001
        logging.exception("b24 testbot: summarize failed")
        return None


def _b24_session_prepare(dialog_id: str, agent_slug: str | None = None) -> tuple[str, str | None]:
    """Return (session_name, seed_summary_for_this_turn) applying idle reset / cap rotation.
    Scoped per agent: the lifecycle row and the hermes session name are keyed by
    (agent_slug, dialog_id), so each agent keeps its own conversation even though the Bitrix
    dialog_id is shared across bots. The history floor is computed from that agent's own turns."""
    scope = _b24_scope(dialog_id, agent_slug)
    safe = re.sub(r"[^A-Za-z0-9_-]", "", str(scope))[:60]
    now = datetime.now(timezone.utc)
    try:
        with pg_connect() as conn:
            with conn.transaction():
                with conn.cursor() as cur:
                    cur.execute(
                        "SELECT epoch, turns, summary, last_at FROM bitrix_bot_sessions WHERE dialog_id=%s FOR UPDATE",
                        (scope,),
                    )
                    row = cur.fetchone()
                    if not row:
                        cur.execute(
                            "INSERT INTO bitrix_bot_sessions (dialog_id, epoch, turns, last_at) VALUES (%s, 1, 0, %s)",
                            (scope, now),
                        )
                        return f"bitrix-{safe}-e1", None
                    epoch, turns, summary, last_at = row["epoch"], row["turns"], row["summary"], row["last_at"]
                    idle = (now - last_at).total_seconds() if last_at else 1e12
                    seed: str | None = None
                    idle_reset = False
                    if idle > B24_IDLE_RESET_SECONDS:
                        # Idle gap → genuinely fresh session: drop carried summary AND raise the
                        # history floor so prior Q/A is no longer injected into the prompt.
                        epoch, turns, summary, idle_reset = epoch + 1, 0, None, True
                    elif turns >= B24_TURN_CAP:
                        seed = _b24_summarize_segment(dialog_id, agent_slug) or summary
                        epoch, turns, summary = epoch + 1, 0, seed
                    elif turns == 0:
                        seed = summary
                    if idle_reset:
                        cur.execute(
                            "SELECT COALESCE(MAX(id), 0) AS floor FROM bitrix_bot_interactions "
                            "WHERE dialog_id=%s AND agent_slug IS NOT DISTINCT FROM %s::text",
                            (str(dialog_id), agent_slug),
                        )
                        floor = cur.fetchone()["floor"]
                        cur.execute(
                            "UPDATE bitrix_bot_sessions SET epoch=%s, turns=%s, summary=%s, last_at=%s, "
                            "history_floor_id=%s WHERE dialog_id=%s",
                            (epoch, turns, summary, now, floor, scope),
                        )
                    else:
                        cur.execute(
                            "UPDATE bitrix_bot_sessions SET epoch=%s, turns=%s, summary=%s, last_at=%s WHERE dialog_id=%s",
                            (epoch, turns, summary, now, scope),
                        )
                    return f"bitrix-{safe}-e{epoch}", seed
    except Exception:  # noqa: BLE001
        logging.exception("b24 testbot: session prepare failed")
        return f"bitrix-{safe}-e1", None


def _b24_session_touch(dialog_id: str, agent_slug: str | None = None) -> None:
    try:
        with pg_connect() as conn:
            with conn.transaction():
                with conn.cursor() as cur:
                    cur.execute(
                        "UPDATE bitrix_bot_sessions SET turns = turns + 1, last_at = now() WHERE dialog_id = %s",
                        (_b24_scope(dialog_id, agent_slug),),
                    )
    except Exception:  # noqa: BLE001
        logging.exception("b24 testbot: session touch failed")


# Reset keywords/commands accepted from chat (typed text or a bot command). Matched against
# the whole trimmed message (case/space-insensitive), so they never fire mid-sentence.
B24_RESET_TRIGGERS = {
    "/new", "/reset", "/clear", "/restart", "/start",
    "/новая", "/сброс", "/сбросить", "/заново",
    "новая сессия", "новый диалог", "новый чат", "начать заново", "начать сначала",
    "сбросить сессию", "сброс сессии", "сбросить контекст", "очистить контекст",
    "сбрось сессию", "сбросить диалог", "сброс",
}


def _b24_is_reset_command(text: str) -> bool:
    """True if the message is purely a 'new session' request (typed or via a bot command)."""
    norm = " ".join(str(text or "").strip().lower().split())
    norm = norm.rstrip("!.,")
    return norm in B24_RESET_TRIGGERS


def _b24_session_reset(dialog_id: str, agent_slug: str | None = None) -> None:
    """Manual 'new session': bump the epoch AND raise the history floor so previously
    injected Q/A from bitrix_bot_interactions is no longer carried into the prompt — i.e.
    a genuinely clean conversation, not just a relabeled session key. Scoped per agent: the
    floor is taken from THIS agent's own turns and the lifecycle row is keyed by (agent, dialog)."""
    scope = _b24_scope(dialog_id, agent_slug)
    now = datetime.now(timezone.utc)
    try:
        with pg_connect() as conn:
            with conn.transaction():
                with conn.cursor() as cur:
                    cur.execute(
                        "SELECT COALESCE(MAX(id), 0) AS floor FROM bitrix_bot_interactions "
                        "WHERE dialog_id=%s AND agent_slug IS NOT DISTINCT FROM %s::text",
                        (str(dialog_id), agent_slug),
                    )
                    floor = cur.fetchone()["floor"]
                    cur.execute(
                        """
                        INSERT INTO bitrix_bot_sessions (dialog_id, epoch, turns, summary, last_at, history_floor_id)
                        VALUES (%s, 2, 0, NULL, %s, %s)
                        ON CONFLICT (dialog_id) DO UPDATE
                          SET epoch = bitrix_bot_sessions.epoch + 1,
                              turns = 0, summary = NULL, last_at = %s,
                              history_floor_id = %s
                        """,
                        (scope, now, floor, now, floor),
                    )
    except Exception:  # noqa: BLE001
        logging.exception("b24 testbot: session reset failed")


# ── Bitrix bot: read screenshots (Groq vision OCR) + reply-to-earlier-message context ──────────
def _b24_groq_api_key() -> str:
    """GROQ key for vision OCR. Albery runs as root, so it can read the Hermes gateway env where
    the key already lives (no secret duplicated into the Albery .env)."""
    k = os.getenv("GROQ_API_KEY", "").strip()
    if k:
        return k
    try:
        for line in Path("/root/.hermes/secure/hermes-gateway.env").read_text(encoding="utf-8").splitlines():
            if line.strip().startswith("GROQ_API_KEY="):
                return line.split("=", 1)[1].strip().strip('"').strip("'")
    except OSError:
        pass
    return ""


def _b24_vision_ocr(image_bytes: bytes, name: str = "") -> str:
    """OCR/describe an image with Groq vision. NOTE: api.groq.com blocks urllib's default
    User-Agent with Cloudflare 1010 — a browser UA passes (same gotcha as the STT path)."""
    key = _b24_groq_api_key()
    if not key or not image_bytes:
        return ""
    ext = (name.rsplit(".", 1)[-1] if "." in name else "png").lower()
    mime = {"jpg": "jpeg", "jpeg": "jpeg", "png": "png", "gif": "gif", "webp": "webp"}.get(ext, "png")
    b64 = base64.b64encode(image_bytes).decode()
    payload = {
        "model": os.getenv("B24_VISION_MODEL", "meta-llama/llama-4-scout-17b-16e-instruct"),
        "max_tokens": 800, "temperature": 0,
        "messages": [{"role": "user", "content": [
            {"type": "text", "text": "Это изображение/скриншот, который пользователь прислал ассистенту "
             "в чате. Извлеки ВЕСЬ текст с него ДОСЛОВНО, а затем кратко (1-2 предложения) опиши, что на "
             "нём и какая может быть проблема/смысл. Ответь по-русски, без лишних вступлений."},
            {"type": "image_url", "image_url": {"url": f"data:image/{mime};base64,{b64}"}}]}],
    }
    try:
        req = urllib.request.Request(
            "https://api.groq.com/openai/v1/chat/completions",
            data=json.dumps(payload).encode(),
            headers={"Authorization": "Bearer " + key, "Content-Type": "application/json",
                     "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)"})
        with urllib.request.urlopen(req, timeout=70) as r:
            d = json.loads(r.read().decode("utf-8", "ignore"))
        return ((d.get("choices") or [{}])[0].get("message", {}) or {}).get("content", "").strip()
    except Exception as exc:  # noqa: BLE001
        logging.warning("b24 vision OCR failed: %s", repr(exc)[:200])
        return ""


def _b24_fetch_bytes(url: str, access_token: str, max_bytes: int = 20 * 1024 * 1024) -> bytes:
    """Download a Bitrix file. Tries the URL as-is, then with ?auth=<token> (REST file URLs need it)."""
    if not url:
        return b""
    candidates = [url]
    if "auth=" not in url and access_token:
        candidates.append(url + ("&" if "?" in url else "?") + "auth=" + access_token)
    for u in candidates:
        try:
            with requests.get(u, stream=True, timeout=60) as resp:
                if resp.status_code != 200:
                    continue
                buf = b""
                for chunk in resp.iter_content(65536):
                    buf += chunk
                    if len(buf) > max_bytes:
                        break
                if buf[:64].lstrip().lower().startswith((b"<!doctype html", b"<html")):
                    continue  # got a login/HTML page, not the file
                return buf
        except Exception:  # noqa: BLE001
            continue
    return b""


_B24_DOC_EXTS = ("pdf", "docx", "doc", "xlsx", "xlsm", "md", "markdown", "txt",
                 "csv", "tsv", "json", "rtf", "htm", "html", "log", "yaml", "yml")
_B24_IMG_EXTS = ("png", "jpg", "jpeg", "gif", "webp", "bmp", "heic")
_B24_DELIVER_RE = re.compile(r"\[\[DELIVER_(PDF|XLSX|EXCEL|DOCX|WORD):\s*([^\]]*)\]\]", re.I)
_B24_DELIVER_FMT = {"pdf": "pdf", "xlsx": "xlsx", "excel": "xlsx", "docx": "docx", "word": "docx"}


def _b24_extract_document(data: bytes, name: str) -> str:
    """Extract readable text from a user-sent document. Pure-python extractors (pypdf / python-docx /
    openpyxl) + native text for md/txt/csv. Legacy binary .doc isn't supported (no libreoffice)."""
    import io as _io
    ext = name.rsplit(".", 1)[-1].lower() if "." in name else ""
    try:
        if ext in ("md", "markdown", "txt", "csv", "tsv", "json", "log", "yaml", "yml", "htm", "html"):
            return data.decode("utf-8", "ignore")
        if ext == "pdf":
            from pypdf import PdfReader
            reader = PdfReader(_io.BytesIO(data))
            return "\n".join((pg.extract_text() or "") for pg in reader.pages).strip()
        if ext == "docx":
            from docx import Document
            doc = Document(_io.BytesIO(data))
            parts = [p.text for p in doc.paragraphs]
            for tbl in doc.tables:
                for row in tbl.rows:
                    parts.append(" | ".join(cell.text for cell in row.cells))
            return "\n".join(parts).strip()
        if ext in ("xlsx", "xlsm"):
            from openpyxl import load_workbook
            wb = load_workbook(_io.BytesIO(data), read_only=True, data_only=True)
            out = []
            for ws in wb.worksheets:
                out.append("# Лист: " + str(ws.title))
                for row in ws.iter_rows(values_only=True):
                    cells = ["" if v is None else str(v) for v in row]
                    if any(c.strip() for c in cells):
                        out.append(" | ".join(cells))
            return "\n".join(out).strip()
        if ext == "doc":
            return ""  # legacy binary .doc — unsupported without libreoffice/antiword
        return data.decode("utf-8", "ignore")
    except Exception as exc:  # noqa: BLE001
        logging.warning("b24 doc extract failed (%s): %s", name, repr(exc)[:160])
        return ""


def _b24_text_to_pdf(title: str, text: str) -> bytes:
    """Render plain/markdown-ish text (supports [b]…[/b], **…**, '- ' bullets, '# ' headings) to a
    Cyrillic-capable PDF via reportlab + the DejaVu fonts already used for owner reports."""
    import io as _io, html as _html
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors
    fr, fb = "Helvetica", "Helvetica-Bold"
    try:
        reg, bold = reportlab_font_paths()
        pdfmetrics.registerFont(TTFont("B24Reg", reg))
        pdfmetrics.registerFont(TTFont("B24Bold", bold))
        fr, fb = "B24Reg", "B24Bold"
    except Exception:  # noqa: BLE001
        pass
    body = ParagraphStyle("b24body", fontName=fr, fontSize=11, leading=15)
    head = ParagraphStyle("b24head", fontName=fb, fontSize=15, leading=19, spaceBefore=4, spaceAfter=8)
    sub = ParagraphStyle("b24sub", fontName=fb, fontSize=12.5, leading=17, spaceBefore=6, spaceAfter=4)

    _emoji = re.compile("[\U0001F000-\U0001FAFF\U00002600-\U000027BF\U0001F1E6-\U0001F1FF\uFE0F\u200D]")

    def fmt(s: str) -> str:
        s = _emoji.sub("", s)  # DejaVu lacks colour emoji -> would render as tofu boxes in the PDF
        s = _html.escape(s)
        s = re.sub(r"\[b\](.+?)\[/b\]", r"<b>\1</b>", s, flags=re.S)
        s = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", s, flags=re.S)
        return s

    flow = []
    if title:
        flow.append(Paragraph(fmt(str(title)), head))
        flow.append(Spacer(1, 4))
    PURPLE = colors.HexColor("#5440F6")
    avail = (210 - 36) * mm  # A4 width minus left/right margins
    for kind, val in _b24_split_doc_blocks(text):
        if kind == "table":
            ncols = max((len(r) for r in val), default=1) or 1
            cst = ParagraphStyle("b24cell", fontName=fr, fontSize=9.5, leading=12)
            chs = ParagraphStyle("b24cellh", fontName=fb, fontSize=9.5, leading=12, textColor=colors.white)
            data = [[Paragraph(fmt(c), chs if ri == 0 else cst) for c in row] for ri, row in enumerate(val)]
            tbl = Table(data, colWidths=[avail / ncols] * ncols, repeatRows=1)
            tbl.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), PURPLE),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F4F2FF")]),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#D9D5F5")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5), ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]))
            flow.append(tbl)
            flow.append(Spacer(1, 8))
            continue
        line = val.rstrip()
        if not line.strip():
            flow.append(Spacer(1, 6)); continue
        st = line.lstrip()
        if st.startswith("# "):
            flow.append(Paragraph(fmt(st[2:]), head))
        elif st.startswith("## "):
            flow.append(Paragraph(fmt(st[3:]), sub))
        elif st.startswith(("- ", "• ", "* ")):
            flow.append(Paragraph("•&nbsp;" + fmt(st[2:]), body))
        else:
            flow.append(Paragraph(fmt(line), body))
    buf = _io.BytesIO()
    SimpleDocTemplate(buf, pagesize=A4, leftMargin=18 * mm, rightMargin=18 * mm,
                      topMargin=16 * mm, bottomMargin=16 * mm,
                      title=str(title or "Документ")).build(flow)
    return buf.getvalue()


def _b24_save_pdf_export(pdf_bytes: bytes, name: str) -> str:
    """Save a PDF to the export dir and return a public, token-protected, time-limited download URL
    (reuses the zoom-export token mechanism; the route serves any file in that dir)."""
    base = re.sub(r"[^\w.\-() а-яёА-ЯЁ]+", "_", str(name or "Документ")).strip() or "Документ"
    if not base.lower().endswith(".pdf"):
        base += ".pdf"
    fname = "%d_%s" % (int(time.time()), base)
    ZOOM_EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    cleanup_zoom_exports()
    (ZOOM_EXPORT_DIR / fname).write_bytes(pdf_bytes)
    url = _zoom_export_public_url(fname)
    if url.startswith("/"):
        url = "https://mcp.m4s.ru" + url
    return url


def _b24_plain(s) -> str:
    """Drop inline formatting markers ([b]…[/b], **…**) for raw cell/run text."""
    s = re.sub(r"\[/?b\]", "", str(s or ""))
    s = re.sub(r"\*\*(.+?)\*\*", r"\1", s, flags=re.S)
    return s


def _b24_bold_segments(s):
    """Split a line into (text, is_bold) runs honoring [b]…[/b] / **…**."""
    s = re.sub(r"\*\*(.+?)\*\*", r"[b]\1[/b]", str(s or ""), flags=re.S)
    out = []
    for i, part in enumerate(re.split(r"\[b\](.+?)\[/b\]", s, flags=re.S)):
        if part:
            out.append((part, i % 2 == 1))
    return out or [("", False)]


def _b24_split_doc_blocks(text: str):
    """Split answer text into ('table', rows) and ('line', str) blocks. A table is a run of
    markdown pipe rows ('| a | b |'); the '| --- | --- |' separator row is dropped and ragged
    rows are padded to the widest row."""
    out = []
    lines = (text or "").split("\n")

    def is_row(s):
        t = s.strip()
        return t.startswith("|") and t.count("|") >= 2

    def is_sep(s):
        t = s.strip().strip("|").replace(" ", "")
        return bool(t) and set(t) <= set("-:|")

    i = 0
    while i < len(lines):
        if is_row(lines[i]):
            rows = []
            while i < len(lines) and is_row(lines[i]):
                if not is_sep(lines[i]):
                    rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            if rows:
                w = max(len(r) for r in rows)
                out.append(("table", [r + [""] * (w - len(r)) for r in rows]))
        else:
            out.append(("line", lines[i]))
            i += 1
    return out


def _b24_text_to_xlsx(title: str, text: str) -> bytes:
    """Build a real .xlsx: pipe tables become rows (bold purple header); non-table lines are
    written into column A above/around them."""
    import io as _io
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    wb = Workbook()
    ws = wb.active
    ws.title = (re.sub(r"[^\w \-]+", "", str(title or "Лист"))[:28] or "Лист")
    hfont = Font(bold=True, color="FFFFFF")
    hfill = PatternFill("solid", fgColor="5440F6")
    wrap = Alignment(wrap_text=True, vertical="top")
    r = 1
    had_table = False
    for kind, val in _b24_split_doc_blocks(text):
        if kind == "table":
            had_table = True
            for ri, row in enumerate(val):
                for ci, c in enumerate(row, 1):
                    cell = ws.cell(row=r, column=ci, value=_b24_plain(c))
                    cell.alignment = wrap
                    if ri == 0:
                        cell.font = hfont
                        cell.fill = hfill
                r += 1
            r += 1
        else:
            s = _b24_plain(val).rstrip()
            if s.strip():
                ws.cell(row=r, column=1, value=s)
                r += 1
    for col in ws.columns:
        try:
            width = min(60, max((len(str(c.value)) for c in col if c.value is not None), default=10) + 2)
            ws.column_dimensions[col[0].column_letter].width = width
        except Exception:  # noqa: BLE001
            pass
    buf = _io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _b24_text_to_docx(title: str, text: str) -> bytes:
    """Build a real .docx.

    Ordinary files keep the compact answer-style layout. Contracts and other official/legal
    documents are formatted as a clean Russian business document: A4, GOST-like margins,
    Times New Roman, 14 pt, 1.5 spacing, justified paragraphs and first-line indent.
    """
    import io as _io
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    def is_official_doc(name: str, body: str) -> bool:
        hay = (str(name or "") + "\n" + str(body or "")).lower()
        return bool(re.search(r"\b(договор|соглашение|акт|контракт|оферт|гост)\b", hay, flags=re.I))

    official = is_official_doc(title, text)
    doc = Document()

    def set_run_font(run, *, size=14, bold=None):
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(0, 0, 0)
        if bold is not None:
            run.bold = bold

    def apply_paragraph_format(p, *, align=None, first_line=True, space_after=0):
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(space_after)
        pf.line_spacing = 1.5 if official else 1.15
        if official:
            pf.first_line_indent = Cm(1.25) if first_line else Cm(0)
        if align is not None:
            p.alignment = align
        elif official:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def add_runs(p, line: str, *, force_bold=False, size=14):
        for seg, isb in _b24_bold_segments(line):
            run = p.add_run(seg)
            set_run_font(run, size=size, bold=(force_bold or isb))

    def set_cell_text(cell, value: str, *, header=False):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.text = ""
        p = cell.paragraphs[0]
        apply_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER if header else WD_ALIGN_PARAGRAPH.LEFT,
                               first_line=False, space_after=0)
        run = p.add_run(_b24_plain(value))
        set_run_font(run, size=12 if official else 11, bold=header)

    def set_table_borders(table):
        tbl = table._tbl
        tblPr = tbl.tblPr
        borders = tblPr.first_child_found_in("w:tblBorders")
        if borders is None:
            borders = OxmlElement("w:tblBorders")
            tblPr.append(borders)
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            tag = "w:" + edge
            elem = borders.find(qn(tag))
            if elem is None:
                elem = OxmlElement(tag)
                borders.append(elem)
            elem.set(qn("w:val"), "single")
            elem.set(qn("w:sz"), "6")
            elem.set(qn("w:space"), "0")
            elem.set(qn("w:color"), "000000")

    if official:
        section = doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        for style_name in ("Normal", "Body Text"):
            try:
                style = doc.styles[style_name]
                style.font.name = "Times New Roman"
                style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                style.font.size = Pt(14)
            except Exception:  # noqa: BLE001
                pass
        if title:
            p = doc.add_paragraph()
            apply_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, space_after=6)
            add_runs(p, _b24_plain(title).upper(), force_bold=True, size=14)
    elif title:
        doc.add_heading(_b24_plain(title), level=0)

    for kind, val in _b24_split_doc_blocks(text):
        if kind == "table":
            ncols = max((len(r) for r in val), default=1) or 1
            t = doc.add_table(rows=0, cols=ncols)
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            if not official:
                try:
                    t.style = "Table Grid"
                except Exception:  # noqa: BLE001
                    pass
            set_table_borders(t)
            for ri, row in enumerate(val):
                cells = t.add_row().cells
                for ci in range(ncols):
                    set_cell_text(cells[ci], row[ci] if ci < len(row) else "", header=(ri == 0))
            doc.add_paragraph("")
            continue
        line = val.rstrip()
        if not line.strip():
            continue
        st = line.lstrip()
        plain = _b24_plain(st)
        if official:
            if re.match(r"^(?:#\s*)?(?:ДОГОВОР|СОГЛАШЕНИЕ|АКТ|КОНТРАКТ)\b", plain, flags=re.I):
                p = doc.add_paragraph()
                apply_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, space_after=6)
                add_runs(p, plain.upper(), force_bold=True, size=14)
            elif re.match(r"^\d+(?:\.\d+)*\.\s+\S", plain):
                is_section_heading = bool(re.match(r"^\d+\.\s+[^.]{3,80}$", plain))
                p = doc.add_paragraph()
                apply_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=not is_section_heading, space_after=0)
                add_runs(p, plain, force_bold=is_section_heading, size=14)
            elif re.match(r"^(г\.|город\s+)\s*\S+", plain, flags=re.I) or "___" in plain and "20" in plain:
                p = doc.add_paragraph()
                apply_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False, space_after=0)
                add_runs(p, plain, size=14)
            elif st.startswith(("# ", "## ")):
                p = doc.add_paragraph()
                apply_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, space_after=6)
                add_runs(p, _b24_plain(re.sub(r"^#+\s*", "", st)).upper(), force_bold=True, size=14)
            elif st.startswith(("- ", "• ", "* ")):
                p = doc.add_paragraph()
                apply_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=True, space_after=0)
                add_runs(p, "— " + _b24_plain(st[2:]), size=14)
            else:
                p = doc.add_paragraph()
                apply_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=True, space_after=0)
                add_runs(p, line, size=14)
        else:
            if st.startswith("# "):
                doc.add_heading(_b24_plain(st[2:]), level=1)
            elif st.startswith("## "):
                doc.add_heading(_b24_plain(st[3:]), level=2)
            elif st.startswith(("- ", "• ", "* ")):
                doc.add_paragraph(_b24_plain(st[2:]), style="List Bullet")
            else:
                p = doc.add_paragraph()
                for seg, isb in _b24_bold_segments(line):
                    run = p.add_run(seg)
                    set_run_font(run, size=11, bold=isb)
    buf = _io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _b24_save_export(data: bytes, name: str, ext: str = "pdf") -> str:
    """Save a generated document (pdf/xlsx/docx) to the export dir and return a public,
    token-protected, time-limited download URL (same mechanism as the PDF export)."""
    base = re.sub(r"[^\w.\-() а-яёА-ЯЁ]+", "_", str(name or "Документ")).strip() or "Документ"
    base = re.sub(r"\.(pdf|xlsx|docx)$", "", base, flags=re.I) + "." + ext
    fname = "%d_%s" % (int(time.time()), base)
    ZOOM_EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    cleanup_zoom_exports()
    (ZOOM_EXPORT_DIR / fname).write_bytes(data)
    url = _zoom_export_public_url(fname)
    if url.startswith("/"):
        url = "https://mcp.m4s.ru" + url
    return url


def _b24_extract_deliver(answer: str):
    """Strip a [[DELIVER_<fmt>: name]] marker; return (clean_text, name|None, fmt|None),
    where fmt is one of 'pdf' | 'xlsx' | 'docx'."""
    m = _B24_DELIVER_RE.search(answer or "")
    if not m:
        return answer, None, None
    fmt = _B24_DELIVER_FMT.get((m.group(1) or "").lower(), "pdf")
    name = (m.group(2) or "").strip() or "Документ"
    clean = _B24_DELIVER_RE.sub("", answer).strip()
    return clean, name, fmt


def _b24_app_download_url(endpoint: str, access_token: str, fid) -> str:
    """Resolve a REST DOWNLOAD_URL for a chat file using the BOT's own OAuth token.
    The bot is a participant of the dialog, so it can read files posted to it once the
    local app has the 'disk' scope (the inbound webhook user is not a chat member and is
    denied). Returns '' on any failure."""
    fid_i = to_int(fid)
    if fid_i is None or not endpoint or not access_token:
        return ""
    try:
        data = _b24_app_call(endpoint, access_token, "disk.file.get", {"id": fid_i})
    except Exception as exc:  # noqa: BLE001
        logging.warning("b24 extras: bot disk.file.get failed fid=%s: %s", fid, repr(exc)[:160])
        return ""
    res = data.get("result") if isinstance(data, dict) else None
    if isinstance(res, dict):
        for key in ("DOWNLOAD_URL", "DOWNLOAD_URL_EXTERNAL", "SHOW_URL"):
            value = res.get(key)
            if isinstance(value, str) and value.strip():
                return absolute_bitrix_url(value, endpoint) or ""
    return ""


def task_comment_files(file_ids: list, task_id: int, max_files: int = 5) -> list[dict]:
    """Read files attached to TASK COMMENTS (params.FILE_ID of the task-chat message): download,
    recognize (images → Groq vision OCR, documents → text extraction), persist to the attachment
    store and return [{attachment_id, name, kind, text}]. Used by mcp get_task_comments and the
    in-task mention flow, so agents can actually see screenshots/documents in comments (Sofia's
    case: a comment that is only a screenshot looked «empty» to the agent).

    Results are cached by the Bitrix disk file id (source_disk_file_id, migration 047): a repeated
    get_task_comments over the same task reuses stored text instead of re-downloading/re-OCRing.
    Download: app OAuth token first (same as chat files), inbound webhook as fallback — both were
    proven to reach task-chat files on this portal. Best-effort per file: a failure yields an
    honest «не удалось прочитать» entry instead of a silent drop."""
    import attachments as _att
    out: list[dict] = []
    wh = (os.getenv("B24_TESTBOT_WEBHOOK_BASE", "") or "").rstrip("/")
    ep = tok = ""
    for fid in list(file_ids or [])[:max_files]:
        fid_i = to_int(fid)
        if fid_i is None:
            continue
        cached = _att.find_by_disk_file_id(fid_i)
        if cached and (cached.get("extracted_text") or "").strip():
            out.append({"attachment_id": cached["token"], "name": cached.get("file_name") or "файл",
                        "kind": cached.get("kind") or "document",
                        "text": cached.get("extracted_text") or ""})
            continue
        # File metadata (NAME) + download URL via the webhook; app-token URL as the second road.
        name, durl = "", ""
        if wh:
            try:
                r = requests.post(f"{wh}/disk.file.get.json", json={"id": fid_i}, timeout=20)
                res = (r.json() or {}).get("result") or {}
                if isinstance(res, dict):
                    name = str(res.get("NAME") or "")
                    for key in ("DOWNLOAD_URL", "DOWNLOAD_URL_EXTERNAL", "SHOW_URL"):
                        if isinstance(res.get(key), str) and res[key].strip():
                            durl = absolute_bitrix_url(res[key], wh) or ""
                            break
            except Exception:  # noqa: BLE001
                logging.warning("b24 task files: webhook disk.file.get failed fid=%s", fid_i)
        data = _b24_fetch_bytes(durl, "") if durl else b""
        if not data:
            if not (ep and tok):
                try:
                    ep, tok = _b24_app_access_token()
                except Exception:  # noqa: BLE001
                    ep = tok = ""
            aurl = _b24_app_download_url(ep, tok, fid_i) if ep and tok else ""
            data = _b24_fetch_bytes(aurl, tok) if aurl else b""
        if not data:
            logging.warning("b24 task files: could not download fid=%s name=%s", fid_i, name)
            out.append({"attachment_id": None, "name": name or f"файл {fid_i}", "kind": "unknown",
                        "text": "(⚠️ не удалось скачать это вложение из комментария — честно скажи "
                                "об этом и попроси прислать файл в чат. НЕ придумывай содержимое.)"})
            continue
        ext = name.rsplit(".", 1)[-1].lower() if "." in name else ""
        is_image = ext in _B24_IMG_EXTS or (not ext and data[:3] in (b"\xff\xd8\xff", b"\x89PN", b"GIF"))
        if is_image:
            kind, text = "image", _b24_vision_ocr(data, name or "image.png")
        else:
            kind, text = "document", _b24_extract_document(data, name or "file")
        token = _att.store_attachment(
            data=data, file_name=name or ("image.png" if kind == "image" else "file"),
            kind=kind, extracted_text=text or "", agent_slug=None,
            dialog_id=f"task-{task_id}", source_disk_file_id=fid_i,
        )
        out.append({"attachment_id": token, "name": name or ("скрин" if kind == "image" else "файл"),
                    "kind": kind,
                    "text": (text or "").strip() or "(не удалось извлечь текст из вложения — "
                            "возможно скан без текстового слоя; попроси прислать PDF/DOCX)"})
    return out


def _b24_message_extras(payload: dict, endpoint: str = "", access_token: str = "",
                        agent_slug: str | None = None, dialog_id: str = "", from_user_id: Any = None):
    """Parse images + documents + reply-context straight from the flattened imbot event payload.
    Files come inline as data[PARAMS][FILES][<id>][field]; the inline _esd urls are session-bound, so
    we resolve a REST DOWNLOAD_URL via disk.file.get (needs the webhook 'disk' scope). Images -> Groq
    vision OCR; documents (pdf/docx/xlsx/md/txt/csv) -> text extraction.

    Every successfully downloaded file is ALSO persisted to the attachment store (full text + raw
    bytes), so (a) the agent can read the WHOLE document later via get_attachment_text — no 12k
    truncation — and (b) it can re-attach the original file to a task/comment/result. The store
    token is returned per attachment for prompt injection.

    Reply text is inline in data[PARAMS][REPLY_MESSAGE][MESSAGE].
    Returns (image_texts, reply_text, doc_blocks, attachments) where doc_blocks items are
    (name, content, token|None) and attachments is a list of {token,name,kind,char_len}."""
    image_texts: list = []
    reply_text = ""
    doc_blocks: list = []
    attachments: list = []
    if not isinstance(payload, dict):
        return image_texts, reply_text, doc_blocks, attachments
    for key in ("data[PARAMS][REPLY_MESSAGE][MESSAGE]", "data[params][REPLY_MESSAGE][MESSAGE]"):
        v = payload.get(key)
        if isinstance(v, list):
            v = v[0] if v else ""
        if v and str(v).strip():
            reply_text = str(v).strip()
            break
    files: dict = {}
    pat = re.compile(r"^data\[PARAMS\]\[FILES\]\[([^\]]+)\]\[([^\]]+)\]$", re.I)
    for k, v in payload.items():
        m = pat.match(str(k))
        if not m:
            continue
        val = v[0] if isinstance(v, list) and v else v
        files.setdefault(m.group(1), {})[m.group(2)] = val
    wh = (os.getenv("B24_TESTBOT_WEBHOOK_BASE", "") or "").rstrip("/")
    for fid, f in list(files.items())[:5]:
        name = str(f.get("name") or "")
        ext = name.rsplit(".", 1)[-1].lower() if "." in name else ""
        is_image = str(f.get("type") or "").lower() == "image" or ext in _B24_IMG_EXTS
        is_doc = ext in _B24_DOC_EXTS
        if not (is_image or is_doc):
            continue
        # Download the chat file. The BOT is a member of this dialog, so it can read the
        # file via its own OAuth token once the local app has the 'disk' scope; the static
        # inbound webhook user is NOT a chat member (403), so the bot token is tried first
        # and the webhook only as a fallback. The inline urlDownload is a session-bound ajax
        # URL (needs a browser cookie) and is intentionally not used server-side.
        durl = _b24_app_download_url(endpoint, access_token, fid)
        if not durl and wh:
            try:
                durl = _refresh_bitrix_download_url(wh, fid) or ""
            except Exception:  # noqa: BLE001
                logging.warning("b24 extras: webhook disk.file.get failed fid=%s", fid)
        data = _b24_fetch_bytes(durl, access_token) if durl else b""
        if not data:
            # Never drop a file silently: tell the brain a file arrived but could not be
            # read, so it asks for a resend / PDF instead of inventing the contents.
            logging.warning("b24 extras: could not download fid=%s name=%s (app may need 'disk' scope)", fid, name)
            if is_doc:
                doc_blocks.append((name or "документ",
                                   "(\u26a0\ufe0f не удалось скачать файл с сервера Bitrix. Честно скажи "
                                   "пользователю, что не смог открыть этот файл, и попроси прислать его "
                                   "ещё раз или в виде PDF. НЕ придумывай и не угадывай содержимое.)", None))
            else:
                image_texts.append("(\u26a0\ufe0f не удалось получить изображение. Скажи пользователю, что "
                                   "не смог его открыть, и попроси прислать ещё раз. Не придумывай, что на нём.)")
            continue
        # Persist the file (raw bytes + extracted text) so the agent can read the WHOLE document
        # later and re-attach the original to a task. Best-effort — never blocks the reply.
        def _store(kind: str, full_text: str) -> str | None:
            try:
                import attachments as _att
                return _att.store_attachment(
                    data=data, file_name=name or ("image" if kind == "image" else "file"),
                    kind=kind, extracted_text=full_text or "", agent_slug=agent_slug,
                    dialog_id=str(dialog_id or ""), bitrix_user_id=from_user_id,
                    mime=str(f.get("type") or "") or None,
                )
            except Exception:  # noqa: BLE001
                logging.warning("b24 extras: attachment store failed name=%s", name)
                return None
        if is_image:
            txt = _b24_vision_ocr(data, name or "image.png")
            token = _store("image", txt)
            if txt:
                image_texts.append(txt)
            if token:
                attachments.append({"token": token, "name": name or "image",
                                    "kind": "image", "char_len": len(txt or "")})
        else:
            txt = _b24_extract_document(data, name or "file")
            token = _store("document", txt if (txt and txt.strip()) else "")
            if txt and txt.strip():
                doc_blocks.append((name or "документ", txt.strip(), token))
            else:
                doc_blocks.append((name or "документ",
                                   "(не удалось извлечь текст: возможно скан без текстового слоя, "
                                   "пустой файл или устаревший формат .doc — попросите прислать PDF/DOCX)", token))
            if token:
                attachments.append({"token": token, "name": name or "документ",
                                    "kind": "document", "char_len": len((txt or "").strip())})
    if os.getenv("B24_DEBUG_PAYLOAD", "0") == "1":
        logging.info("b24 extras: images=%d docs=%d reply=%s attach=%d",
                     len(image_texts), len(doc_blocks), bool(reply_text), len(attachments))
    return image_texts, reply_text, doc_blocks, attachments


# How many characters of an extracted document to inline into the prompt. Beyond this the agent
# reads the rest with get_attachment_text(token, offset). Generous by default so short/medium
# docs (most contracts fit) are fully inline; the tool guarantees the complete text regardless.
_B24_DOC_INLINE_CHARS = int(os.getenv("B24_DOC_INLINE_CHARS", "30000") or "30000")


def _b24_compose_user_text(text: str, image_texts: list, reply_text: str, doc_blocks: list = None,
                           attachments: list = None) -> str:
    """Fold reply-context, image OCR and extracted document text into the message for the brain."""
    text = (text or "").strip()
    blocks: list = []
    if reply_text:
        blocks.append("[Пользователь ОТВЕЧАЕТ на это более раннее сообщение (возможно, из прошлой "
                      "сессии) — учитывай его как контекст]:\n«" + reply_text[:1500] + "»")
    if text:
        blocks.append(text)
    multi = len(image_texts) > 1
    for i, ocr in enumerate(image_texts, 1):
        label = (f"[Изображение №{i}. Распознанное содержимое:]" if multi
                 else "[Пользователь прислал изображение. Распознанное содержимое:]")
        blocks.append(label + "\n" + ocr[:2500])
    for entry in (doc_blocks or []):
        # doc_blocks entries are (name, content, token|None). Tolerate the legacy 2-tuple too.
        name, content = entry[0], entry[1]
        token = entry[2] if len(entry) > 2 else None
        cap = _B24_DOC_INLINE_CHARS
        body = content[:cap]
        head = f"[Пользователь прислал документ «{name}»"
        if token:
            head += f", attachment_id={token}"
        head += ". Извлечённое содержимое"
        if len(content) > cap:
            body += ("\n…[показано начало документа. ПОЛНЫЙ текст читай инструментом "
                     f"get_attachment_text(attachment_id='{token}', offset=…) — ничего не обрезано, "
                     "документ доступен целиком.]")
            head += " (начало; полный текст — get_attachment_text)"
        blocks.append(head + ":]\n" + body)
    # Tell the agent it can re-send/attach these exact files (screenshots, documents) to a task,
    # comment, or result — by passing their attachment_id to the task tools.
    forwardable = [a for a in (attachments or []) if a.get("token")]
    if forwardable:
        listing = "; ".join(f"{a['token']} — «{a['name']}» ({a['kind']})" for a in forwardable)
        blocks.append(
            "[ВЛОЖЕНИЯ от пользователя, которые ты можешь переслать/приложить как есть "
            "(передай attachment_id в create_bitrix_task, add_bitrix_task_comment или "
            "attach_files_to_task; полный текст документа — get_attachment_text): " + listing + "]"
        )
    if not text and (image_texts or doc_blocks):
        blocks.append("(Текста в сообщении не было — отвечай по содержимому вложения.)")
    return "\n\n".join(blocks) if blocks else text


def _b24_recent_history(dialog_id: str, limit: int = 6,
                        agent_slug: str | None = None) -> list[tuple[str, str]]:
    """Last successful Q/A pairs for THIS agent in this dialog — we inject them into the prompt
    because each turn runs in a FRESH hermes session (a unique per-run --continue name; hermes
    >=0.17 would otherwise resume the named session and double the memory), so the agent
    has no memory of its own. Filtered by agent_slug so an agent only recalls its OWN history
    (the Bitrix dialog_id is shared across bots); the floor comes from this agent's session row."""
    try:
        with pg_connect() as conn:
            with conn.cursor() as cur:
                cur.execute(
                    "SELECT question, answer FROM bitrix_bot_interactions "
                    "WHERE dialog_id=%s AND agent_slug IS NOT DISTINCT FROM %s::text "
                    "AND status='ok' AND question <> '' "
                    "AND id > COALESCE((SELECT history_floor_id FROM bitrix_bot_sessions WHERE dialog_id=%s), 0) "
                    "ORDER BY id DESC LIMIT %s",
                    (str(dialog_id), agent_slug, _b24_scope(dialog_id, agent_slug), int(limit)),
                )
                return [(r["question"], r["answer"]) for r in reversed(cur.fetchall())]
    except Exception:  # noqa: BLE001
        logging.exception("b24 testbot: history fetch failed")
        return []


# --- Brain-run guard rails: concurrency cap, one retry, owner alerts ------------------------
# Each turn spawns a separate `hermes` CLI process (~250MB) on a 2GB box, so an unbounded burst
# of simultaneous users would swap/OOM the whole server.
_HERMES_MAX_CONCURRENCY = max(1, int(os.getenv("B24_HERMES_MAX_CONCURRENCY", "3")))
_HERMES_QUEUE_WAIT_S = int(os.getenv("B24_HERMES_QUEUE_WAIT_S", "180"))
_HERMES_RUN_SLOTS = threading.BoundedSemaphore(_HERMES_MAX_CONCURRENCY)
# Live turn registry for cancellation: «Новая сессия» must STOP a running brain turn of that
# (agent, dialog) — kill its hermes subprocess — instead of letting it finish into a session
# the user already reset. Keyed by _b24_scope; the worker sees `cancelled` and drops the turn.
_LIVE_TURNS_LOCK = threading.Lock()
_LIVE_TURNS: dict[str, list[dict]] = {}


def _b24_hermes_popen_kwargs(cmd: list) -> dict[str, Any]:
    kwargs: dict[str, Any] = {
        "args": cmd,
        "stdout": subprocess.PIPE,
        "stderr": subprocess.PIPE,
        "text": True,
        "cwd": "/root",
        "env": {**os.environ, "HOME": "/root"},
    }
    if os.name == "posix":
        # Hermes can spawn children. Give the turn its own process group so reset/timeout
        # can kill the whole tree and not hang on inherited stdout/stderr pipes.
        kwargs["start_new_session"] = True
    return kwargs


def _b24_kill_process_tree(proc: Any) -> None:
    if proc is None or proc.poll() is not None:
        return
    try:
        if os.name == "posix":
            os.killpg(os.getpgid(proc.pid), getattr(signal, "SIGKILL", signal.SIGTERM))
        else:
            proc.kill()
    except OSError:
        try:
            proc.kill()
        except OSError:
            pass


def _b24_cancel_live_turns(scope: str) -> int:
    """Cancel every running brain turn of this scope: mark + kill the hermes subprocess.
    Returns how many turns were cancelled (0 = nothing was running)."""
    cancelled = 0
    with _LIVE_TURNS_LOCK:
        for entry in _LIVE_TURNS.get(scope, []):
            entry["cancelled"] = True
            proc = entry.get("proc")
            _b24_kill_process_tree(proc)
            cancelled += 1
    return cancelled


_OPS_ALERT_COOLDOWN_S = 600
_ops_alert_last_sent: dict[str, float] = {}
_ops_alert_lock = threading.Lock()


def _b24_ops_alert(kind: str, dialog_id: Any, tier: str, from_user_id: Any, detail: str) -> None:
    """Fire-and-forget Telegram alert to the Albery notifications group about a failed bot turn.
    Per-kind cooldown so one incident hitting many users does not flood the group."""
    now = time.monotonic()
    with _ops_alert_lock:
        if now - _ops_alert_last_sent.get(kind, -_OPS_ALERT_COOLDOWN_S) < _OPS_ALERT_COOLDOWN_S:
            return
        _ops_alert_last_sent[kind] = now
    text = (
        f"🚨 ИИ-агент (Bitrix): {kind}\n"
        f"Диалог {dialog_id}, tier={tier}, user={from_user_id}\n{detail}\n"
        f"(повторы этого типа ближайшие {_OPS_ALERT_COOLDOWN_S // 60} мин не дублируются; "
        f"детали: journalctl -u albery)"
    )

    def _do() -> None:
        ok, err = _albery_tg_notify(text)
        if not ok:
            logging.error("b24 testbot: ops alert delivery failed: %s", err)

    threading.Thread(target=_do, daemon=True).start()


# The hermes CLI prints its own failure notice to stdout AND exits rc=0 when an LLM call gives up
# (e.g. "API call failed after 3 retries. Connection error." / "[Errno 32] Broken pipe" / HTTP 429
# usage limit / context-window overflow). Left unchecked, b24bot posts that raw diagnostic straight
# to the employee — this is exactly what surfaced as the "AI-lawyer error". Detect it so the turn is
# retried and, if it still fails, the user sees a friendly message instead of the raw stack text.
_HERMES_ERROR_RE = re.compile(r"^\s*API call failed after\s+\d+\s+retr", re.IGNORECASE)


def _hermes_answer_is_error(answer: str) -> bool:
    return bool(answer) and bool(_HERMES_ERROR_RE.match(answer))


def _b24_brain_error_message(answer: str) -> str:
    low = (answer or "").lower()
    if "usage limit" in low or "429" in low:
        return ("Сейчас мой ИИ-мозг перегружен лимитами запросов 🙏 Дай минуту-другую и отправь "
                "сообщение ещё раз — я отвечу.")
    if "context window" in low or "exceeds the context" in low or "too long" in low:
        return ("Запрос получился слишком большим для одного сообщения 😅 Попробуй разбить его на "
                "части или начни новую сессию кнопкой «🆕 Новая сессия» и повтори покороче.")
    return ("Что-то временно сбоит на стороне ИИ 😔 Обычно это разовый сетевой сбой — попробуй, "
            "пожалуйста, повторить запрос через минуту.")


def _hermes_run_guarded(cmd: list, timeout_s: int, dialog_id: Any, tier: str,
                        from_user_id: Any, prompt_chars: int, scope: str = "",
                        retry_prompt_suffix: str = ""):
    """Run the hermes CLI under the concurrency semaphore, retrying once on a quick failure
    (non-zero rc / empty stdout / an LLM error sentinel printed as the answer). Returns
    (proc, None), (None, 'busy'), (None, 'timeout') or (None, 'cancelled') — the latter when
    «Новая сессия» killed this turn mid-flight.

    On the retry we (a) pause a short backoff so a transient provider blip (Broken pipe to the
    single Codex account, which is what kills heavy contract turns) can clear, and (b) optionally
    append `retry_prompt_suffix` to the -z prompt so the second attempt runs LEANER (fewer web
    pages, one export call) — heavy multi-step turns are exactly what breaks the connection."""
    backoff_s = float(os.getenv("B24_HERMES_RETRY_BACKOFF_S", "6") or "6")
    if not _HERMES_RUN_SLOTS.acquire(timeout=_HERMES_QUEUE_WAIT_S):
        logging.warning("b24 testbot: hermes slot wait exceeded %ss dialog_id=%s tier=%s user_id=%s",
                        _HERMES_QUEUE_WAIT_S, dialog_id, tier, from_user_id)
        _b24_ops_alert(
            "очередь переполнена", dialog_id, tier, from_user_id,
            f"Свободный слот не появился за {_HERMES_QUEUE_WAIT_S}с "
            f"(лимит {_HERMES_MAX_CONCURRENCY} одновременных прогонов).",
        )
        return None, "busy"
    entry: dict[str, Any] = {"proc": None, "cancelled": False}
    with _LIVE_TURNS_LOCK:
        _LIVE_TURNS.setdefault(scope, []).append(entry)
    try:
        proc = None
        for attempt in (1, 2):
            attempt_cmd = cmd
            if attempt == 2:
                # Give a transient provider blip time to clear, then rerun LEANER.
                if backoff_s > 0 and not entry["cancelled"]:
                    time.sleep(backoff_s)
                if retry_prompt_suffix:
                    attempt_cmd = list(cmd)
                    try:
                        zi = attempt_cmd.index("-z")
                        attempt_cmd[zi + 1] = str(attempt_cmd[zi + 1]) + retry_prompt_suffix
                    except (ValueError, IndexError):
                        pass
            # Spawn under the lock so a concurrent cancel can't slip between check and start.
            with _LIVE_TURNS_LOCK:
                if entry["cancelled"]:
                    return None, "cancelled"
                child = subprocess.Popen(**_b24_hermes_popen_kwargs(attempt_cmd))
                entry["proc"] = child
            try:
                out, err = child.communicate(timeout=timeout_s)
            except subprocess.TimeoutExpired:
                _b24_kill_process_tree(child)
                child.communicate()
                if entry["cancelled"]:
                    return None, "cancelled"
                logging.warning(
                    "b24 testbot: hermes timed out after %ss dialog_id=%s tier=%s user_id=%s prompt_chars=%s",
                    timeout_s, dialog_id, tier, from_user_id, prompt_chars,
                )
                _b24_ops_alert("таймаут хода", dialog_id, tier, from_user_id,
                               f"Мозг не ответил за {timeout_s}с; пользователь получил вежливый отказ.")
                return None, "timeout"
            if entry["cancelled"]:
                return None, "cancelled"
            proc = subprocess.CompletedProcess(cmd, child.returncode, out or "", err or "")
            if (proc.returncode == 0 and (proc.stdout or "").strip()
                    and not _hermes_answer_is_error(proc.stdout.strip())):
                return proc, None
            logging.error("b24 testbot: hermes run failed (attempt %s/2): rc=%s err=%s answer=%s",
                          attempt, proc.returncode, (proc.stderr or "")[:200],
                          (proc.stdout or "")[:120])
        return proc, None  # both attempts bad -> caller reports the empty answer
    finally:
        with _LIVE_TURNS_LOCK:
            entries = _LIVE_TURNS.get(scope, [])
            if entry in entries:
                entries.remove(entry)
            if not entries:
                _LIVE_TURNS.pop(scope, None)
        _HERMES_RUN_SLOTS.release()


def hermes_brain_answer(user_text: str, dialog_id: str, tier: str = "faq", from_user_id: Any = "",
                        agent: dict[str, Any] | None = None) -> str | None:
    """Run one turn through the local Hermes brain. Toolset is chosen by access tier:
    admin → full MCP `albery` (incl. instruction/settings edits + deletes); ops → `albery-ops`
    (operational, no admin tools); everyone else → read-only `albery-faq`. A subagent turn
    (`agent` profile from agent_center) uses the agent's OWN connector `agent-<slug>` —
    its tier toolset plus personal self-learning tools, scoped to that agent only.
    Session lifecycle (idle reset + cap rotation + carried summary) is handled by
    _b24_session_prepare. All session/history keying is scoped to this agent (agent_slug)."""
    agent_slug = (agent or {}).get("slug")
    session, seed = _b24_session_prepare(dialog_id, agent_slug)
    toolset = {"admin": "albery", "ops": "albery-ops"}.get(tier, "albery-faq")
    core_toolset = tier in ("admin", "ops") and os.getenv("B24_CORE_TOOLSET", "").strip() == "1"
    if core_toolset:
        # Two-stage tools: the bot registers a curated core + find_tool/call_tool (fast turns,
        # small context); the full connectors stay untouched for cron agents.
        toolset = {"admin": "albery-core", "ops": "albery-ops-core"}[tier]
    # Universal (main) agent: when enabled, the main bot runs on ONE configurable набор
    # (its own /mcp-agent/main connector) for everyone allowed, instead of per-user tiers —
    # capped at ops (no admin). Guarded + falls back to the classic connectors if not ready.
    if agent is None:
        try:
            from agent_center import universal_main_connector
            universal = universal_main_connector()
        except Exception:  # noqa: BLE001
            universal = None
        if universal:
            tier = "ops"
            core_toolset = False
            toolset = universal
    if agent is not None:
        toolset = f"agent-{agent['slug']}"
    timeout_s = int(os.getenv("B24_TESTBOT_HERMES_TIMEOUT", "170"))
    fmt = (
        " СТИЛЬ И ФОРМАТ ОТВЕТА (важно): пиши КРАТКО и по делу, но КРАСИВО и удобно для чтения — "
        "глазу должно быть приятно. Выделяй ГЛАВНОЕ жирным через [b]...[/b] (ключевые слова, итоги, "
        "важные имена/цифры). Разбивай ответ на короткие абзацы по смыслу с пустой строкой между "
        "ними; перечисления оформляй списком — каждый пункт с новой строки и с «- » в начале. "
        "Уместно используй эмодзи, чтобы оживить ответ (1–2, иногда 3 — но без перебора и не в каждой "
        "строке). БЕЗ воды: не пиши вводных («Понимаю вопрос», «Отличный вопрос»), не повторяй вопрос, "
        "не добавляй финальных резюме «То есть коротко…». Не используй Markdown (#, **, ` или таблицы) "
        "— Битрикс их не отображает; жирный ТОЛЬКО через [b]...[/b]. Итог: коротко по содержанию, но с "
        "аккуратным оформлением, акцентами жирным и живыми эмодзи."
        " ФАЙЛ-ОТВЕТ: если нужен ОФИЦИАЛЬНЫЙ ДОКУМЕНТ Word (договор, соглашение, акт, оферта, "
        "официальное письмо) и в твоих инструментах есть export_document — используй ЕГО: собери "
        "полный HTML документа (оформление целиком под твоим контролем: разрывы страниц, реквизиты, "
        "выравнивание; без эмодзи и [b]-кодов) и пришли пользователю ссылку из ответа инструмента. "
        "Для остальных файлов: если пользователь просит ответ ДОКУМЕНТОМ/ФАЙЛОМ (\"пришли PDF\", "
        "\"оформи в PDF\", \"сделай файл\"), сформируй ПОЛНОЕ содержимое документа прямо в ответе "
        "(заголовки [b]…[/b], списки \"- \", абзацы) и в САМОМ конце добавь служебный маркер РОВНО так: "
        "[[DELIVER_PDF: Имя]] — PDF; [[DELIVER_XLSX: Имя]] — Excel; [[DELIVER_DOCX: Имя]] — Word. "
        "Формат выбирай по просьбе пользователя: «эксель/xlsx/таблицей» → XLSX, «ворд/docx» → DOCX, иначе PDF. "
        "Если данные табличные (задачи, отчёты, колонки) — В ФАЙЛЕ оформляй их Markdown-таблицей через «|»: "
        "строка заголовков, затем строка «| --- | --- |», затем строки данных; система превратит это в "
        "НАСТОЯЩУЮ таблицу в PDF/Excel/Word. Реальные данные бери из инструментов (задачи — search_tasks и "
        "т.п.), НЕ выдумывай содержимое. Система сама соберёт файл и пришлёт ссылку; маркер "
        "пользователь не увидит. Без явной просьбы о файле маркер НЕ добавляй."
    )
    access_rule = (
        " ПРАВИЛО ПРО ДОСТУП (обязательно): если пользователь просит ДЕЙСТВИЕ, которого нет в твоём "
        "текущем уровне доступа (например, самому зайти и отредактировать Google-таблицу, поставить "
        "задачу, отправить сообщение), НЕ юли, НЕ говори обтекаемо и НЕ предлагай обходные пути ВМЕСТО "
        "прямого ответа. Ответь чётко и коротко по образцу: «В целом я это умею, но в вашем текущем "
        "уровне доступа это недоступно. Доступ выдаёт Александр Никитенко — хотите, я передам ему ваш "
        "запрос? 🙌». Ответственный за доступ — ТОЛЬКО Александр Никитенко; НИКОГДА не упоминай других "
        "людей (в том числе Евгения). Если пользователь соглашается передать запрос — ответь «Готово, "
        "передал Александру 🙌» и в САМОМ конце добавь скрытый служебный маркер РОВНО так: "
        "[[ESCALATE: краткая суть запроса]] (пользователь его не увидит, система сама уберёт и отправит "
        "выжимку Александру). Если в обучении показаны умения, которых нет в твоём доступе — поясни, что "
        "обучение показывает возможности в целом, а доступно именно этому пользователю — по его уровню "
        "доступа. "
        "ДОСТУП К GOOGLE-ОБЪЕКТАМ (КРИТИЧНО): любую Google-таблицу, документ, папку, Apps Script или "
        "веб-приложение, которые ты СОЗДАЁШЬ или ссылку на которые ДАЁШЬ человеку, ОБЯЗАТЕЛЬНО открывай "
        "«для всех по ссылке» (редактор). НИКОГДА не присылай ссылку на объект, к которому нет доступа по "
        "ссылке — иначе человек видит «Нет доступа». Таблицы создавай через create_google_sheet (он сразу "
        "открывает доступ); для ЛЮБОГО другого объекта — включая таблицу, которую создало само "
        "приложение/скрипт, — вызови share_drive_item_for_everyone(item=<id или url>) ПЕРЕД тем как дать "
        "ссылку. Если пишешь Apps Script, создающий файл/таблицу — добавь в код открытие доступа: "
        "DriveApp.getFileById(id).setSharing(DriveApp.Access.ANYONE_WITH_LINK, DriveApp.Permission.EDIT). "
        "ВЕБ-ПРИЛОЖЕНИЯ НА ТАБЛИЦАХ (чтобы открывались У ВСЕХ без входа Google): НЕЛЬЗЯ обращаться к "
        "SpreadsheetApp/DriveApp в коде Apps Script — тогда Google требует вход и отдаёт 403 анонимам. "
        "Делай так: doGet возвращает ТОЛЬКО HTML/JS; для данных вызови инструмент "
        "make_sheet_applet(spreadsheet_id=<id таблицы>) и встрой его html_snippet/applet_url в страницу "
        "(JS читает строки appletRows() и добавляет appletAdd([...]) через Albery — это работает "
        "анонимно, без авторизации Google). Публикуй ТОЛЬКО ANYONE_ANONYMOUS, никогда не ANYONE. ДИЗАЙН (важно): ВСЕГДА бери основу веб-приложения "
        "через инструмент get_webapp_template — фирменный стиль Albery (светлый фон, белые карточки, фиолетовый #5440F6, шрифт Inter) — и собирай интерфейс на его CSS-классах "
        "(.card/.btn-primary/.input/.field/table/.badge/.stat); приложения должны выглядеть единообразно и красиво, как наш прод-сайт, а не случайно."
    )
    if agent is not None:
        role = (agent.get("role_prompt") or "").strip() or "специализированный помощник компании"
        head = ("[Канал: Битрикс24. Ты — специализированный агент «" + str(agent.get("name") or "Агент")
                + "». ТВОЯ РОЛЬ: " + role + " Работай СТРОГО в рамках своей роли: по вопросам вне её "
                "вежливо скажи, что этим занимается Основной агент Албери, и не пытайся выполнить сам. "
                "Отвечай по-русски." + access_rule + fmt + "]")
    elif tier == "admin":
        head = ("[Канал: Битрикс24. Уровень доступа пользователя: ПОЛНЫЙ — доступны все инструменты, "
                "включая изменение настроек/инструкций. Любое изменение данных сначала подтверждай. "
                "Отвечай по-русски." + access_rule + fmt + "]")
    elif tier == "ops":
        head = ("[Канал: Битрикс24. Уровень доступа пользователя: ВСЕ ФУНКЦИИ — можешь искать/ставить/"
                "закрывать задачи, готовить отчёты, писать сотрудникам, а также СОЗДАВАТЬ и "
                "РЕДАКТИРОВАТЬ Google-таблицы и Google-документы (вносить данные, формулы, оформление, "
                "автоматизацию — используй инструменты и навык работы с Google Sheets через Google "
                "Sheets/Apps Script API). Любое изменение данных сначала подтверждай. Недоступно "
                "только изменение собственных настроек/инструкций. Отвечай по-русски."
                + access_rule + fmt + "]")
    else:  # faq
        head = ("[Канал: Битрикс24. Уровень доступа пользователя: СПРАВКА (только знания и чтение). "
                "Тебе доступно: отвечать на вопросы по компании, регламентам, оргструктуре, базе знаний, "
                "разбирать Zoom-созвоны. Тебе НЕдоступны любые ДЕЙСТВИЯ: редактировать Google-таблицы, "
                "ставить/закрывать задачи, отправлять сообщения, что-либо менять в системах. "
                "Отвечай по-русски." + access_rule + fmt + "]")
    parts = [head]
    # Self-orientation, same for every agent: how to see what it can do and act fast + correctly.
    # Keeps agents from either faking a capability they lack or refusing one they have. The model
    # natively sees its enabled tools (the connector serves tools/list) — this tells it to trust
    # that set as the source of truth and gives a tight operating procedure.
    parts.append(
        "КАК ТЫ РАБОТАЕШЬ (ориентируйся по этому, отвечай быстро и точно): "
        "1) ТВОИ ВОЗМОЖНОСТИ — это РОВНО те инструменты, что доступны тебе в этом ходе (система "
        "показывает их как доступные функции), плюс перечисленные ниже подключённые навыки и "
        "инструкции. Это и есть точный список того, что ты умеешь. "
        "2) ПОРЯДОК: пойми, что нужно человеку → выбери подходящий инструмент из своего набора → "
        "выполни за МИНИМУМ шагов. Не делай пробных вызовов без обязательных полей и не "
        "переспрашивай лишний раз: недостающее (исполнитель, срок, критерий и т.п.) собери ОДНИМ "
        "уточняющим сообщением, затем действуй. "
        "3) ЕСЛИ НУЖНОГО ИНСТРУМЕНТА ИЛИ НАВЫКА У ТЕБЯ НЕТ — не выдумывай и не имитируй работу: "
        "коротко и честно скажи, что именно вне твоих возможностей, и направь — профильные вопросы "
        "не по твоей роли ведёт Основной агент Албери, вопросы доступа/настроек решает Александр "
        "Никитенко. "
        "4) Прежде чем сказать «не могу/не помню» — проверь свои инструменты и историю диалога: "
        "часто нужное уже под рукой."
    )
    if agent is not None:
        # Skills the owner connected to THIS subagent (manifest). Injecting only the
        # selected ones is what enforces the selection at the prompt level. (Instructions
        # are delivered through the scoped start_here tool, not injected here again.)
        try:
            from agent_center import agent_selected_knowledge
            selected = agent_selected_knowledge(agent)
        except Exception:  # noqa: BLE001
            logging.exception("subagent %s: selected knowledge load failed", agent.get("slug"))
            selected = {"instructions": [], "skills": []}
        if selected["skills"]:
            parts.append(
                "ТВОИ НАВЫКИ (подключены владельцем): "
                + "; ".join(f"«{s['title']}» — {s['description']}" for s in selected["skills"])
                + ". Пользуйся ТОЛЬКО этими навыками. Если для задачи нужен навык, которого нет в "
                "этом списке, не выдумывай и не применяй сторонние приёмы — скажи, что этим "
                "занимается Основной агент Албери."
            )
            for s in selected["skills"]:
                if s.get("content"):
                    parts.append("ПОЛНЫЙ ТЕКСТ НАВЫКА «" + s["title"] + "» — следуй ему буквально:\n"
                                 + s["content"])
    # Personal instructions + self-learning apply to ANY agent that runs on its own
    # connector: subagents AND the universal (main) agent when it is on agent-main. The
    # main bot keeps its general head; here we only add its accumulated skills + the
    # "ask after a good task, then form a skill" nudge.
    learning_agent = agent
    if agent is None and universal:
        try:
            from agent_center import MAIN_AGENT_SLUG, _agent_by_slug
            learning_agent = _agent_by_slug(MAIN_AGENT_SLUG)
        except Exception:  # noqa: BLE001
            logging.exception("main agent: profile load for self-learning failed")
            learning_agent = None
    if learning_agent is not None:
        learned = learning_agent.get("instructions") or []
        if learned:
            parts.append(
                "ТВОИ ЛИЧНЫЕ ИНСТРУКЦИИ И НАВЫКИ (накоплены обучением, применяй обязательно):\n"
                + "\n\n".join(f"— {i['name']}:\n{i['content']}" for i in learned)
            )
        parts.append(
            "САМООБУЧЕНИЕ И ОБРАТНАЯ СВЯЗЬ (важно): когда ты ЗАВЕРШИЛ ощутимую задачу для человека "
            "(подготовил отчёт/файл, решил вопрос, что-то настроил) и работа явно удалась — КОРОТКО "
            "спроси, всё ли устроило, и предложи закрепить это как правило: «Всё ок? 🙌 Оформить это "
            "как постоянное правило работы, чтобы делать так и дальше?». НЕ спрашивай в каждом "
            "сообщении и не по мелочам — только после реально завершённой задачи и не чаще одного раза "
            "за тему. Если человек подтвердил, что понравилось, — выдели из этой работы МНОГОРАЗОВОЕ "
            "правило/приём (формат, порядок действий, предпочтение команды) и сохрани его себе через "
            "upsert_my_instruction (коротко и по делу) — так ты формируешь новый навык из удачной "
            "работы. Устаревшую свою инструкцию обнови тем же инструментом или удали через "
            "delete_my_instruction. НЕ сохраняй разовые факты, персональные данные и содержимое "
            "конкретных диалогов. Это ТВОИ личные навыки — глобальные правила и других агентов ты "
            "изменить не можешь."
        )
        parts.append(
            "АВТОМАТИЗАЦИИ ПО РАСПИСАНИЮ: если сотрудник просит делать что-то РЕГУЛЯРНО («каждый "
            "день/неделю», «по расписанию», «присылай сводку», «напоминай») — настрой это инструментом "
            "schedule_my_automation (cron, время МСК; deliver_to='" + str(dialog_id) + "' — текущий "
            "диалог; requested_by — ИМЯ собеседника, который просит: владелец видит, кто поставил). "
            "СНАЧАЛА проверь свои инструменты: если их для задачи не хватает — ЧЕСТНО скажи, "
            "чего именно не хватает, и автоматизацию НЕ создавай. Перед созданием подтверди у "
            "пользователя расписание и суть задачи. Посмотреть/удалить: list_my_automations / "
            "delete_my_automation (поставленные владельцем не удаляй — это делает владелец в приложении)."
        )
    parts.append(
        "Текущие дата и время: " + msk_now().strftime("%d.%m.%Y %H:%M")
        + " МСК (Europe/Moscow) — это «сегодня/сейчас» для любых расчётов сроков и дат."
    )
    if core_toolset:
        parts.append(
            "ИНСТРУМЕНТЫ — ДВУХСТУПЕНЧАТАЯ СХЕМА (важно): в твоём списке — ядро самых нужных "
            "инструментов. Если нужного действия в списке НЕТ — не отвечай «не умею/нет доступа»: "
            "сначала вызови find_tool (query — короткие английские ключевые слова, например "
            "'delete task', 'zoom report', 'drive folder'), возьми из результата точное имя и "
            "схему аргументов и выполни действие через call_tool(name=..., arguments={...})."
        )
    if tier in ("admin", "ops"):
        parts.append(
            "ПАМЯТЬ И КОНТЕКСТ (важно). Это диалог Битрикса dialog_id=`" + str(dialog_id) + "`. У тебя "
            "ЕСТЬ полный доступ ко всей твоей прошлой работе в ЭТОМ диалоге, включая ПРОШЛЫЕ и уже "
            "сброшенные сессии. Инструменты: get_bitrix_bot_chat(dialog_id='" + str(dialog_id) + "') — "
            "возвращает полный транскрипт твоих прошлых шагов (что ты делал, какие ссылки и "
            "идентификаторы создавал: script_id, web_app_url, editor_url, id таблиц); "
            "list_bitrix_bot_sessions — обзор сессий этого диалога. "
            "ЖЕЛЕЗНОЕ ПРАВИЛО: прежде чем сказать «не помню», «пришлите ссылку/script_id», «повторите», "
            "или прежде чем создавать заново то, что ты, возможно, уже делал, — СНАЧАЛА вызови "
            "get_bitrix_bot_chat и найди нужное в своей истории. Не перекладывай на пользователя то, что "
            "можешь вспомнить сам. "
            "Когда создаёшь/правишь Apps Script приложение или таблицу — ВСЕГДА явно указывай в ответе "
            "его web_app_url, script_id И editor_url (ссылку на редактор), чтобы в будущем легко найти и "
            "переопубликовать именно его (а не плодить новые копии). Проверяя доступность ссылки, бери "
            "АКТУАЛЬНЫЙ web_app_url из своей истории, а не угадывай."
        )
    # Document + forwarded-message recall — applies to EVERY agent (incl. tier=faq юрист), so the
    # agent can always reach earlier context, including Word/PDF sent in a past/reset session.
    if agent is not None or tier in ("admin", "ops"):
        parts.append(
            "ПАМЯТЬ О ДОКУМЕНТАХ И ПЕРЕСЛАННЫХ СООБЩЕНИЯХ (важно). Это диалог dialog_id=`"
            + str(dialog_id) + "`. Ты можешь поднять ВСЁ, что было в этом диалоге раньше — включая "
            "прошлые и уже сброшенные сессии — инструментом get_bitrix_bot_chat(dialog_id='"
            + str(dialog_id) + "'). Если пользователь ссылается на прежний документ/сообщение, "
            "ПЕРЕСЫЛАЕТ или ОТВЕЧАЕТ на старое сообщение (в т.ч. в новой сессии), а ты не видишь его "
            "полностью — СНАЧАЛА вызови get_bitrix_bot_chat и найди нужный момент. У КАЖДОГО присланного "
            "ранее файла (Word/PDF/Excel/скан) в истории есть его attachment_id (att_…): прочитай "
            "документ ЦЕЛИКОМ через get_attachment_text(attachment_id='att_…', offset=…), по частям, "
            "пока has_more не станет false. Если файл прикреплён к текущему сообщению заново — ты видишь "
            "его сразу; если он только процитирован/переслан — подними его из истории по attachment_id. "
            "НИКОГДА не проси прислать заново то, что уже было в этом диалоге."
        )
    if tier in ("ops", "admin") and str(from_user_id).strip():
        parts.append(
            "ПОСТАНОВЩИК ЗАДАЧ (важно): по умолчанию постановщик создаваемой задачи — "
            "ТЕКУЩИЙ собеседник, его Bitrix id=" + str(from_user_id) + ". При вызове "
            "create_bitrix_task передавай creator_bitrix_user_id=" + str(from_user_id) + ", "
            "КРОМЕ случая, когда пользователь явно просит сделать постановщиком другого "
            "человека — тогда укажи creator_name этого человека. У каждой задачи "
            "результат обязателен (завершить без результата нельзя). "
            "СРОК — ПРОВЕРКА НА АДЕКВАТНОСТЬ (важно): перед постановкой сравни срок (deadline) с "
            "текущими датой/временем МСК (см. выше). Если срок уже В ПРОШЛОМ (раньше «сейчас») — НЕ "
            "ставь задачу молча: подсвети пользователю «Срок <дата> уже прошёл (сегодня <сегодня>). "
            "Поставить задачу с этим сроком как есть или укажете новый?» и дождись ответа. Если "
            "пользователь подтвердил («да», «как есть», «ставь так») — создавай задачу с ТЕМ ЖЕ "
            "сроком; если дал новый срок — ставь с новым; если срок сегодня или в будущем — действуй "
            "как обычно. Подтверждение «как есть» = повтори create_bitrix_task с "
            "confirm_past_deadline=true (без этого флага инструмент НЕ поставит прошедший срок). "
            "Никогда не меняй прошедший срок сам — только по решению пользователя. "
            "РЕЗУЛЬТАТ ОБЯЗАТЕЛЕН У КАЖДОЙ ЗАДАЧИ (важно): всегда уточняй у пользователя, какой результат/"
            "критерий выполнения у задачи — по чему поймём, что сделано, и чем подтверждается (скрин/"
            "ссылка/файл). Передавай это в параметр result_criteria. Если пользователь результат не "
            "назвал — СПРОСИ, не выдумывай и не ставь задачу без него (инструмент откажет). Лучше "
            "спрашивай срок и результат вместе, одним сообщением. "
            "СКОРОСТЬ (важно): ставь задачу за МИНИМУМ шагов. Собери всё недостающее (исполнитель, "
            "срок, результат) за ОДНО уточнение, и при подтверждении вызови create_bitrix_task ОДИН "
            "раз со ВСЕМИ полями сразу: title, responsible_name/responsible_bitrix_user_id, deadline, "
            "result_criteria, creator_bitrix_user_id. НЕ делай пробных вызовов create_bitrix_task без "
            "обязательных полей (получишь отказ и потеряешь целый ход). НЕ вызывай get_employee_absences "
            "и прочие инструменты при обычной постановке — только если пользователь сам упомянул отпуск/"
            "занятость. Не переспрашивай и не перепроверяй лишний раз — действуй сразу, одним вызовом. "
            "ВЕДЕНИЕ ЗАДАЧ ОТ ЛИЦА СОБЕСЕДНИКА (важно): по умолчанию всё, что ты делаешь в задаче, — "
            "ОТ ИМЕНИ текущего собеседника (id=" + str(from_user_id) + "). "
            "• Комментарий: add_bitrix_task_comment(bitrix_task_id, comment_text, "
            "author_bitrix_user_id=" + str(from_user_id) + ") — комментарий покажется от его имени. "
            "Другого автора ставь ТОЛЬКО если он явно просит прокомментировать от другого человека. "
            "• Завершить задачу: complete_bitrix_task(bitrix_task_id, on_behalf_bitrix_user_id=" + str(from_user_id) + "); "
            "если у задачи обязателен результат — передай result_text и/или attachment_ids (файл-подтверждение). "
            "• Отметить результат / приложить скрин как результат: add_bitrix_task_comment(..., as_result=true, "
            "attachment_ids=[...]) или attach_files_to_task(bitrix_task_id, attachment_ids=[...], as_result=true). "
            "• Возобновить задачу: reopen_bitrix_task(bitrix_task_id, reason, confirm=true); с новым сроком — "
            "добавь new_deadline. "
            "ВЛОЖЕНИЯ (важно): файлы, которые прислал пользователь, приходят тебе с токенами attachment_id "
            "(att_…) в блоке [ВЛОЖЕНИЯ …]. Чтобы переслать/приложить их — передавай эти attachment_id в "
            "create_bitrix_task, add_bitrix_task_comment, complete_bitrix_task или attach_files_to_task "
            "(и скрин, и документ). Чтобы прочитать ДЛИННЫЙ документ целиком (договор и т.п.) — вызывай "
            "get_attachment_text(attachment_id=..., offset=...) по частям, пока has_more не станет false; "
            "в промпте показан только предпросмотр, полный текст бери инструментом, ничего не обрезано. "
            "Правку документа отдавай обратно через export_document (свой HTML → docx)."
        )
    if seed:
        parts.append("Сводка более ранней части разговора: " + seed)
    history = _b24_recent_history(dialog_id, int(os.getenv("B24_TESTBOT_HISTORY_TURNS", "10")), agent_slug)
    if history:
        # Long answers (contracts, tables) would balloon the prompt of every following turn —
        # clip each history item; the rolling session summary carries the older context anyway.
        def _clip(text: str, cap: int) -> str:
            return text if len(text) <= cap else text[:cap] + " …[обрезано]"
        convo = "\n".join(f"Пользователь: {_clip(q, 500)}\nАссистент: {_clip(a, 1500)}" for q, a in history)
        parts.append("История этого диалога (предыдущие реплики, помни их):\n" + convo)
    parts.append("Текущее сообщение пользователя:\n" + user_text)
    prompt = "\n\n".join(parts)
    # Hermes >=0.17 actually RESUMES a named --continue session (0.14 one-shots never did).
    # Our only memory channel is the prompt-injected history above; a resumed session would
    # duplicate it AND replay every past tool result into each turn's context (compression is
    # disabled on this box). A unique per-run suffix keeps one fresh session per turn — the
    # pre-0.17 behaviour — while the epoch name stays as a searchable prefix in state.db.
    run_session = f"{session}-r{uuid.uuid4().hex[:8]}"
    # Every agent gets the internet by default: hermes's built-in `web` toolset (search + page
    # fetch) rides along with the agent's own MCP connector. `-t` REPLACES the toolset list, so
    # without this agents are cut off from the web entirely. Terminal/file/exec stay OFF —
    # the isolation model is unchanged.
    extra_toolsets = os.getenv("B24_EXTRA_TOOLSETS", "web").strip().strip(",")
    toolset_arg = f"{toolset},{extra_toolsets}" if extra_toolsets else toolset
    cmd = ["hermes", "-z", prompt, "--continue", run_session, "-t", toolset_arg, "--yolo"]
    # If the first attempt fails (typically a Broken pipe to the single Codex account on a HEAVY
    # turn — big document HTML + lots of web browsing), the retry runs leaner so it actually
    # completes: one short web lookup at most, assemble the file in ONE export_document call.
    retry_lean = (
        "\n\n[СИСТЕМА: предыдущая попытка не завершилась. Доведи ответ до конца надёжно. Если "
        "готовишь документ Word — собирай его ИНКРЕМЕНТАЛЬНО по частям (export_document с "
        "section/doc_token/finalize), НЕ одним огромным вызовом: один большой вывод обрывает связь. "
        "Каждая секция небольшая. Качество и полнота остаются прежними.]"
    )
    proc, run_fail = _hermes_run_guarded(cmd, timeout_s, dialog_id, tier, from_user_id, len(prompt),
                                         scope=_b24_scope(dialog_id, agent_slug),
                                         retry_prompt_suffix=retry_lean)
    if run_fail == "cancelled":
        return None  # «Новая сессия» остановила этот ход — reset уже ответил пользователю
    if run_fail == "busy":
        return ("Сейчас я обрабатываю много запросов одновременно 🙏 Подожди минуту-другую и "
                "отправь сообщение ещё раз — я отвечу.")
    if run_fail == "timeout":
        sheets_like = re.search(
            "(google|гугл|таблиц|spreadsheet|sheet|лист|артикул|баркод|barcode|размер|остатк|заказ)",
            user_text or "",
            re.IGNORECASE,
        )
        if sheets_like:
            return (
                ("[b]Операция с Google-таблицами превысила лимит %s сек.[/b]\\n\\n" % timeout_s)
                + "Это не проблема формулировки: задача слишком большая для одного синхронного сообщения "
                + "в Битриксе (много строк/листов/формул или долгие Google API-вызовы). Я не буду "
                + "писать, что файл готов, пока он реально не создан.\\n\\n"
                + "Нужно запускать такую задачу пакетно/в фоне или сначала собрать тест на 1-2 артикула, "
                + "проверить результат и затем масштабировать на остальные."
            )
        return (
            ("[b]Операция превысила лимит %s сек.[/b]\\n\\n" % timeout_s)
            + "Я остановил ожидание, чтобы не держать чат бесконечно. Если это длинная операция, ее нужно "
            + "запускать пакетно/в фоне или разбить на короткие этапы."
        )
    answer = (proc.stdout or "").strip()
    if not answer:
        logging.error("hermes brain empty after retry: rc=%s err=%s",
                      proc.returncode, (proc.stderr or "")[:300])
        _b24_ops_alert("пустой ответ мозга", dialog_id, tier, from_user_id,
                       f"Два прогона подряд без ответа (rc={proc.returncode}).")
        return "Мозг временно недоступен, попробуй ещё раз чуть позже."
    if _hermes_answer_is_error(answer):
        # Both attempts came back as the CLI's own LLM-failure notice (transient network / Broken
        # pipe / limit / context overflow). Never post that raw text to the employee.
        logging.error("hermes brain LLM error sentinel dialog_id=%s tier=%s: %s",
                      dialog_id, tier, answer[:200])
        _b24_ops_alert("ошибка LLM в ходе", dialog_id, tier, from_user_id, answer[:300])
        return _b24_brain_error_message(answer)
    return answer


def _b24_log_interaction(dialog_id: str, from_user_id: Any, tier: str, question: str, answer: str,
                         latency_ms: int, status: str, error: str | None,
                         agent_slug: str | None = None) -> None:
    """Best-effort analytics row (never breaks the reply path)."""
    try:
        session = "bitrix-" + re.sub(r"[^A-Za-z0-9_-]", "", str(dialog_id))[:40]
        with pg_connect() as conn:
            with conn.transaction():
                with conn.cursor() as cur:
                    cur.execute(
                        """
                        INSERT INTO bitrix_bot_interactions
                          (dialog_id, bitrix_user_id, tier, session_name, question, answer, latency_ms, status, error, agent_slug)
                        VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                        """,
                        (str(dialog_id), to_int(from_user_id), tier, session,
                         question, answer, latency_ms, status, error, agent_slug),
                    )
    except Exception:  # noqa: BLE001
        logging.exception("b24 testbot: interaction log failed")


# --- Access-upgrade escalation: when the model marks a reply with [[ESCALATE: <суть>]] (after the
# user agrees to forward an access request), we strip the marker and notify the owner in Telegram.
_B24_ESCALATE_RE = re.compile(r"\[\[\s*ESCALATE\s*:?\s*(.*?)\]\]", re.IGNORECASE | re.DOTALL)


def _b24_extract_escalation(answer: str) -> tuple[str, str | None]:
    """Strip a [[ESCALATE: ...]] marker from the reply; return (clean_text, request_summary|None)."""
    match = _B24_ESCALATE_RE.search(answer or "")
    if not match:
        return answer, None
    summary = (match.group(1) or "").strip()
    clean = _B24_ESCALATE_RE.sub("", answer).strip()
    return clean, summary


def _b24_log_access_request(dialog_id: str, from_user_id: Any, requester_name: str,
                            request_text: str, delivered: bool, delivery_error: str | None) -> None:
    try:
        with pg_connect() as conn:
            with conn.transaction():
                with conn.cursor() as cur:
                    cur.execute(
                        "INSERT INTO access_requests (dialog_id, bitrix_user_id, requester_name, "
                        "request_text, delivered, delivery_error) VALUES (%s, %s, %s, %s, %s, %s)",
                        (str(dialog_id), to_int(from_user_id), requester_name or None,
                         request_text, bool(delivered), delivery_error),
                    )
    except Exception:  # noqa: BLE001
        logging.exception("b24 testbot: access-request log failed")


def _b24_forward_access_request(dialog_id: str, from_user_id: Any, request_text: str) -> None:
    """Notify the owner (Telegram) that a user requests more access, and log it for the record."""
    name = _b24_requester_name(from_user_id)
    text = (f"🔐 {name} просит расширить доступ к ИИ-агенту.\n\n"
            f"Запрос: {request_text}\n\n"
            f"Сейчас прав не хватает — уровень доступа можно изменить в «Настройки Агента».")
    ok, err = _albery_tg_notify(text, os.getenv("ALBERY_ACCESS_REQUEST_TG_CHAT", "").strip() or None)
    _b24_log_access_request(dialog_id, from_user_id, name, request_text, ok, err)
    if not ok:
        logging.error("b24 testbot: access-request TG delivery failed: %s", err)


# --- Live progress message: one bot message edited in place while the brain works ------------
# Toggled by B24_STATUS_MESSAGE=1. Every call is best-effort: a failed status update must never
# break the answer path.
import random

_B24_STATUS_GREETING = [
    "👋 Принял! Уже смотрю...",
    "👌 Взял в работу, секунду...",
    "🫡 Есть! Уже разбираюсь...",
    "⚡ Поймал запрос, приступаю...",
]

# What the agent is ACTUALLY doing right now: canonical MCP tool name -> human phrase.
# Fed live by mcp.context_server.recent_core_tool_calls (same Flask process).
_B24_TOOL_STATUS = {
    "search_tasks": "🔍 Ищу по задачам в Bitrix...",
    "get_task_comments": "💬 Читаю комментарии к задаче...",
    "add_bitrix_task_comment": "💬 Пишу комментарий в задаче...",
    "create_bitrix_task": "🎯 Ставлю задачу в Bitrix...",
    "reopen_bitrix_task": "🔁 Возобновляю задачу...",
    "delete_bitrix_task": "🗑 Удаляю задачу...",
    "search_company_knowledge": "📚 Ищу в базе знаний компании...",
    "list_company_files": "🗂 Просматриваю документы компании...",
    "get_company_file": "📖 Читаю документ из базы знаний...",
    "get_org_structure": "🧑‍🤝‍🧑 Сверяюсь с оргструктурой...",
    "get_employee_absences": "🏖 Проверяю график отсутствий...",
    "list_zoom_calls": "🎥 Просматриваю список созвонов...",
    "get_zoom_call_transcript": "🎧 Читаю транскрипт созвона...",
    "search_zoom_transcripts": "🎥 Ищу по транскриптам созвонов...",
    "get_bitrix_bot_chat": "🧠 Вспоминаю нашу переписку...",
    "list_bitrix_bot_sessions": "🧠 Поднимаю историю диалога...",
    "send_bitrix_message": "✉️ Отправляю сообщение...",
    "fetch_url": "🌐 Открываю ссылку...",
    "create_google_sheet": "📊 Создаю Google-таблицу...",
    "get_google_sheet_meta": "📊 Открываю таблицу, смотрю листы...",
    "write_google_sheet_values": "✏️ Вношу данные в таблицу...",
    "write_company_sheet": "✏️ Вношу данные в таблицу...",
    "format_google_sheet": "🎨 Навожу красоту в таблице...",
    "share_drive_item_for_everyone": "🔗 Открываю доступ по ссылке...",
    "move_drive_file_to_folder": "🗂 Раскладываю файлы по папкам...",
    "list_drive_folder_items": "🗂 Просматриваю папку на Диске...",
    "create_drive_folder": "🗂 Создаю папку на Диске...",
    "organize_drive_folder": "🗂 Навожу порядок в папке...",
    "manage_apps_script": "🛠 Пишу код автоматизации...",
    "get_webapp_template": "🧩 Собираю веб-приложение...",
    "make_sheet_applet": "🧩 Подключаю таблицу к приложению...",
    "find_tool": "🧰 Подбираю подходящий инструмент...",
    "start_here_always_read_ai_instructions": "📋 Сверяюсь со своими инструкциями...",
    "get_ai_instructions": "📋 Сверяюсь со своими инструкциями...",
    "get_ai_capabilities": "📋 Уточняю свои возможности...",
    "get_context_guide": "🗺 Смотрю карту данных...",
    "search_messages": "💬 Ищу по сообщениям...",
    "get_chat_transcript": "💬 Читаю переписку чата...",
    "process_chat_ocr": "🖼 Распознаю изображения из чата...",
    "get_compact_export": "📦 Готовлю выгрузку данных...",
}

_B24_TOOL_STATUS_DEFAULT = "⚙️ Работаю с данными..."

# Fillers between real tool events — rotate every tick so the message never looks frozen.
_B24_STATUS_THINKING = [
    "🧠 Обдумываю, что с этим делать...",
    "🤔 Прикидываю варианты...",
    "🧩 Складываю всё воедино...",
    "✍️ Формулирую мысль...",
    "🔎 Перепроверяю детали...",
    "📐 Сверяю цифры и факты...",
    "☕ Минутку, довожу до ума...",
]


def _b24_status_text(_elapsed_s: float = 0) -> str:
    return random.choice(_B24_STATUS_GREETING)


def _b24_status_for_tool(tool_name: str) -> str:
    return _B24_TOOL_STATUS.get(str(tool_name), _B24_TOOL_STATUS_DEFAULT)


def _b24_status_thinking(previous: str) -> str:
    pool = [p for p in _B24_STATUS_THINKING if not str(previous).startswith(p)]
    return random.choice(pool or _B24_STATUS_THINKING)


def _b24_status_send(client_endpoint: str, access_token: str, bot_id: Any, dialog_id: str) -> Any:
    """Post the initial progress message; returns its message id (None = feature off or failed)."""
    if os.getenv("B24_STATUS_MESSAGE", "").strip() != "1":
        return None
    try:
        res = _b24_app_call(client_endpoint, access_token, "imbot.message.add", {
            "BOT_ID": bot_id, "DIALOG_ID": dialog_id, "MESSAGE": _b24_status_text(0),
        })
        return res.get("result")
    except Exception:  # noqa: BLE001
        logging.debug("b24 testbot: status message send failed", exc_info=True)
        return None


def _b24_status_update(client_endpoint: str, access_token: str, bot_id: Any,
                       message_id: Any, text: str) -> None:
    try:
        _b24_app_call(client_endpoint, access_token, "imbot.message.update", {
            "BOT_ID": bot_id, "MESSAGE_ID": message_id, "MESSAGE": text,
        })
    except Exception:  # noqa: BLE001
        logging.debug("b24 testbot: status message update failed", exc_info=True)


def _b24_status_finish(client_endpoint: str, access_token: str, bot_id: Any, message_id: Any) -> None:
    """Remove the progress message right before the real answer lands. If the portal refuses the
    delete, degrade to editing it into a short 'done' pointer instead of leaving a stale status."""
    try:
        _b24_app_call(client_endpoint, access_token, "imbot.message.delete", {
            "BOT_ID": bot_id, "MESSAGE_ID": message_id, "COMPLETE": "Y",
        })
    except Exception:  # noqa: BLE001
        logging.debug("b24 testbot: status delete failed — editing to done-pointer", exc_info=True)
        _b24_status_update(client_endpoint, access_token, bot_id, message_id, "✅ Готово — ответ ниже 👇")


def _b24_app_process(client_endpoint: str, access_token: str, bot_id: Any, dialog_id: str,
                     user_text: str, message_id: Any = "", from_user_id: Any = "",
                     agent: dict[str, Any] | None = None) -> None:
    started = time.monotonic()
    tier = str(agent["tier"]) if agent is not None else _b24_tier_for(from_user_id)
    status, error = "ok", None
    # Live progress: one status message edited in place while the brain works, plus the native
    # 'typing…' indicator (it fades after ~30s on its own and the bot would look frozen).
    stop_typing = threading.Event()
    status_message_id = _b24_status_send(client_endpoint, access_token, bot_id, dialog_id)

    def _typing_keepalive() -> None:
        keepalive_started = time.monotonic()
        turn_started_ts = time.time()
        last_shown = ""
        last_tool_ts = 0.0
        while not stop_typing.wait(12):
            _b24_app_typing(client_endpoint, access_token, bot_id, dialog_id)
            if not status_message_id:
                continue
            text = None
            try:
                # Same process as the MCP server: show what the agent is REALLY doing right now.
                # With several concurrent turns the attribution is approximate — acceptable for
                # a cosmetic status line.
                from mcp.context_server import recent_core_tool_calls
                calls = recent_core_tool_calls(turn_started_ts)
                if calls and calls[-1][0] > last_tool_ts:
                    last_tool_ts = calls[-1][0]
                    text = _b24_status_for_tool(calls[-1][1])
            except Exception:  # noqa: BLE001
                pass
            if text is None:
                text = _b24_status_thinking(last_shown)
            minutes = int((time.monotonic() - keepalive_started) // 60)
            if minutes >= 1:
                text = f"{text} · {minutes} мин"
            if text != last_shown:
                last_shown = text
                _b24_status_update(client_endpoint, access_token, bot_id,
                                   status_message_id, text)

    threading.Thread(target=_typing_keepalive, daemon=True).start()
    # Durable in-flight marker: if this process is killed mid-turn (deploy restart, OOM, crash),
    # the row survives and boot recovery tells the user to resend — never an eternal 'typing…'.
    inflight_id = _b24_inflight_register(bot_id, dialog_id, (agent or {}).get("slug"), from_user_id,
                                         message_id, status_message_id, user_text)
    try:
        answer = hermes_brain_answer(user_text, dialog_id, tier, from_user_id, agent=agent)
    except Exception as exc:  # noqa: BLE001
        logging.exception("b24 testbot: hermes brain failed")
        status, error = "error", str(exc)[:500]
        _b24_ops_alert("исключение в ходе", dialog_id, tier, from_user_id, str(exc)[:300])
        answer = ("Что-то пошло не так на моей стороне 😔 Я уже отправил отчёт разработчикам. "
                  "Попробуй повторить запрос через пару минут.")
    finally:
        stop_typing.set()
        _b24_inflight_clear(inflight_id)
    if answer is None:
        # Turn cancelled by «Новая сессия»: the reset flow already replied to the user —
        # drop the progress message, release the 👀 reaction and post nothing.
        if status_message_id:
            _b24_status_finish(client_endpoint, access_token, bot_id, status_message_id)
        _b24_app_react(client_endpoint, access_token, message_id, "eyes", add=False)
        return
    latency_ms = int((time.monotonic() - started) * 1000)
    answer, escalation_request = _b24_extract_escalation(answer)
    answer, _deliver_name, _deliver_fmt = _b24_extract_deliver(answer)
    if _deliver_name:
        try:
            if _deliver_fmt == "xlsx":
                _data, _ext, _label = _b24_text_to_xlsx(_deliver_name, answer or _deliver_name), "xlsx", "Excel"
            elif _deliver_fmt == "docx":
                _data, _ext, _label = _b24_text_to_docx(_deliver_name, answer or _deliver_name), "docx", "Word"
            else:
                _data, _ext, _label = _b24_text_to_pdf(_deliver_name, answer or _deliver_name), "pdf", "PDF"
            _link = _b24_save_export(_data, _deliver_name, _ext)
            answer = "📎 Готово, оформил ответ в " + _label + ":\n" + _link + "\n\n(ссылка действует ~30 минут)"
        except Exception:  # noqa: BLE001
            logging.exception("b24 testbot: file deliver failed — sending text instead")
    if status_message_id:
        _b24_status_finish(client_endpoint, access_token, bot_id, status_message_id)
    _b24_app_reply(client_endpoint, access_token, bot_id, dialog_id, answer,
                   keyboard=_b24_keyboard())
    if escalation_request is not None:
        threading.Thread(
            target=_b24_forward_access_request,
            args=(dialog_id, from_user_id, escalation_request or user_text),
            daemon=True,
        ).start()
    # done: swap 👀 (read) → 👍 (done) on the user's message.
    _b24_app_react(client_endpoint, access_token, message_id, "eyes", add=False)
    _b24_app_react(client_endpoint, access_token, message_id, "like", add=True)
    _b24_session_touch(dialog_id, (agent or {}).get("slug"))
    _b24_ensure_command_registered(client_endpoint, access_token, bot_id)
    _b24_log_interaction(dialog_id, from_user_id, tier, user_text, answer, latency_ms, status, error,
                         agent_slug=(agent or {}).get("slug"))


def _b24_do_reset(client_endpoint: str, access_token: str, bot_id: Any, dialog_id: str,
                  message_id: Any = "", agent_slug: str | None = None) -> None:
    """Reset the dialog's session and confirm — runs WITHOUT calling the model. Scoped per agent
    so 'New session' only clears THIS agent's conversation, not another bot's in the same chat.
    A brain turn still running in this scope is KILLED first: a reset means the user no longer
    wants that answer, and letting it finish would post into the fresh session."""
    stopped = _b24_cancel_live_turns(_b24_scope(dialog_id, agent_slug))
    _b24_session_reset(dialog_id, agent_slug)
    if message_id:
        _b24_app_react(client_endpoint, access_token, message_id, "eyes", add=False)
        _b24_app_react(client_endpoint, access_token, message_id, "like", add=True)
    _b24_app_reply(client_endpoint, access_token, bot_id, dialog_id,
                   ("⏹ Остановил текущую работу и начал новую сессию — предыдущий контекст очищен. "
                    "Спрашивайте!") if stopped else
                   "🆕 Начал новую сессию — предыдущий контекст очищен. Спрашивайте!",
                   keyboard=_b24_keyboard())
    _b24_ensure_command_registered(client_endpoint, access_token, bot_id)


# --- "Сообщить об ошибке": ask → capture next message → forward to Telegram + log -----------
# The notifications group is a Telegram group ("Albery_Уведомления"); generic Bitrix REST is the
# wrong channel anyway. We deliver via the Telegram Bot API using @albery_ai_bot's token and log
# every report to bitrix_error_reports (delivered flag + any delivery error) for an audit trail.

def _albery_tg_bot_token() -> str:
    """Telegram bot token for Albery notifications. Prefer an explicit ALBERY_TG_BOT_TOKEN;
    otherwise reuse the Hermes gateway bot (@albery_ai_bot) token from /root/.hermes/.env — the
    albery service runs as root on the same box, so no secret needs to be duplicated."""
    token = os.getenv("ALBERY_TG_BOT_TOKEN", "").strip()
    if token:
        return token
    try:
        for line in Path("/root/.hermes/.env").read_text(encoding="utf-8").splitlines():
            if line.startswith("TELEGRAM_BOT_TOKEN="):
                return line.split("=", 1)[1].strip().strip('"').strip("'")
    except OSError:
        pass
    return ""


def _albery_tg_notify(text: str, chat_id: str | None = None) -> tuple[bool, str | None]:
    """Send a plain-text message to a Telegram chat (default: the Albery notifications group).
    Returns (ok, error)."""
    token = _albery_tg_bot_token()
    chat_id = (chat_id or os.getenv("ALBERY_ERROR_REPORT_TG_CHAT", "-5283789593")).strip()
    if not token or not chat_id:
        return False, "telegram token/chat not configured"
    try:
        resp = requests.post(
            f"https://api.telegram.org/bot{token}/sendMessage",
            json={"chat_id": chat_id, "text": text, "disable_web_page_preview": True},
            timeout=20,
        )
        data = resp.json() if resp.content else {}
        if not (isinstance(data, dict) and data.get("ok")):
            detail = data.get("description") if isinstance(data, dict) else resp.text
            return False, str(detail)[:300]
        return True, None
    except Exception as exc:  # noqa: BLE001
        return False, str(exc)[:300]


def _albery_bitrix_notify(text: str, dialog_id: str | None = None, *,
                          client_endpoint: str = "", access_token: str = "",
                          bot_id: Any = None) -> tuple[bool, str | None]:
    """Post a notification AS THE BOT into the Bitrix24 notifications chat ("Albery Уведомления")
    via imbot.message.add. During a live event the caller passes the event's fresh token; otherwise
    (a cron) we obtain one via _b24_app_access_token() (refresh_token grant — never expires while the
    weekly job keeps rotating it). The bot must be a member of the chat. Default chat = chat728
    (override via ALBERY_BITRIX_NOTIFY_CHAT). Returns (ok, error)."""
    dialog_id = (dialog_id or os.getenv("ALBERY_BITRIX_NOTIFY_CHAT", "chat728")).strip()
    if not bot_id:
        bot_id = _b24_load_state().get("bot_id")
    if not (client_endpoint and access_token):
        try:
            client_endpoint, access_token = _b24_app_access_token()
        except Exception as exc:  # noqa: BLE001
            return False, str(exc)[:300]
    if not (client_endpoint and access_token and bot_id and dialog_id):
        return False, "bitrix bot token/chat not bootstrapped yet"
    try:
        _b24_app_call(client_endpoint, access_token, "imbot.message.add",
                      {"BOT_ID": bot_id, "DIALOG_ID": dialog_id, "MESSAGE": text})
        return True, None
    except Exception as exc:  # noqa: BLE001
        return False, str(exc)[:300]


def _b24_user_name(client_endpoint: str, access_token: str, user_id: Any) -> str:
    """Best-effort '<First Last>' for a Bitrix user id (empty string if unavailable)."""
    uid = to_int(user_id)
    if not uid:
        return ""
    try:
        data = _b24_app_call(client_endpoint, access_token, "user.get", {"ID": uid})
        rows = data.get("result") or []
        if isinstance(rows, list) and rows and isinstance(rows[0], dict):
            u = rows[0]
            return " ".join(p for p in (u.get("NAME"), u.get("LAST_NAME")) if p).strip()
    except Exception:  # noqa: BLE001
        logging.debug("b24 testbot: user.get name lookup failed", exc_info=True)
    return ""


def _b24_log_error_report(dialog_id: str, from_user_id: Any, reporter_name: str,
                          report_text: str, delivered: bool, delivery_error: str | None,
                          agent_slug: str | None = None) -> None:
    try:
        with pg_connect() as conn:
            with conn.transaction():
                with conn.cursor() as cur:
                    cur.execute(
                        """
                        INSERT INTO bitrix_error_reports
                          (dialog_id, bitrix_user_id, reporter_name, report_text, delivered, delivery_error, agent_slug)
                        VALUES (%s, %s, %s, %s, %s, %s, %s)
                        """,
                        (str(dialog_id), to_int(from_user_id), reporter_name or None,
                         report_text, bool(delivered), delivery_error, agent_slug),
                    )
    except Exception:  # noqa: BLE001
        logging.exception("b24 testbot: error-report log failed")


def _b24_start_error_report(client_endpoint: str, access_token: str, bot_id: Any,
                            dialog_id: str, message_id: Any = "") -> None:
    """User pressed '⚠️ Сообщить об ошибке' — ask for the description and arm capture of their
    next message (handled in ONIMBOTMESSAGEADD via _b24_pop_awaiting_error). Armed per (bot, dialog)
    so pressing it on one bot never captures a message the user later sends to a DIFFERENT bot."""
    _b24_set_awaiting_error(bot_id, dialog_id)
    if message_id:
        _b24_app_react(client_endpoint, access_token, message_id, "eyes", add=False)
        _b24_app_react(client_endpoint, access_token, message_id, "like", add=True)
    _b24_app_reply(
        client_endpoint, access_token, bot_id, dialog_id,
        "Объясните, пожалуйста, в чём заключается ошибка? Опишите одним сообщением — "
        "я сразу передам Александру.",
        keyboard=_b24_keyboard(),
    )
    _b24_ensure_command_registered(client_endpoint, access_token, bot_id)


def _b24_handle_error_report(client_endpoint: str, access_token: str, bot_id: Any, dialog_id: str,
                             report_text: str, from_user_id: Any, reporter_name: str,
                             message_id: Any = "") -> None:
    """Forward a user's error description to the Albery notifications Telegram group and log it."""
    name = (reporter_name or "").strip() or "Сотрудник"
    tg_text = f"⚠️ {name} отправил отчёт об ошибке, текст: {report_text}. Я уже иду разбираться дальше."
    ok, err = _albery_tg_notify(tg_text)
    report_agent_slug = None  # NULL = the main bot, same convention as interactions
    try:
        from agent_center import agent_for_bot_id
        report_agent = agent_for_bot_id(bot_id)
        report_agent_slug = report_agent.get("slug") if report_agent else None
    except Exception:  # noqa: BLE001
        logging.warning("b24 testbot: error-report agent resolve failed", exc_info=True)
    _b24_log_error_report(dialog_id, from_user_id, name, report_text, ok, err, report_agent_slug)
    # Mirror to the Bitrix24 notifications chat AS THE BOT (best-effort; never blocks the TG path).
    # We are inside a live event, so pass its fresh token — no refresh needed here.
    b24_ok, b24_err = _albery_bitrix_notify(
        tg_text, client_endpoint=client_endpoint, access_token=access_token, bot_id=bot_id)
    if not b24_ok:
        logging.warning("b24 testbot: error report Bitrix mirror failed: %s", b24_err)
    if message_id:
        _b24_app_react(client_endpoint, access_token, message_id, "eyes", add=False)
        _b24_app_react(client_endpoint, access_token, message_id, "like", add=True)
    if ok:
        confirm = "Спасибо! Передал ваше сообщение Александру — уже разбираемся."
    else:
        logging.error("b24 testbot: error report TG delivery failed: %s", err)
        confirm = "Спасибо! Записал вашу ошибку и передам Александру."
    _b24_app_reply(client_endpoint, access_token, bot_id, dialog_id, confirm, keyboard=_b24_keyboard())


# ================= Agent works INSIDE a task: reply in the comment when called =================
# When an employee writes a task comment naming an agent ("Албери, ..."), and that employee has
# access to the agent, the agent replies IN the task comment with the FULL task context and can act
# on the task (comment / close / result / deadline) through its tools. Delivered via the existing
# OnTaskCommentAdd events are bound PROGRAMMATICALLY via the app OAuth token (event.bind) — the
# original design relied on a manual portal step that was never done, so event.get stayed empty and
# mentions silently never fired (found 2026-07-09). _b24_ensure_task_comment_event_bound() is
# self-healing: it runs on process start and on live events, checks event.get and binds only what
# is missing, so the binding survives app reinstalls/token rotations without manual steps.
#
# Guards make company-wide comment traffic safe — the event fires on EVERY comment on EVERY task:
#   * dedupe by comment id (bitrix_task_comment_seen) — each comment handled at most once;
#   * skip comments authored by the technical webhook user or any bot (no self-trigger loops);
#   * act ONLY when a configured agent trigger phrase is present AND the author has access;
#   * kill-switch B24_TASK_MENTION_ENABLED=0.

def _b24_task_mention_enabled() -> bool:
    return os.getenv("B24_TASK_MENTION_ENABLED", "1").strip().lower() not in {"0", "false", "no", "off"}


_B24_TASK_EVENT_BIND_CHECKED = False


def _b24_task_events_handler_url() -> str:
    secret = os.getenv("BITRIX_EVENT_SECRET", "").strip()
    return f"https://mcp.m4s.ru/bitrix/events/tasks/{secret}" if secret else ""


def _b24_ensure_task_comment_event_bound(client_endpoint: str, access_token: str) -> None:
    """Bind ONTASKCOMMENTADD/ONTASKCOMMENTUPDATE to our tasks-events endpoint via the app token.
    Idempotent and cheap: once per process; reads event.get and binds only the missing events.
    A duplicate delivery is harmless anyway — the comment claim is atomic (first sight wins)."""
    global _B24_TASK_EVENT_BIND_CHECKED
    if _B24_TASK_EVENT_BIND_CHECKED or not _b24_task_mention_enabled():
        return
    handler = _b24_task_events_handler_url()
    if not (handler and client_endpoint and access_token):
        return
    try:
        existing = _b24_app_call(client_endpoint, access_token, "event.get", {})
        bound = {
            str(e.get("event") or "").strip().upper()
            for e in (existing.get("result") or [])
            if isinstance(e, dict) and str(e.get("handler") or "").strip() == handler
        }
        for event_name in ("ONTASKCOMMENTADD", "ONTASKCOMMENTUPDATE"):
            if event_name in bound:
                continue
            try:
                _b24_app_call(client_endpoint, access_token, "event.bind",
                              {"event": event_name, "handler": handler})
                logging.info("b24 task-mention: bound %s -> tasks events endpoint", event_name)
            except Exception as exc:  # noqa: BLE001
                # "Handler already binded" is fine — someone raced us; anything else is loud.
                if "already" in str(exc).lower():
                    continue
                raise
        _B24_TASK_EVENT_BIND_CHECKED = True
    except Exception:  # noqa: BLE001
        logging.warning("b24 task-mention: event bind failed (will retry on next event)", exc_info=True)


def _b24_task_bot_author_ids() -> set[int]:
    """User ids whose comments must NEVER trigger the agent (its own replies go out as these)."""
    ids = {to_int(x) for x in os.getenv("B24_TASK_BOT_AUTHOR_IDS", "22").split(",")}
    return {i for i in ids if i is not None}


def _b24_task_targets() -> list[dict[str, Any]]:
    """Agents that can be summoned in a task comment, each with its trigger phrases.
    main (the universal agent) + every active subagent with a Bitrix bot. Triggers derive from the
    agent name; B24_TASK_MENTION_TRIGGERS_<SLUG> / _MAIN can add extras (comma-separated)."""
    targets: list[dict[str, Any]] = []
    main_bot = to_int(_b24_load_state().get("bot_id"))
    main_trigs = {"албери", "агент албери", "@албери", "albery"}
    main_trigs |= {t.strip().lower() for t in os.getenv("B24_TASK_MENTION_TRIGGERS_MAIN", "").split(",") if t.strip()}
    targets.append({"slug": None, "bot_id": main_bot, "name": "Агент Албери",
                    "triggers": main_trigs, "is_main": True})
    try:
        with pg_connect() as conn:
            with conn.cursor() as cur:
                cur.execute("SELECT slug, name, bitrix_bot_id FROM agents "
                            "WHERE is_active AND bitrix_bot_id IS NOT NULL")
                for row in cur.fetchall():
                    name = str(row["name"] or "").strip()
                    trigs = {name.lower()} if name else set()
                    for w in re.findall(r"[а-яёa-z]{4,}", name.lower()):
                        if w not in {"агент"}:
                            trigs.add(w)
                    env_key = "B24_TASK_MENTION_TRIGGERS_" + str(row["slug"] or "").upper().replace("-", "_")
                    trigs |= {t.strip().lower() for t in os.getenv(env_key, "").split(",") if t.strip()}
                    targets.append({"slug": row["slug"], "bot_id": to_int(row["bitrix_bot_id"]),
                                    "name": name, "triggers": {t for t in trigs if t}, "is_main": False})
    except Exception:  # noqa: BLE001
        logging.warning("b24 task-mention: target list failed", exc_info=True)
    return targets


def _b24_task_pick_agent(text: str) -> dict[str, Any] | None:
    """Return the summoned agent target whose trigger phrase appears in the comment, longest first
    (so «агент-юрист» wins over the generic «албери»). None if no agent is named."""
    low = " " + re.sub(r"\s+", " ", (text or "").lower()) + " "
    best = None
    best_len = 0
    for tgt in _b24_task_targets():
        for trig in tgt["triggers"]:
            if not trig:
                continue
            # word-ish boundary: trigger surrounded by non-letters (handles «Албери,» / «@албери»)
            if re.search(r"(?<![а-яёa-z])" + re.escape(trig) + r"(?![а-яёa-z])", low):
                if len(trig) > best_len:
                    best, best_len = tgt, len(trig)
    return best


def _b24_task_comment_claim(comment_id: int, task_id: int, agent_slug: str | None, author_id: Any) -> bool:
    """Atomically mark a comment as seen. Returns True only on FIRST sight (safe to process)."""
    try:
        with pg_connect() as conn:
            with conn.cursor() as cur:
                cur.execute(
                    "INSERT INTO bitrix_task_comment_seen (comment_id, task_id, agent_slug, author_id) "
                    "VALUES (%s,%s,%s,%s) ON CONFLICT (comment_id) DO NOTHING RETURNING comment_id",
                    (int(comment_id), int(task_id), agent_slug, to_int(author_id)))
                return cur.fetchone() is not None
    except Exception:  # noqa: BLE001
        logging.warning("b24 task-mention: claim failed comment=%s", comment_id, exc_info=True)
        return False


def _b24_fetch_task_comment(task_id: int, comment_id: int) -> dict[str, Any] | None:
    """Read a task comment (author id + text) via the task's IM chat (the modern task card stores
    comments as chat messages — the legacy forum getlist is empty on this portal)."""
    wh = (os.getenv("B24_TESTBOT_WEBHOOK_BASE", "") or "").rstrip("/")

    def _wh(method, payload):
        if not wh:
            return {}
        try:
            r = requests.post(f"{wh}/{method}.json", json=payload, timeout=20)
            return r.json() if r.content else {}
        except Exception:  # noqa: BLE001
            return {}

    task = _wh("tasks.task.get", {"taskId": task_id, "select": ["ID", "CHAT_ID", "TITLE"]})
    t = (task.get("result") or {}).get("task") or {} if isinstance(task, dict) else {}
    chat_id = t.get("chatId") or t.get("CHAT_ID")
    if not chat_id:
        return None
    msgs = _wh("im.dialog.messages.get", {"DIALOG_ID": f"chat{chat_id}", "LIMIT": 30})
    for m in ((msgs.get("result") or {}).get("messages") or []) if isinstance(msgs, dict) else []:
        if str(m.get("id")) == str(comment_id):
            params = m.get("params") if isinstance(m.get("params"), dict) else {}
            file_ids = params.get("FILE_ID") if isinstance(params.get("FILE_ID"), list) else []
            return {"author_id": to_int(m.get("author_id")), "text": str(m.get("text") or ""),
                    "chat_id": chat_id, "title": t.get("title") or t.get("TITLE"),
                    "file_ids": file_ids}
    return None


_B24_TASK_STATUS_LABELS = {
    1: "новая", 2: "ждёт выполнения", 3: "выполняется", 4: "ждёт контроля",
    5: "завершена", 6: "отложена", 7: "отклонена",
}


def _b24_strip_task_bbcode(text: str) -> str:
    """Readable comment text: [USER=24]Имя[/USER] -> Имя, other short BB tags dropped."""
    out = re.sub(r"\[USER=\d+\]([^\[]*)\[/USER\]", r"\1", str(text or ""), flags=re.IGNORECASE)
    out = re.sub(r"\[/?[A-Za-z][^\]]{0,60}\]", "", out)
    return re.sub(r"\s+", " ", out).strip()


def _b24_task_context_text(task_id: int) -> str:
    """The FULL task context block for the agent: title, description, status, deadline,
    responsible, creator and the recent comment thread — so the agent «обязательно видит
    контекст задачи, в которой его зовут» (требование задачи 1152)."""
    wh = (os.getenv("B24_TESTBOT_WEBHOOK_BASE", "") or "").rstrip("/")
    lines = [f"Задача №{task_id}."]
    chat_id = None
    try:
        r = requests.post(f"{wh}/tasks.task.get.json",
                          json={"taskId": task_id,
                                "select": ["ID", "TITLE", "DESCRIPTION", "STATUS", "DEADLINE",
                                           "RESPONSIBLE_ID", "CREATED_BY", "CHAT_ID"]}, timeout=20)
        t = (r.json().get("result") or {}).get("task") or {}
        chat_id = t.get("chatId") or t.get("CHAT_ID")
        if t.get("title"):
            lines.append("Название: " + str(t.get("title")))
        status = to_int(t.get("status") or t.get("STATUS"))
        if status in _B24_TASK_STATUS_LABELS:
            lines.append("Статус: " + _B24_TASK_STATUS_LABELS[status])
        if t.get("description"):
            lines.append("Описание: " + str(t.get("description"))[:3000])
        if t.get("deadline"):
            lines.append("Срок: " + str(t.get("deadline")))
        directory = _b24_portal_user_directory()
        rid = to_int(t.get("responsibleId") or t.get("RESPONSIBLE_ID"))
        if rid:
            lines.append("Ответственный: " + (directory.get(rid, {}).get("name") or str(rid)))
        cid = to_int(t.get("createdBy") or t.get("CREATED_BY"))
        if cid:
            lines.append("Постановщик: " + (directory.get(cid, {}).get("name") or str(cid)))
    except Exception:  # noqa: BLE001
        logging.warning("b24 task-mention: context fetch failed task=%s", task_id, exc_info=True)

    # Recent comment thread (the modern task card keeps comments in the task IM chat).
    if chat_id:
        try:
            r = requests.post(f"{wh}/im.dialog.messages.get.json",
                              json={"DIALOG_ID": f"chat{chat_id}", "LIMIT": 20}, timeout=20)
            msgs = (r.json().get("result") or {}).get("messages") or []
            directory = _b24_portal_user_directory()
            thread = []
            for m in sorted(msgs, key=lambda x: to_int(x.get("id")) or 0):
                author = to_int(m.get("author_id"))
                if not author:  # system notices (task created/status changed) — noise
                    continue
                text = _b24_strip_task_bbcode(m.get("text"))
                if not text:
                    continue
                name = directory.get(author, {}).get("name") or f"id {author}"
                thread.append(f"- {name}: {text[:400]}")
            if thread:
                lines.append("Последние комментарии (старые выше):\n" + "\n".join(thread[-10:]))
        except Exception:  # noqa: BLE001
            logging.warning("b24 task-mention: comments fetch failed task=%s", task_id, exc_info=True)
    return "\n".join(lines)


def _b24_post_task_comment(task_id: int, text: str, agent_name: str) -> bool:
    """Post the agent's reply as a task comment (author = technical webhook user, marked with the
    agent name so employees see who answered). Loop-safe: our handler skips this author id."""
    wh = (os.getenv("B24_TESTBOT_WEBHOOK_BASE", "") or "").rstrip("/")
    if not wh:
        return False
    body = f"🤖 {agent_name}: {text}"
    try:
        r = requests.post(f"{wh}/task.commentitem.add.json",
                          json={"TASKID": task_id, "FIELDS": {"POST_MESSAGE": body[:20000]}}, timeout=30)
        data = r.json() if r.content else {}
        return bool(data.get("result"))
    except Exception:  # noqa: BLE001
        logging.warning("b24 task-mention: reply post failed task=%s", task_id, exc_info=True)
        return False


def _b24_handle_task_comment_event(task_id: int, comment_id: int) -> dict[str, Any]:
    """Core of the in-task agent. Returns a small status dict (also used by the smoke test)."""
    if not _b24_task_mention_enabled():
        return {"handled": False, "reason": "disabled"}
    comment = _b24_fetch_task_comment(task_id, comment_id)
    if not comment:
        return {"handled": False, "reason": "comment_not_found"}
    author_id = comment.get("author_id")
    text = comment.get("text") or ""
    # Loop guard: never react to the bot's own replies or the technical webhook user.
    if to_int(author_id) in _b24_task_bot_author_ids() or text.lstrip().startswith("🤖"):
        return {"handled": False, "reason": "own_comment"}
    target = _b24_task_pick_agent(text)
    if not target:
        return {"handled": False, "reason": "no_agent_named"}
    # Access gate: the author must have access to THIS agent (same rules as the chat).
    allowed = _b24_main_allows(author_id) if target["is_main"] else _b24_task_subagent_allows(target["slug"], author_id)
    if not allowed:
        return {"handled": False, "reason": "no_access", "agent": target["name"]}
    # First-sight claim (dedupe) — do this AFTER the cheap filters so we don't burn the id on noise.
    if not _b24_task_comment_claim(comment_id, task_id, target["slug"], author_id):
        return {"handled": False, "reason": "already_seen"}

    agent = None
    if not target["is_main"] and target["slug"]:
        try:
            from agent_center import agent_for_bot_id
            agent = agent_for_bot_id(target["bot_id"])
        except Exception:  # noqa: BLE001
            logging.warning("b24 task-mention: agent resolve failed slug=%s", target["slug"], exc_info=True)
    tier = "ops" if target["is_main"] else str((agent or {}).get("tier") or "faq")

    ctx = _b24_task_context_text(task_id)
    requester = _b24_portal_user_directory().get(to_int(author_id), {}).get("name") or f"id {author_id}"
    # Files attached to the triggering comment (screenshots/documents): recognize and inject,
    # so «Албери, посмотри скрин» inside a task actually sees the screenshot.
    files_block = ""
    if comment.get("file_ids"):
        try:
            parts = []
            for f in task_comment_files(comment["file_ids"], task_id):
                head = f"[Вложение в комментарии: «{f['name']}» ({'скрин/изображение' if f['kind'] == 'image' else 'документ'}"
                if f.get("attachment_id"):
                    head += f", attachment_id={f['attachment_id']}"
                head += "). Распознанное/извлечённое содержимое:]"
                parts.append(head + "\n" + (f["text"] or "")[:4000])
            if parts:
                files_block = "\n\n" + "\n\n".join(parts) + (
                    "\n\n(Полный текст вложения — get_attachment_text(attachment_id=…); "
                    "переслать файл в задачу/комментарий — attachment_ids.)")
        except Exception:  # noqa: BLE001
            logging.warning("b24 task-mention: comment files read failed task=%s", task_id, exc_info=True)
    user_text = (
        "Тебя позвали ПРЯМО В ЗАДАЧЕ Bitrix (в комментарии). Работай с ЭТОЙ задачей.\n\n"
        + ctx + "\n\n"
        + f"Сотрудник {requester} (id={author_id}) написал в комментарии к задаче №{task_id}:\n«"
        + text.strip() + "»"
        + files_block + "\n\n"
        "Ответь по существу и, если он просит действие с задачей, выполни его своими инструментами "
        f"(комментарий — add_bitrix_task_comment(bitrix_task_id={task_id}, author_bitrix_user_id={author_id}); "
        f"завершить — complete_bitrix_task(bitrix_task_id={task_id}, on_behalf_bitrix_user_id={author_id}); "
        f"результат/скрин как результат — as_result=true; возобновить — reopen_bitrix_task; "
        f"новый срок — new_deadline). Пиши кратко: твой ответ уйдёт в этот же комментарий задачи. "
        "НЕ дублируй постановку — задача уже есть, работай в её контексте."
    )
    # Per-task memory so the agent keeps the thread of THIS task, separate from private chat.
    dialog_id = f"task-{task_id}"
    answer = hermes_brain_answer(user_text, dialog_id, tier, author_id, agent=agent)
    if not answer or _hermes_answer_is_error(answer):
        return {"handled": False, "reason": "brain_error", "agent": target["name"]}
    posted = _b24_post_task_comment(task_id, answer, target["name"])
    try:
        _b24_log_interaction(dialog_id, author_id, tier, user_text, answer, 0,
                             "ok" if posted else "post_failed", "" if posted else "reply post failed",
                             agent_slug=target["slug"])
    except Exception:  # noqa: BLE001
        pass
    return {"handled": bool(posted), "agent": target["name"], "task_id": task_id}


def _b24_task_subagent_allows(slug: str | None, author_id: Any) -> bool:
    """Subagent access for the in-task path: explicit member list, empty = open to non-'none'."""
    if not slug:
        return _b24_main_allows(author_id)
    try:
        from agent_center import agent_for_bot_id  # noqa: F401
        with pg_connect() as conn:
            with conn.cursor() as cur:
                cur.execute("SELECT a.id FROM agents a WHERE a.slug=%s", (slug,))
                row = cur.fetchone()
                if not row:
                    return False
                cur.execute("SELECT bitrix_user_id FROM agent_members WHERE agent_id=%s", (row["id"],))
                members = {to_int(r["bitrix_user_id"]) for r in cur.fetchall()}
    except Exception:  # noqa: BLE001
        logging.warning("b24 task-mention: subagent access lookup failed slug=%s", slug, exc_info=True)
        return False
    if members:
        return to_int(author_id) in members
    return _b24_tier_for(author_id) != "none"


def _bitrix_imbot_app_event():
    if request.method == "GET":
        return jsonify({"ok": True, "message": "Bitrix imbot app endpoint is ready."})

    payload = flatten_request_payload()
    event_name = str(first_non_empty(payload.get("event"), payload.get("EVENT")) or "").upper()
    app_token = _imbot_auth(payload, "application_token")
    access_token = _imbot_auth(payload, "access_token")
    client_endpoint = _imbot_auth(payload, "client_endpoint")
    state = _b24_load_state()
    # Persist the app OAuth tokens from this event so cron digests can post as the bot later.
    _b24_capture_tokens(payload, state)

    if event_name in ("ONAPPINSTALL", "ONAPPUPDATE"):
        try:
            bot_id = state.get("bot_id") or _b24_app_register_bot(client_endpoint, access_token)
            state.update({"application_token": app_token, "client_endpoint": client_endpoint, "bot_id": bot_id})
            _b24_save_state(state)
            return jsonify({"ok": True, "event": event_name, "bot_id": bot_id})
        except Exception as exc:  # noqa: BLE001
            logging.exception("b24 app install failed")
            return jsonify({"ok": False, "event": event_name, "error": str(exc)[:300]}), 200

    stored = state.get("application_token", "")
    if not stored and app_token:
        # Self-heal after an accidental state wipe (the pre-fix ONIMBOTDELETE bug erased the
        # whole state when a SUBAGENT bot was deleted): every event of this install keeps
        # carrying the same application_token, so trust it once, persist, and keep serving.
        logging.warning("b24 testbot: empty stored application_token — self-healing from event")
        state["application_token"] = stored = app_token
        if client_endpoint:
            state["client_endpoint"] = client_endpoint
        _b24_save_state(state)
    if not stored or not app_token or not hmac.compare_digest(app_token, stored):
        return jsonify({"error": "forbidden"}), 403

    if event_name in ("ONAPPUNINSTALL", "ONIMBOTDELETE"):
        # Wipe the install state ONLY when the whole app is removed or the MAIN bot is
        # deleted; a deleted SUBAGENT bot must never take the main bot's tokens with it.
        deleted_bot = str(_imbot_scan(payload, "BOT_ID") or "").strip()
        main_bot = str(state.get("bot_id") or "")
        if event_name == "ONAPPUNINSTALL" or (deleted_bot and deleted_bot == main_bot):
            _b24_save_state({})
        else:
            logging.info("b24 testbot: subagent bot %s deleted — main state kept", deleted_bot or "?")
            try:
                from agent_center import _agent_cache_bust
                _agent_cache_bust()
            except Exception:  # noqa: BLE001
                pass
        return jsonify({"ok": True, "event": event_name})

    bot_id = state.get("bot_id")
    endpoint = client_endpoint or state.get("client_endpoint", "")
    dialog_id = str(_imbot_event_param(payload, "DIALOG_ID") or "")

    # Subagent routing: every bot registered by this application posts to this same
    # handler; the event's BOT_ID tells us WHICH bot was addressed. Unknown ids fall
    # back to the main agent so a stale cache can never silence the bot entirely.
    agent = None
    event_bot_id = str(_imbot_scan(payload, "BOT_ID") or "").strip()
    if event_bot_id and str(bot_id or "") != event_bot_id:
        try:
            from agent_center import agent_for_bot_id
            agent = agent_for_bot_id(event_bot_id)
        except Exception:  # noqa: BLE001
            logging.exception("b24 testbot: subagent resolve failed for bot %s", event_bot_id)
    if agent is not None:
        bot_id = event_bot_id  # reply as the addressed bot, not the main one
    elif not bot_id and event_bot_id:
        # Second half of the self-heal: restore the main bot id lost with the wiped state.
        state["bot_id"] = bot_id = event_bot_id
        _b24_save_state(state)
        logging.warning("b24 testbot: restored main bot_id=%s from event", event_bot_id)

    def _agent_allows(user_id: Any) -> bool:
        """Subagent access: its explicit member list; empty list = agent open to everyone
        who is not globally denied ('none' tier)."""
        if agent is None:
            return True
        members = agent.get("members") or set()
        uid = to_int(user_id)
        if members:
            return uid in members
        return _b24_tier_for(user_id) != "none"

    # Self-heal: ensure the /new command is registered (background, one-time, no-op after) for the
    # addressed bot, AND once per process for EVERY bot (main + subagents) so subagent buttons are
    # live without needing a priming message.
    if access_token and bot_id and endpoint and event_name in ("ONIMBOTMESSAGEADD", "ONIMCOMMANDADD"):
        _b24_ensure_command_registered(endpoint, access_token, bot_id)
        _b24_bootstrap_all_commands(endpoint, access_token, state.get("bot_id"))
        _b24_ensure_task_comment_event_bound(endpoint, access_token)

    if event_name == "ONIMBOTJOINCHAT":
        if dialog_id:
            if agent is not None:
                welcome = (f"👋 Я — {agent['name']}, специализированный ИИ-агент Албери.\n\n"
                           f"{(agent.get('role_prompt') or '').strip()[:300]}\n\n"
                           "Просто напишите свой вопрос — поехали! 🚀").replace("\n\n\n\n", "\n\n")
                _b24_app_reply(endpoint, access_token, bot_id, dialog_id, welcome)
            else:
                _b24_app_reply(
                    endpoint, access_token, bot_id, dialog_id,
                    "👋 Я — ИИ-агент Албери, и я могу сильно упростить вашу работу: подскажу по компании "
                    "и регламентам, разберу Zoom-созвоны, помогу с задачами и отчётами.\n\n"
                    "Чтобы пользоваться мной эффективно, пройдите короткое обучение (1 минута) — или "
                    "сразу задайте свой вопрос. Поехали! 🚀",
                    keyboard=_b24_welcome_keyboard(),
                )
            _b24_ensure_command_registered(endpoint, access_token, bot_id)
        return jsonify({"ok": True, "event": event_name})

    # A keyboard button / slash command click arrives as ONIMCOMMANDADD. The command-event field
    # layout differs from messages, so read fields robustly: `report_error` starts the error-report
    # flow; only an explicit `new/reset` command resets the session. Empty/unknown command
    # events are ignored to avoid accidental context loss after Bitrix keyboard glitches.
    if event_name == "ONIMCOMMANDADD":
        try:
            logging.info("b24 testbot: ONIMCOMMANDADD keys=%s", [k for k in payload.keys()][:40])
        except Exception:  # noqa: BLE001
            pass
        command = (_imbot_scan(payload, "COMMAND") or "").strip().lstrip("/").lower()
        message_id = _imbot_scan(payload, "MESSAGE_ID")
        command_id = _imbot_scan(payload, "COMMAND_ID")
        # A keyboard-button command event carries the clicker in data[PARAMS][FROM_USER_ID] and,
        # crucially, NO DIALOG_ID for a private 1-1 chat. In Bitrix a private dialog is keyed by
        # the peer user id, so derive cmd_dialog from the clicker's id when DIALOG_ID is absent —
        # otherwise every dispatch branch below (they all require cmd_dialog) is skipped and the
        # button silently does nothing (this was THE bug: buttons acknowledged but no-op'd).
        cmd_user = (str(_imbot_event_param(payload, "FROM_USER_ID") or "")
                    or _imbot_scan(payload, "FROM_USER_ID")
                    or _imbot_scan(payload, "USER_ID"))
        cmd_dialog = dialog_id or _imbot_scan(payload, "DIALOG_ID") or str(cmd_user or "")
        cmd_denied = (
            (not _b24_main_allows(cmd_user)) if agent is None else not _agent_allows(cmd_user)
        )
        if cmd_dialog and cmd_user and cmd_denied:
            try:
                _b24_app_call(endpoint, access_token, "imbot.message.add", {
                    "BOT_ID": bot_id, "DIALOG_ID": cmd_dialog,
                    "MESSAGE": "😔 К сожалению, у вас нет доступа к агенту.\n\n"
                               "Пожалуйста, обратитесь к вашему руководителю или к Александру Никитенко 🙌",
                })
            except Exception:  # noqa: BLE001
                logging.exception("b24 testbot: no-access notice (command) failed")
        elif cmd_dialog and command in ("report_error", "error", "ошибка"):
            _b24_start_error_report(endpoint, access_token, bot_id, cmd_dialog, message_id)
        elif cmd_dialog and command in ("help", "обучение", "onboarding", "помощь"):
            _b24_send_onboarding(endpoint, access_token, bot_id, cmd_dialog, 1,
                                 _b24_tier_for(cmd_user), message_id)
        elif cmd_dialog and command == "onb_next":
            _b24_send_onboarding(endpoint, access_token, bot_id, cmd_dialog,
                                 _b24_onboarding_step(bot_id, cmd_dialog) + 1, _b24_tier_for(cmd_user), message_id)
        elif cmd_dialog and command in ("new", "reset", "новая", "сброс"):
            _b24_do_reset(endpoint, access_token, bot_id, cmd_dialog, message_id,
                          agent_slug=(agent or {}).get("slug"))
        elif cmd_dialog:
            logging.info("b24 testbot: ignored unknown/empty command event: %r", command)
        # Acknowledge the command so Bitrix stops retrying / showing 'typing…' (best-effort).
        if command_id:
            try:
                _b24_app_call(endpoint, access_token, "imbot.command.answer",
                              {"COMMAND_ID": command_id, "MESSAGE": ""})
            except Exception:  # noqa: BLE001
                logging.debug("b24 testbot: command.answer failed", exc_info=True)
        return jsonify({"ok": True, "event": event_name, "command": command})

    if event_name == "ONIMBOTMESSAGEADD":
        message_text = str(_imbot_event_param(payload, "MESSAGE") or "").strip()
        message_id = _imbot_event_param(payload, "MESSAGE_ID") or ""
        from_user_id = _imbot_event_param(payload, "FROM_USER_ID") or ""
        if os.getenv("B24_DEBUG_PAYLOAD", "0") == "1":
            try:
                _red = {k: ("<redacted>" if re.search(r"TOKEN|AUTH|SECRET|REFRESH", k, re.I) else v)
                        for k, v in payload.items()}
                logging.info("b24 MSGADD keys=%s payload=%s",
                             list(payload.keys()), json.dumps(_red, ensure_ascii=False)[:3500])
            except Exception:  # noqa: BLE001
                logging.exception("b24 MSGADD debug log failed")
        # The bot must SEE screenshots and understand replies to earlier (possibly reset) messages.
        image_texts, reply_text, doc_blocks, msg_attachments = [], "", [], []
        try:
            image_texts, reply_text, doc_blocks, msg_attachments = _b24_message_extras(
                payload, endpoint, access_token,
                agent_slug=(agent or {}).get("slug"), dialog_id=dialog_id, from_user_id=from_user_id)
        except Exception:  # noqa: BLE001
            logging.exception("b24 testbot: image/reply/doc extras failed")
        if not dialog_id or (not message_text and not image_texts and not reply_text and not doc_blocks):
            return jsonify({"ok": True, "ignored": True, "reason": "empty"}), 200
        # Access gate: the universal (main) agent answers ONLY its team (allowlist =
        # non-'none' agent_access grants); a subagent answers only its member list. Everyone
        # else gets a plain system notice (no model, no disclaimer) so they know to ask for access.
        denied = (not _b24_main_allows(from_user_id)) if agent is None else not _agent_allows(from_user_id)
        if agent is not None and not agent.get("is_active", True):
            denied = True
        if denied:
            try:
                _b24_app_call(endpoint, access_token, "imbot.message.add", {
                    "BOT_ID": bot_id, "DIALOG_ID": dialog_id,
                    "MESSAGE": ("😔 Этот агент сейчас выключен." if agent is not None and not agent.get("is_active", True)
                                else "😔 К сожалению, у вас нет доступа к агенту.") + "\n\n"
                               "Пожалуйста, обратитесь к вашему руководителю или к Александру Никитенко 🙌",
                })
            except Exception:  # noqa: BLE001
                logging.exception("b24 testbot: no-access notice failed")
            return jsonify({"ok": True, "event": event_name, "no_access": True}), 200
        # Pending error report: this message is the user's error description — capture, forward to
        # the Albery notifications Telegram group + log it, WITHOUT calling the model. Scoped to THIS
        # bot so an error-report armed on another bot can't swallow a normal message here.
        if _b24_pop_awaiting_error(bot_id, dialog_id):
            reporter = _b24_user_name(endpoint, access_token, from_user_id)
            _b24_app_react(endpoint, access_token, message_id, "eyes", add=True)
            threading.Thread(
                target=_b24_handle_error_report,
                args=(endpoint, access_token, bot_id, dialog_id, message_text, from_user_id, reporter, message_id),
                daemon=True,
            ).start()
            return jsonify({"ok": True, "event": event_name, "error_report": True})
        # 'New session' request (typed /new or 'новая сессия'): reset WITHOUT calling the model.
        if _b24_is_reset_command(message_text):
            _b24_do_reset(endpoint, access_token, bot_id, dialog_id, message_id,
                          agent_slug=(agent or {}).get("slug"))
            return jsonify({"ok": True, "event": event_name, "reset": True})
        # Return to Bitrix immediately. Even cosmetic pre-processing calls (reaction/typing)
        # can hang during Bitrix/network degradation; if they run before the HTTP response,
        # Bitrix retries or drops the webhook. Do those signals inside the background worker.
        def _process_message_async() -> None:
            _b24_app_react(endpoint, access_token, message_id, "eyes", add=True)
            _b24_app_typing(endpoint, access_token, bot_id, dialog_id)
            _b24_app_process(
                endpoint, access_token, bot_id, dialog_id,
                _b24_compose_user_text(message_text, image_texts, reply_text, doc_blocks, msg_attachments),
                message_id, from_user_id, agent=agent,
            )

        threading.Thread(target=_process_message_async, daemon=True).start()
        return jsonify({"ok": True, "event": event_name, "accepted": True})

    return jsonify({"ok": True, "event": event_name, "ignored": True}), 200


@app.route("/bitrix/imbot/<secret>", methods=["GET", "POST"])
def bitrix_imbot_webhook(secret: str):
    if secret == "app":
        return _bitrix_imbot_app_event()
    if not b24_testbot_secret_valid(secret):
        return jsonify({"error": "forbidden"}), 403
    if request.method == "GET":
        return jsonify({"ok": True, "message": "Bitrix imbot endpoint is ready."})

    payload = flatten_request_payload()
    event_name = str(first_non_empty(payload.get("event"), payload.get("EVENT")) or "").upper()
    if event_name == "ONIMBOTJOINCHAT":
        dialog_id = str(_imbot_event_param(payload, "DIALOG_ID") or "")
        if dialog_id:
            _b24_testbot_reply(dialog_id, "Привет! Я тестовый ассистент. Могу искать сотрудников, ставить и закрывать задачи, писать людям. Что сделать?")
        return jsonify({"ok": True, "event": event_name})
    if event_name != "ONIMBOTMESSAGEADD":
        return jsonify({"ok": True, "event": event_name, "ignored": True}), 200

    dialog_id = str(_imbot_event_param(payload, "DIALOG_ID") or "")
    message_text = str(_imbot_event_param(payload, "MESSAGE") or "").strip()
    # Ignore messages authored by a bot to avoid loops.
    if str(_imbot_event_param(payload, "FROM_USER_ID") or "") and _imbot_event_param(payload, "MESSAGE") is None:
        return jsonify({"ok": True, "ignored": True, "reason": "no_message"}), 200
    if not dialog_id or not message_text:
        return jsonify({"ok": True, "ignored": True, "reason": "empty"}), 200

    # Respond to Bitrix immediately; run the (slower) agent loop in the background.
    threading.Thread(target=_b24_testbot_process, args=(dialog_id, message_text), daemon=True).start()
    return jsonify({"ok": True, "event": event_name, "accepted": True})


