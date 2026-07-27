from __future__ import annotations

# PDF → редактируемый Word и обратно: что именно переживает конвертацию.

import io

import pytest

import docconvert
import pdfedit

fitz = pytest.importorskip("fitz", reason="конвертация PDF требует pymupdf")


def _pdf_page(draw) -> bytes:
    doc = fitz.open()
    page = doc.new_page(width=595, height=842)
    page.insert_font(fontname="probe", fontfile=pdfedit.font_path("Times", bold=False, italic=False))
    draw(page)
    buffer = io.BytesIO()
    doc.save(buffer)
    doc.close()
    return buffer.getvalue()


def _document(data: bytes):
    from docx import Document

    return Document(io.BytesIO(data))


def test_text_and_reading_order_survive_the_conversion():
    source = _pdf_page(lambda page: [
        page.insert_text(fitz.Point(72, 100), "ДОГОВОР КОМИССИИ", fontname="probe", fontsize=14),
        page.insert_text(fitz.Point(72, 140), "1.1. Комиссионер обязуется реализовать Товар.", fontname="probe", fontsize=12),
        page.insert_text(fitz.Point(72, 160), "1.2. Комитент передаёт Товар на реализацию.", fontname="probe", fontsize=12),
    ])

    result, _ = docconvert.pdf_to_docx(source)

    paragraphs = [p.text.strip() for p in _document(result).paragraphs if p.text.strip()]
    assert "ДОГОВОР КОМИССИИ" in paragraphs[0]
    joined = "\n".join(paragraphs)
    assert joined.index("1.1.") < joined.index("1.2.")


def test_centred_heading_stays_centred_and_body_stays_justified():
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    body_left = 72
    font = fitz.Font(fontfile=pdfedit.font_path("Times", bold=False, italic=False))
    heading = "ПУБЛИЧНАЯ ОФЕРТА"
    first_line = "Комиссионер обязуется по поручению Комитента совершать сделки по продаже"
    second_line = "Товара третьим лицам через торговую площадку Wildberries на условиях Оферты."
    # Правый край текстовой области — по фактической ширине строк абзаца; заголовок
    # ставится ровно посередине между границами, как в настоящем документе.
    body_right = body_left + max(
        font.text_length(first_line, fontsize=12), font.text_length(second_line, fontsize=12)
    )

    def draw(page):
        heading_x = body_left + (body_right - body_left - font.text_length(heading, fontsize=14)) / 2
        page.insert_text(fitz.Point(heading_x, 100), heading, fontname="probe", fontsize=14)
        page.insert_text(fitz.Point(body_left, 140), first_line, fontname="probe", fontsize=12)
        page.insert_text(fitz.Point(body_left, 158), second_line, fontname="probe", fontsize=12)

    result, _ = docconvert.pdf_to_docx(_pdf_page(draw))

    paragraphs = [p for p in _document(result).paragraphs if p.text.strip()]
    assert paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert paragraphs[1].alignment in {WD_ALIGN_PARAGRAPH.JUSTIFY, WD_ALIGN_PARAGRAPH.LEFT}


def test_bold_run_stays_bold_and_size_is_kept():
    def draw(page):
        page.insert_font(
            fontname="probe-bold",
            fontfile=pdfedit.font_path("Times", bold=True, italic=False),
        )
        page.insert_text(fitz.Point(72, 100), "РЕКВИЗИТЫ", fontname="probe-bold", fontsize=13)

    result, _ = docconvert.pdf_to_docx(_pdf_page(draw))

    runs = [run for p in _document(result).paragraphs for run in p.runs if run.text.strip()]
    assert runs, "текст обязан доехать в Word"
    assert runs[0].bold is True
    assert abs(runs[0].font.size.pt - 13) < 0.6


def test_table_becomes_a_word_table_not_a_wall_of_text():
    doc = fitz.open()
    page = doc.new_page(width=595, height=842)
    page.insert_font(fontname="probe", fontfile=pdfedit.font_path("Times", bold=False, italic=False))
    # Рамки + текст в ячейках: так таблицу видит поиск таблиц PyMuPDF.
    rows = [(100, 130), (130, 160), (160, 190)]
    columns = [(72, 250), (250, 430)]
    for top, bottom in rows:
        for left, right in columns:
            page.draw_rect(fitz.Rect(left, top, right, bottom), color=(0, 0, 0), width=0.7)
    cells = [["Наименование", "Сумма"], ["Товар А", "1 000"], ["Товар Б", "2 000"]]
    for row_index, (top, _bottom) in enumerate(rows):
        for column_index, (left, _right) in enumerate(columns):
            page.insert_text(
                fitz.Point(left + 4, top + 20), cells[row_index][column_index],
                fontname="probe", fontsize=11,
            )
    buffer = io.BytesIO()
    doc.save(buffer)
    doc.close()

    result, _ = docconvert.pdf_to_docx(buffer.getvalue())

    document = _document(result)
    assert document.tables, "таблица обязана остаться таблицей, иначе её нельзя редактировать"
    table = document.tables[0]
    assert len(table.rows) == 3 and len(table.columns) == 2
    assert "Наименование" in table.rows[0].cells[0].text
    assert "2 000" in table.rows[2].cells[1].text
    # Текст таблицы не должен продублироваться ещё и абзацами.
    assert "Наименование" not in "\n".join(p.text for p in document.paragraphs)


def test_images_are_reported_as_lost_rather_than_silently_dropped():
    doc = fitz.open()
    page = doc.new_page(width=595, height=842)
    pixmap = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 40, 40))
    pixmap.set_rect(pixmap.irect, (10, 120, 200))
    page.insert_image(fitz.Rect(72, 72, 112, 112), pixmap=pixmap)
    buffer = io.BytesIO()
    doc.save(buffer)
    doc.close()

    _, warnings = docconvert.pdf_to_docx(buffer.getvalue())

    assert warnings and "изображен" in warnings[0].lower()


def test_password_protected_pdf_is_refused():
    doc = fitz.open()
    doc.new_page()
    buffer = io.BytesIO()
    doc.save(buffer, encryption=fitz.PDF_ENCRYPT_AES_256, owner_pw="owner", user_pw="user")
    doc.close()

    with pytest.raises(docconvert.ConvertError):
        docconvert.pdf_to_docx(buffer.getvalue())


def test_word_to_pdf_says_plainly_when_libreoffice_is_absent(monkeypatch):
    monkeypatch.setattr(docconvert, "soffice_path", lambda: None)

    with pytest.raises(docconvert.ConvertError) as excinfo:
        docconvert.docx_to_pdf(b"PK\x03\x04 fake docx")

    assert "libreoffice" in str(excinfo.value).lower()


@pytest.mark.skipif(docconvert.soffice_path() is None, reason="LibreOffice есть только на сервере")
def test_word_to_pdf_produces_a_readable_pdf():
    from docx import Document

    document = Document()
    document.add_heading("ДОГОВОР", level=1)
    document.add_paragraph("Комиссионер обязуется реализовать Товар Комитента.")
    buffer = io.BytesIO()
    document.save(buffer)

    result = docconvert.docx_to_pdf(buffer.getvalue())

    assert result.startswith(b"%PDF")
    assert "Комиссионер обязуется реализовать" in pdfedit.extract_text(result)
