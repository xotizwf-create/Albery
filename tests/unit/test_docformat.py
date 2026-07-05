"""docformat — the agent-controlled HTML → DOCX renderer behind export_document."""
from __future__ import annotations

import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docformat import html_to_docx


def render(html: str, **kw) -> Document:
    return Document(io.BytesIO(html_to_docx(html, **kw)))


def texts(doc: Document) -> list[str]:
    return [p.text for p in doc.paragraphs if p.text.strip()]


class TestBlocks:
    def test_title_centered_bold(self):
        doc = render('<h1 style="text-align:center">ДОГОВОР ПОСТАВКИ</h1><p>Тело.</p>')
        title = doc.paragraphs[0]
        assert title.text == "ДОГОВОР ПОСТАВКИ"
        assert title.alignment == WD_ALIGN_PARAGRAPH.CENTER
        assert all(r.bold for r in title.runs)

    def test_inline_formatting(self):
        doc = render("<p>обычный <b>жирный</b> <i>курсив</i> <u>подчёркнутый</u></p>")
        runs = doc.paragraphs[0].runs
        flags = {r.text.strip(): (bool(r.bold), bool(r.italic), bool(r.underline)) for r in runs if r.text.strip()}
        assert flags["жирный"] == (True, False, False)
        assert flags["курсив"] == (False, True, False)
        assert flags["подчёркнутый"] == (False, False, True)

    def test_page_break_before_for_appendix(self):
        doc = render('<p>Основной текст.</p>'
                     '<h2 style="page-break-before:always">Приложение № 1</h2><p>Спецификация.</p>')
        appendix = next(p for p in doc.paragraphs if p.text == "Приложение № 1")
        assert appendix.paragraph_format.page_break_before is True
        body = next(p for p in doc.paragraphs if p.text == "Основной текст.")
        assert not body.paragraph_format.page_break_before

    def test_text_indent_and_font_size(self):
        doc = render('<p style="text-indent:1.25cm; font-size:14pt">Абзац с отступом.</p>')
        p = doc.paragraphs[0]
        assert round(p.paragraph_format.first_line_indent.cm, 2) == 1.25
        assert p.runs[0].font.size.pt == 14

    def test_emoji_stripped(self):
        doc = render("<p>Договор готов ⚖️🙌 к подписанию → сегодня</p>")
        assert "Договор готов" in doc.paragraphs[0].text
        for bad in ("⚖", "🙌", "→"):
            assert bad not in doc.paragraphs[0].text


class TestTables:
    def test_bordered_data_table(self):
        doc = render("<table><tr><th>Товар</th><th>Цена</th></tr><tr><td>Ткань</td><td>100</td></tr></table>")
        t = doc.tables[0]
        assert t.cell(0, 0).text == "Товар"
        assert t.cell(1, 1).text == "100"
        assert t._tbl.tblPr.find(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblBorders") is not None

    def test_borderless_requisites(self):
        doc = render('<table border="0"><tr><td>Поставщик:<br>ООО «Ромашка»</td>'
                     '<td>Покупатель:<br>ООО «Лютик»</td></tr></table>')
        t = doc.tables[0]
        assert "Ромашка" in t.cell(0, 0).text
        assert t._tbl.tblPr.find(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblBorders") is None


class TestDocumentDefaults:
    def test_gost_margins_and_font(self):
        doc = render("<p>Текст.</p>")
        s = doc.sections[0]
        assert round(s.left_margin.cm, 1) == 3.0
        assert round(s.right_margin.cm, 1) == 1.5
        run = doc.paragraphs[0].runs[0]
        assert run.font.name == "Times New Roman"

    def test_line_spacing_option(self):
        doc = render("<p>Текст.</p>", line_spacing=1.5)
        assert doc.paragraphs[0].paragraph_format.line_spacing == 1.5

    def test_unknown_tags_text_flows(self):
        doc = render("<article><section><p>Внутри неизвестных тегов.</p></section></article>")
        assert "Внутри неизвестных тегов." in texts(doc)
