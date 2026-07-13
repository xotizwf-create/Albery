"""Unit tests for docedit — targeted in-place document edits (the anti-«пересобрал с нуля» tool)."""
import io
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[2]))

import pytest

import docedit


def _xlsx_bytes():
    from openpyxl import load_workbook, Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Регламент"
    ws["A1"] = "Контрольная встреча"
    ws["B1"] = "пятница, 15:00"
    ws["C1"] = "=SUM(1,2)"  # formula must survive untouched
    ws2 = wb.create_sheet("Прочее")
    ws2["A1"] = "пятница"
    out = io.BytesIO()
    wb.save(out)
    return out.getvalue()


def test_xlsx_edit_replaces_only_strings_and_keeps_formulas():
    data = _xlsx_bytes()
    new, counts, warns = docedit.apply_edits(data, "r.xlsx", [("пятница", "понедельник")])
    assert counts == [2] and not warns
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(new))
    assert wb["Регламент"]["B1"].value == "понедельник, 15:00"
    assert wb["Регламент"]["C1"].value == "=SUM(1,2)"
    assert wb["Прочее"]["A1"].value == "понедельник"


def test_xlsx_sheet_scope_limits_edit():
    data = _xlsx_bytes()
    new, counts, _ = docedit.apply_edits(data, "r.xlsx", [("пятница", "понедельник")], sheet="Прочее")
    assert counts == [1]
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(new))
    assert wb["Регламент"]["B1"].value == "пятница, 15:00"  # untouched
    assert wb["Прочее"]["A1"].value == "понедельник"


def test_xlsx_unknown_sheet_is_a_clear_error():
    with pytest.raises(docedit.UnsupportedFormat, match="Доступные листы"):
        docedit.apply_edits(_xlsx_bytes(), "r.xlsx", [("x", "y")], sheet="Нет такого")


def test_docx_edit_survives_split_runs_and_tables():
    from docx import Document
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Перенос с ")
    bold = p.add_run("пят")
    bold.bold = True
    p.add_run("ницы")
    tbl = doc.add_table(rows=1, cols=1)
    tbl.rows[0].cells[0].text = "день: пятница"
    buf = io.BytesIO()
    doc.save(buf)
    new, counts, _ = docedit.apply_edits(buf.getvalue(), "d.docx",
                                         [("пятницы", "понедельника"), ("пятница", "понедельник")])
    assert counts == [1, 1]
    d2 = Document(io.BytesIO(new))
    assert d2.paragraphs[0].text == "Перенос с понедельника"
    assert d2.paragraphs[0].runs[1].bold  # formatting of the run the match started in survives
    assert d2.tables[0].rows[0].cells[0].text == "день: понедельник"


def test_text_zero_match_counts_and_replacement():
    new, counts, _ = docedit.apply_edits("a;пятница\n".encode(), "t.csv",
                                         [("пятница", "понедельник"), ("нету", "x")])
    assert counts == [1, 0]
    assert new.decode() == "a;понедельник\n"


def test_unsupported_format_raises():
    with pytest.raises(docedit.UnsupportedFormat):
        docedit.apply_edits(b"%PDF-1.4", "scan.pdf", [("a", "b")])
