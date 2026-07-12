from __future__ import annotations

import json
import calendar
import csv
import io
import importlib
import logging
import os
import re
import sys
import time
import urllib.error
import urllib.request
from datetime import date, datetime, timedelta, timezone
from decimal import Decimal
from pathlib import Path
from typing import Any
from urllib.parse import urlparse, urlunparse
from uuid import UUID
from zoneinfo import ZoneInfo

from psycopg.types.json import Jsonb

from shared.db import connect as pg_connection, load_env_value as shared_load_env_value, normalize_postgres_url as shared_normalize_postgres_url


logging.basicConfig(
    level=os.getenv("LOG_LEVEL", "INFO").upper(),
    format="%(asctime)s %(levelname)s %(name)s %(message)s",
    stream=sys.stderr,
)
logger = logging.getLogger(__name__)

ROOT = Path(__file__).resolve().parents[1]
ENV_PATH = ROOT / ".env"
SERVER_NAME = "employee-analytics-context"
SERVER_VERSION = "0.18.0"
PROTOCOL_VERSION = "2024-11-05"
MAX_LIMIT = 500
ZOOM_TRANSCRIPT_MAX_LIMIT = 2000
TOOL_USAGE_CONTRACT = (
    "Call start_here_always_read_ai_instructions first; if scope is unclear, "
    "ask one short clarifying question before guessing. "
)
REFERENCE_CACHE_TTL_SECONDS = int(os.getenv("MCP_REFERENCE_CACHE_TTL_SECONDS", "60") or "60")
TOOL_LATENCY_LOG_MS = float(os.getenv("MCP_TOOL_LATENCY_LOG_MS", "250") or "250")
_TTL_CACHE: dict[tuple[Any, ...], tuple[float, Any]] = {}


class McpError(Exception):
    def __init__(self, code: int, message: str) -> None:
        super().__init__(message)
        self.code = code
        self.message = message


def ttl_cache_get(key: tuple[Any, ...]) -> Any | None:
    cached = _TTL_CACHE.get(key)
    if not cached:
        return None
    expires_at, value = cached
    if expires_at < time.time():
        _TTL_CACHE.pop(key, None)
        return None
    return value


def ttl_cache_set(key: tuple[Any, ...], value: Any, ttl_seconds: int = REFERENCE_CACHE_TTL_SECONDS) -> Any:
    _TTL_CACHE[key] = (time.time() + ttl_seconds, value)
    return value


def ttl_cache_delete_prefix(prefix: tuple[Any, ...]) -> None:
    for key in list(_TTL_CACHE):
        if key[: len(prefix)] == prefix:
            _TTL_CACHE.pop(key, None)


def load_database_url() -> str:
    value = shared_load_env_value("DATABASE_URL")
    if not value:
        raise McpError(-32000, "DATABASE_URL is not set in environment or .env")
    return normalize_postgres_url(value)


def load_env_value(key: str) -> str:
    return shared_load_env_value(key)


def normalize_postgres_url(database_url: str) -> str:
    return shared_normalize_postgres_url(database_url)


def connect() -> Any:
    try:
        return pg_connection()
    except RuntimeError as exc:
        raise McpError(-32000, str(exc)) from exc


_MSK_TZ = ZoneInfo("Europe/Moscow")


def _to_msk(value: datetime) -> datetime:
    """Render datetimes in Moscow time (UTC+3) for all MCP output. The DB stores timestamptz and
    the connection runs in UTC, so a bare .isoformat() would leak UTC to the model (e.g. Zoom call
    times). Naive datetimes are assumed to already be UTC (the server runs in UTC)."""
    if value.tzinfo is None:
        value = value.replace(tzinfo=timezone.utc)
    return value.astimezone(_MSK_TZ)


def json_default(value: Any) -> Any:
    if isinstance(value, datetime):
        return _to_msk(value).isoformat()
    if isinstance(value, date):
        return value.isoformat()
    if isinstance(value, Decimal):
        return float(value)
    if isinstance(value, UUID):
        return str(value)
    return str(value)


def text_response(payload: Any) -> dict[str, Any]:
    structured = json_safe(payload)
    return {
        "structuredContent": structured,
        "content": [
            {
                "type": "text",
                "text": json.dumps(structured, ensure_ascii=False, default=json_default, indent=2),
            }
        ]
    }


def local_app_base_url() -> str:
    value = os.getenv("MCP_LOCAL_APP_BASE_URL", "").strip() or os.getenv("APP_BASE_URL", "").strip()
    return value.rstrip("/") if value else "http://127.0.0.1:5002"


# app.py is being split module-by-module (move-only refactor), so a workflow may now
# live in an extracted module; resolve across all of them, app first.
WORKFLOW_MODULES = ("app", "bitrix", "gdrive", "zoom", "b24bot", "llm", "utils", "agent_center")


def app_workflow_function(name: str) -> Any:
    import_errors: list[str] = []
    for module_name in WORKFLOW_MODULES:
        try:
            module = importlib.import_module(module_name)
        except Exception as exc:  # noqa: BLE001
            import_errors.append(f"{module_name}: {exc}")
            continue
        workflow = getattr(module, name, None)
        if callable(workflow):
            return workflow
    detail = (" (import errors: " + "; ".join(import_errors) + ")") if import_errors else ""
    raise McpError(-32000, f"Local app workflow is not available: {name}{detail}")


def json_safe(value: Any) -> Any:
    if isinstance(value, dict):
        return {str(key): json_safe(item) for key, item in value.items()}
    if isinstance(value, list):
        return [json_safe(item) for item in value]
    if isinstance(value, datetime):
        return _to_msk(value).isoformat()
    if isinstance(value, date):
        return value.isoformat()
    if isinstance(value, Decimal):
        return float(value)
    if isinstance(value, UUID):
        return str(value)
    return value


def parse_date_arg(args: dict[str, Any], name: str, required: bool = True) -> date | None:
    raw = args.get(name)
    if raw in (None, ""):
        if required:
            raise McpError(-32602, f"Missing required argument: {name}")
        return None
    try:
        return date.fromisoformat(str(raw))
    except ValueError as exc:
        raise McpError(-32602, f"{name} must use YYYY-MM-DD format") from exc


def parse_limit(args: dict[str, Any], default: int = 100, max_limit: int = MAX_LIMIT) -> int:
    raw = args.get("limit", default)
    try:
        limit = int(raw)
    except (TypeError, ValueError) as exc:
        raise McpError(-32602, "limit must be an integer") from exc
    return max(1, min(limit, max_limit))


def parse_offset(args: dict[str, Any]) -> int:
    raw = args.get("offset", 0)
    try:
        offset = int(raw)
    except (TypeError, ValueError) as exc:
        raise McpError(-32602, "offset must be an integer") from exc
    return max(0, offset)


def is_ocr_supported_file(file_name: Any, file_type: Any, mime_type: Any) -> bool:
    markers = [file_type, mime_type, file_name]
    for marker in markers:
        value = str(marker or "").lower().strip()
        if not value:
            continue
        if value == "image" or value.startswith("image/"):
            return True
        if value == "pdf" or value == "application/pdf" or value.endswith(".pdf"):
            return True
        if value.endswith((".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp", ".tif", ".tiff")):
            return True
    return False


def safe_table_exists(cur: Any, table_name: str) -> bool:
    cur.execute("SELECT to_regclass(%s) IS NOT NULL AS exists", (f"public.{table_name}",))
    return bool(cur.fetchone()["exists"])


def column_exists(cur: Any, table_name: str, column_name: str) -> bool:
    cur.execute(
        """
        SELECT EXISTS (
            SELECT 1 FROM information_schema.columns
            WHERE table_schema = 'public' AND table_name = %s AND column_name = %s
        ) AS exists
        """,
        (table_name, column_name),
    )
    return bool(cur.fetchone()["exists"])


def tool_health(_: dict[str, Any]) -> dict[str, Any]:
    url = load_database_url()
    parsed = urlparse(url)
    database_name = parsed.path.lstrip("/")
    safe_url = urlunparse((parsed.scheme, parsed.netloc.split("@")[-1], parsed.path, "", "", ""))
    with connect() as conn:
        with conn.cursor() as cur:
            cur.execute("SELECT now() AS server_time, current_database() AS database")
            row = cur.fetchone()
    return {
        "status": "ok",
        "server": SERVER_NAME,
        "database": row["database"] or database_name,
        "database_url": safe_url,
        "server_time": row["server_time"],
    }


def tool_get_runtime_status(_: dict[str, Any]) -> dict[str, Any]:
    url = load_database_url()
    parsed = urlparse(url)
    safe_url = urlunparse((parsed.scheme, parsed.netloc.split("@")[-1], parsed.path, "", "", ""))
    with connect() as conn:
        with conn.cursor() as cur:
            cur.execute("SELECT current_database() AS database, current_schema() AS schema, now() AS server_time")
            row = cur.fetchone()
    return {
        "mode": "mcp_first_postgresql_only",
        "database": row["database"],
        "schema": row["schema"],
        "database_url": safe_url,
        "server_time": row["server_time"],
        "legacy_http_api_enabled": os.getenv("ALLOW_LEGACY_HTTP_API", "").strip() == "1",
        "path_token_auth_enabled": os.getenv("MCP_ALLOW_PATH_TOKEN", "").strip() == "1",
        "reference_cache_ttl_seconds": REFERENCE_CACHE_TTL_SECONDS,
        "rules": [
            "Use MCP tools as the primary interface for AI agents.",
            "Read and write company data through PostgreSQL-backed MCP tools.",
            "Do not depend on /api/* HTTP routes for MCP workflows.",
        ],
    }


def tool_list_available_sources(_: dict[str, Any]) -> dict[str, Any]:
    table_names = [
        "users",
        "departments",
        "user_departments",
        "bitrix_tasks",
        "chats",
        "chat_messages",
        "chat_message_files",
        "chat_file_ocr",
        "owner_daily_reports",
        "owner_weekly_reports",
        "owner_manager_recommendations",
        "owner_recommendation_dispatches",
        "owner_recommendation_events",
        "company_profile",
        "company_folders",
        "company_drive_sources",
        "ai_instruction_folders",
        "zoom_calls",
        "zoom_call_participants",
        "zoom_call_transcript_segments",
    ]
    with connect() as conn:
        with conn.cursor() as cur:
            sources: dict[str, dict[str, Any]] = {}
            for table in table_names:
                exists = safe_table_exists(cur, table)
                count = None
                if exists:
                    cur.execute(f"SELECT count(*) AS count FROM {table}")
                    count = cur.fetchone()["count"]
                sources[table] = {"exists": exists, "rows": count}
    return {"sources": sources}


def load_ai_instructions(
    path_prefix: str | None = None,
    allowed_paths: set[str] | list[str] | None = None,
) -> list[dict[str, Any]]:
    """Live instructions. Source of truth is the GitHub registry
    (agent_knowledge/instructions); falls back to the ai_instruction_folders DB tree
    when the registry is absent.

    ``allowed_paths`` (per-agent scoping) — when given, only instructions whose full
    path is in the set are returned. This is what makes a per-agent instruction
    selection real: an agent scoped to a subset literally cannot read the rest via
    start_here / get_ai_instructions. ``None`` = no scoping (the full tree), used by
    the admin/tier connectors.

    ``path_prefix`` (case-insensitive) — narrow to one folder subtree, so the
    assistant can re-read a single instruction instead of the whole set.
    """
    rows = _all_instruction_rows()
    if allowed_paths is not None:
        allowed = set(allowed_paths)
        rows = [row for row in rows if str(row.get("path") or "") in allowed]
    if not path_prefix:
        return rows
    needle = path_prefix.strip().lower()
    if not needle:
        return rows
    return [row for row in rows if str(row.get("path") or "").lower().startswith(needle)]


def _all_instruction_rows() -> list[dict[str, Any]]:
    """Registry-first instruction rows (git canonical, DB fallback). Rows carry the
    same ``path`` / ``content`` / ``name`` keys in both modes."""
    try:
        from agent_knowledge import load_instructions
        reg = load_instructions()
    except Exception:  # noqa: BLE001
        reg = None
    if reg is not None:
        return reg
    return _load_ai_instruction_rows()


def load_ai_instruction_index() -> list[dict[str, Any]]:
    """Compact map of instruction folders without the heavy ``content`` body.

    Lets the assistant see which instructions exist and fetch only the relevant
    one via get_ai_instructions(path=...) instead of re-reading the full tree.
    """
    index: list[dict[str, Any]] = []
    for row in _all_instruction_rows():
        content = row.get("content")
        index.append(
            {
                "id": row.get("id"),
                "path": row.get("path"),
                "name": row.get("name"),
                "content_chars": len(content) if isinstance(content, str) else 0,
                "updated_at": row.get("updated_at"),
            }
        )
    return index


def _load_ai_instruction_rows() -> list[dict[str, Any]]:
    cache_key = ("ai_instruction_rows",)
    cached = ttl_cache_get(cache_key)
    if cached is not None:
        return cached
    with connect() as conn:
        with conn.cursor() as cur:
            if not safe_table_exists(cur, "ai_instruction_folders"):
                return []
            cur.execute(
                """
                WITH RECURSIVE folder_tree AS (
                    SELECT id, parent_id, name, content, sort_order, ARRAY[name]::text[] AS path, updated_at
                    FROM ai_instruction_folders
                    WHERE parent_id IS NULL
                    UNION ALL
                    SELECT child.id, child.parent_id, child.name, child.content, child.sort_order,
                           folder_tree.path || child.name, child.updated_at
                    FROM ai_instruction_folders child
                    JOIN folder_tree ON folder_tree.id = child.parent_id
                )
                SELECT id, parent_id, name, array_to_string(path, ' / ') AS path, content, updated_at
                FROM folder_tree
                ORDER BY path
                """
            )
            return ttl_cache_set(cache_key, cur.fetchall())


INTENT_SOURCE_MAP: dict[str, list[str]] = {
    "company_rule_question": ["company_knowledge", "organization"],
    "employee_period_question": ["organization", "bitrix_tasks", "bitrix_chats", "zoom_calls", "owner_reports"],
    "chat_event_question": ["bitrix_chats", "organization"],
    "bitrix_task_creation": ["bitrix_tasks", "organization"],
    "recommendation_answer": ["owner_reports", "company_knowledge", "bitrix_tasks", "bitrix_chats", "zoom_calls"],
    "owner_daily_report_creation": ["owner_reports", "bitrix_chats", "zoom_calls", "company_knowledge", "bitrix_tasks", "organization"],
    "owner_weekly_report_creation": ["owner_reports", "bitrix_chats", "zoom_calls", "company_knowledge", "bitrix_tasks", "organization"],
}

# Aliases so a loose intent string still routes to the right workflow.
INTENT_ALIASES: dict[str, str] = {
    "company_rule": "company_rule_question",
    "rules": "company_rule_question",
    "regulation": "company_rule_question",
    "employee": "employee_period_question",
    "period": "employee_period_question",
    "analytics": "employee_period_question",
    "chat": "chat_event_question",
    "chat_event": "chat_event_question",
    "create_task": "bitrix_task_creation",
    "bitrix_task": "bitrix_task_creation",
    "recommendation": "recommendation_answer",
    "advice": "recommendation_answer",
    "owner_daily": "owner_daily_report_creation",
    "company_daily_report": "owner_daily_report_creation",
    "owner_daily_report": "owner_daily_report_creation",
    "owner_weekly": "owner_weekly_report_creation",
    "owner_weekly_report": "owner_weekly_report_creation",
}


def resolve_intent(raw: str) -> str | None:
    key = (raw or "").strip().lower()
    if not key:
        return None
    if key in INTENT_SOURCE_MAP:
        return key
    return INTENT_ALIASES.get(key)


def tool_get_context_guide(args: dict[str, Any] | None = None) -> dict[str, Any]:
    args = args or {}
    instructions_index = load_ai_instruction_index()
    full_guide = {
        "purpose": "Navigation guide for using this MCP server systematically instead of guessing where data lives.",
        "ai_instructions_index": instructions_index,
        "ai_instructions_note": (
            "Full live instruction text was already returned by start_here_always_read_ai_instructions. "
            "To re-read one folder by path, call get_ai_instructions(path='Folder / Subfolder'). "
            "Pass intent to get_context_guide to receive only the workflow and sources for the current task."
        ),
        "operating_rules": [
            "Before any substantive answer or data work, call start_here_always_read_ai_instructions and follow its live instructions exactly.",
            "For unfamiliar questions, call get_context_guide after start_here_always_read_ai_instructions, then list_available_sources if source freshness or row counts matter.",
            "Act as an internal company AI agent: answers must be based on company context, regulations, reports, Bitrix tasks, chats, Zoom, and live AI instructions, not generic advice.",
            "For company rules, regulations, document mirrors, and persistent business knowledge, use search_company_knowledge first.",
            "For recommendations, management advice, or owner-facing conclusions, read recent owner reports and raw chat transcripts/OCR before concluding what is done, open, overdue, or repeated.",
            "If the request is vague, ambiguous, underspecified, or can be interpreted in several ways, stop and ask a short clarifying question before using data tools or answering.",
            "Ask what exact date/period, chat, person, task, source, output format, target decision, or save/write action is needed. Do not guess missing scope.",
            "Always use concrete names and task titles. Do not write only 'task 318099' or 'Natalia task'; write 'task 318099: Сформировать реестр платежей' with responsible person, status, deadline, and source when available.",
            "For employee identity, managers, departments, and Bitrix user ids, use get_org_structure before person-specific filters.",
            "For a date period, call get_period_index before reading messages/tasks; it shows where data exists and which chats are active.",
            "For task status, ownership, deadlines, and Bitrix work items, use search_tasks; when a task row shows comments_human_count > 0, read the actual discussion with get_task_comments(bitrix_task_id).",
            "For creating Bitrix tasks, use create_bitrix_task only when the user provided a task title, one responsible person, and a deadline. If any of these are missing, ask for the missing field instead of creating the task.",
            "For task comments, use add_bitrix_task_comment only with an exact bitrix_task_id. If the task reference is ambiguous, search_tasks first.",
            "For reopening a completed Bitrix task, first read the task/result/comments, explain why the result is unsatisfactory, and call reopen_bitrix_task with reason and confirm=true only after explicit confirmation or a standing review instruction.",
            "For deleting Bitrix tasks, first identify exactly one bitrix_task_id, show the user the task title/status/responsible/deadline, and ask for explicit confirmation. Only after the user confirms deletion, call delete_bitrix_task with confirm=true.",
            "For discussion evidence, commitments, blockers, decisions, and OCR from chat images, use list_chats then search_messages or get_chat_transcript.",
            "For meeting evidence, use list_zoom_calls first, then search_zoom_transcripts with topic keywords, then get_zoom_call_transcript for matching calls, and get_org_structure before generating a Zoom report.",
            "For cross-source reports over a bounded date range, use get_compact_export, then deepen with source-specific tools.",
            "This project is MCP-first and PostgreSQL-only: do not call or depend on legacy /api/* HTTP routes for AI workflows.",
            "Prefer narrow queries with date ranges, dialog_id, responsible_bitrix_user_id, or search text. Page with offset instead of requesting everything.",
            "Do not treat absence of one search result as proof until the relevant source index/count was checked.",
            "When evidence is incomplete, separate confirmed facts from assumptions and ask the user for the missing input needed to continue.",
        ],
        "source_map": {
            "company_knowledge": {
                "tools": ["search_company_knowledge", "list_company_files", "get_company_file", "get_company_profile"],
                "tables": ["company_folders", "company_drive_sources", "company_profile"],
                "use_for": ["company rules", "regulations", "Google Drive mirrored documents", "static process knowledge"],
            },
            "organization": {
                "tools": ["get_org_structure"],
                "tables": ["users", "departments", "user_departments"],
                "use_for": ["people", "roles", "managers", "departments", "Bitrix user ids"],
            },
            "bitrix_tasks": {
                "tools": ["search_tasks", "get_task_comments", "add_bitrix_task_comment", "create_bitrix_task", "reopen_bitrix_task", "delete_bitrix_task"],
                "tables": ["bitrix_tasks", "bitrix_task_members", "bitrix_task_snapshots"],
                "use_for": ["task ownership", "deadlines", "statuses", "overdue work", "responsibility", "task discussion and comments", "adding Bitrix task comments", "creating Bitrix tasks with required title/responsible/deadline", "reopening completed tasks with reason/comment", "deleting Bitrix tasks only after exact id lookup and explicit confirmation"],
            },
            "bitrix_chats": {
                "tools": ["get_report_readiness", "list_chats", "search_messages", "get_chat_transcript", "get_chat_ocr_status", "process_chat_ocr"],
                "tables": ["chats", "chat_messages", "chat_message_files", "chat_file_ocr"],
                "use_for": ["conversation evidence", "commitments", "decisions", "questions", "OCR from screenshots", "raw chat transcript retrieval"],
                "rules": [
                    "Daily chat reports are disabled. Do not create, request, or wait for chat_daily_reports.",
                    "Use get_chat_transcript with include_ocr=true after OCR is ready.",
                    "process_chat_ocr calls the local PostgreSQL workflow directly from MCP; it does not require a local HTTP API route.",
                ],
            },
            "owner_reports": {
                "tools": ["get_report_readiness", "get_previous_owner_daily_context", "get_owner_reports", "save_owner_daily_report", "save_owner_weekly_report", "list_recommendations", "get_recommendation_feedback_context", "save_recommendation_event"],
                "tables": ["owner_daily_reports", "owner_weekly_reports", "owner_manager_recommendations", "owner_recommendation_dispatches", "owner_recommendation_events"],
                "use_for": ["recent owner context", "general daily reports for owner", "general weekly reports for owner", "owner-level report storage", "recommendation lifecycle", "recommendation feedback and statuses"],
            },
            "zoom_calls": {
                "tools": ["list_zoom_calls", "get_zoom_call_transcript", "search_zoom_transcripts", "save_zoom_call_report", "delete_zoom_call_report"],
                "tables": ["zoom_calls", "zoom_call_participants", "zoom_call_transcript_segments"],
                "use_for": ["meeting transcripts", "call participants from Zoom API", "mentioned people in transcript", "spoken decisions", "facts of task execution", "standalone Zoom report storage"],
                "rules": [
                    "For standalone Zoom reports, always include factual participants, mentioned people, a strict task block with owner/deadline/success-criteria gaps, and behavioral factors.",
                    "For chat context, Zoom relevance is based on transcript content, participants, and topics from chat/OCR/tasks, not only call title.",
                    "If the report date has exactly one Zoom call, read its transcript before saying there are no relevant Zoom facts.",
                    "Search transcript keywords extracted from OCR/tasks/risks, for example payment calendar, motivation, margins, project names, overdue work, Bitrix, and owner names.",
                ],
            },
        },
        "recommended_workflows": {
            "company_rule_question": ["search_company_knowledge(query)", "list_company_files() to inspect all documents", "get_company_file(folder_id/google_file_id) to read full content", "get_company_profile() if full tree is needed"],
            "employee_period_question": [
                "get_org_structure(include_inactive=false)",
                "get_owner_reports(report_kind='daily', limit=7) and get_owner_reports(report_kind='weekly', limit=4) if the answer includes recommendations or management conclusions",
                "get_period_index(date_from,date_to)",
                "search_tasks(date_from,date_to,responsible_bitrix_user_id)",
                "search_messages(date_from,date_to,query/person name)",
                "search_zoom_transcripts(date_from,date_to,query/person name)",
            ],
            "chat_event_question": ["list_chats(date_from,date_to,query)", "get_chat_transcript(dialog_id,date_from,date_to)"],
            "bitrix_task_creation": [
                "Verify the user provided title, responsible person, and deadline.",
                "If responsible person is a name, create_bitrix_task resolves it through org structure; if ambiguous, ask for responsible_bitrix_user_id.",
                "Call create_bitrix_task(title, responsible_name/responsible_bitrix_user_id, deadline, description).",
                "Return created task_id, responsible, deadline, and title.",
                "For comments: add_bitrix_task_comment(bitrix_task_id, comment_text) only after resolving the exact task.",
                "For reopening: search_tasks + get_task_comments/result first, explain the unsatisfactory result, then reopen_bitrix_task(bitrix_task_id, reason, confirm=true) after confirmation/standing review instruction.",
                "For deletion: search_tasks(bitrix_task_id=...) first, show the exact task title/status/responsible/deadline, ask the user to confirm deletion, then call delete_bitrix_task(bitrix_task_id=..., confirm=true). Never delete by name, search text, or ambiguous reference.",
            ],
            "recommendation_answer": [
                "instructions already arrived from start_here_always_read_ai_instructions; re-read a specific folder only if needed via get_ai_instructions(path=...)",
                "search_company_knowledge(query) for company rules and regulations",
                "get_owner_reports(report_kind='daily', limit=7)",
                "get_owner_reports(report_kind='weekly', limit=4)",
                "get_period_index(date_from,date_to) for the requested period",
                "list_recommendations(status='open', date_from,date_to) and recommendation events for feedback continuity",
                "search_tasks/search_messages/search_zoom_transcripts for concrete evidence",
                "answer with specific task titles, owners, statuses, deadlines, and sources; ask clarifying questions when evidence is missing",
            ],
            "owner_daily_report_creation": [
                "instructions already arrived from start_here_always_read_ai_instructions; re-read only if needed via get_ai_instructions(path='Формирование отчетов / Ежедневный отчет по компании')",
                "open the active AI prompt in Сводная аналитика / Настройка промтов / ежедневный общий отчет для собственника and follow it as the report contract",
                "get_report_readiness(date_from=report_date,date_to=report_date) ONCE to get, in a single call: active chats with messages/OCR readiness, same-day Zoom calls and which already have an analytical_note, and whether the previous owner daily report exists — use this instead of probing each chat/Zoom one by one",
                "read company regulations with search_company_knowledge before writing recommendations; compare Bitrix/Zoom/chat facts against regulated owners, process roles, payment calendar ownership, approval rules, meeting rhythm, and SLA",
                "for every active chat with messages, read get_chat_transcript(..., include_ocr=true); daily chat reports are disabled and must not be generated",
                "for every Zoom call in missing_zoom_reports from get_report_readiness, use zoom_call_report instructions and save_zoom_call_report before continuing",
                "if previous_owner_daily_report_exists is false in readiness, stop or create the missing previous day first; otherwise get_previous_owner_daily_context(report_date) for continuity",
                "read recommendation feedback from raw chat transcripts and recommendation event context; every addressable recommendation must account for the recipient's previous reply, objection, delegation, unclear answer, or missing response — but recommendations are now delivered as one Bitrix task per recipient (title 'Рекомендации DD.MM', deadline 10:00 next day, non-movable), so write each person's recommendations as a clean numbered list WITHOUT any greeting or salutation",
                "for every addressable recommendation, explicitly use regulation comparison when relevant: if the actual owner/executor/deadline/process differs from company regulation, mention the regulated owner/process and propose delegation, confirmation, Bitrix fixation, or regulation update",
                "only after every needed raw chat transcript/OCR, every Zoom analytical report, and previous owner_daily_report are ready, create owner_daily_report",
                "if any required source is missing or failed, stop and return the missing chat/Zoom/OCR source list instead of writing owner_daily_report",
            ],
            "owner_weekly_report_creation": [
                "get_report_readiness(date_from=week_start,date_to=week_end) ONCE to see per-day readiness across the whole week (chats/Zoom/owner daily) in a single call before deepening",
                "for each day in the week, run owner_daily_report_creation until the daily chain is complete",
                "use raw chat transcripts/OCR for weekly chat context; daily chat reports are disabled",
                "create or refresh chat_overall_weekly_report",
                "only then create owner_weekly_report",
                "if any daily source chain is incomplete, stop and return the incomplete days and missing sources",
            ],
        },
    }

    intent = resolve_intent(str(args.get("intent") or ""))
    if intent is None:
        return full_guide

    # Task-scoped view: only the workflow and sources relevant to this intent,
    # so the assistant reads one route instead of the whole guide.
    source_keys = INTENT_SOURCE_MAP.get(intent, [])
    return {
        "purpose": full_guide["purpose"],
        "intent": intent,
        "operating_rules": full_guide["operating_rules"],
        "source_map": {key: full_guide["source_map"][key] for key in source_keys if key in full_guide["source_map"]},
        "workflow": full_guide["recommended_workflows"].get(intent, []),
        "ai_instructions_index": full_guide["ai_instructions_index"],
        "ai_instructions_note": full_guide["ai_instructions_note"],
        "note": "Scoped to the requested intent. Call get_context_guide without intent to see all workflows and sources.",
    }


def tool_start_here_always_read_ai_instructions(args: dict[str, Any]) -> dict[str, Any]:
    scope = args.get("_allowed_instruction_paths")
    instructions = load_ai_instructions(allowed_paths=set(scope) if scope is not None else None)
    available_tools = list(args.get("_connector_tools") or sorted(TOOLS.keys()))
    hidden_tools = list(args.get("_connector_hidden_tools") or [])
    connector_id = args.get("_connector_id") or "full"
    is_faq = connector_id == "faq" or set(available_tools) == FAQ_TOOL_NAMES
    connector_label = "FAQ MCP (read-only)" if is_faq else "Full MCP"
    return {
        "mandatory_status": "READ_FIRST_AND_OBEY_EXACTLY",
        "purpose": "This is the mandatory entry tool for this MCP server. The assistant must read these live settings before any analysis, report, recommendation, database lookup plan, or final answer.",
        "source": "Настройки -> Инструкции для ИИ (table ai_instruction_folders)",
        "connector_scope": {
            "connector_id": connector_id,
            "connector_label": connector_label,
            "available_tools": sorted(available_tools),
            **({
                "more_tools_via_call_tool": hidden_tools,
                "two_stage_note": (
                    "ВАЖНО: инструменты из more_tools_via_call_tool тебе ТОЖЕ доступны — найди "
                    "нужный через find_tool и вызывай через call_tool(name=..., arguments={...}). "
                    "Правило «нет доступа» применяй ТОЛЬКО к тому, чего нет ни в available_tools, "
                    "ни в more_tools_via_call_tool."
                ),
            } if hidden_tools else {}),
            "rules": [
                "STEP 0 (before any planning): inspect this connector's available_tools list above. Treat it as the complete and only set of capabilities you have right now.",
                "Do not assume access to any tool, data source, table, file, integration, web search, or external service that is not in available_tools. Do not invoke or describe tools by name that are not in the list.",
                "Do not look outside this MCP connector for information: no general web knowledge, no other connectors, no prior conversations, no offline assumptions about what 'usually' exists.",
                "Plan the order of actions ONLY after Step 0: pick the tool from available_tools that matches the user's task, then follow live_ai_instructions and the execution_contract below.",
                "If the user asks for data that requires a tool not present in available_tools (for example Bitrix tasks/chats, OCR, report saving, or any source not exposed here), reply briefly in the user's language: 'Нет доступа к этой информации в текущем коннекторе.' Then list which tools you do have so the user can refine the request or switch connectors. Never guess or fabricate the missing data.",
                "If multiple tools could match, prefer the more specific one. If none match, follow the previous rule.",
                "Even when available_tools is small (FAQ connector), do not apologize, do not invent capabilities, and do not offer to do work you cannot do here.",
            ],
        },
        "live_ai_instructions": instructions,
        "execution_contract": [
            "Treat every non-empty instruction as binding for the current user request.",
            "Do exactly what the relevant instruction says: source order, report format, required checks, save/read workflow, and stopping conditions.",
            "If a specific report contract is required by the instructions, call get_report_contract for that category before generating the report.",
            "If instructions require missing source checks, OCR, Zoom analysis, previous reports, or Bitrix refresh, complete those checks before conclusions.",
            "If the request conflicts with these instructions, explain the conflict and ask the user to update Настройки -> Инструкции для ИИ or confirm a one-off exception.",
            "If the user request is vague, ambiguous, or missing the needed scope, ask one concise clarifying question first. Do not infer dates, chats, people, report type, or whether to save/write unless the user said it clearly.",
            "For Bitrix task creation, never call create_bitrix_task unless the user provided all three required fields: task title, exactly one responsible person, and a deadline. If any field is missing or the responsible person is ambiguous, ask for clarification and do not create anything.",
            "For Bitrix task comments, never call add_bitrix_task_comment without an exact bitrix_task_id and non-empty comment_text.",
            "For Bitrix task reopening, never call reopen_bitrix_task without an exact bitrix_task_id, a reason, and confirm=true after checking the result/comments.",
            "For Bitrix task deletion, never call delete_bitrix_task unless the user has already confirmed deletion of one exact bitrix_task_id after seeing its title/status/responsible/deadline.",
            "When instructions are incomplete for the task, continue with get_context_guide and the relevant source tools instead of guessing.",
            "If an instruction names a tool that is not in connector_scope.available_tools, skip that step and tell the user that the action requires a connector that exposes that tool. Do not pretend the step succeeded.",
        ],
        "next_tool_guidance": [
            "The live_ai_instructions above are the full text for this request; do not re-fetch them with get_ai_instructions unless you need a folder by path after the conversation grew long.",
            "Use get_context_guide(intent='owner_daily_report_creation'|'recommendation_answer'|...) to get only the workflow and sources for the current task instead of the whole guide.",
            "Use get_report_contract when generating configured reports.",
            "Use get_report_readiness(date_from,date_to) before building daily/weekly/owner reports to learn in one call what is missing.",
            "Use list_available_sources when freshness, availability, or row counts matter.",
            "When the user asks what you can do / your capabilities ('что ты умеешь', 'твои возможности'), call get_ai_capabilities — it returns the human-readable list for the current connector/tool set; answer only within it, and keep it current with update_ai_capabilities when that tool is available.",
        ],
    }

def tool_get_ai_instructions(args: dict[str, Any] | None = None) -> dict[str, Any]:
    args = args or {}
    path = str(args.get("path") or "").strip()
    scope = args.get("_allowed_instruction_paths")
    instructions = load_ai_instructions(path or None, allowed_paths=set(scope) if scope is not None else None)
    note = "These instructions are loaded live from ai_instruction_folders. Edit them in the UI: Settings -> AI instructions."
    if path:
        note = (
            f"Filtered to folders whose path starts with '{path}'. "
            "Omit path to read the full tree, or use get_context_guide for the index of available paths."
        )
    return {
        "instructions": instructions,
        "path": path or None,
        "note": note,
    }


def tool_get_report_contract(args: dict[str, Any]) -> dict[str, Any]:
    category_key = str(args.get("category_key") or "").strip()
    if not category_key:
        raise McpError(-32602, "Missing required argument: category_key")

    with connect() as conn:
        with conn.cursor() as cur:
            if not safe_table_exists(cur, "ai_prompt_categories") or not safe_table_exists(cur, "ai_prompts"):
                return {
                    "category_key": category_key,
                    "contract": None,
                    "message": "Report contract tables do not exist yet.",
                }
            cur.execute(
                """
                SELECT
                    c.category_key,
                    c.title AS category_title,
                    c.description AS category_description,
                    p.id AS contract_id,
                    p.prompt_key AS contract_key,
                    p.title AS contract_title,
                    p.prompt_text AS contract_text,
                    p.version,
                    p.created_at
                FROM ai_prompt_categories c
                JOIN ai_prompts p ON p.category_id = c.id
                WHERE c.is_active = TRUE
                  AND p.is_active = TRUE
                  AND c.category_key = %s
                ORDER BY p.version DESC, p.created_at DESC
                LIMIT 1
                """,
                (category_key,),
            )
            row = cur.fetchone()

    return {
        "category_key": category_key,
        "contract": row,
        "note": (
            "Use contract_text as the exact report-generation contract. "
            "Daily/weekly chat reports are disabled; use raw chat transcript tools instead."
        ),
    }


def tool_get_company_profile(_: dict[str, Any]) -> dict[str, Any]:
    cache_key = ("company_profile",)
    cached = ttl_cache_get(cache_key)
    if cached is not None:
        return cached
    with connect() as conn:
        with conn.cursor() as cur:
            folders: list[dict[str, Any]] = []
            if safe_table_exists(cur, "company_folders"):
                cur.execute(
                    """
                    SELECT id, parent_id, name, content, sort_order, updated_at
                    FROM company_folders
                    ORDER BY parent_id NULLS FIRST, sort_order, lower(name), created_at
                    """
                )
                folders = cur.fetchall()
            if not safe_table_exists(cur, "company_profile"):
                return {
                    "title": "О компании",
                    "content": "",
                    "folders": folders,
                    "updated_at": None,
                    "message": "company_profile table does not exist yet.",
                }
            cur.execute(
                """
                SELECT title, content, updated_at
                FROM company_profile
                WHERE profile_key = 'main'
                """
            )
            row = cur.fetchone()
    if not row:
        return ttl_cache_set(cache_key, {"title": "О компании", "content": "", "folders": folders, "updated_at": None})
    return ttl_cache_set(cache_key, {
        "title": row["title"] or "О компании",
        "content": row["content"] or "",
        "folders": folders,
        "updated_at": row["updated_at"],
    })


def tool_list_company_files(args: dict[str, Any]) -> dict[str, Any]:
    include_empty = bool(args.get("include_empty", False))
    limit = parse_limit(args, 500)
    offset = parse_offset(args)

    filters = []
    if not include_empty:
        filters.append("COALESCE(f.content, '') <> ''")
    where_sql = "WHERE " + " AND ".join(filters) if filters else ""

    with connect() as conn:
        with conn.cursor() as cur:
            if not safe_table_exists(cur, "company_folders"):
                return {"items": [], "limit": limit, "offset": offset, "message": "company_folders table does not exist yet."}

            drive_join = ""
            drive_select = """
                NULL::text AS google_file_id,
                NULL::text AS source_url,
                NULL::text AS mime_type,
                NULL::timestamptz AS google_updated_at
            """
            if safe_table_exists(cur, "company_drive_sources"):
                drive_join = "LEFT JOIN company_drive_sources ds ON ds.folder_id = f.id"
                drive_select = """
                    ds.google_file_id,
                    ds.source_url,
                    ds.mime_type,
                    ds.google_updated_at
                """

            cur.execute(
                f"""
                WITH RECURSIVE folder_tree AS (
                    SELECT id, parent_id, name, ARRAY[name]::text[] AS path
                    FROM company_folders
                    WHERE parent_id IS NULL
                    UNION ALL
                    SELECT child.id, child.parent_id, child.name, folder_tree.path || child.name
                    FROM company_folders child
                    JOIN folder_tree ON folder_tree.id = child.parent_id
                )
                SELECT
                    f.id AS folder_id,
                    f.parent_id,
                    f.name,
                    array_to_string(ft.path, ' / ') AS path,
                    char_length(COALESCE(f.content, '')) AS content_length,
                    COALESCE(f.content, '') <> '' AS has_content,
                    f.updated_at,
                    {drive_select}
                FROM company_folders f
                JOIN folder_tree ft ON ft.id = f.id
                {drive_join}
                {where_sql}
                ORDER BY ft.path
                LIMIT %s OFFSET %s
                """,
                (limit, offset),
            )
            rows = cur.fetchall()

    return {
        "items": rows,
        "limit": limit,
        "offset": offset,
        "note": "Use get_company_file with folder_id or google_file_id to read full content.",
    }


def tool_get_company_file(args: dict[str, Any]) -> dict[str, Any]:
    folder_id = str(args.get("folder_id") or "").strip()
    google_file_id = str(args.get("google_file_id") or "").strip()
    if not folder_id and not google_file_id:
        raise McpError(-32602, "Provide folder_id or google_file_id")

    with connect() as conn:
        with conn.cursor() as cur:
            if not safe_table_exists(cur, "company_folders"):
                raise McpError(-32004, "company_folders table does not exist yet.")

            has_drive_sources = safe_table_exists(cur, "company_drive_sources")
            drive_join = ""
            drive_select = """
                NULL::text AS google_file_id,
                NULL::text AS source_url,
                NULL::text AS mime_type,
                NULL::timestamptz AS google_updated_at,
                '{}'::jsonb AS drive_raw_json
            """
            if has_drive_sources:
                drive_join = "LEFT JOIN company_drive_sources ds ON ds.folder_id = f.id"
                drive_select = """
                    ds.google_file_id,
                    ds.source_url,
                    ds.mime_type,
                    ds.google_updated_at,
                    ds.raw_json AS drive_raw_json
                """

            if google_file_id:
                if not has_drive_sources:
                    raise McpError(-32004, "company_drive_sources table does not exist yet.")
                filter_sql = "ds.google_file_id = %s"
                params: list[Any] = [google_file_id]
            else:
                filter_sql = "f.id = %s"
                params = [folder_id]

            cur.execute(
                f"""
                WITH RECURSIVE folder_tree AS (
                    SELECT id, parent_id, name, ARRAY[name]::text[] AS path
                    FROM company_folders
                    WHERE parent_id IS NULL
                    UNION ALL
                    SELECT child.id, child.parent_id, child.name, folder_tree.path || child.name
                    FROM company_folders child
                    JOIN folder_tree ON folder_tree.id = child.parent_id
                )
                SELECT
                    f.id AS folder_id,
                    f.parent_id,
                    f.name,
                    array_to_string(ft.path, ' / ') AS path,
                    f.content,
                    char_length(COALESCE(f.content, '')) AS content_length,
                    f.updated_at,
                    {drive_select}
                FROM company_folders f
                JOIN folder_tree ft ON ft.id = f.id
                {drive_join}
                WHERE {filter_sql}
                LIMIT 1
                """,
                params,
            )
            row = cur.fetchone()

    if not row:
        raise McpError(-32004, "Company file not found.")
    return row


def tool_search_company_knowledge(args: dict[str, Any]) -> dict[str, Any]:
    """RAG-style retrieval over company knowledge.

    With a ``query`` this returns the few most relevant PASSAGES (chunks ~400 tokens),
    not whole documents — whole files averaged ~24 KB and burned the model's context.
    The chunk index self-refreshes (``ensure_fresh``) so edits and Drive syncs are
    reflected automatically. Read a full document with ``get_company_file(folder_id)``.
    Without a query it returns a lightweight document list (names/paths + short preview).
    """
    query = str(args.get("query") or "").strip()
    limit = parse_limit(args, 6, max_limit=30)
    offset = parse_offset(args)

    with connect() as conn:
        with conn.cursor() as cur:
            if not safe_table_exists(cur, "company_folders"):
                return {"items": [], "limit": limit, "offset": offset, "message": "company_folders table does not exist yet."}
            chunks_ready = safe_table_exists(cur, "company_knowledge_chunks")

        # No query → lightweight listing (names, paths, previews) — never full bodies.
        if not query:
            with conn.cursor() as cur:
                cur.execute(
                    """
                    WITH RECURSIVE folder_tree AS (
                        SELECT id, parent_id, name, ARRAY[name]::text[] AS path
                        FROM company_folders WHERE parent_id IS NULL
                        UNION ALL
                        SELECT child.id, child.parent_id, child.name, folder_tree.path || child.name
                        FROM company_folders child JOIN folder_tree ON folder_tree.id = child.parent_id
                    )
                    SELECT
                        f.id AS folder_id, f.parent_id, f.name,
                        array_to_string(ft.path, ' / ') AS path,
                        left(COALESCE(f.content, ''), 200) AS preview,
                        length(COALESCE(f.content, '')) AS content_length,
                        f.updated_at
                    FROM company_folders f
                    JOIN folder_tree ft ON ft.id = f.id
                    ORDER BY f.updated_at DESC NULLS LAST, lower(f.name)
                    LIMIT %(limit)s OFFSET %(offset)s
                    """,
                    {"limit": limit, "offset": offset},
                )
                items = cur.fetchall()
            return {
                "items": items,
                "limit": limit,
                "offset": offset,
                "note": "Document list (names, paths, previews). Pass a `query` for focused passages, "
                        "or read a full document with get_company_file(folder_id).",
            }

        # Query → chunked passage retrieval (preferred), with query-expansion fallback.
        if chunks_ready:
            from shared.knowledge_chunks import ensure_fresh, search_expanded

            ensure_fresh(conn)  # cheap signature check; re-chunks only changed docs
            rows, mode = search_expanded(conn, query, limit=limit, offset=offset)
            items = [
                {
                    "folder_id": r["folder_id"],
                    "name": r["name"],
                    "path": r["path"],
                    "chunk_index": r["chunk_index"],
                    "content": r["content"],  # the focused passage
                    "snippet": r["content"],
                    "score": round(float(r["score"]), 4),
                }
                for r in rows
            ]
            if not items:
                return {
                    "items": [],
                    "limit": limit,
                    "offset": offset,
                    "note": "Ничего не найдено по этому запросу. ПЕРЕФОРМУЛИРУЙТЕ синонимами и вызовите "
                            "search_company_knowledge ещё раз перед выводом «не нашёл» (напр.: "
                            "«созвоны»→«встречи/планёрки/Zoom/созвоны», «график»→«расписание/периодичность/ритм», "
                            "«зарплата»→«оплата труда/штатное расписание/ФОТ»). Либо вызовите без query, "
                            "чтобы увидеть список документов, и откройте нужный через get_company_file(folder_id).",
                }
            note = (
                "Найдено по отдельным словам (широкий поиск) — проверьте релевантность; если не то, "
                "переформулируйте синонимами и поищите ещё раз. "
                if mode == "broad"
                else "Focused passages (chunks), not full documents. "
            )
            return {
                "items": items,
                "limit": limit,
                "offset": offset,
                "note": note + "Read the whole document with get_company_file(folder_id).",
            }

        # Fallback: chunks table not built yet (brief mid-deploy window). Legacy whole-doc
        # hybrid search, but content is truncated so it can't blow up the context.
        return _legacy_search_company_knowledge(conn, query, limit, offset)


def _legacy_search_company_knowledge(conn: Any, query: str, limit: int, offset: int) -> dict[str, Any]:
    like = f"%{query}%"
    with conn.cursor() as cur:
        has_tsv = column_exists(cur, "company_folders", "content_tsv")
        if has_tsv:
            where_sql = (
                "WHERE f.content_tsv @@ websearch_to_tsquery('russian', %(query)s) "
                "OR f.name ILIKE %(like)s OR COALESCE(f.content, '') ILIKE %(like)s "
                "OR similarity(f.name, %(query)s) >= 0.3"
            )
            order_sql = (
                "ORDER BY (ts_rank_cd(f.content_tsv, websearch_to_tsquery('russian', %(query)s)) "
                "+ 0.5 * similarity(f.name, %(query)s)) DESC, f.updated_at DESC NULLS LAST, lower(f.name)"
            )
        else:
            where_sql = "WHERE f.name ILIKE %(like)s OR COALESCE(f.content, '') ILIKE %(like)s"
            order_sql = "ORDER BY f.updated_at DESC NULLS LAST, lower(f.name)"
        cur.execute(
            f"""
            WITH RECURSIVE folder_tree AS (
                SELECT id, parent_id, name, ARRAY[name]::text[] AS path
                FROM company_folders WHERE parent_id IS NULL
                UNION ALL
                SELECT child.id, child.parent_id, child.name, folder_tree.path || child.name
                FROM company_folders child JOIN folder_tree ON folder_tree.id = child.parent_id
            )
            SELECT f.id AS folder_id, f.parent_id, f.name,
                   array_to_string(ft.path, ' / ') AS path,
                   left(COALESCE(f.content, ''), 1500) AS content,
                   f.updated_at
            FROM company_folders f JOIN folder_tree ft ON ft.id = f.id
            {where_sql}
            {order_sql}
            LIMIT %(limit)s OFFSET %(offset)s
            """,
            {"query": query, "like": like, "limit": limit, "offset": offset},
        )
        items = cur.fetchall()
    return {
        "items": items,
        "limit": limit,
        "offset": offset,
        "note": "Fallback search (chunk index not built yet); content truncated to 1500 chars. "
                "Read a full document with get_company_file(folder_id).",
    }


def tool_list_periods(args: dict[str, Any]) -> dict[str, Any]:
    limit = parse_limit(args, 90)
    with connect() as conn:
        with conn.cursor() as cur:
            cur.execute(
                """
                SELECT message_day AS period, count(*) AS messages_count
                FROM chat_messages
                GROUP BY message_day
                ORDER BY message_day DESC
                LIMIT %s
                """,
                (limit,),
            )
            chat_periods = cur.fetchall()
            report_periods: list[dict[str, Any]] = []
            zoom_periods: list[dict[str, Any]] = []
            if safe_table_exists(cur, "zoom_calls"):
                cur.execute(
                    """
                    SELECT call_date AS period, count(*) AS zoom_calls_count
                    FROM zoom_calls
                    GROUP BY call_date
                    ORDER BY call_date DESC
                    LIMIT %s
                    """,
                    (limit,),
                )
                zoom_periods = cur.fetchall()
    return {
        "chat_message_periods": chat_periods,
        "chat_report_periods": report_periods,
        "chat_reports_enabled": False,
        "zoom_call_periods": zoom_periods,
    }


def tool_get_period_index(args: dict[str, Any]) -> dict[str, Any]:
    date_from = parse_date_arg(args, "date_from")
    date_to = parse_date_arg(args, "date_to")
    if date_to < date_from:
        raise McpError(-32602, "date_to must be greater than or equal to date_from")

    with connect() as conn:
        with conn.cursor() as cur:
            cur.execute(
                """
                SELECT count(*) AS messages_count, count(DISTINCT chat_id) AS chats_count
                FROM chat_messages
                WHERE message_day BETWEEN %s AND %s
                """,
                (date_from, date_to),
            )
            messages = cur.fetchone()
            cur.execute(
                """
                SELECT count(*) AS tasks_count
                FROM bitrix_tasks
                WHERE COALESCE(updated_at_bitrix, created_at_bitrix, deadline_at, created_at)::date
                      BETWEEN %s AND %s
                """,
                (date_from, date_to),
            )
            tasks = cur.fetchone()
            zoom_counts = {"zoom_calls_count": 0, "zoom_transcript_segments_count": 0}
            if safe_table_exists(cur, "zoom_calls"):
                cur.execute(
                    """
                    SELECT
                        count(*) AS zoom_calls_count,
                        COALESCE(sum(segment_counts.segments_count), 0)::bigint AS zoom_transcript_segments_count
                    FROM zoom_calls zc
                    LEFT JOIN (
                        SELECT call_id, count(*) AS segments_count
                        FROM zoom_call_transcript_segments
                        GROUP BY call_id
                    ) segment_counts ON segment_counts.call_id = zc.id
                    WHERE zc.call_date BETWEEN %s AND %s
                    """,
                    (date_from, date_to),
                )
                zoom_counts = cur.fetchone()
            cur.execute(
                """
                SELECT c.dialog_id, c.chat_title, count(m.id) AS messages_count
                FROM chat_messages m
                JOIN chats c ON c.id = m.chat_id
                WHERE m.message_day BETWEEN %s AND %s
                GROUP BY c.dialog_id, c.chat_title
                ORDER BY messages_count DESC, c.chat_title NULLS LAST
                LIMIT 100
                """,
                (date_from, date_to),
            )
            chats = cur.fetchall()
    return {
        "period": {"date_from": date_from, "date_to": date_to},
        "counts": {**messages, **tasks, **zoom_counts},
        "top_chats": chats,
        "available_tools": [
            "get_org_structure",
            "search_tasks",
            "list_chats",
            "search_messages",
            "get_chat_transcript",
            "list_zoom_calls",
            "get_zoom_call_transcript",
            "search_zoom_transcripts",
            "get_company_profile",
            "get_compact_export",
            "get_report_readiness",
        ],
    }


def tool_get_report_readiness(args: dict[str, Any]) -> dict[str, Any]:
    """One-call source-readiness check for report building.

    Daily chat reports are disabled. For each date in the range this reports
    which active chats have raw messages, which Zoom calls have an analytical_note,
    and whether the current and previous owner daily reports exist.
    """
    date_from = parse_date_arg(args, "date_from")
    date_to = parse_date_arg(args, "date_to", required=False) or date_from
    if date_to < date_from:
        raise McpError(-32602, "date_to must be greater than or equal to date_from")

    with connect() as conn:
        with conn.cursor() as cur:
            cur.execute(
                """
                SELECT m.message_day::date AS day, c.id AS chat_id,
                       c.dialog_id, c.chat_title, count(m.id) AS messages_count
                FROM chat_messages m
                JOIN chats c ON c.id = m.chat_id
                WHERE c.is_excluded = FALSE AND m.message_day BETWEEN %s AND %s
                GROUP BY m.message_day::date, c.id, c.dialog_id, c.chat_title
                ORDER BY day, messages_count DESC, c.chat_title NULLS LAST
                """,
                (date_from, date_to),
            )
            message_rows = cur.fetchall()

            zoom_rows: list[dict[str, Any]] = []
            if safe_table_exists(cur, "zoom_calls"):
                cur.execute(
                    """
                    SELECT id AS call_id, call_date, topic, technical_topic,
                           (analytical_note IS NOT NULL AND length(btrim(analytical_note)) > 0) AS has_report
                    FROM zoom_calls
                    WHERE call_date BETWEEN %s AND %s
                    ORDER BY call_date, start_time_msk NULLS LAST
                    """,
                    (date_from, date_to),
                )
                zoom_rows = cur.fetchall()

            owner_dates: set[Any] = set()
            if safe_table_exists(cur, "owner_daily_reports"):
                cur.execute(
                    """
                    SELECT report_date
                    FROM owner_daily_reports
                    WHERE is_current = TRUE AND report_date BETWEEN %s AND %s
                    """,
                    (date_from - timedelta(days=1), date_to),
                )
                owner_dates = {row["report_date"] for row in cur.fetchall()}

    days: list[dict[str, Any]] = []
    total_zoom_missing = 0
    days_ready = 0
    current = date_from
    while current <= date_to:
        day_chats = [row for row in message_rows if row["day"] == current]
        chat_transcripts = [
            {
                "dialog_id": row["dialog_id"],
                "chat_title": row["chat_title"],
                "messages_count": row["messages_count"],
            }
            for row in day_chats
        ]
        day_zoom = [row for row in zoom_rows if row["call_date"] == current]
        missing_zoom = [
            {
                "call_id": str(row["call_id"]),
                "topic": row.get("topic") or row.get("technical_topic"),
            }
            for row in day_zoom
            if not row["has_report"]
        ]
        owner_exists = current in owner_dates
        prev_owner_exists = (current - timedelta(days=1)) in owner_dates
        ready_for_owner = not missing_zoom and prev_owner_exists

        total_zoom_missing += len(missing_zoom)
        if ready_for_owner:
            days_ready += 1

        days.append(
            {
                "date": current,
                "chats": {
                    "with_messages": len(day_chats),
                    "transcripts": chat_transcripts,
                    "daily_reports_enabled": False,
                },
                "zoom": {
                    "calls": len(day_zoom),
                    "reports_ready": len(day_zoom) - len(missing_zoom),
                    "missing_zoom_reports": missing_zoom,
                },
                "owner_daily_report_exists": owner_exists,
                "previous_owner_daily_report_exists": prev_owner_exists,
                "ready_for_owner_daily": ready_for_owner,
            }
        )
        current += timedelta(days=1)

    return {
        "period": {"date_from": date_from, "date_to": date_to},
        "days": days,
        "summary": {
            "days": len(days),
            "daily_chat_reports_enabled": False,
            "total_zoom_missing_reports": total_zoom_missing,
            "days_ready_for_owner_daily": days_ready,
        },
        "next_actions": [
            "Do not generate chat_daily_reports; they are disabled.",
            "Read raw chat context with get_chat_transcript(dialog_id,date_from,date_to,include_ocr=true).",
            "Generate a Zoom report only for the calls in missing_zoom_reports.",
            "Build owner_daily_report only for days where ready_for_owner_daily is true; otherwise close the missing sources first.",
        ],
    }


def tool_get_org_structure(args: dict[str, Any]) -> dict[str, Any]:
    include_inactive = bool(args.get("include_inactive", False))
    cache_key = ("org_structure", include_inactive)
    cached = ttl_cache_get(cache_key)
    if cached is not None:
        return cached
    with connect() as conn:
        with conn.cursor() as cur:
            cur.execute(
                """
                SELECT
                    d.id, d.bitrix_department_id, d.name,
                    pd.bitrix_department_id AS parent_bitrix_department_id,
                    pd.name AS parent_name,
                    hu.bitrix_user_id AS head_bitrix_user_id,
                    hu.full_name AS head_name
                FROM departments d
                LEFT JOIN departments pd ON pd.id = d.parent_id
                LEFT JOIN users hu ON hu.id = d.head_id
                ORDER BY COALESCE(pd.name, ''), d.name
                """
            )
            departments = cur.fetchall()
            cur.execute(
                """
                SELECT
                    u.id, u.bitrix_user_id, u.full_name, u.email, u.work_position,
                    u.is_active,
                    mu.bitrix_user_id AS manager_bitrix_user_id,
                    mu.full_name AS manager_name,
                    array_remove(array_agg(DISTINCT d.name), NULL) AS departments
                FROM users u
                LEFT JOIN users mu ON mu.id = u.manager_id
                LEFT JOIN user_departments ud ON ud.user_id = u.id
                LEFT JOIN departments d ON d.id = ud.department_id
                WHERE (%s = TRUE OR u.is_active = TRUE)
                GROUP BY u.id, mu.bitrix_user_id, mu.full_name
                ORDER BY u.full_name NULLS LAST, u.bitrix_user_id
                """,
                (include_inactive,),
            )
            users = cur.fetchall()
    return ttl_cache_set(cache_key, {"departments": departments, "users": users})


def _name_tokens(value: Any) -> list[str]:
    aliases = {
        "настя": "анастасия",
        "дима": "дмитрий",
        "саша": "александр",
        "женя": "евгений",
    }
    normalized = str(value or "").strip().lower().replace("ё", "е")
    return [aliases.get(token, token) for token in re.findall(r"[a-zа-я0-9]+", normalized, flags=re.IGNORECASE)]


def _person_names_match(left: Any, right: Any) -> bool:
    left_tokens = _name_tokens(left)
    right_tokens = _name_tokens(right)
    if not left_tokens or not right_tokens:
        return False
    if left_tokens == right_tokens:
        return True
    smaller, larger = (left_tokens, right_tokens) if len(left_tokens) <= len(right_tokens) else (right_tokens, left_tokens)
    return all(any(candidate.startswith(token) if len(token) == 1 else candidate == token for candidate in larger) for token in smaller)


def _resolve_active_bitrix_user(bitrix_user_id: Any = None, name: Any = None) -> dict[str, Any]:
    user_id = None
    if bitrix_user_id not in (None, ""):
        try:
            user_id = int(bitrix_user_id)
        except (TypeError, ValueError) as exc:
            raise McpError(-32602, "responsible_bitrix_user_id must be an integer.") from exc

    responsible_name = str(name or "").strip()
    with connect() as conn:
        with conn.cursor() as cur:
            if user_id is not None:
                cur.execute(
                    """
                    SELECT bitrix_user_id, full_name, email, work_position, is_active
                    FROM users
                    WHERE bitrix_user_id = %s AND is_active = TRUE
                    LIMIT 1
                    """,
                    (user_id,),
                )
                row = cur.fetchone()
                if not row:
                    raise McpError(
                        -32602,
                        f"Не найден активный сотрудник Bitrix с id {user_id}. Не перебирай id — "
                        "уточни у пользователя правильного сотрудника.",
                    )
                return dict(row)

            if not responsible_name:
                raise McpError(-32602, "Нужно указать исполнителя: responsible_name или responsible_bitrix_user_id.")

            cur.execute(
                """
                SELECT bitrix_user_id, full_name, email, work_position, is_active
                FROM users
                WHERE is_active = TRUE AND bitrix_user_id IS NOT NULL
                ORDER BY full_name
                """
            )
            rows = [dict(row) for row in cur.fetchall()]

    exact = [row for row in rows if str(row.get("full_name") or "").strip().lower() == responsible_name.lower()]
    matches = exact or [row for row in rows if _person_names_match(row.get("full_name"), responsible_name)]
    if not matches:
        raise McpError(
            -32602,
            f"Не удалось найти исполнителя в оргструктуре: {responsible_name}. "
            "НЕ повторяй вызов с этим же именем и НЕ подбирай замену сам: если задач несколько — "
            "создай остальные, а по этому человеку спроси у пользователя, кто это "
            "(или попроси точный responsible_bitrix_user_id).",
        )
    if len(matches) > 1:
        candidates = [
            {
                "bitrix_user_id": row.get("bitrix_user_id"),
                "full_name": row.get("full_name"),
                "work_position": row.get("work_position"),
            }
            for row in matches[:10]
        ]
        raise McpError(-32602, "Исполнитель найден неоднозначно. Укажите responsible_bitrix_user_id. Кандидаты: " + json.dumps(candidates, ensure_ascii=False))
    return matches[0]


def _resolve_active_bitrix_users(
    bitrix_user_ids: Any,
    names: Any,
    *,
    role_label: str,
    id_field: str,
    name_field: str,
) -> list[dict[str, Any]]:
    ids_input = bitrix_user_ids if isinstance(bitrix_user_ids, list) else []
    names_input = names if isinstance(names, list) else []
    if not ids_input and not names_input:
        return []

    parsed_ids: list[int] = []
    for raw in ids_input:
        if raw in (None, ""):
            continue
        try:
            parsed_ids.append(int(raw))
        except (TypeError, ValueError) as exc:
            raise McpError(-32602, f"{id_field}: каждый элемент должен быть integer.") from exc

    clean_names = [str(n).strip() for n in names_input if str(n or "").strip()]

    resolved: list[dict[str, Any]] = []
    seen: set[int] = set()

    with connect() as conn:
        with conn.cursor() as cur:
            for uid in parsed_ids:
                if uid in seen:
                    continue
                cur.execute(
                    """
                    SELECT bitrix_user_id, full_name, email, work_position, is_active
                    FROM users
                    WHERE bitrix_user_id = %s AND is_active = TRUE
                    LIMIT 1
                    """,
                    (uid,),
                )
                row = cur.fetchone()
                if not row:
                    raise McpError(
                        -32602,
                        f"{role_label} не найден: активный сотрудник Bitrix с id {uid} отсутствует.",
                    )
                seen.add(uid)
                resolved.append(dict(row))

            if clean_names:
                cur.execute(
                    """
                    SELECT bitrix_user_id, full_name, email, work_position, is_active
                    FROM users
                    WHERE is_active = TRUE AND bitrix_user_id IS NOT NULL
                    ORDER BY full_name
                    """
                )
                all_active = [dict(r) for r in cur.fetchall()]
                for name in clean_names:
                    exact = [
                        r for r in all_active
                        if str(r.get("full_name") or "").strip().lower() == name.lower()
                    ]
                    matches = exact or [
                        r for r in all_active if _person_names_match(r.get("full_name"), name)
                    ]
                    if not matches:
                        raise McpError(
                            -32602,
                            f"{role_label} не найден в оргструктуре: {name}. "
                            "НЕ повторяй вызов с этим именем — пропусти этого человека, сделай "
                            "остальное и уточни у пользователя, кто это.",
                        )
                    if len(matches) > 1:
                        candidates = [
                            {
                                "bitrix_user_id": r.get("bitrix_user_id"),
                                "full_name": r.get("full_name"),
                                "work_position": r.get("work_position"),
                            }
                            for r in matches[:10]
                        ]
                        raise McpError(
                            -32602,
                            f"{role_label} '{name}' найден неоднозначно. Уточните через {id_field}. Кандидаты: "
                            + json.dumps(candidates, ensure_ascii=False),
                        )
                    uid = int(matches[0]["bitrix_user_id"])
                    if uid in seen:
                        continue
                    seen.add(uid)
                    resolved.append(matches[0])

    return resolved


_BITRIX_WEEKDAY_CODES = ("MO", "TU", "WE", "TH", "FR", "SA", "SU")


def _build_bitrix_regular_parameters(periodic: dict[str, Any]) -> dict[str, Any]:
    type_raw = str(periodic.get("type") or "").strip().lower()
    if type_raw not in {"daily", "weekly", "monthly"}:
        raise McpError(-32602, "periodic.type must be one of: daily, weekly, monthly.")

    interval_raw = periodic.get("interval")
    if interval_raw in (None, ""):
        interval = 1
    else:
        try:
            interval = int(interval_raw)
        except (TypeError, ValueError) as exc:
            raise McpError(-32602, "periodic.interval must be a positive integer.") from exc
        if interval < 1:
            raise McpError(-32602, "periodic.interval must be >= 1.")

    params: dict[str, Any] = {"REPEAT_EVERY": interval}

    if type_raw == "daily":
        daily_mode = str(periodic.get("daily_mode") or "all").strip().lower()
        if daily_mode not in {"all", "workdays"}:
            raise McpError(-32602, "periodic.daily_mode must be 'all' or 'workdays'.")
        params["REPEAT_TYPE"] = "daily"
        params["DAILY_MODE"] = daily_mode
    elif type_raw == "weekly":
        weekdays_raw = periodic.get("weekdays")
        if not isinstance(weekdays_raw, list) or not weekdays_raw:
            raise McpError(
                -32602,
                "periodic.weekdays must be a non-empty list when type=weekly (e.g. ['MO','WE','FR']).",
            )
        normalized_weekdays: list[str] = []
        for code in weekdays_raw:
            code_str = str(code or "").strip().upper()
            if code_str not in _BITRIX_WEEKDAY_CODES:
                raise McpError(
                    -32602,
                    f"periodic.weekdays: '{code}' is not a valid weekday code. Use MO/TU/WE/TH/FR/SA/SU.",
                )
            if code_str not in normalized_weekdays:
                normalized_weekdays.append(code_str)
        params["REPEAT_TYPE"] = "weekly"
        params["REPEAT_WEEKDAYS"] = normalized_weekdays
    else:  # monthly
        dom_raw = periodic.get("day_of_month")
        if dom_raw in (None, ""):
            raise McpError(-32602, "periodic.day_of_month is required when type=monthly (1-31).")
        try:
            dom = int(dom_raw)
        except (TypeError, ValueError) as exc:
            raise McpError(-32602, "periodic.day_of_month must be an integer 1-31.") from exc
        if not 1 <= dom <= 31:
            raise McpError(-32602, "periodic.day_of_month must be between 1 and 31.")
        params["REPEAT_TYPE"] = "monthlydays"
        params["REPEAT_MONTHDAY"] = dom

    until_raw = str(periodic.get("until") or "").strip()
    if until_raw:
        if not re.match(r"^\d{4}-\d{2}-\d{2}$", until_raw):
            raise McpError(-32602, "periodic.until must be YYYY-MM-DD.")
        params["REPEAT_TILL"] = until_raw

    return params


def _normalize_bitrix_datetime(value: Any, *, field: str = "deadline",
                               missing_msg: str | None = None) -> str:
    """Accept the natural date/datetime formats the model produces, not only ISO-with-T.
    Supported: YYYY-MM-DD and DD.MM.YYYY (date-only -> 19:00 MSK); the same dates
    followed by ' HH:MM[:SS]' (space OR 'T'); and full ISO with optional tz.
    A space-separated date+time (e.g. '2026-06-28 15:00') was previously REJECTED,
    which made the model loop on format retries and report a phantom tool timeout.
    `field` names the field in error messages so plan dates / deadline read naturally."""
    raw = str(value or "").strip()
    if not raw:
        raise McpError(-32602, missing_msg or f"Нужно указать {field}.")
    # DD.MM.YYYY  optionally followed by  (space|T) HH:MM[:SS]
    m = re.match(r"^(\d{2})\.(\d{2})\.(\d{4})(?:[ T](\d{1,2}):(\d{2})(?::(\d{2}))?)?$", raw)
    if m:
        d, mo, y, hh, mm, ss = m.groups()
        if hh is None:
            return f"{y}-{mo}-{d}T19:00:00+03:00"
        return f"{y}-{mo}-{d}T{int(hh):02d}:{mm}:{ss or '00'}+03:00"
    # YYYY-MM-DD  optionally followed by  (space|T) HH:MM[:SS]  optional tz
    m = re.match(r"^(\d{4})-(\d{2})-(\d{2})(?:[ T](\d{1,2}):(\d{2})(?::(\d{2}))?(Z|[+-]\d{2}:?\d{2})?)?$", raw)
    if m:
        y, mo, d, hh, mm, ss, tz = m.groups()
        if hh is None:
            return f"{y}-{mo}-{d}T19:00:00+03:00"
        tzs = "+03:00" if not tz else ("+00:00" if tz == "Z" else tz)
        return f"{y}-{mo}-{d}T{int(hh):02d}:{mm}:{ss or '00'}{tzs}"
    raise McpError(-32602, f"{field} должен быть в формате YYYY-MM-DD[ HH:MM], DD.MM.YYYY[ HH:MM] или ISO datetime.")


def _normalize_bitrix_deadline(value: Any) -> str:
    """Task deadline normalizer (see _normalize_bitrix_datetime). Behaviour preserved."""
    return _normalize_bitrix_datetime(
        value, field="deadline",
        missing_msg="Нужно указать крайний срок задачи: deadline.")

def _bitrix_call_with_fallback(
    method: str,
    payload: dict[str, Any],
    prefer_api: bool = True,
    fallback: bool = True,
) -> dict[str, Any]:
    workflow = app_workflow_function("bitrix_method_call")
    try:
        return workflow(method, payload, prefer_api, fallback)
    except ValueError as exc:
        raise McpError(-32000, str(exc)) from exc
    except Exception as exc:  # noqa: BLE001
        raise McpError(-32010, f"Bitrix API call failed: {exc}") from exc


def _positive_bitrix_task_id(value: Any) -> int:
    if value in (None, ""):
        raise McpError(-32602, "Нужно указать точный номер задачи: bitrix_task_id.")
    try:
        task_id = int(value)
    except (TypeError, ValueError) as exc:
        raise McpError(-32602, "bitrix_task_id must be an integer.") from exc
    if task_id <= 0:
        raise McpError(-32602, "bitrix_task_id must be a positive integer.")
    return task_id


def _confirmed(args: dict[str, Any]) -> bool:
    raw = args.get("confirm")
    return raw is True or str(raw or "").strip().lower() in {"true", "1", "yes", "да"}


def _indexed_task_for_action(task_id: int) -> dict[str, Any] | None:
    with connect() as conn:
        with conn.cursor() as cur:
            cur.execute(
                """
                SELECT
                    t.bitrix_task_id,
                    t.title,
                    t.status,
                    t.status_name,
                    t.deadline_at,
                    ru.bitrix_user_id AS responsible_bitrix_user_id,
                    ru.full_name AS responsible_name
                FROM bitrix_tasks t
                LEFT JOIN users ru ON ru.id = t.responsible_id
                WHERE t.bitrix_task_id = %s
                LIMIT 1
                """,
                (task_id,),
            )
            row = cur.fetchone()
    return dict(row) if row else None


def _assert_expected_task_title(task: dict[str, Any] | None, expected_title: Any) -> None:
    expected = str(expected_title or "").strip()
    if not expected or not task:
        return
    actual = str(task.get("title") or "").strip()
    if expected.lower() != actual.lower():
        raise McpError(
            -32602,
            "expected_title не совпадает с найденной задачей. Действие остановлено, чтобы не изменить не ту задачу.",
        )


def _task_payload(task: dict[str, Any] | None, task_id: int) -> dict[str, Any]:
    if not task:
        return {"bitrix_task_id": task_id}
    return {
        "bitrix_task_id": task.get("bitrix_task_id"),
        "title": task.get("title"),
        "status": task.get("status"),
        "status_name": task.get("status_name"),
        "deadline_at": task.get("deadline_at").isoformat() if hasattr(task.get("deadline_at"), "isoformat") else task.get("deadline_at"),
        "responsible_bitrix_user_id": task.get("responsible_bitrix_user_id"),
        "responsible_name": task.get("responsible_name"),
    }


def _portal_base_url() -> str:
    """The Bitrix portal web host (https://<portal>), derived from the webhook base."""
    base = (os.getenv("BITRIX_WEBHOOK_BASE", "") or "").strip()
    if base:
        p = urlparse(base)
        if p.scheme and p.netloc:
            return f"{p.scheme}://{p.netloc}"
    return (os.getenv("BITRIX_PORTAL_URL", "") or "").rstrip("/")


def _task_deep_link(task_id: Any) -> str | None:
    """Clickable Bitrix task deep link — the number a user taps to open the task."""
    base = _portal_base_url()
    try:
        tid = int(task_id)
    except (TypeError, ValueError):
        return None
    if not base or tid <= 0:
        return None
    return f"{base}/company/personal/user/0/tasks/task/view/{tid}/"


def _task_deadline_change_count(task_id: Any) -> int | None:
    """How many times the DEADLINE was changed («перенесена X раз»), via live task history.
    Best-effort — None on any failure. Uses the v2 webhook directly (history.list works there)."""
    try:
        tid = int(task_id)
    except (TypeError, ValueError):
        return None
    try:
        r = _webhook_raw("tasks.task.history.list", {"taskId": tid})
    except Exception:  # noqa: BLE001
        return None
    res = r.get("result") if isinstance(r, dict) else None
    lst = res.get("list") if isinstance(res, dict) else res
    if not isinstance(lst, list):
        return None
    return sum(1 for it in lst if isinstance(it, dict)
               and str(it.get("field") or it.get("FIELD") or "").upper() == "DEADLINE")


def _deadline_in_past(deadline: str) -> "datetime | None":
    """Return the parsed deadline (MSK-aware) if it is at/before 'now' in Europe/Moscow, else None.
    Deterministic backstop so an already-overdue task is never created without explicit confirmation."""
    s = str(deadline or "").strip()
    dt = None
    for cand in (s, s + ":00"):
        try:
            dt = datetime.fromisoformat(cand)
            break
        except ValueError:
            continue
    if dt is None:
        return None
    if dt.tzinfo is None:
        dt = dt.replace(tzinfo=_MSK_TZ)
    now = datetime.now(_MSK_TZ)
    return dt if dt <= now else None


# --- Full task field support (соисполнители, план, CRM, пользовательские поля, ...) ----------
# The Bitrix task editor exposes many field groups (see the "add field" chips). These helpers let
# the agent set the REST-addable ones on create AND on update. The action-based entities
# (checklists, time tracking, dependencies, reminders) have their own dedicated tools below.

_UF_KEY_RE = re.compile(r"^UF_[A-Z0-9_]+$")
_CRM_PREFIX_MAP = {"DEAL": "D", "LEAD": "L", "CONTACT": "C", "COMPANY": "CO"}


def _clean_crm_elements(value: Any) -> list[str]:
    """Normalize CRM bindings for UF_CRM_TASK — accept ['D_12','L_3',...] or long forms
    (deal/lead/contact/company). Returns ['D_12', ...] or []. Malformed entries are refused."""
    if value in (None, ""):
        return []
    items = value if isinstance(value, (list, tuple)) else [value]
    out: list[str] = []
    for it in items:
        s = str(it or "").strip().upper().replace(" ", "")
        if not s:
            continue
        m = re.match(r"^([A-Z]+)_?(\d+)$", s)
        if not m:
            raise McpError(-32602, f"crm_elements: '{it}' — ожидается вид 'D_123' (сделка), "
                                   "'L_' (лид), 'C_' (контакт), 'CO_' (компания).")
        pref = _CRM_PREFIX_MAP.get(m.group(1), m.group(1))
        ref = f"{pref}_{m.group(2)}"
        if ref not in out:
            out.append(ref)
    return out


def _clean_custom_fields(value: Any) -> dict[str, Any]:
    """Validate a dict of custom task fields. Keys must be UF_* (task user fields) so the model
    cannot set arbitrary system fields. Field codes come from list_task_userfields."""
    if value in (None, ""):
        return {}
    if not isinstance(value, dict):
        raise McpError(-32602, "custom_fields должен быть объектом вида {\"UF_...\": значение}.")
    out: dict[str, Any] = {}
    for k, v in value.items():
        key = str(k or "").strip().upper()
        if not _UF_KEY_RE.match(key):
            raise McpError(-32602, f"custom_fields: ключ '{k}' должен начинаться с UF_ "
                                   "(пользовательское поле задачи). Коды полей — list_task_userfields.")
        out[key] = v
    return out


def _resolve_person_ids(args: dict[str, Any], id_field: str, name_field: str, role_label: str) -> list[int]:
    """Resolve a list of people (ids and/or names) to a de-duplicated list of active user ids."""
    users = _resolve_active_bitrix_users(
        args.get(id_field), args.get(name_field),
        role_label=role_label, id_field=id_field, name_field=name_field)
    return [int(u["bitrix_user_id"]) for u in users]


def _assemble_task_fields(
    *, title: str, description: str, responsible_id: int, deadline_iso: str,
    priority: int = 1, auditor_ids: list[int] | None = None,
    accomplice_ids: list[int] | None = None, creator_id: int | None = None,
    tags: list[str] | None = None, parent_task_id: int | None = None, group_id: int | None = None,
    start_plan: str | None = None, end_plan: str | None = None,
    time_estimate_seconds: int | None = None, crm_elements: list[str] | None = None,
    custom_fields: dict[str, Any] | None = None,
) -> dict[str, Any]:
    """Build the Bitrix tasks.task.add FIELDS dict from already-resolved values. Shared by the
    create tool and the recurring scheduler so both produce identical tasks. Only provided optional
    fields are set; a result is always required (SE_PARAMETER code 3)."""
    fields: dict[str, Any] = {
        "TITLE": title,
        "DESCRIPTION": description,
        "RESPONSIBLE_ID": int(responsible_id),
        "DEADLINE": deadline_iso,
        "PRIORITY": priority,
        "SE_PARAMETER": [{"CODE": 3, "VALUE": "Y"}],
    }
    if auditor_ids:
        fields["AUDITORS"] = [int(x) for x in auditor_ids]
    if accomplice_ids:
        fields["ACCOMPLICES"] = [int(x) for x in accomplice_ids]
    if creator_id is not None:
        fields["CREATED_BY"] = int(creator_id)
    if tags:
        fields["TAGS"] = list(tags)
    if parent_task_id:
        fields["PARENT_ID"] = int(parent_task_id)
    if group_id:
        fields["GROUP_ID"] = int(group_id)
    if start_plan:
        fields["START_DATE_PLAN"] = start_plan
    if end_plan:
        fields["END_DATE_PLAN"] = end_plan
    if time_estimate_seconds:
        fields["TIME_ESTIMATE"] = int(time_estimate_seconds)
    if crm_elements:
        fields["UF_CRM_TASK"] = list(crm_elements)
    if custom_fields:
        for k, v in custom_fields.items():
            fields[str(k)] = v
    return fields


def _add_checklist_items(task_id: int, items: Any) -> list[dict[str, Any]]:
    """Add checklist items to a task (task.checklistitem.add, v2 webhook). `items` = list of
    strings or {title, complete?}. Best-effort per item; a per-item error never aborts the rest."""
    summary: list[dict[str, Any]] = []
    if not isinstance(items, (list, tuple)):
        return summary
    for it in items:
        if isinstance(it, dict):
            title = str(it.get("title") or it.get("text") or "").strip()
            complete = it.get("complete") is True or str(it.get("complete") or "").strip().lower() in {"true", "1", "yes", "да"}
        else:
            title = str(it or "").strip()
            complete = False
        if not title:
            continue
        try:
            resp = _webhook_raw("task.checklistitem.add", {"TASKID": task_id, "FIELDS": {"TITLE": title}})
            item_id = resp.get("result") if isinstance(resp, dict) else None
            if complete and item_id:
                try:
                    _webhook_raw("task.checklistitem.complete", {"TASKID": task_id, "ITEMID": item_id})
                except Exception:  # noqa: BLE001
                    pass
            summary.append({"title": title, "checklist_item_id": item_id, "complete": complete})
        except Exception as exc:  # noqa: BLE001
            summary.append({"title": title, "error": str(exc)[:160]})
    return summary


def tool_create_bitrix_task(args: dict[str, Any]) -> dict[str, Any]:
    title = str(args.get("title") or "").strip()
    if not title:
        raise McpError(-32602, "Нужно указать название задачи: title.")
    result_criteria = str(args.get("result_criteria") or "").strip()
    if not result_criteria:
        raise McpError(
            -32602,
            "У задачи ОБЯЗАН быть результат. Спроси пользователя: какой результат/критерий выполнения "
            "у этой задачи — по чему поймём, что сделано, и чем подтверждается (скрин/ссылка/файл)? "
            "Затем повтори вызов create_bitrix_task с параметром result_criteria=...",
        )

    # Recurring tasks are NOT created here. The old Bitrix REPLICATE template mechanism needs a
    # paid plan (which this portal does not have) and silently never spawns tasks; recurring tasks
    # are now driven by the agent's own scheduler. Redirect to the dedicated tool.
    periodic_arg = args.get("periodic")
    if isinstance(periodic_arg, dict) and periodic_arg:
        raise McpError(
            -32602,
            "Для ПОВТОРЯЮЩЕЙСЯ (регулярной) задачи используй инструмент create_recurring_task "
            "(create_bitrix_task делает только разовые). Передай period (daily/weekly/monthly), для "
            "weekly — weekdays, create_time (во сколько создавать) и срок каждой задачи (deadline_time "
            "или deadline_after_hours).",
        )

    deadline = _normalize_bitrix_deadline(args.get("deadline"))
    _confirm_past = args.get("confirm_past_deadline")
    _confirm_past = _confirm_past is True or str(_confirm_past or "").strip().lower() in {"true", "1", "yes", "да"}
    if not _confirm_past:
        _past = _deadline_in_past(deadline)
        if _past is not None:
            _now = datetime.now(_MSK_TZ)
            raise McpError(
                -32602,
                "Срок " + _past.strftime("%d.%m.%Y %H:%M") + " уже в прошлом (сейчас "
                + _now.strftime("%d.%m.%Y %H:%M") + " МСК). НЕ ставь задачу молча: спроси пользователя — "
                "поставить с этим сроком как есть или указать новый. Если подтвердил как есть — повтори "
                "вызов с confirm_past_deadline=true; если дал новый срок — вызови с новым deadline.",
            )
    responsible = _resolve_active_bitrix_user(args.get("responsible_bitrix_user_id"), args.get("responsible_name"))
    description = str(args.get("description") or "").strip() or title
    if "Критерий результата" not in description:
        description = description + "\n\nКритерий результата: " + result_criteria
    priority_raw = str(args.get("priority") or "normal").strip().lower()
    priority = 2 if priority_raw in {"high", "critical", "2", "важно", "высокий"} else 1

    auditors = _resolve_active_bitrix_users(
        args.get("auditor_bitrix_user_ids"), args.get("auditor_names"),
        role_label="Наблюдатель", id_field="auditor_bitrix_user_ids", name_field="auditor_names")
    accomplices = _resolve_active_bitrix_users(
        args.get("accomplice_bitrix_user_ids"), args.get("accomplice_names"),
        role_label="Соисполнитель", id_field="accomplice_bitrix_user_ids", name_field="accomplice_names")

    creator_info = None
    if args.get("creator_bitrix_user_id") not in (None, "") or args.get("creator_name"):
        creator_info = _resolve_active_bitrix_user(args.get("creator_bitrix_user_id"), args.get("creator_name"))

    # Optional scalar fields (all safe to omit; only set when provided).
    parent_task_id = _positive_bitrix_task_id(args.get("parent_task_id")) if args.get("parent_task_id") not in (None, "") else None
    group_id = None
    if args.get("group_id") not in (None, ""):
        try:
            group_id = int(args.get("group_id"))
        except (TypeError, ValueError) as exc:
            raise McpError(-32602, "group_id (проект/рабочая группа) должен быть числом.") from exc
    start_plan = _normalize_bitrix_datetime(args.get("start_plan"), field="start_plan") if args.get("start_plan") not in (None, "") else None
    end_plan = _normalize_bitrix_datetime(args.get("end_plan"), field="end_plan") if args.get("end_plan") not in (None, "") else None
    time_estimate_seconds = None
    if args.get("time_estimate_hours") not in (None, ""):
        try:
            time_estimate_seconds = int(float(args.get("time_estimate_hours")) * 3600)
        except (TypeError, ValueError) as exc:
            raise McpError(-32602, "time_estimate_hours должно быть числом (часы).") from exc
    crm_elements = _clean_crm_elements(args.get("crm_elements"))
    custom_fields = _clean_custom_fields(args.get("custom_fields"))
    tags = [str(t).strip() for t in args.get("tags") if str(t or "").strip()] if isinstance(args.get("tags"), list) else None

    fields = _assemble_task_fields(
        title=title, description=description, responsible_id=int(responsible["bitrix_user_id"]),
        deadline_iso=deadline, priority=priority,
        auditor_ids=[int(u["bitrix_user_id"]) for u in auditors],
        accomplice_ids=[int(u["bitrix_user_id"]) for u in accomplices],
        creator_id=int(creator_info["bitrix_user_id"]) if creator_info else None,
        tags=tags, parent_task_id=parent_task_id, group_id=group_id,
        start_plan=start_plan, end_plan=end_plan, time_estimate_seconds=time_estimate_seconds,
        crm_elements=crm_elements or None, custom_fields=custom_fields or None,
    )

    response = _bitrix_call_with_fallback("tasks.task.add", {"fields": fields})
    result = response.get("result") if isinstance(response, dict) else {}
    task_id = None
    if isinstance(result, dict):
        task = result.get("task") if isinstance(result.get("task"), dict) else {}
        task_id = task.get("id") or result.get("id")
    else:
        task_id = result

    # Post-create actions that need the task id: attach forwarded files + checklist items.
    attach_summary: list[dict[str, Any]] = []
    if task_id and isinstance(args.get("attachment_ids"), (list, tuple)) and any(str(a or "").strip() for a in args["attachment_ids"]):
        try:
            refs, attach_summary = _resolve_attachment_disk_refs(args.get("attachment_ids"))
            _attach_disk_refs_to_task(int(task_id), refs)
        except Exception as exc:  # noqa: BLE001
            logging.warning("create_task attach failed task=%s: %s", task_id, repr(exc)[:120])
    checklist_summary: list[dict[str, Any]] = []
    if task_id and args.get("checklist"):
        checklist_summary = _add_checklist_items(int(task_id), args.get("checklist"))

    # Agent-created task → offer-to-help comment from the most suitable agent (задача 1300).
    # Fire-and-forget background thread; can never fail or delay the creation itself.
    if task_id:
        try:
            from task_offers import schedule_offer
            schedule_offer(task_id, title=title, description=description,
                           checklist=args.get("checklist"),
                           responsible_id=responsible.get("bitrix_user_id"),
                           creator_id=(creator_info or {}).get("bitrix_user_id"))
        except Exception:  # noqa: BLE001
            logging.warning("create_task: offer scheduling failed task=%s", task_id, exc_info=True)

    return {
        "created": True,
        "task_id": task_id,
        "title": title,
        "description": description,
        "deadline": deadline,
        "responsible": {
            "bitrix_user_id": responsible.get("bitrix_user_id"),
            "full_name": responsible.get("full_name"),
            "work_position": responsible.get("work_position"),
        },
        "auditors": [{"bitrix_user_id": u.get("bitrix_user_id"), "full_name": u.get("full_name")} for u in auditors],
        "accomplices": [{"bitrix_user_id": u.get("bitrix_user_id"), "full_name": u.get("full_name")} for u in accomplices],
        "creator": (
            {"bitrix_user_id": creator_info.get("bitrix_user_id"), "full_name": creator_info.get("full_name")}
            if creator_info else None
        ),
        "parent_task_id": parent_task_id,
        "group_id": group_id,
        "planning": {"start_plan": start_plan, "end_plan": end_plan,
                     "time_estimate_hours": (time_estimate_seconds / 3600) if time_estimate_seconds else None},
        "crm_elements": crm_elements or None,
        "custom_fields": custom_fields or None,
        "attachments": attach_summary or None,
        "checklist": checklist_summary or None,
        "require_result": True,
        "bitrix_response": response.get("result") if isinstance(response, dict) else response,
        "rule": "Task creation requires title, responsible_name/responsible_bitrix_user_id, and deadline. "
                "Optional: соисполнители (accomplice_*), наблюдатели (auditor_*), теги, родительская задача "
                "(parent_task_id), проект (group_id), планирование (start_plan/end_plan/time_estimate_hours), "
                "элементы CRM (crm_elements), пользовательские поля (custom_fields), файлы (attachment_ids), "
                "чек-лист (checklist). Missing or ambiguous data blocks creation.",
    }


def tool_delete_bitrix_task(args: dict[str, Any]) -> dict[str, Any]:
    task_id = _positive_bitrix_task_id(args.get("bitrix_task_id"))

    if not _confirmed(args):
        raise McpError(
            -32602,
            "Удаление задачи требует явного подтверждения. Сначала покажите пользователю точную задачу "
            "(номер, название, статус, ответственный, дедлайн) и спросите подтверждение. После подтверждения "
            "повторите вызов с confirm=true.",
        )

    task = _indexed_task_for_action(task_id)
    if not task:
        raise McpError(
            -32602,
            f"Задача Bitrix {task_id} не найдена в локальном индексе. Сначала проверьте номер через search_tasks.",
        )

    _assert_expected_task_title(task, args.get("expected_title"))

    try:
        response = _bitrix_call_with_fallback("tasks.task.delete", {"taskId": task_id})
    except McpError:
        # Bitrix answers «Нет доступа к удалению задачи» (1048582) for tasks that DON'T EXIST —
        # our local index keeps rows for tasks deleted earlier, so the agent sees a ghost,
        # tries to delete it and gets a misleading permission error (live incident 11.07:
        # 7 of 8 «неудаляемых» задач просто не существовали). Probe live before failing.
        if _task_exists_live(task_id):
            raise  # the task is real — this IS a genuine error/permission problem
        _purge_task_from_index(task_id)
        return {
            "deleted": True,
            "already_gone": True,
            "task": _task_payload(task, task_id),
            "note": "Задачи уже не было в Bitrix (устаревший локальный индекс) — запись вычищена "
                    "из индекса, в списках она больше не появится.",
        }
    _purge_task_from_index(task_id)
    return {
        "deleted": True,
        "task": _task_payload(task, task_id),
        "bitrix_response": response.get("result") if isinstance(response, dict) else response,
        "rule": "Deletion requires exact bitrix_task_id and confirm=true after the user has seen the exact task and explicitly confirmed deletion.",
    }


def _task_exists_live(task_id: int) -> bool:
    """Does the task exist in Bitrix RIGHT NOW (not in the local snapshot)? Best-effort:
    unknown (call failed) counts as existing so a delete error is never masked by a probe error."""
    try:
        r = _webhook_raw("tasks.task.list", {"filter": {"ID": int(task_id)}, "select": ["ID"]})
        tasks = (r.get("result") or {}).get("tasks")
        return bool(tasks)
    except Exception:  # noqa: BLE001
        return True


def _purge_task_from_index(task_id: int) -> None:
    """Drop the task's row from the local bitrix_tasks index after deletion, so ghosts don't
    linger in search_tasks lists (the sync never removes deleted tasks)."""
    try:
        with connect() as conn:
            with conn.cursor() as cur:
                cur.execute("DELETE FROM bitrix_tasks WHERE bitrix_task_id = %s", (int(task_id),))
    except Exception:  # noqa: BLE001
        logging.warning("task index purge failed for %s", task_id, exc_info=True)


# --- Attachment forwarding: re-upload a stored inbound file to Bitrix and reference it -------
# The attachment store (module `attachments`) keeps the raw bytes of every file an employee sent
# the bot. To put such a file on a task/comment/result it must be a Bitrix disk object; we upload
# it once to the webhook user's storage and cache the disk id so repeated attaches are cheap.

def _bitrix_upload_bytes(file_name: str, data: bytes) -> int:
    import base64
    me = _bitrix_call_with_fallback("user.current", {}, prefer_api=False)
    uid = (me.get("result") or {}).get("ID") if isinstance(me, dict) else None
    storages = []
    if uid is not None:
        st = _bitrix_call_with_fallback(
            "disk.storage.getlist", {"filter": {"ENTITY_TYPE": "user", "ENTITY_ID": uid}}, prefer_api=False)
        storages = st.get("result") or [] if isinstance(st, dict) else []
    if not storages:
        st = _bitrix_call_with_fallback("disk.storage.getlist", {}, prefer_api=False)
        rows = st.get("result") or [] if isinstance(st, dict) else []
        storages = [r for r in rows if str(r.get("ENTITY_TYPE") or "").lower() == "user"] or rows
    sid = storages[0].get("ID") if storages else None
    if not sid:
        raise McpError(-32010, "Не удалось найти хранилище Диска для загрузки файла в Bitrix.")
    content = base64.b64encode(data).decode()
    up = _bitrix_call_with_fallback(
        "disk.storage.uploadfile",
        {"id": sid, "data": {"NAME": file_name or "file"},
         "fileContent": [file_name or "file", content], "generateUniqueName": True},
        prefer_api=False,
    )
    fid = (up.get("result") or {}).get("ID") if isinstance(up, dict) else None
    if not fid:
        raise McpError(-32010, "Bitrix отклонил загрузку файла на Диск.")
    return int(fid)


def _attachment_disk_id(token: str) -> int:
    """Resolve a stored attachment token to a FRESH Bitrix disk file id. A disk object is consumed
    once attached to an entity (comment/task) — so we upload a new copy per attach target rather
    than reusing a cached id (reuse yields «Не удалось найти файл»)."""
    import attachments as _att
    row = _att.get_attachment(token)
    if not row:
        raise McpError(-32602, f"Вложение {token} не найдено. Проверь attachment_id (его выдаёт "
                               "система при получении файла от пользователя).")
    blob = _att.attachment_bytes(token)
    if not blob:
        raise McpError(-32010, f"Файл вложения {token} больше недоступен для пересылки "
                               "(истёк срок хранения). Попроси пользователя прислать файл заново.")
    data, name = blob
    fid = _bitrix_upload_bytes(row.get("file_name") or name, data)
    _att.set_disk_id(token, fid)  # last-used id, for reference only (not reused)
    return fid


def _resolve_attachment_disk_refs(attachment_ids: Any) -> tuple[list[str], list[dict[str, Any]]]:
    """Turn a list of attachment tokens into ['n<fid>', ...] disk refs + a human summary."""
    refs: list[str] = []
    summary: list[dict[str, Any]] = []
    if not isinstance(attachment_ids, (list, tuple)):
        return refs, summary
    import attachments as _att
    for tok in attachment_ids:
        tok = str(tok or "").strip()
        if not tok:
            continue
        fid = _attachment_disk_id(tok)
        refs.append(f"n{fid}")
        row = _att.get_attachment(tok) or {}
        summary.append({"attachment_id": tok, "file_name": row.get("file_name"), "disk_id": fid})
    return refs, summary


def _webhook_raw(method: str, payload: dict[str, Any]) -> dict[str, Any]:
    """Call the plain Bitrix incoming webhook (REST v2) directly. Needed for legacy task fields
    like UF_TASK_WEBDAV_FILES that the v3 TaskDto (used by BitrixClient) rejects."""
    base = (os.getenv("BITRIX_WEBHOOK_BASE", "") or "").rstrip("/")
    if not base:
        raise McpError(-32000, "BITRIX_WEBHOOK_BASE is not configured.")
    data = json.dumps(payload or {}).encode()
    req = urllib.request.Request(f"{base}/{method}.json", data=data,
                                 headers={"Content-Type": "application/json"})
    try:
        with urllib.request.urlopen(req, timeout=30) as r:
            return json.loads(r.read().decode())
    except urllib.error.HTTPError as exc:  # noqa: BLE001
        body = ""
        try:
            body = exc.read().decode()[:300]
        except Exception:  # noqa: BLE001
            pass
        raise McpError(-32010, f"Bitrix webhook {method} failed: HTTP {exc.code} {body}") from exc
    except Exception as exc:  # noqa: BLE001
        raise McpError(-32010, f"Bitrix webhook {method} failed: {exc}") from exc


def _attach_disk_refs_to_task(task_id: int, refs: list[str]) -> None:
    """Append disk refs to a task's UF_TASK_WEBDAV_FILES without dropping existing files.
    Uses the v2 webhook directly — the v3 TaskDto (BitrixClient) rejects UF_TASK_WEBDAV_FILES."""
    if not refs:
        return
    existing: list[str] = []
    try:
        cur = _webhook_raw("tasks.task.get", {"taskId": task_id, "select": ["ID", "UF_TASK_WEBDAV_FILES"]})
        task = (cur.get("result") or {}).get("task") or {} if isinstance(cur, dict) else {}
        raw = task.get("ufTaskWebdavFiles") or task.get("UF_TASK_WEBDAV_FILES") or []
        for item in raw if isinstance(raw, list) else []:
            s = str(item)
            existing.append(s if s.startswith("n") else f"n{s}")
    except Exception:  # noqa: BLE001
        existing = []
    merged = list(dict.fromkeys(existing + refs))
    _webhook_raw("tasks.task.update", {"taskId": task_id, "fields": {"UF_TASK_WEBDAV_FILES": merged}})


def _resolve_task_actor(args: dict[str, Any], id_key: str, name_key: str) -> dict[str, Any] | None:
    """Resolve the person an action is performed 'on behalf of' (comment author / task closer).
    Returns the org-structure user dict or None when nothing was passed (Bitrix keeps the webhook user)."""
    if args.get(id_key) in (None, "") and not args.get(name_key):
        return None
    return _resolve_active_bitrix_user(args.get(id_key), args.get(name_key))


def _preclaim_task_comment(comment_id: Any, task_id: int, author_id: Any = None) -> None:
    """A comment created by OUR OWN tools must never trigger the in-task mention handler — the
    OnTaskCommentAdd event fires for it too. Pre-claiming the id in the dedupe table makes the
    handler see it as already handled. This precise guard is what allows LIVE humans writing
    through the technical webhook user (author 22 — наш веб-интерфейс) to summon the agent:
    the old blanket «skip author 22» rule silently ate their mentions (2026-07-09, «Рекомендации
    09.07»)."""
    try:
        cid = int(comment_id)
    except (TypeError, ValueError):
        return
    try:
        with connect() as conn:
            with conn.cursor() as cur:
                cur.execute(
                    "INSERT INTO bitrix_task_comment_seen (comment_id, task_id, agent_slug, author_id, handled) "
                    "VALUES (%s, %s, 'self-tool', %s, TRUE) ON CONFLICT (comment_id) DO NOTHING",
                    (cid, int(task_id), _int_or_none(author_id)))
    except Exception:  # noqa: BLE001
        logging.warning("task-comment preclaim failed comment=%s", comment_id, exc_info=True)


def _int_or_none(value: Any):
    try:
        return int(value)
    except (TypeError, ValueError):
        return None


def tool_add_bitrix_task_comment(args: dict[str, Any]) -> dict[str, Any]:
    task_id = _positive_bitrix_task_id(args.get("bitrix_task_id"))
    text = str(args.get("comment_text") or args.get("message") or "").strip()
    as_result = args.get("as_result") is True or str(args.get("as_result") or "").strip().lower() in {"true", "1", "yes", "да"}
    attachment_ids = args.get("attachment_ids")
    has_files = isinstance(attachment_ids, (list, tuple)) and any(str(a or "").strip() for a in attachment_ids)
    if not text and not has_files:
        raise McpError(-32602, "Нужно указать текст комментария (comment_text) или вложения (attachment_ids).")
    if len(text) > 20000:
        raise McpError(-32602, "comment_text is too long (max 20000 characters).")

    task = _indexed_task_for_action(task_id)
    _assert_expected_task_title(task, args.get("expected_title"))

    # On behalf of the requesting employee (verified: task.commentitem.add honours AUTHOR_ID on
    # this portal — the comment shows as authored by that person, not the technical webhook user).
    author = _resolve_task_actor(args, "author_bitrix_user_id", "author_name")

    post_text = text
    if as_result:
        # Native "Результат задачи" (tasks.task.result.addFromComment) does NOT work through the
        # REST webhook on this portal (comments are IM-chat based, the legacy forum id is rejected
        # as "Comment not found"). So a result is delivered as an unmistakably-labelled comment
        # PLUS the file(s) attached to the task — functionally the result, visible to everyone.
        post_text = ("✅ РЕЗУЛЬТАТ ЗАДАЧИ:\n" + text) if text else "✅ РЕЗУЛЬТАТ ЗАДАЧИ (см. вложение)"

    # Re-upload any forwarded attachments and reference them on the comment.
    refs, attach_summary = _resolve_attachment_disk_refs(attachment_ids)

    fields: dict[str, Any] = {"POST_MESSAGE": post_text or "(вложение)"}
    if author:
        fields["AUTHOR_ID"] = int(author["bitrix_user_id"])
    if refs:
        fields["UF_FORUM_MESSAGE_DOC"] = refs
    response = _bitrix_call_with_fallback(
        "task.commentitem.add", {"TASKID": task_id, "FIELDS": fields}, prefer_api=False, fallback=False)
    comment_id = response.get("result") if isinstance(response, dict) else None
    _preclaim_task_comment(comment_id, task_id, (author or {}).get("bitrix_user_id"))

    # For a result, also pin the file(s) to the task itself so they show in the task's files.
    # A disk object is consumed by the comment attach above, so upload FRESH copies for the task.
    if as_result and has_files:
        try:
            task_refs, _ = _resolve_attachment_disk_refs(attachment_ids)
            _attach_disk_refs_to_task(task_id, task_refs)
        except Exception as exc:  # noqa: BLE001
            logging.warning("as_result task attach failed task=%s: %s", task_id, repr(exc)[:120])

    return {
        "comment_added": True,
        "task": _task_payload(task, task_id),
        "comment_text": post_text,
        "comment_id": comment_id,
        "as_result": as_result,
        "author": ({"bitrix_user_id": author.get("bitrix_user_id"), "full_name": author.get("full_name")}
                   if author else None),
        "attachments": attach_summary,
        "method": "task.commentitem.add",
        "rule": ("Comments require exact bitrix_task_id. By default author = the current chat user "
                 "(author_bitrix_user_id). as_result posts a labelled result comment + pins files to the task "
                 "(the native Bitrix result badge is not settable via REST on this portal)."),
    }


def tool_complete_bitrix_task(args: dict[str, Any]) -> dict[str, Any]:
    """Complete (close) a Bitrix task, optionally on behalf of the current chat user and with a
    result comment + attached files. The status change is attributed to the person via
    STATUS_CHANGED_BY (verified settable); CLOSED_BY is computed by Bitrix and stays the webhook user."""
    task_id = _positive_bitrix_task_id(args.get("bitrix_task_id"))
    task = _indexed_task_for_action(task_id)
    _assert_expected_task_title(task, args.get("expected_title"))

    actor = _resolve_task_actor(args, "on_behalf_bitrix_user_id", "on_behalf_name")
    result_text = str(args.get("result_text") or args.get("comment_text") or "").strip()
    attachment_ids = args.get("attachment_ids")
    has_files = isinstance(attachment_ids, (list, tuple)) and any(str(a or "").strip() for a in attachment_ids)

    # If a result was provided, record it FIRST (as a labelled result comment + files pinned to the
    # task), so a task with "результат обязателен" has its result visible before completion.
    result_info = None
    if result_text or has_files:
        result_info = tool_add_bitrix_task_comment({
            "bitrix_task_id": task_id,
            "comment_text": result_text,
            "as_result": True,
            "attachment_ids": attachment_ids,
            "author_bitrix_user_id": args.get("on_behalf_bitrix_user_id"),
            "author_name": args.get("on_behalf_name"),
            "expected_title": args.get("expected_title"),
        })

    fields: dict[str, Any] = {"STATUS": 5}
    if actor:
        fields["STATUS_CHANGED_BY"] = int(actor["bitrix_user_id"])
    # tasks.task.complete is the canonical close; it also bypasses the "result required" gate for
    # the webhook admin user. We use update(STATUS=5) so we can attribute STATUS_CHANGED_BY.
    # v2 webhook: the v3 TaskDto rejects the uppercase STATUS/STATUS_CHANGED_BY fields.
    last_error: McpError | None = None
    for method, payload in (
        ("tasks.task.update", {"taskId": task_id, "fields": fields}),
        ("tasks.task.complete", {"taskId": task_id}),
    ):
        try:
            response = _webhook_raw(method, payload)
            if isinstance(response, dict) and response.get("error") and not response.get("result"):
                raise McpError(-32010, f"Bitrix {method}: {response.get('error_description') or response.get('error')}")
            return {
                "completed": True,
                "task": _task_payload(task, task_id),
                "on_behalf": ({"bitrix_user_id": actor.get("bitrix_user_id"), "full_name": actor.get("full_name")}
                              if actor else None),
                "result": result_info,
                "method": method,
                "rule": ("Completion closes the task. By default the closer = the current chat user "
                         "(on_behalf_bitrix_user_id → STATUS_CHANGED_BY). Attach the result via result_text/"
                         "attachment_ids so a result-required task keeps its proof."),
            }
        except McpError as exc:
            last_error = exc
            continue
    raise last_error or McpError(-32010, "Bitrix API call failed: could not complete task.")


def tool_attach_files_to_task(args: dict[str, Any]) -> dict[str, Any]:
    """Attach one or more stored inbound attachments (screenshots/documents the user sent the bot)
    to a Bitrix task — either to the task's files, or as a comment/result carrying the files."""
    task_id = _positive_bitrix_task_id(args.get("bitrix_task_id"))
    attachment_ids = args.get("attachment_ids")
    if not (isinstance(attachment_ids, (list, tuple)) and any(str(a or "").strip() for a in attachment_ids)):
        raise McpError(-32602, "Нужно указать attachment_ids — токены вложений от пользователя.")
    task = _indexed_task_for_action(task_id)
    _assert_expected_task_title(task, args.get("expected_title"))

    as_result = args.get("as_result") is True or str(args.get("as_result") or "").strip().lower() in {"true", "1", "yes", "да"}
    as_comment = args.get("as_comment") is True or str(args.get("as_comment") or "").strip().lower() in {"true", "1", "yes", "да"}
    note = str(args.get("note") or "").strip()

    if as_result or as_comment or note:
        # Deliver as a (result) comment carrying the files.
        return tool_add_bitrix_task_comment({
            "bitrix_task_id": task_id,
            "comment_text": note or ("Файлы приложены" if not as_result else "Результат приложен"),
            "as_result": as_result,
            "attachment_ids": attachment_ids,
            "author_bitrix_user_id": args.get("author_bitrix_user_id"),
            "author_name": args.get("author_name"),
            "expected_title": args.get("expected_title"),
        })

    # Default: pin the files to the task's files section.
    refs, summary = _resolve_attachment_disk_refs(attachment_ids)
    _attach_disk_refs_to_task(task_id, refs)
    return {
        "attached": True,
        "task": _task_payload(task, task_id),
        "attachments": summary,
        "target": "task_files",
        "rule": "Files are pinned to the task. Pass as_result=true (labelled result comment + files) "
                "or as_comment=true (files in a discussion comment) to deliver differently.",
    }


def tool_get_attachment_text(args: dict[str, Any]) -> dict[str, Any]:
    """Return the FULL extracted text of a stored attachment (chunked). This is how an agent reads a
    long document (e.g. a contract) end to end — nothing is truncated; the prompt only ever holds a
    preview, the complete text is served here on demand."""
    import attachments as _att
    token = str(args.get("attachment_id") or args.get("token") or "").strip()
    if not token:
        raise McpError(-32602, "Нужно указать attachment_id (токен вложения из промпта).")
    row = _att.get_attachment(token)
    if not row:
        raise McpError(-32602, f"Вложение {token} не найдено. Проверь attachment_id.")
    full = row.get("extracted_text") or ""
    total = len(full)
    try:
        offset = max(0, int(args.get("offset") or 0))
    except (TypeError, ValueError):
        offset = 0
    try:
        max_chars = int(args.get("max_chars") or 40000)
    except (TypeError, ValueError):
        max_chars = 40000
    max_chars = max(1000, min(max_chars, 120000))
    chunk = full[offset:offset + max_chars]
    next_offset = offset + len(chunk)
    return {
        "attachment_id": token,
        "file_name": row.get("file_name"),
        "kind": row.get("kind"),
        "total_chars": total,
        "offset": offset,
        "returned_chars": len(chunk),
        "has_more": next_offset < total,
        "next_offset": next_offset if next_offset < total else None,
        "text": chunk,
        "rule": "Read the whole document by calling again with offset=next_offset until has_more is false.",
    }


def tool_reopen_bitrix_task(args: dict[str, Any]) -> dict[str, Any]:
    task_id = _positive_bitrix_task_id(args.get("bitrix_task_id"))
    reason = str(args.get("reason") or args.get("comment_text") or "").strip()
    if not reason:
        raise McpError(-32602, "Нужно указать причину возобновления: reason.")
    if not _confirmed(args):
        raise McpError(
            -32602,
            "Возобновление задачи требует confirm=true. Сначала проверь задачу через search_tasks/get_task_comments, "
            "объясни пользователю, почему результат неудовлетворительный, и получи подтверждение на возобновление.",
        )

    task = _indexed_task_for_action(task_id)
    _assert_expected_task_title(task, args.get("expected_title"))

    # Optional new deadline when reopening ("возобнови с новым сроком до …").
    new_deadline = None
    if args.get("new_deadline"):
        new_deadline = _normalize_bitrix_deadline(args.get("new_deadline"))
        confirm_past = args.get("confirm_past_deadline")
        confirm_past = confirm_past is True or str(confirm_past or "").strip().lower() in {"true", "1", "yes", "да"}
        if not confirm_past and _deadline_in_past(new_deadline) is not None:
            raise McpError(
                -32602,
                "Новый срок " + str(args.get("new_deadline")) + " уже в прошлом. Уточни у пользователя новый "
                "срок в будущем, либо, если нужно оставить как есть, повтори с confirm_past_deadline=true.",
            )

    comment = "Результат требует доработки: " + reason
    if new_deadline:
        comment += f"\nНовый срок: {new_deadline}."
    comment_result = tool_add_bitrix_task_comment({
        "bitrix_task_id": task_id,
        "comment_text": comment,
        "author_bitrix_user_id": args.get("on_behalf_bitrix_user_id"),
        "author_name": args.get("on_behalf_name"),
        "expected_title": args.get("expected_title"),
    })

    attempts = (
        ("tasks.task.renew", {"taskId": task_id}),
        ("task.item.renew", {"TASKID": task_id}),
    )
    last_error: McpError | None = None
    reopened_response = None
    reopened_method = None
    for method, payload in attempts:
        try:
            reopened_response = _webhook_raw(method, payload)
            if isinstance(reopened_response, dict) and reopened_response.get("error") and not reopened_response.get("result"):
                raise McpError(-32010, str(reopened_response.get("error_description") or reopened_response.get("error")))
            reopened_method = method
            break
        except McpError as exc:
            last_error = exc
            continue
    if reopened_method is None:
        raise last_error or McpError(-32010, "Bitrix API call failed: could not reopen task.")

    if new_deadline:
        try:
            _webhook_raw("tasks.task.update", {"taskId": task_id, "fields": {"DEADLINE": new_deadline}})
        except McpError as exc:
            logging.warning("reopen: new deadline set failed task=%s: %s", task_id, str(exc)[:120])

    return {
        "reopened": True,
        "task": _task_payload(task, task_id),
        "reason": reason,
        "new_deadline": new_deadline,
        "comment": comment_result,
        "method": reopened_method,
        "bitrix_response": reopened_response.get("result") if isinstance(reopened_response, dict) else reopened_response,
        "rule": "Reopen requires exact bitrix_task_id, reason, and confirm=true after checking the result/comments. "
                "Pass new_deadline to renew with a new due date.",
    }


# --- General task editor + action tools (the "add field" chips of the Bitrix task) -----------
# One editor (update_bitrix_task) covers the field-group chips settable via REST; the action
# entities (checklists, time tracking, related tasks/Gantt, reminders) get dedicated tools.
# Everything goes through the v2 webhook (_webhook_raw): the v3 TaskDto rejects legacy fields
# like ACCOMPLICES/AUDITORS/TAGS/PARENT_ID/UF_* on update.

def _webhook_ok(resp: dict[str, Any], method: str) -> Any:
    """Return resp['result'] or raise a clear McpError if the webhook reported an error.
    Never a silent drop (the lesson from the REPLICATE recurring-task failure)."""
    if isinstance(resp, dict) and resp.get("error") and not resp.get("result"):
        raise McpError(-32010, f"Bitrix {method}: {resp.get('error_description') or resp.get('error')}")
    return resp.get("result") if isinstance(resp, dict) else resp


def tool_update_bitrix_task(args: dict[str, Any]) -> dict[str, Any]:
    """Edit an existing Bitrix task: any subset of the field-group chips (соисполнители,
    наблюдатели, теги, родительская задача, проект, планирование сроков, элементы CRM,
    пользовательские поля, срок, приоритет, ответственный, название/описание)."""
    task_id = _positive_bitrix_task_id(args.get("bitrix_task_id"))
    task = _indexed_task_for_action(task_id)
    _assert_expected_task_title(task, args.get("expected_title"))

    fields: dict[str, Any] = {}
    changed: list[str] = []

    if str(args.get("title") or "").strip():
        fields["TITLE"] = str(args["title"]).strip(); changed.append("название")
    if args.get("description") not in (None, ""):
        fields["DESCRIPTION"] = str(args["description"]).strip(); changed.append("описание")
    if args.get("priority") not in (None, ""):
        pr = str(args.get("priority")).strip().lower()
        fields["PRIORITY"] = 2 if pr in {"high", "critical", "2", "важно", "высокий"} else 1
        changed.append("приоритет")
    if args.get("deadline") not in (None, ""):
        dl = _normalize_bitrix_deadline(args.get("deadline"))
        confirm_past = args.get("confirm_past_deadline") is True or str(args.get("confirm_past_deadline") or "").strip().lower() in {"true", "1", "yes", "да"}
        if not confirm_past and _deadline_in_past(dl) is not None:
            raise McpError(-32602, "Новый срок уже в прошлом. Уточни у пользователя срок в будущем, "
                                   "либо повтори с confirm_past_deadline=true, чтобы оставить как есть.")
        fields["DEADLINE"] = dl; changed.append("срок")
    if args.get("responsible_bitrix_user_id") not in (None, "") or args.get("responsible_name"):
        resp_user = _resolve_active_bitrix_user(args.get("responsible_bitrix_user_id"), args.get("responsible_name"))
        fields["RESPONSIBLE_ID"] = int(resp_user["bitrix_user_id"]); changed.append("ответственный")
    if args.get("accomplice_bitrix_user_ids") not in (None, "") or args.get("accomplice_names"):
        fields["ACCOMPLICES"] = _resolve_person_ids(args, "accomplice_bitrix_user_ids", "accomplice_names", "Соисполнитель")
        changed.append("соисполнители")
    if args.get("auditor_bitrix_user_ids") not in (None, "") or args.get("auditor_names"):
        fields["AUDITORS"] = _resolve_person_ids(args, "auditor_bitrix_user_ids", "auditor_names", "Наблюдатель")
        changed.append("наблюдатели")
    if isinstance(args.get("tags"), list):
        fields["TAGS"] = [str(t).strip() for t in args["tags"] if str(t or "").strip()]; changed.append("теги")
    if args.get("parent_task_id") not in (None, ""):
        fields["PARENT_ID"] = _positive_bitrix_task_id(args.get("parent_task_id")); changed.append("родительская задача")
    if args.get("group_id") not in (None, ""):
        try:
            fields["GROUP_ID"] = int(args.get("group_id"))
        except (TypeError, ValueError) as exc:
            raise McpError(-32602, "group_id должен быть числом.") from exc
        changed.append("проект")
    if args.get("start_plan") not in (None, ""):
        fields["START_DATE_PLAN"] = _normalize_bitrix_datetime(args.get("start_plan"), field="start_plan"); changed.append("план: старт")
    if args.get("end_plan") not in (None, ""):
        fields["END_DATE_PLAN"] = _normalize_bitrix_datetime(args.get("end_plan"), field="end_plan"); changed.append("план: финиш")
    if args.get("time_estimate_hours") not in (None, ""):
        try:
            fields["TIME_ESTIMATE"] = int(float(args.get("time_estimate_hours")) * 3600)
        except (TypeError, ValueError) as exc:
            raise McpError(-32602, "time_estimate_hours должно быть числом.") from exc
        changed.append("оценка времени")
    crm = _clean_crm_elements(args.get("crm_elements"))
    if crm:
        fields["UF_CRM_TASK"] = crm; changed.append("элементы CRM")
    custom = _clean_custom_fields(args.get("custom_fields"))
    if custom:
        fields.update(custom); changed.append("пользовательские поля")

    if not fields:
        raise McpError(-32602, "Нечего менять: передай хотя бы одно поле (напр. accomplice_names, tags, "
                               "deadline, parent_task_id, group_id, start_plan/end_plan, crm_elements, custom_fields).")

    resp = _webhook_raw("tasks.task.update", {"taskId": task_id, "fields": fields})
    _webhook_ok(resp, "tasks.task.update")
    return {
        "updated": True,
        "task": _task_payload(task, task_id),
        "changed": changed,
        "fields_set": sorted(fields.keys()),
        "rule": "Правит существующую задачу по точному bitrix_task_id. Списочные поля (соисполнители/"
                "наблюдатели/теги) ЗАМЕНЯЮТСЯ переданным списком — передавай полный набор. Смену "
                "ответственного/родительской задачи подтверди у пользователя перед вызовом.",
    }


def tool_add_task_checklist(args: dict[str, Any]) -> dict[str, Any]:
    """Add checklist items (чек-лист) to an existing task. items = list of strings or
    {title, complete?}."""
    task_id = _positive_bitrix_task_id(args.get("bitrix_task_id"))
    items = args.get("items") or args.get("checklist")
    if not (isinstance(items, (list, tuple)) and any(
            (str((it.get("title") if isinstance(it, dict) else it) or "").strip()) for it in items)):
        raise McpError(-32602, "Нужно указать items — список пунктов чек-листа (строки или {title, complete}).")
    task = _indexed_task_for_action(task_id)
    _assert_expected_task_title(task, args.get("expected_title"))
    summary = _add_checklist_items(task_id, items)
    added = [s for s in summary if not s.get("error")]
    if not added and summary:
        raise McpError(-32010, "Bitrix не принял пункты чек-листа: " + str(summary[0].get("error") or "")[:180])
    return {"checklist_added": True, "task": _task_payload(task, task_id), "items": summary,
            "rule": "Пункты чек-листа добавлены (task.checklistitem.add). complete=true отмечает пункт выполненным."}


def tool_log_task_time(args: dict[str, Any]) -> dict[str, Any]:
    """Log spent time (учёт времени) on a task via task.elapseditem.add. Time from hours/minutes/
    seconds; optional comment and on-behalf user."""
    task_id = _positive_bitrix_task_id(args.get("bitrix_task_id"))
    total = 0
    for key, mult in (("hours", 3600), ("minutes", 60), ("seconds", 1)):
        if args.get(key) not in (None, ""):
            try:
                total += int(float(args.get(key)) * mult)
            except (TypeError, ValueError) as exc:
                raise McpError(-32602, f"{key} должно быть числом.") from exc
    if total <= 0:
        raise McpError(-32602, "Укажи затраченное время: hours и/или minutes (или seconds).")
    task = _indexed_task_for_action(task_id)
    _assert_expected_task_title(task, args.get("expected_title"))
    fields: dict[str, Any] = {"SECONDS": total}
    comment = str(args.get("comment") or args.get("comment_text") or "").strip()
    if comment:
        fields["COMMENT_TEXT"] = comment
    actor = _resolve_task_actor(args, "on_behalf_bitrix_user_id", "on_behalf_name")
    if actor:
        fields["USER_ID"] = int(actor["bitrix_user_id"])
    resp = _webhook_raw("task.elapseditem.add", {"TASKID": task_id, "FIELDS": fields})
    item_id = _webhook_ok(resp, "task.elapseditem.add")
    return {"time_logged": True, "task": _task_payload(task, task_id), "seconds": total,
            "minutes": round(total / 60, 1), "elapsed_item_id": item_id,
            "on_behalf": ({"bitrix_user_id": actor.get("bitrix_user_id"), "full_name": actor.get("full_name")} if actor else None),
            "rule": "Записан учёт времени по задаче (task.elapseditem.add)."}


_TASK_LINK_TYPES = {"finish_start": 2, "start_start": 0, "start_finish": 1, "finish_finish": 3}


def tool_link_tasks(args: dict[str, Any]) -> dict[str, Any]:
    """Link two tasks (связанные задачи / зависимость для Ганта) via task.dependence.add.
    task_id_from depends on task_id_to; link_type finish_start (default) / start_start /
    start_finish / finish_finish, or an int 0..3."""
    src = _positive_bitrix_task_id(args.get("task_id_from") or args.get("bitrix_task_id"))
    dst = _positive_bitrix_task_id(args.get("task_id_to") or args.get("related_task_id"))
    if src == dst:
        raise McpError(-32602, "task_id_from и task_id_to должны быть разными задачами.")
    lt_raw = args.get("link_type")
    if lt_raw in (None, ""):
        link_type = 2
    elif isinstance(lt_raw, int) or str(lt_raw).isdigit():
        link_type = int(lt_raw)
        if link_type not in (0, 1, 2, 3):
            raise McpError(-32602, "link_type (число) должен быть 0..3.")
    else:
        key = str(lt_raw).strip().lower()
        if key not in _TASK_LINK_TYPES:
            raise McpError(-32602, "link_type: finish_start | start_start | start_finish | finish_finish (или 0..3).")
        link_type = _TASK_LINK_TYPES[key]
    resp = _webhook_raw("task.dependence.add", {"taskIdFrom": src, "taskIdTo": dst, "linkType": link_type})
    result = _webhook_ok(resp, "task.dependence.add")
    return {"linked": True, "task_id_from": src, "task_id_to": dst, "link_type": link_type,
            "bitrix_response": result,
            "rule": "Создана связь/зависимость между задачами (task.dependence.add) — она же строит Гант."}


def tool_add_task_reminder(args: dict[str, Any]) -> dict[str, Any]:
    """Add a reminder (напоминание) for a task at an absolute time. Best-effort: reminders REST
    varies by portal — on rejection the tool returns a clear error rather than failing silently."""
    task_id = _positive_bitrix_task_id(args.get("bitrix_task_id"))
    if args.get("remind_at") in (None, ""):
        raise McpError(-32602, "Укажи remind_at — когда напомнить (YYYY-MM-DD HH:MM или DD.MM.YYYY HH:MM).")
    remind_iso = _normalize_bitrix_datetime(args.get("remind_at"), field="remind_at")
    # Whom to remind: explicit user, else the responsible person of the indexed task.
    user = _resolve_task_actor(args, "user_bitrix_user_id", "user_name")
    uid = int(user["bitrix_user_id"]) if user else None
    task = _indexed_task_for_action(task_id)
    if uid is None and task and task.get("responsible_bitrix_user_id"):
        uid = int(task["responsible_bitrix_user_id"])
    if uid is None:
        raise McpError(-32602, "Кому напомнить не определено: передай user_name/user_bitrix_user_id.")
    # task.reminder.add wants ONE `data` object with TASK_ID/USER_ID/TYPE + DATE in Bitrix format
    # (DD.MM.YYYY HH:MM:SS). Some portals block this action for the integration webhook user
    # (ACTION_FAILED) — degrade to a clear, actionable message instead of a raw error.
    bx_date = datetime.fromisoformat(remind_iso).strftime("%d.%m.%Y %H:%M:%S")
    try:
        resp = _webhook_raw("task.reminder.add", {
            "data": {"TASK_ID": task_id, "USER_ID": uid, "TYPE": "date", "DATE": bx_date}})
        result = _webhook_ok(resp, "task.reminder.add")
    except McpError as exc:
        raise McpError(
            -32010,
            "Этот портал Bitrix не разрешает ставить напоминания через REST (task.reminder.add). "
            "Надёжные альтернативы: у задачи есть срок — Bitrix сам напомнит о дедлайне; для отдельного "
            "напоминания в нужное время используй schedule_my_automation (агентское напоминание). "
            "Тех.детали: " + str(exc.message)[:150]) from exc
    return {"reminder_added": True, "task_id": task_id, "user_bitrix_user_id": uid, "remind_at": remind_iso,
            "bitrix_response": result, "rule": "Напоминание по задаче добавлено (task.reminder.add)."}


def tool_list_task_userfields(args: dict[str, Any]) -> dict[str, Any]:
    """List the custom task fields (пользовательские поля, UF_*) defined on the portal, so the agent
    uses real UF_* codes in create/update instead of guessing. Read-only. Uses tasks.task.getFields
    (task.item.userfield.getlist is blocked for the integration webhook on this portal)."""
    resp = _webhook_raw("tasks.task.getFields", {})
    result = resp.get("result") if isinstance(resp, dict) else None
    fmap = result.get("fields") if isinstance(result, dict) and isinstance(result.get("fields"), dict) else result
    fields = []
    if isinstance(fmap, dict):
        for code, meta in fmap.items():
            if not str(code).upper().startswith("UF_"):
                continue
            label, ftype = code, None
            if isinstance(meta, dict):
                label = meta.get("title") or meta.get("TITLE") or meta.get("EDIT_FORM_LABEL") or code
                ftype = meta.get("type") or meta.get("TYPE")
            fields.append({"code": code, "label": label, "type": ftype})
    return {"count": len(fields), "fields": fields,
            "rule": "Коды пользовательских полей задач (UF_*). Их значения передавай в custom_fields "
                    "инструментов create_bitrix_task / update_bitrix_task. Пусто = на портале нет своих "
                    "полей задач (кроме системных)."}


# --- Recurring (regular) tasks — fired by the agent's OWN scheduler, not Bitrix ---------------
# The portal has no paid subscription and Bitrix's automatic task replication (task.template.add
# REPLICATE) is a paid feature: the template is created but never spawns tasks ("не создаётся
# нормально"). So the recurrence schedule lives in bitrix_recurring_tasks (migrations 045+046) and
# the Albery app's recurring_scheduler.py creates a plain one-off task on time — plain tasks via
# REST work without a subscription. This module defines the schedule maths + the create-from-spec
# path both the tool and the scheduler share; the scheduler thread itself lives in
# recurring_scheduler.py (started from agent_center.py, same process as these tools).

_WEEKDAY_NUM = {"MO": 1, "TU": 2, "WE": 3, "TH": 4, "FR": 5, "SA": 6, "SU": 7}
_WEEKDAY_RU = {1: "понедельник", 2: "вторник", 3: "среда", 4: "четверг", 5: "пятница", 6: "суббота", 7: "воскресенье"}


def _parse_hhmm(value: Any, field: str, default: str | None = None) -> str:
    s = str(value or "").strip()
    if not s and default is not None:
        return default
    m = re.match(r"^(\d{1,2}):(\d{2})$", s)
    if not m or not (0 <= int(m.group(1)) <= 23) or not (0 <= int(m.group(2)) <= 59):
        raise McpError(-32602, f"{field} должно быть в формате ЧЧ:ММ (например 10:00).")
    return f"{int(m.group(1)):02d}:{m.group(2)}"


def _recurring_schedule_desc(period: str, interval: int, weekday_nums: list[int],
                             day_of_month: int | None, create_time: str, deadline_desc: str) -> str:
    every = "" if interval == 1 else f"каждые {interval} "
    if period == "daily":
        base = "каждый день" if interval == 1 else f"каждые {interval} дн."
    elif period == "weekly":
        days = ", ".join(_WEEKDAY_RU.get(n, str(n)) for n in weekday_nums)
        base = f"еженедельно ({days})" if interval == 1 else f"каждые {interval} нед. ({days})"
    else:
        base = f"ежемесячно, {day_of_month} числа" if interval == 1 else f"каждые {interval} мес., {day_of_month} числа"
    txt = f"{base}, создание в {create_time}"
    if deadline_desc:
        txt += f", дедлайн {deadline_desc}"
    return txt


def _recurring_day_matches(period: str, interval: int, weekday_nums: list[int],
                           day_of_month: int | None, d: "date", anchor: "date") -> bool:
    """Does calendar day `d` (>= anchor) match the recurrence? interval>1 counts from `anchor`."""
    if d < anchor:
        return False
    if period == "daily":
        return interval == 1 or ((d - anchor).days % interval == 0)
    if period == "weekly":
        if d.isoweekday() not in weekday_nums:  # Mon=1..Sun=7 (== Bitrix WEEK_DAYS)
            return False
        if interval == 1:
            return True
        a_mon = anchor - timedelta(days=anchor.isoweekday() - 1)
        d_mon = d - timedelta(days=d.isoweekday() - 1)
        weeks = (d_mon - a_mon).days // 7
        return weeks >= 0 and weeks % interval == 0
    # monthly: fire on day_of_month, clamped to the month's length (31 -> last day of a short month)
    eff = min(int(day_of_month or 1), calendar.monthrange(d.year, d.month)[1])
    if d.day != eff:
        return False
    if interval == 1:
        return True
    months = (d.year - anchor.year) * 12 + (d.month - anchor.month)
    return months >= 0 and months % interval == 0


def _recurring_next_run(period: str, interval: int, weekday_nums: list[int],
                        day_of_month: int | None, create_time: str, *,
                        after: "datetime", anchor: "date | None" = None) -> "datetime | None":
    """Next MSK-aware datetime at create_time on a matching day, strictly AFTER `after`.
    None if nothing within ~2 years (shouldn't happen for valid schedules)."""
    hh, mm = (int(x) for x in create_time.split(":"))
    after = after.astimezone(_MSK_TZ) if after.tzinfo else after.replace(tzinfo=_MSK_TZ)
    anchor = anchor or after.date()
    for offset in range(0, 800):
        d = after.date() + timedelta(days=offset)
        cand = datetime(d.year, d.month, d.day, hh, mm, tzinfo=_MSK_TZ)
        if cand <= after:
            continue
        if _recurring_day_matches(period, interval, weekday_nums, day_of_month, d, anchor):
            return cand
    return None


def create_oneoff_task_from_spec(spec: dict[str, Any], deadline_iso: str) -> dict[str, Any]:
    """Create a plain one-off Bitrix task from a stored recurring spec (used by the scheduler).
    People are already resolved to ids in the spec. Returns {task_id, checklist}. Raises on failure
    (the scheduler records last_error and retries)."""
    title = str(spec.get("title") or "").strip()
    description = str(spec.get("description") or "").strip() or title
    responsible_id = int(spec["responsible_bitrix_id"])
    priority = 2 if str(spec.get("priority") or "").strip().lower() in {"high", "critical", "2", "важно", "высокий"} else 1
    fields = _assemble_task_fields(
        title=title, description=description, responsible_id=responsible_id, deadline_iso=deadline_iso,
        priority=priority,
        auditor_ids=spec.get("auditor_ids") or None,
        accomplice_ids=spec.get("accomplice_ids") or None,
        creator_id=int(spec["creator_bitrix_id"]) if spec.get("creator_bitrix_id") else responsible_id,
        tags=spec.get("tags") or None, group_id=spec.get("group_id") or None,
        crm_elements=spec.get("crm_elements") or None, custom_fields=spec.get("custom_fields") or None,
    )
    response = _bitrix_call_with_fallback("tasks.task.add", {"fields": fields})
    result = response.get("result") if isinstance(response, dict) else {}
    task_id = None
    if isinstance(result, dict):
        t = result.get("task") if isinstance(result.get("task"), dict) else {}
        task_id = t.get("id") or result.get("id")
    else:
        task_id = result
    if not task_id:
        raise McpError(-32010, f"Bitrix не создал разовую задачу по расписанию: {response.get('error_description') if isinstance(response, dict) else response}")
    checklist = _add_checklist_items(int(task_id), spec.get("checklist")) if spec.get("checklist") else None
    # Scheduler-created instances get the same offer-to-help comment as one-off agent tasks.
    try:
        from task_offers import schedule_offer
        schedule_offer(task_id, title=title, description=description,
                       checklist=spec.get("checklist"),
                       responsible_id=responsible_id,
                       creator_id=spec.get("creator_bitrix_id"))
    except Exception:  # noqa: BLE001
        logging.warning("recurring instance: offer scheduling failed task=%s", task_id, exc_info=True)
    return {"task_id": int(task_id), "checklist": checklist}


def _ddmmyyyy_to_date(value: Any) -> "date | None":
    """Parse a DD.MM.YYYY string (the registry's until format) to a date; None on failure."""
    try:
        d, m, y = str(value).split(".")
        return date(int(y), int(m), int(d))
    except Exception:  # noqa: BLE001
        return None


def tool_create_recurring_task(args: dict[str, Any]) -> dict[str, Any]:
    """Create a RECURRING (regular) task that is fired by the agent's OWN scheduler — the app creates
    a plain one-off Bitrix task on schedule (e.g. every Friday at 10:00, deadline 19:00 the same day).
    No Bitrix subscription needed (unlike Bitrix's own recurring-task templates)."""
    title = str(args.get("title") or "").strip()
    if not title:
        raise McpError(-32602, "Нужно указать название задачи: title.")
    result_criteria = str(args.get("result_criteria") or "").strip()
    if not result_criteria:
        raise McpError(-32602, "У повторяющейся задачи ОБЯЗАН быть результат. Спроси пользователя, "
                               "по чему поймём, что сделано, и передай result_criteria.")
    period = str(args.get("period") or "").strip().lower()
    if period not in {"daily", "weekly", "monthly"}:
        raise McpError(-32602, "period должно быть daily, weekly или monthly.")
    try:
        interval = int(args.get("interval") or 1)
    except (TypeError, ValueError) as exc:
        raise McpError(-32602, "interval должно быть целым числом ≥ 1.") from exc
    if interval < 1:
        raise McpError(-32602, "interval должно быть ≥ 1.")

    weekday_nums: list[int] = []
    day_of_month: int | None = None
    if period == "weekly":
        raw = args.get("weekdays")
        if not isinstance(raw, list) or not raw:
            raise McpError(-32602, "Для weekly укажи weekdays — список дней (MO/TU/WE/TH/FR/SA/SU), например ['FR'].")
        for code in raw:
            c = str(code or "").strip().upper()
            if c not in _WEEKDAY_NUM:
                raise McpError(-32602, f"weekdays: '{code}' — неверный код. Используй MO/TU/WE/TH/FR/SA/SU.")
            if _WEEKDAY_NUM[c] not in weekday_nums:
                weekday_nums.append(_WEEKDAY_NUM[c])
    elif period == "monthly":
        try:
            day_of_month = int(args.get("day_of_month"))
        except (TypeError, ValueError) as exc:
            raise McpError(-32602, "Для monthly укажи day_of_month (1-31).") from exc
        if not 1 <= day_of_month <= 31:
            raise McpError(-32602, "day_of_month должно быть 1-31.")

    create_time = _parse_hhmm(args.get("create_time"), "create_time", default="10:00")
    # Deadline: either an explicit deadline_time (same day; wraps to next day if earlier than create)
    # or deadline_after_hours. Default: end of the creation day is not assumed — require one.
    deadline_after_seconds: int | None = None
    deadline_desc = ""
    if args.get("deadline_time") not in (None, ""):
        dl = _parse_hhmm(args.get("deadline_time"), "deadline_time")
        ch, cm = (int(x) for x in create_time.split(":"))
        dh, dm = (int(x) for x in dl.split(":"))
        diff = (dh * 60 + dm) - (ch * 60 + cm)
        if diff <= 0:
            diff += 24 * 60  # deadline on the next day
            deadline_desc = f"{dl} следующего дня"
        else:
            deadline_desc = f"{dl} того же дня"
        deadline_after_seconds = diff * 60
    elif args.get("deadline_after_hours") not in (None, ""):
        try:
            hrs = float(args.get("deadline_after_hours"))
        except (TypeError, ValueError) as exc:
            raise McpError(-32602, "deadline_after_hours должно быть числом.") from exc
        if hrs <= 0:
            raise McpError(-32602, "deadline_after_hours должно быть > 0.")
        deadline_after_seconds = int(hrs * 3600)
        deadline_desc = f"через {hrs:g} ч после создания"
    else:
        raise McpError(-32602, "Укажи срок каждой задачи: deadline_time (ЧЧ:ММ, напр. 19:00) или deadline_after_hours.")

    until = None
    if args.get("until"):
        until = str(args.get("until")).strip()
        if not re.match(r"^\d{4}-\d{2}-\d{2}$", until) and not re.match(r"^\d{2}\.\d{2}\.\d{4}$", until):
            raise McpError(-32602, "until должно быть YYYY-MM-DD или DD.MM.YYYY.")
        if re.match(r"^\d{4}-\d{2}-\d{2}$", until):
            y, m, d = until.split("-")
            until = f"{d}.{m}.{y}"  # Bitrix wants DD.MM.YYYY

    responsible = _resolve_active_bitrix_user(args.get("responsible_bitrix_user_id"), args.get("responsible_name"))
    creator = None
    if args.get("creator_bitrix_user_id") not in (None, "") or args.get("creator_name"):
        creator = _resolve_active_bitrix_user(args.get("creator_bitrix_user_id"), args.get("creator_name"))

    description = str(args.get("description") or "").strip() or title
    if "Критерий результата" not in description:
        description += "\n\nКритерий результата: " + result_criteria

    # Optional extra fields carried into every generated instance (same palette as create_bitrix_task).
    priority_raw = str(args.get("priority") or "normal").strip().lower()
    priority = "high" if priority_raw in {"high", "critical", "2", "важно", "высокий"} else "normal"
    auditors = _resolve_active_bitrix_users(
        args.get("auditor_bitrix_user_ids"), args.get("auditor_names"),
        role_label="Наблюдатель", id_field="auditor_bitrix_user_ids", name_field="auditor_names")
    accomplices = _resolve_active_bitrix_users(
        args.get("accomplice_bitrix_user_ids"), args.get("accomplice_names"),
        role_label="Соисполнитель", id_field="accomplice_bitrix_user_ids", name_field="accomplice_names")
    group_id = None
    if args.get("group_id") not in (None, ""):
        try:
            group_id = int(args.get("group_id"))
        except (TypeError, ValueError) as exc:
            raise McpError(-32602, "group_id должен быть числом.") from exc
    tags = [str(t).strip() for t in args.get("tags") if str(t or "").strip()] if isinstance(args.get("tags"), list) else None
    crm_elements = _clean_crm_elements(args.get("crm_elements")) or None
    custom_fields = _clean_custom_fields(args.get("custom_fields")) or None
    checklist = args.get("checklist") if isinstance(args.get("checklist"), list) else None

    schedule_desc = _recurring_schedule_desc(period, interval, weekday_nums, day_of_month, create_time, deadline_desc)
    until_date = until  # DD.MM.YYYY or None (normalized above)

    now_msk = datetime.now(_MSK_TZ)
    next_run = _recurring_next_run(period, interval, weekday_nums, day_of_month, create_time, after=now_msk)
    until_d = _ddmmyyyy_to_date(until_date) if until_date else None
    if next_run and until_d and next_run.date() > until_d:
        next_run = None  # end date already passed — nothing to schedule

    spec = {
        "title": title,
        "description": description,
        "responsible_bitrix_id": int(responsible["bitrix_user_id"]),
        "creator_bitrix_id": int(creator["bitrix_user_id"]) if creator else None,
        "auditor_ids": [int(u["bitrix_user_id"]) for u in auditors] or None,
        "accomplice_ids": [int(u["bitrix_user_id"]) for u in accomplices] or None,
        "tags": tags, "group_id": group_id, "crm_elements": crm_elements,
        "custom_fields": custom_fields, "checklist": checklist,
        "priority": priority, "deadline_after_seconds": deadline_after_seconds,
    }

    # Which agent's «Автоматизации» tab shows this row: the per-agent MCP connector injects
    # its slug as _agent_slug (see handle_request); legacy connectors fall back to 'main'.
    agent_slug = str(args.get("_agent_slug") or "").strip() or "main"

    recurring_id = None
    try:
        with connect() as conn:
            with conn.cursor() as cur:
                cur.execute(
                    "INSERT INTO bitrix_recurring_tasks (title, description, responsible_bitrix_id, "
                    "responsible_name, creator_bitrix_id, period, interval_every, weekdays, day_of_month, "
                    "create_time, deadline_after_seconds, deadline_desc, schedule_desc, until_date, "
                    "result_criteria, priority, spec, next_run_at, source, agent_slug) "
                    "VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s::jsonb,%s,%s,%s) RETURNING id",
                    (title, description, int(responsible["bitrix_user_id"]), responsible.get("full_name"),
                     (int(creator["bitrix_user_id"]) if creator else None), period, interval,
                     weekday_nums or None, day_of_month, create_time, deadline_after_seconds, deadline_desc,
                     schedule_desc, until_date, result_criteria, priority,
                     json.dumps(spec, ensure_ascii=False), next_run, "agent_scheduler", agent_slug))
                row = cur.fetchone()
                recurring_id = (row or {}).get("id") if isinstance(row, dict) else (row[0] if row else None)
    except Exception as exc:  # noqa: BLE001
        logging.warning("recurring registry insert failed: %s", repr(exc)[:200])
        raise McpError(-32010, "Не удалось сохранить повторяющуюся задачу в реестр. Повтори позже.") from exc

    return {
        "created": True,
        "recurring": True,
        "recurring_id": recurring_id,
        "title": title,
        "responsible": {"bitrix_user_id": responsible.get("bitrix_user_id"), "full_name": responsible.get("full_name")},
        "creator": ({"bitrix_user_id": creator.get("bitrix_user_id"), "full_name": creator.get("full_name")} if creator else None),
        "schedule": schedule_desc,
        "next_run": next_run.isoformat() if next_run else None,
        "period": period, "interval": interval, "weekdays": weekday_nums or None, "day_of_month": day_of_month,
        "create_time": create_time, "deadline": deadline_desc, "until": until_date,
        "rule": "Повторяющаяся задача поставлена в СОБСТВЕННЫЙ планировщик агента (не в Bitrix — там нет "
                "подписки). Приложение само создаёт обычную разовую задачу в момент next_run по расписанию. "
                "Скажи пользователю: запись видна в Центре Агента → Агенты → вкладка «Автоматизации» "
                "(чип «регулярная задача»), там же её можно выключить/удалить. "
                "Смотреть/останавливать из чата: list_recurring_tasks / delete_recurring_task.",
    }


def tool_list_recurring_tasks(args: dict[str, Any]) -> dict[str, Any]:
    """List recurring (regular) tasks — all of them or for one person. Reads the agent's registry;
    each row shows the schedule, the next auto-creation time, and the last created instance."""
    filt_user = None
    if args.get("responsible_bitrix_user_id") not in (None, "") or args.get("responsible_name"):
        u = _resolve_active_bitrix_user(args.get("responsible_bitrix_user_id"), args.get("responsible_name"))
        filt_user = int(u["bitrix_user_id"])
    include_inactive = bool(args.get("include_inactive", False))

    where = [] if include_inactive else ["active"]
    params: list[Any] = []
    if filt_user is not None:
        where.append("responsible_bitrix_id = %s")
        params.append(filt_user)
    where_sql = ("WHERE " + " AND ".join(where)) if where else ""
    with connect() as conn:
        with conn.cursor() as cur:
            if not safe_table_exists(cur, "bitrix_recurring_tasks"):
                return {"items": [], "note": "Таблица реестра ещё не создана."}
            cur.execute(
                f"SELECT id, title, responsible_bitrix_id, responsible_name, period, interval_every, "
                f"weekdays, day_of_month, create_time, deadline_desc, schedule_desc, until_date, active, "
                f"next_run_at, last_created_at, last_task_id, last_error, created_at "
                f"FROM bitrix_recurring_tasks {where_sql} ORDER BY responsible_name NULLS LAST, created_at DESC",
                params,
            )
            rows = cur.fetchall()

    items = []
    for r in rows:
        row = dict(r)
        for col in ("next_run_at", "last_created_at", "created_at"):
            v = row.get(col)
            row[col] = _to_msk(v).isoformat() if hasattr(v, "isoformat") else v
        items.append(row)

    return {
        "count": len(items),
        "responsible_filter": filt_user,
        "items": items,
        "note": ("Повторяющиеся задачи из реестра агента. Их создаёт СОБСТВЕННЫЙ планировщик приложения "
                 "(не Bitrix — там нет подписки): next_run_at — ближайшее авто-создание, last_task_id — "
                 "последняя созданная задача. Остановить: delete_recurring_task."),
    }


def tool_delete_recurring_task(args: dict[str, Any]) -> dict[str, Any]:
    """Stop a recurring task: deactivate it in the registry so the scheduler no longer creates it.
    Identify it by recurring_id (from list_recurring_tasks). Already-created task instances are not
    touched. Requires confirm=true."""
    rec_id = args.get("recurring_id")
    if rec_id in (None, ""):
        raise McpError(-32602, "Укажи recurring_id — id повторяющейся задачи из list_recurring_tasks.")
    try:
        rec_id = int(rec_id)
    except (TypeError, ValueError) as exc:
        raise McpError(-32602, "recurring_id должен быть числом.") from exc
    if not _confirmed(args):
        raise McpError(-32602, "Остановка повторяющейся задачи требует confirm=true. Сначала покажи "
                               "пользователю задачу (list_recurring_tasks) и получи подтверждение.")
    with connect() as conn:
        with conn.cursor() as cur:
            cur.execute("SELECT id, title, schedule_desc, responsible_name, active, bitrix_template_id "
                        "FROM bitrix_recurring_tasks WHERE id = %s", (rec_id,))
            row = cur.fetchone()
            if not row:
                raise McpError(-32602, f"Повторяющаяся задача {rec_id} не найдена в реестре.")
            row = dict(row)
            cur.execute("UPDATE bitrix_recurring_tasks SET active=false, updated_at=now() WHERE id = %s", (rec_id,))
    # Best-effort: also drop the dead Bitrix template for legacy (pre-scheduler) rows.
    tpl = row.get("bitrix_template_id")
    if tpl:
        try:
            _webhook_raw("task.template.delete", {"id": int(tpl)})
        except Exception:  # noqa: BLE001
            pass
    return {"stopped": True, "recurring_id": rec_id, "title": row.get("title"),
            "schedule": row.get("schedule_desc"), "responsible_name": row.get("responsible_name"),
            "rule": "Повторяющаяся задача остановлена (active=false) — планировщик её больше не создаёт. "
                    "Уже созданные задачи остаются."}


_DL_DESC_RE = re.compile(r"^(\d{1,2}:\d{2}) (того же|следующего) дня$")


def _normalize_weekday_list(raw: Any) -> list[int]:
    """Accepts MO..SU codes and/or ints 1..7 (Mon=1); returns a sorted unique int list."""
    if not isinstance(raw, list) or not raw:
        raise McpError(-32602, "weekdays должен быть непустым списком дней: MO/TU/WE/TH/FR/SA/SU или числа 1-7 (Пн=1).")
    out: list[int] = []
    for v in raw:
        if isinstance(v, str) and v.strip().upper() in _WEEKDAY_NUM:
            n = _WEEKDAY_NUM[v.strip().upper()]
        else:
            try:
                n = int(v)
            except (TypeError, ValueError) as exc:
                raise McpError(-32602, f"weekdays: '{v}' — не день недели (MO..SU или 1-7).") from exc
        if not 1 <= n <= 7:
            raise McpError(-32602, f"weekdays: {n} вне диапазона 1-7 (Пн=1..Вс=7).")
        if n not in out:
            out.append(n)
    return sorted(out)


def apply_recurring_update(rec_id: int, changes: dict[str, Any]) -> dict[str, Any]:
    """Update a recurring-task registry row: schedule (weekdays / create_time / deadline_time /
    day_of_month / until) and/or content (title / description / checklist / result_criteria /
    priority). Recomputes deadline offset, human schedule text and next_run_at, and keeps the
    jsonb spec in sync so the next created instance reflects the edit. Shared by the
    «Автоматизации» tab editor (PATCH endpoint) and the update_recurring_task MCP tool.

    Weekday semantics: a weekdays list of all 7 days == period 'daily'; any subset == 'weekly'
    with those days (so «каждый день, но не в выходные» is weekdays=[1..5]). day_of_month
    switches the row to 'monthly'."""
    try:
        rec_id = int(rec_id)
    except (TypeError, ValueError) as exc:
        raise McpError(-32602, "recurring_id должен быть числом.") from exc
    with connect() as conn:
        with conn.cursor() as cur:
            cur.execute("SELECT * FROM bitrix_recurring_tasks WHERE id = %s", (rec_id,))
            row = cur.fetchone()
    if not row:
        raise McpError(-32602, f"Повторяющаяся задача {rec_id} не найдена (list_recurring_tasks покажет id).")
    row = dict(row)

    period = row.get("period") or "daily"
    interval = int(row.get("interval_every") or 1)
    weekday_nums = list(row.get("weekdays") or [])
    day_of_month = row.get("day_of_month")
    create_time = row.get("create_time") or "10:00"
    dl_secs = row.get("deadline_after_seconds")
    deadline_desc = row.get("deadline_desc") or ""

    if changes.get("weekdays") is not None:
        wd = _normalize_weekday_list(changes["weekdays"])
        if len(wd) == 7:
            period, weekday_nums, day_of_month = "daily", [], None
        else:
            period, weekday_nums, day_of_month = "weekly", wd, None
        interval = 1  # a day-of-week edit always means "every listed day"
    if changes.get("day_of_month") is not None:
        try:
            dom = int(changes["day_of_month"])
        except (TypeError, ValueError) as exc:
            raise McpError(-32602, "day_of_month должно быть числом 1-31.") from exc
        if not 1 <= dom <= 31:
            raise McpError(-32602, "day_of_month должно быть 1-31.")
        period, weekday_nums, day_of_month, interval = "monthly", [], dom, 1
    if period == "weekly" and not weekday_nums:
        raise McpError(-32602, "Для weekly-расписания нужен непустой список weekdays.")

    old_create_time = create_time
    if changes.get("create_time") not in (None, ""):
        create_time = _parse_hhmm(changes["create_time"], "create_time")

    # Deadline: explicit edit wins; otherwise, when the creation time moved and the old deadline
    # was anchored to a clock time («18:00 того же дня»), keep that clock time by recomputing the
    # offset — moving creation 09:00→10:00 must not silently move the deadline 18:00→19:00.
    def _offset_from(dl_hhmm: str) -> tuple[int, str]:
        ch, cm = (int(x) for x in create_time.split(":"))
        dh, dm = (int(x) for x in dl_hhmm.split(":"))
        diff = (dh * 60 + dm) - (ch * 60 + cm)
        if diff <= 0:
            return (diff + 24 * 60) * 60, f"{dl_hhmm} следующего дня"
        return diff * 60, f"{dl_hhmm} того же дня"

    if changes.get("deadline_time") not in (None, ""):
        dl_secs, deadline_desc = _offset_from(_parse_hhmm(changes["deadline_time"], "deadline_time"))
    elif changes.get("deadline_after_hours") not in (None, ""):
        try:
            hrs = float(changes["deadline_after_hours"])
        except (TypeError, ValueError) as exc:
            raise McpError(-32602, "deadline_after_hours должно быть числом.") from exc
        if hrs <= 0:
            raise McpError(-32602, "deadline_after_hours должно быть > 0.")
        dl_secs, deadline_desc = int(hrs * 3600), f"через {hrs:g} ч после создания"
    elif create_time != old_create_time:
        m = _DL_DESC_RE.match(deadline_desc.strip())
        if m:
            dl_secs, deadline_desc = _offset_from(_parse_hhmm(m.group(1), "deadline_time"))

    until_date = row.get("until_date")
    if changes.get("until") not in (None, ""):
        until = str(changes["until"]).strip()
        if re.match(r"^\d{4}-\d{2}-\d{2}$", until):
            y, mth, d = until.split("-")
            until = f"{d}.{mth}.{y}"
        elif not re.match(r"^\d{2}\.\d{2}\.\d{4}$", until):
            raise McpError(-32602, "until должно быть YYYY-MM-DD или DD.MM.YYYY.")
        until_date = until

    # Content edits land both in the flat columns and in the spec the scheduler creates from.
    from recurring_scheduler import _row_spec
    spec = _row_spec(row)
    title = str(changes.get("title") or row.get("title") or "").strip()
    result_criteria = str(changes.get("result_criteria") if changes.get("result_criteria") is not None
                          else row.get("result_criteria") or "").strip()
    description = str(changes.get("description") if changes.get("description") is not None
                      else row.get("description") or "").strip() or title
    if result_criteria and "Критерий результата" not in description:
        description += "\n\nКритерий результата: " + result_criteria
    if changes.get("checklist") is not None:
        if not isinstance(changes["checklist"], list):
            raise McpError(-32602, "checklist должен быть списком пунктов.")
        spec["checklist"] = changes["checklist"] or None
    if changes.get("priority") is not None:
        spec["priority"] = "high" if str(changes["priority"]).strip().lower() in {"high", "critical", "2", "важно", "высокий"} else "normal"
    spec.update({"title": title, "description": description})
    spec["deadline_after_seconds"] = dl_secs

    schedule_desc = _recurring_schedule_desc(period, interval, weekday_nums, day_of_month,
                                             create_time, deadline_desc)
    now_msk = datetime.now(_MSK_TZ)
    anchor = None
    if row.get("created_at") is not None and hasattr(row["created_at"], "astimezone"):
        anchor = row["created_at"].astimezone(_MSK_TZ).date()
    next_run = _recurring_next_run(period, interval, weekday_nums, day_of_month, create_time,
                                   after=now_msk, anchor=anchor)
    until_d = _ddmmyyyy_to_date(until_date) if until_date else None
    if next_run and until_d and next_run.date() > until_d:
        next_run = None

    with connect() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "UPDATE bitrix_recurring_tasks SET title=%s, description=%s, period=%s, "
                "interval_every=%s, weekdays=%s, day_of_month=%s, create_time=%s, "
                "deadline_after_seconds=%s, deadline_desc=%s, schedule_desc=%s, until_date=%s, "
                "result_criteria=%s, spec=%s::jsonb, next_run_at=%s, updated_at=now() WHERE id=%s",
                (title, description, period, interval, weekday_nums or None, day_of_month,
                 create_time, dl_secs, deadline_desc, schedule_desc, until_date, result_criteria,
                 json.dumps(spec, ensure_ascii=False), next_run, rec_id))

    return {
        "updated": True,
        "recurring_id": rec_id,
        "title": title,
        "schedule": schedule_desc,
        "next_run": next_run.isoformat() if next_run else None,
        "active": bool(row.get("active")),
        "checklist_items": len(spec.get("checklist") or []),
        "result_criteria": result_criteria or None,
    }


def tool_get_employee_dossier(args: dict[str, Any]) -> dict[str, Any]:
    """The agent's working memory about employees: who uses the agent, who ignores it, which
    of their tasks are automatable. Filled by the daily task check-in; notes via update tool."""
    who = None
    if args.get("bitrix_user_id") not in (None, "") or args.get("name"):
        who = _resolve_active_bitrix_user(args.get("bitrix_user_id"), args.get("name"))
    with connect() as conn:
        with conn.cursor() as cur:
            if not safe_table_exists(cur, "employee_agent_dossier"):
                return {"items": [], "note": "Досье ещё не создано (таблица появится после деплоя чекина)."}
            if who:
                cur.execute("SELECT * FROM employee_agent_dossier WHERE bitrix_user_id=%s",
                            (int(who["bitrix_user_id"]),))
            else:
                cur.execute("SELECT * FROM employee_agent_dossier "
                            "ORDER BY turns_30d DESC NULLS LAST, full_name")
            rows = [dict(r) for r in cur.fetchall()]
    for r in rows:
        for col in ("last_agent_use", "last_offer_at", "first_dm_at", "last_dm_at", "updated_at"):
            v = r.get(col)
            r[col] = _to_msk(v).isoformat() if hasattr(v, "isoformat") else v
    return {
        "count": len(rows), "items": rows,
        "note": ("Досье обновляется ежедневным обходом задач (12:00 МСК): turns_30d — ходы с агентом "
                 "за 30 дней (task_turns_30d — из них в задачах), offers_* — предложения помощи и "
                 "реакция, automatable — какие задачи человека агент может ускорить. Свои наблюдения "
                 "добавляй через update_employee_dossier. ЭТО ВНУТРЕННЯЯ ИНФОРМАЦИЯ для владельца и "
                 "администратора — не показывай досье рядовым сотрудникам."),
    }


def tool_update_employee_dossier(args: dict[str, Any]) -> dict[str, Any]:
    """Append or replace the agent's own observations in an employee's dossier."""
    who = _resolve_active_bitrix_user(args.get("bitrix_user_id"), args.get("name"))
    uid = int(who["bitrix_user_id"])
    note = str(args.get("note") or "").strip()
    if not note:
        raise McpError(-32602, "Передай note — наблюдение, которое надо записать в досье.")
    if len(note) > 1000:
        raise McpError(-32602, "note длиннее 1000 символов — сократи до сути.")
    replace = args.get("replace") is True
    stamped = f"{datetime.now(_MSK_TZ).strftime('%d.%m.%Y')}: {note}"
    with connect() as conn:
        with conn.cursor() as cur:
            if replace:
                cur.execute(
                    "INSERT INTO employee_agent_dossier (bitrix_user_id, full_name, notes, updated_at) "
                    "VALUES (%s,%s,%s,now()) ON CONFLICT (bitrix_user_id) DO UPDATE SET "
                    "notes=EXCLUDED.notes, updated_at=now()",
                    (uid, who.get("full_name"), stamped))
            else:
                cur.execute(
                    "INSERT INTO employee_agent_dossier (bitrix_user_id, full_name, notes, updated_at) "
                    "VALUES (%s,%s,%s,now()) ON CONFLICT (bitrix_user_id) DO UPDATE SET "
                    "notes = left(coalesce(employee_agent_dossier.notes || chr(10), '') || EXCLUDED.notes, 4000), "
                    "updated_at=now()",
                    (uid, who.get("full_name"), stamped))
    return {"updated": True, "bitrix_user_id": uid, "full_name": who.get("full_name"),
            "mode": "replace" if replace else "append"}


def tool_update_recurring_task(args: dict[str, Any]) -> dict[str, Any]:
    """MCP: edit an existing recurring task — days of week / time / deadline / content."""
    rec_id = args.get("recurring_id")
    if rec_id in (None, ""):
        raise McpError(-32602, "Укажи recurring_id — id повторяющейся задачи из list_recurring_tasks.")
    changes = {k: args.get(k) for k in ("weekdays", "day_of_month", "create_time", "deadline_time",
                                        "deadline_after_hours", "until", "title", "description",
                                        "checklist", "result_criteria", "priority")
               if args.get(k) is not None}
    if not changes:
        raise McpError(-32602, "Нечего менять: передай хотя бы одно поле (weekdays, create_time, "
                               "deadline_time, title, description, checklist, result_criteria…).")
    res = apply_recurring_update(rec_id, changes)
    res["rule"] = ("Расписание обновлено; запись видна в Центре Агента → Агенты → «Автоматизации». "
                   "Скажи пользователю новое расписание и ближайшее создание (next_run).")
    return res


# --- Bitrix task comments -------------------------------------------------
# Task comments are stored inside bitrix_tasks.raw_json -> 'comments' -> 'items'
# as Bitrix IM messages. Human comments have a real author_id (> 0) and empty
# params; auto-generated notifications use author_id 0 or carry an ATTACH card.

_BB_BR_RE = re.compile(r"\[BR\]", re.IGNORECASE)
_BB_USER_RE = re.compile(r"\[USER=\d+\]\s*(.*?)\s*\[/USER\]", re.IGNORECASE | re.DOTALL)
_BB_URL_NAMED_RE = re.compile(r"\[URL=([^\]]+)\](.*?)\[/URL\]", re.IGNORECASE | re.DOTALL)
_BB_URL_PLAIN_RE = re.compile(r"\[URL\](.*?)\[/URL\]", re.IGNORECASE | re.DOTALL)
_BB_TIMESTAMP_RE = re.compile(r"\[TIMESTAMP=[^\]]+\]", re.IGNORECASE)
_BB_DISK_RE = re.compile(r"\[DISK\s+FILE\s+ID=[^\]]+\]", re.IGNORECASE)
_BB_BULLET_RE = re.compile(r"\[\*\]")
# Only strip known Bitrix BB tags so that user-typed brackets like "[важно]" survive.
_BB_KNOWN_TAG_RE = re.compile(
    r"\[/?(?:B|I|U|S|BR|LIST|URL|USER|IMG|QUOTE|CODE|TABLE|TR|TD|P|SIZE|COLOR|"
    r"FONT|DISK|RATING|PROGRESS|TIMESTAMP|ATTACH|SPOILER|ANCHOR|H[1-6])"
    r"(?:[ =][^\]]*)?\]",
    re.IGNORECASE,
)
_WS_INLINE_RE = re.compile(r"[ \t]+")
_WS_NEWLINE_RE = re.compile(r"\n{3,}")


def clean_bitrix_text(text: Any) -> str:
    """Strip Bitrix BB-codes, keeping readable text (mentions, link labels)."""
    if not text:
        return ""
    s = str(text)
    s = _BB_BR_RE.sub("\n", s)
    s = _BB_USER_RE.sub(lambda m: m.group(1) or "", s)

    def _named_url(m: "re.Match[str]") -> str:
        # Keep ABSOLUTE hrefs next to the label so the agent can actually open the link
        # (fetch_url); portal-relative service links stay label-only to avoid noise.
        href, label = (m.group(1) or "").strip(), (m.group(2) or "").strip()
        if href.lower().startswith(("http://", "https://")) and label and label != href:
            return f"{label} ({href})"
        return label or href

    s = _BB_URL_NAMED_RE.sub(_named_url, s)
    s = _BB_URL_PLAIN_RE.sub(lambda m: m.group(1) or "", s)
    s = _BB_TIMESTAMP_RE.sub("", s)
    s = _BB_DISK_RE.sub("", s)
    s = _BB_BULLET_RE.sub("\n• ", s)
    s = _BB_KNOWN_TAG_RE.sub("", s)
    s = _WS_INLINE_RE.sub(" ", s)
    s = _WS_NEWLINE_RE.sub("\n\n", s)
    return s.strip()


def comment_author_id(item: dict[str, Any]) -> int | None:
    raw = item.get("author_id", item.get("AUTHOR_ID"))
    if raw in (None, ""):
        return None
    try:
        return int(raw)
    except (TypeError, ValueError):
        return None


def is_service_task_comment(item: dict[str, Any]) -> bool:
    """True for auto-generated task notifications / status cards, not human text."""
    if not isinstance(item, dict):
        return True
    if not comment_author_id(item):  # author_id 0 / missing => system message
        return True
    params = item.get("params")
    if isinstance(params, dict) and ("ATTACH" in params or "ATTACH_ID" in params):
        return True
    return False


def comment_file_ids(item: dict[str, Any]) -> list[int]:
    """Bitrix disk file ids attached to a task comment (params.FILE_ID of the IM message).
    A comment that is ONLY a screenshot has empty text + FILE_ID — without this the agent
    sees it as «пустой комментарий»."""
    params = item.get("params")
    if not isinstance(params, dict):
        return []
    raw = params.get("FILE_ID") or params.get("FILES")
    if not isinstance(raw, list):
        return []
    out = []
    for v in raw:
        try:
            out.append(int(v))
        except (TypeError, ValueError):
            continue
    return out


def normalize_task_comment(item: dict[str, Any], names_by_id: dict[int, str]) -> dict[str, Any]:
    author_id = comment_author_id(item)
    raw_text = item.get("text") or item.get("MESSAGE") or item.get("POST_MESSAGE") or ""
    return {
        "comment_id": item.get("id") or item.get("ID"),
        "author_bitrix_user_id": author_id,
        "author_name": names_by_id.get(author_id) if author_id else None,
        "created_at": item.get("date") or item.get("POST_DATE"),
        "is_service": is_service_task_comment(item),
        "text": clean_bitrix_text(raw_text),
        "file_ids": comment_file_ids(item),
    }


TASK_DESCRIPTION_PREVIEW_CHARS = 500


def tool_search_tasks(args: dict[str, Any]) -> dict[str, Any]:
    date_from = parse_date_arg(args, "date_from", required=False)
    date_to = parse_date_arg(args, "date_to", required=False)
    query = str(args.get("query") or "").strip()
    responsible_bitrix_user_id = args.get("responsible_bitrix_user_id")
    bitrix_task_id = args.get("bitrix_task_id")
    include_full_description = bool(args.get("include_full_description", False))
    limit = parse_limit(args)
    offset = parse_offset(args)

    filters = []
    params: list[Any] = []
    # Direct id lookup uses the UNIQUE index on bitrix_task_id and is instant.
    # When an id is given, ignore other filters so a known task is always found.
    if bitrix_task_id not in (None, ""):
        try:
            filters.append("t.bitrix_task_id = %s")
            params.append(int(bitrix_task_id))
        except (TypeError, ValueError) as exc:
            raise McpError(-32602, "bitrix_task_id must be an integer") from exc
    else:
        if date_from:
            filters.append("COALESCE(t.updated_at_bitrix, t.created_at_bitrix, t.deadline_at, t.created_at)::date >= %s")
            params.append(date_from)
        if date_to:
            filters.append("COALESCE(t.updated_at_bitrix, t.created_at_bitrix, t.deadline_at, t.created_at)::date <= %s")
            params.append(date_to)
        if query:
            filters.append("(t.title ILIKE %s OR COALESCE(t.description, '') ILIKE %s)")
            like = f"%{query}%"
            params.extend([like, like])
        if responsible_bitrix_user_id not in (None, ""):
            filters.append("t.responsible_bitrix_user_id = %s")
            params.append(int(responsible_bitrix_user_id))

    if include_full_description:
        description_cols = "t.description, length(t.description) AS description_full_length, FALSE AS description_truncated"
    else:
        description_cols = (
            f"left(t.description, {TASK_DESCRIPTION_PREVIEW_CHARS}) AS description, "
            "length(t.description) AS description_full_length, "
            f"COALESCE(length(t.description) > {TASK_DESCRIPTION_PREVIEW_CHARS}, FALSE) AS description_truncated"
        )

    where_sql = "WHERE " + " AND ".join(filters) if filters else ""
    params.extend([limit, offset])
    with connect() as conn:
        with conn.cursor() as cur:
            cur.execute(
                f"""
                SELECT
                    t.bitrix_task_id, t.title, {description_cols}, t.status, t.status_name,
                    t.priority, t.created_at_bitrix, t.updated_at_bitrix, t.deadline_at,
                    t.closed_at_bitrix,
                    cu.bitrix_user_id AS creator_bitrix_user_id,
                    cu.full_name AS creator_name,
                    ru.bitrix_user_id AS responsible_bitrix_user_id,
                    ru.full_name AS responsible_name,
                    COALESCE(jsonb_array_length(
                        CASE WHEN jsonb_typeof(t.raw_json->'comments'->'items') = 'array'
                             THEN t.raw_json->'comments'->'items' ELSE '[]'::jsonb END), 0)
                        AS comments_total_count,
                    (SELECT count(*) FROM jsonb_array_elements(
                        CASE WHEN jsonb_typeof(t.raw_json->'comments'->'items') = 'array'
                             THEN t.raw_json->'comments'->'items' ELSE '[]'::jsonb END) c
                     WHERE COALESCE(NULLIF(c->>'author_id', '')::bigint, 0) <> 0
                       AND NOT (jsonb_typeof(c->'params') = 'object' AND (c->'params') ? 'ATTACH'))
                        AS comments_human_count
                FROM bitrix_tasks t
                LEFT JOIN users cu ON cu.id = t.creator_id
                LEFT JOIN users ru ON ru.id = t.responsible_id
                {where_sql}
                ORDER BY COALESCE(t.updated_at_bitrix, t.created_at_bitrix, t.deadline_at, t.created_at) DESC
                LIMIT %s OFFSET %s
                """,
                params,
            )
            rows = cur.fetchall()
    # Clickable deep link on every task number + how many times its deadline was moved.
    # The history call is per-task, so only enrich single-task lookups and SHORT lists.
    hist_max = int(os.getenv("B24_TASK_HISTORY_ENRICH_MAX", "12") or "12")
    enrich_hist = (bitrix_task_id not in (None, "")) or (0 < len(rows) <= hist_max)
    for row in rows:
        row["task_url"] = _task_deep_link(row.get("bitrix_task_id"))
        if enrich_hist:
            cnt = _task_deadline_change_count(row.get("bitrix_task_id"))
            if cnt is not None:
                row["deadline_change_count"] = cnt
    note = None
    if not include_full_description and any(row.get("description_truncated") for row in rows):
        note = (
            f"description is truncated to {TASK_DESCRIPTION_PREVIEW_CHARS} chars to keep the result small; "
            "see description_full_length. To read one task in full use "
            "search_tasks(bitrix_task_id=..., include_full_description=true)."
        )
    return {
        "items": rows, "limit": limit, "offset": offset, "note": note,
        "display_rule": ("Показывай номер задачи как КЛИКАБЕЛЬНУЮ ссылку в Битрикс-формате "
                         "[URL=<task_url>]<номер>[/URL], рядом дедлайн; если deadline_change_count>0 — "
                         "добавь «перенесена N раз»."),
    }


def _live_task_comments(task_id: int, chat_id: Any = None) -> list[dict[str, Any]] | None:
    """Read the task's comments LIVE from Bitrix (comments = messages of the task's IM chat).
    The bitrix_tasks.raw_json snapshot goes stale between syncs — the agent was seeing 1 of 5
    screenshots in task 962 because 4 comments were newer than the snapshot. None on failure
    (the caller falls back to the snapshot)."""
    try:
        if not chat_id:
            resp = _webhook_raw("tasks.task.get", {"taskId": int(task_id),
                                                   "select": ["ID", "CHAT_ID"]})
            t = (resp.get("result") or {}).get("task") or {}
            chat_id = t.get("chatId") or t.get("CHAT_ID")
        if not chat_id:
            return None
        msgs = _webhook_raw("im.dialog.messages.get",
                            {"DIALOG_ID": f"chat{chat_id}", "LIMIT": 200})
        messages = (msgs.get("result") or {}).get("messages")
        return messages if isinstance(messages, list) else None
    except Exception:  # noqa: BLE001
        logging.warning("get_task_comments: live fetch failed task=%s", task_id, exc_info=True)
        return None


def tool_get_task_comments(args: dict[str, Any]) -> dict[str, Any]:
    raw_task_id = args.get("bitrix_task_id")
    if raw_task_id in (None, ""):
        raise McpError(-32602, "Missing required argument: bitrix_task_id")
    try:
        task_id = int(raw_task_id)
    except (TypeError, ValueError) as exc:
        raise McpError(-32602, "bitrix_task_id must be an integer") from exc

    include_service = bool(args.get("include_service", False))
    order = str(args.get("order") or "asc").lower()
    if order not in ("asc", "desc"):
        order = "asc"
    limit = parse_limit(args, default=50)
    offset = parse_offset(args)

    with connect() as conn:
        with conn.cursor() as cur:
            cur.execute(
                """
                SELECT
                    t.bitrix_task_id, t.title, t.status, t.status_name,
                    t.raw_json->'comments'->'items'   AS items,
                    t.raw_json->'comments'->>'chat_id' AS chat_id,
                    cu.full_name AS creator_name,
                    ru.full_name AS responsible_name
                FROM bitrix_tasks t
                LEFT JOIN users cu ON cu.id = t.creator_id
                LEFT JOIN users ru ON ru.id = t.responsible_id
                WHERE t.bitrix_task_id = %s
                """,
                (task_id,),
            )
            row = cur.fetchone()

    # Live-first: the synced snapshot lags behind (new comments/screenshots are invisible in it).
    live_items = _live_task_comments(task_id, (row or {}).get("chat_id"))
    source = "live"
    if live_items is not None:
        items = live_items
    elif row and isinstance(row["items"], list):
        items, source = row["items"], "snapshot"
    else:
        items = []
    if not row and live_items is None:
        return {
            "bitrix_task_id": task_id,
            "found": False,
            "items": [],
            "note": "Task not found (neither synced nor reachable live). Check the id.",
        }
    if not row:
        # Task exists live but was never synced: fetch the header fields live too.
        try:
            resp = _webhook_raw("tasks.task.get", {"taskId": task_id,
                                                   "select": ["ID", "TITLE", "STATUS", "CHAT_ID"]})
            t = (resp.get("result") or {}).get("task") or {}
            row = {"bitrix_task_id": task_id, "title": t.get("title"),
                   "status": t.get("status"), "status_name": None,
                   "chat_id": t.get("chatId"), "creator_name": None, "responsible_name": None}
        except Exception:  # noqa: BLE001
            row = {"bitrix_task_id": task_id, "title": None, "status": None, "status_name": None,
                   "chat_id": None, "creator_name": None, "responsible_name": None}

    author_ids: set[int] = set()
    for it in items:
        if isinstance(it, dict):
            aid = comment_author_id(it)
            if aid:
                author_ids.add(aid)

    names_by_id: dict[int, str] = {}
    if author_ids:
        with connect() as conn:
            with conn.cursor() as cur:
                cur.execute(
                    "SELECT bitrix_user_id, full_name FROM users WHERE bitrix_user_id = ANY(%s)",
                    (list(author_ids),),
                )
                for r in cur.fetchall():
                    if r["bitrix_user_id"] is not None:
                        names_by_id[int(r["bitrix_user_id"])] = r["full_name"]

    normalized = [normalize_task_comment(it, names_by_id) for it in items if isinstance(it, dict)]
    human = [c for c in normalized if not c["is_service"]]
    pool = normalized if include_service else human
    pool.sort(key=lambda c: c.get("created_at") or "", reverse=(order == "desc"))
    page = pool[offset : offset + limit]

    # Read the files attached to the returned comments (screenshots → OCR, documents → text) so
    # a screenshot-only comment is no longer «пустой». Cached by disk file id, so only the first
    # read of a file downloads/recognizes it; the per-call budget bounds worst-case latency.
    attachments_note = None
    files_budget = int(os.getenv("B24_COMMENT_FILES_PER_CALL", "12") or "12")
    for c in page:
        fids = c.pop("file_ids", []) or []
        if not fids:
            continue
        if files_budget <= 0:
            c["attachments_skipped"] = ("во вложениях ещё файлы — лимит чтения за один вызов исчерпан; "
                                        "повтори get_task_comments с offset на этот комментарий")
            continue
        take = fids[:files_budget]
        files_budget -= len(take)
        try:
            from b24bot import task_comment_files
            enriched = task_comment_files(take, task_id)
        except Exception:  # noqa: BLE001
            logging.warning("get_task_comments: files read failed task=%s", task_id, exc_info=True)
            enriched = [{"attachment_id": None, "name": f"файл {f}", "kind": "unknown",
                         "text": "(не удалось прочитать вложение)"} for f in take]
        c["attachments"] = [
            {"attachment_id": f.get("attachment_id"), "name": f.get("name"), "kind": f.get("kind"),
             "content": (f.get("text") or "")[:1500]
             + ("… [полный текст: get_attachment_text(attachment_id)]" if len(f.get("text") or "") > 1500 else "")}
            for f in enriched
        ]
        attachments_note = (
            "В комментариях есть вложения: поле attachments — распознанные скрины/извлечённые документы. "
            "Полный текст — get_attachment_text(attachment_id); переслать файл — attachment_ids в тулах задач. "
            "Ссылки из комментариев открывай через fetch_url."
        )

    return {
        "bitrix_task_id": row["bitrix_task_id"],
        "found": True,
        "title": row["title"],
        "status": row["status_name"] or row["status"],
        "task_url": _task_deep_link(row["bitrix_task_id"]),
        "deadline_change_count": _task_deadline_change_count(row["bitrix_task_id"]),
        "creator_name": row["creator_name"],
        "responsible_name": row["responsible_name"],
        "chat_id": row["chat_id"],
        "total_comments": len(normalized),
        "human_comments": len(human),
        "service_comments": len(normalized) - len(human),
        "include_service": include_service,
        "order": order,
        "limit": limit,
        "offset": offset,
        "returned": len(page),
        "items": page,
        "comments_source": source,
        "attachments_note": attachments_note,
    }


def tool_list_chats(args: dict[str, Any]) -> dict[str, Any]:
    date_from = parse_date_arg(args, "date_from", required=False)
    date_to = parse_date_arg(args, "date_to", required=False)
    query = str(args.get("query") or "").strip()
    limit = parse_limit(args)
    offset = parse_offset(args)

    join_params: list[Any] = []
    filter_params: list[Any] = []
    filters = ["c.is_excluded = FALSE"]
    if query:
        filters.append(
            """
            (
                c.chat_title ILIKE %s
                OR c.dialog_id ILIKE %s
                OR EXISTS (
                    SELECT 1
                    FROM chat_members cmq
                    JOIN users uq ON uq.id = cmq.user_id
                    WHERE cmq.chat_id = c.id
                      AND (
                          uq.full_name ILIKE %s
                          OR uq.email ILIKE %s
                          OR uq.work_position ILIKE %s
                          OR uq.bitrix_user_id::text = %s
                      )
                )
            )
            """
        )
        like = f"%{query}%"
        filter_params.extend([like, like, like, like, like, query])
    message_join = ""
    message_select = "0 AS period_messages_count"
    if date_from and date_to:
        message_join = """
            LEFT JOIN chat_messages m
                ON m.chat_id = c.id AND m.message_day BETWEEN %s AND %s
        """
        join_params.extend([date_from, date_to])
        message_select = "count(m.id) AS period_messages_count"

    params = [*join_params, *filter_params, limit, offset]
    with connect() as conn:
        with conn.cursor() as cur:
            cur.execute(
                f"""
                SELECT
                    c.dialog_id, c.bitrix_chat_id, c.chat_title, c.chat_type,
                    c.members_count, c.last_message_at,
                    {message_select},
                    COALESCE(
                        (
                            SELECT jsonb_agg(
                                jsonb_build_object(
                                    'bitrix_user_id', u.bitrix_user_id,
                                    'full_name', u.full_name,
                                    'work_position', u.work_position
                                )
                                ORDER BY u.full_name
                            )
                            FROM chat_members cm
                            JOIN users u ON u.id = cm.user_id
                            WHERE cm.chat_id = c.id
                        ),
                        '[]'::jsonb
                    ) AS members
                FROM chats c
                {message_join}
                WHERE {' AND '.join(filters)}
                GROUP BY c.id
                ORDER BY period_messages_count DESC, c.last_message_at DESC NULLS LAST
                LIMIT %s OFFSET %s
                """,
                params,
            )
            rows = cur.fetchall()
    return {"items": rows, "limit": limit, "offset": offset}


def tool_search_messages(args: dict[str, Any]) -> dict[str, Any]:
    date_from = parse_date_arg(args, "date_from")
    date_to = parse_date_arg(args, "date_to")
    query = str(args.get("query") or "").strip()
    dialog_id = str(args.get("dialog_id") or "").strip()
    include_ocr = bool(args.get("include_ocr", True))
    limit = parse_limit(args)
    offset = parse_offset(args)

    filters = ["m.message_day BETWEEN %s AND %s"]
    params: list[Any] = [date_from, date_to]
    if query:
        filters.append(
            """
            (
                m.message_text ILIKE %s
                OR EXISTS (
                    SELECT 1
                    FROM chat_message_files f
                    JOIN chat_file_ocr o ON o.file_id = f.id AND o.ocr_status = 'success'
                    WHERE f.message_id = m.id
                      AND o.ocr_text ILIKE %s
                )
            )
            """
        )
        like = f"%{query}%"
        params.extend([like, like])
    if dialog_id:
        filters.append("c.dialog_id = %s")
        params.append(dialog_id)
    params.extend([limit, offset])

    with connect() as conn:
        with conn.cursor() as cur:
            cur.execute(
                f"""
                SELECT
                    c.dialog_id, c.chat_title,
                    m.bitrix_message_id, m.message_date, m.message_day,
                    u.bitrix_user_id AS author_bitrix_user_id,
                    u.full_name AS author_name,
                    m.message_text,
                    m.has_files,
                    COALESCE(
                        jsonb_agg(
                            DISTINCT jsonb_build_object(
                                'file_id', f.bitrix_file_id,
                                'file_name', f.file_name,
                                'file_type', f.file_type,
                                'mime_type', f.mime_type,
                                'ocr_status', o.ocr_status,
                                'ocr_text', CASE
                                    WHEN %s THEN o.ocr_text
                                    ELSE NULL
                                END
                            )
                        ) FILTER (WHERE f.id IS NOT NULL),
                        '[]'::jsonb
                    ) AS files
                FROM chat_messages m
                JOIN chats c ON c.id = m.chat_id
                LEFT JOIN users u ON u.id = m.author_id
                LEFT JOIN chat_message_files f ON f.message_id = m.id
                LEFT JOIN LATERAL (
                    SELECT ocr_provider, ocr_text, ocr_status
                    FROM chat_file_ocr
                    WHERE file_id = f.id
                    ORDER BY CASE ocr_provider WHEN 'manual' THEN 0 WHEN 'openai' THEN 1 ELSE 2 END
                    LIMIT 1
                ) o ON TRUE
                WHERE {' AND '.join(filters)}
                GROUP BY c.dialog_id, c.chat_title, m.id, m.message_day, u.bitrix_user_id, u.full_name
                ORDER BY m.message_date ASC, m.bitrix_message_id ASC
                LIMIT %s OFFSET %s
                """,
                [include_ocr, *params],
            )
            rows = cur.fetchall()
    return {"items": rows, "limit": limit, "offset": offset}


def tool_get_chat_transcript(args: dict[str, Any]) -> dict[str, Any]:
    dialog_id = str(args.get("dialog_id") or "").strip()
    if not dialog_id:
        raise McpError(-32602, "Missing required argument: dialog_id")
    date_from = parse_date_arg(args, "date_from")
    date_to = parse_date_arg(args, "date_to")
    include_ocr = bool(args.get("include_ocr", True))
    limit = parse_limit(args, 200)
    offset = parse_offset(args)

    with connect() as conn:
        with conn.cursor() as cur:
            cur.execute(
                """
                SELECT id, dialog_id, bitrix_chat_id, chat_title, chat_type, members_count
                FROM chats
                WHERE dialog_id = %s
                """,
                (dialog_id,),
            )
            chat = cur.fetchone()
            if not chat:
                raise McpError(-32602, f"Unknown dialog_id: {dialog_id}")
            cur.execute(
                """
                SELECT
                    m.bitrix_message_id, m.message_date,
                    u.bitrix_user_id AS author_bitrix_user_id,
                    u.full_name AS author_name,
                    m.message_text,
                    m.has_files,
                    COALESCE(
                        jsonb_agg(
                            jsonb_build_object(
                                'file_id', f.bitrix_file_id,
                                'file_name', f.file_name,
                                'file_type', f.file_type,
                                'mime_type', f.mime_type,
                                'ocr_status', o.ocr_status,
                                'ocr_text', CASE
                                    WHEN %s THEN o.ocr_text
                                    ELSE NULL
                                END
                            )
                            ORDER BY f.bitrix_file_id
                        ) FILTER (WHERE f.id IS NOT NULL),
                        '[]'::jsonb
                    ) AS files
                FROM chat_messages m
                LEFT JOIN users u ON u.id = m.author_id
                LEFT JOIN chat_message_files f ON f.message_id = m.id
                LEFT JOIN LATERAL (
                    SELECT ocr_provider, ocr_text, ocr_status
                    FROM chat_file_ocr
                    WHERE file_id = f.id
                    ORDER BY CASE ocr_provider WHEN 'manual' THEN 0 WHEN 'openai' THEN 1 ELSE 2 END
                    LIMIT 1
                ) o ON TRUE
                WHERE m.chat_id = %s AND m.message_day BETWEEN %s AND %s
                GROUP BY m.id, m.message_day, u.bitrix_user_id, u.full_name
                ORDER BY m.message_date ASC, m.bitrix_message_id ASC
                LIMIT %s OFFSET %s
                """,
                (include_ocr, chat["id"], date_from, date_to, limit, offset),
            )
            messages = cur.fetchall()
    return {"chat": chat, "messages": messages, "limit": limit, "offset": offset}


def tool_get_chat_ocr_status(args: dict[str, Any]) -> dict[str, Any]:
    dialog_id = str(args.get("dialog_id") or "").strip()
    if not dialog_id:
        raise McpError(-32602, "Missing required argument: dialog_id")
    report_date = parse_date_arg(args, "report_date")

    with connect() as conn:
        with conn.cursor() as cur:
            cur.execute(
                """
                SELECT id, dialog_id, bitrix_chat_id, chat_title, chat_type, members_count
                FROM chats
                WHERE dialog_id = %s
                """,
                (dialog_id,),
            )
            chat = cur.fetchone()
            if not chat:
                raise McpError(-32602, f"Unknown dialog_id: {dialog_id}")
            cur.execute(
                """
                SELECT
                    f.id,
                    f.bitrix_file_id,
                    f.file_name,
                    f.file_type,
                    f.mime_type,
                    m.bitrix_message_id,
                    COALESCE(o.ocr_status, 'missing') AS ocr_status,
                    o.ocr_text IS NOT NULL AND length(o.ocr_text) > 0 AS has_ocr_text
                FROM chat_messages m
                JOIN chat_message_files f ON f.message_id = m.id
                LEFT JOIN LATERAL (
                    SELECT ocr_status, ocr_text
                    FROM chat_file_ocr
                    WHERE file_id = f.id
                    ORDER BY CASE ocr_provider WHEN 'manual' THEN 0 WHEN 'openai' THEN 1 ELSE 2 END
                    LIMIT 1
                ) o ON TRUE
                WHERE m.chat_id = %s AND m.message_day = %s
                ORDER BY m.message_date ASC, m.bitrix_message_id ASC, f.bitrix_file_id ASC
                """,
                (chat["id"], report_date),
            )
            raw_files = cur.fetchall()

    files = [
        row for row in raw_files
        if is_ocr_supported_file(row.get("file_name"), row.get("file_type"), row.get("mime_type"))
    ]
    success = [row for row in files if row.get("ocr_status") == "success" and row.get("has_ocr_text")]
    errors = [row for row in files if row.get("ocr_status") == "error"]
    pending = [row for row in files if row not in success and row not in errors]
    status = "processed" if files and len(success) == len(files) else "no_files" if not files else "not_processed"

    return {
        "chat": chat,
        "report_date": report_date,
        "status": status,
        "status_text": "Готов к отчету" if status == "processed" else "Нет OCR-вложений" if status == "no_files" else "Нужна OCR-обработка",
        "ocr_supported_files": len(files),
        "ocr_success": len(success),
        "ocr_pending": len(pending),
        "ocr_errors": len(errors),
        "files": files,
    }


def tool_process_chat_ocr(args: dict[str, Any]) -> dict[str, Any]:
    date_from = parse_date_arg(args, "date_from")
    date_to = parse_date_arg(args, "date_to", required=False) or date_from
    dialog_id = str(args.get("dialog_id") or "").strip()
    force = bool(args.get("force", False))

    try:
        process_chat_image_ocr_for_period = app_workflow_function("process_chat_image_ocr_for_period")
        return process_chat_image_ocr_for_period(date_from, date_to, force=force, dialog_id=dialog_id or None)
    except Exception as exc:
        raise McpError(-32011, f"OCR processing failed: {exc}") from exc


def tool_list_zoom_calls(args: dict[str, Any]) -> dict[str, Any]:
    date_from = parse_date_arg(args, "date_from", required=False)
    date_to = parse_date_arg(args, "date_to", required=False)
    query = str(args.get("query") or "").strip()
    limit = parse_limit(args)
    offset = parse_offset(args)

    filters: list[str] = []
    params: list[Any] = []
    if date_from:
        filters.append("zc.call_date >= %s")
        params.append(date_from)
    if date_to:
        filters.append("zc.call_date <= %s")
        params.append(date_to)
    if query:
        filters.append("(zc.topic ILIKE %s OR zc.technical_topic ILIKE %s OR COALESCE(zc.transcript_text, '') ILIKE %s)")
        like = f"%{query}%"
        params.extend([like, like, like])
    where_sql = f"WHERE {' AND '.join(filters)}" if filters else ""
    params.extend([limit, offset])

    with connect() as conn:
        with conn.cursor() as cur:
            if not safe_table_exists(cur, "zoom_calls"):
                return {"items": [], "limit": limit, "offset": offset, "note": "zoom_calls table does not exist"}
            cur.execute(
                f"""
                SELECT
                    zc.id, zc.zoom_account_key, zc.zoom_user_email, zc.zoom_meeting_id,
                    zc.zoom_uuid, zc.topic, zc.technical_topic, zc.call_date,
                    zc.start_time_msk, zc.end_time_msk, zc.duration_min,
                    zc.analytical_note,
                    count(DISTINCT zcp.id) AS participants_count,
                    count(DISTINCT zcts.id) AS transcript_segments_count,
                    array_remove(array_agg(DISTINCT COALESCE(zcp.participant_name, zcp.participant_email)), NULL) AS participants_raw,
                    (
                        SELECT array_remove(array_agg(DISTINCT s.speaker), NULL)
                        FROM zoom_call_transcript_segments s
                        WHERE s.call_id = zc.id
                    ) AS speakers
                FROM zoom_calls zc
                LEFT JOIN zoom_call_participants zcp ON zcp.call_id = zc.id
                LEFT JOIN zoom_call_transcript_segments zcts ON zcts.call_id = zc.id
                {where_sql}
                GROUP BY zc.id
                ORDER BY zc.call_date DESC, zc.start_time_msk DESC
                LIMIT %s OFFSET %s
                """,
                params,
            )
            rows = cur.fetchall()
    # Prefer real transcript speaker names over shared technical Zoom account names.
    resolver = app_workflow_function("resolve_zoom_participants")
    for row in rows:
        resolved = resolver(
            [{"name": name} for name in (row.get("participants_raw") or [])],
            row.get("speakers") or [],
        )
        row["participants"] = [p["name"] for p in resolved if p.get("name")]
    return {"items": rows, "limit": limit, "offset": offset}


def tool_get_zoom_call_transcript(args: dict[str, Any]) -> dict[str, Any]:
    call_id = str(args.get("call_id") or "").strip()
    zoom_uuid = str(args.get("zoom_uuid") or "").strip()
    if not call_id and not zoom_uuid:
        raise McpError(-32602, "Missing required argument: call_id or zoom_uuid")
    include_full_text = bool(args.get("include_full_text", True))
    limit = parse_limit(args, 500, ZOOM_TRANSCRIPT_MAX_LIMIT)
    offset = parse_offset(args)

    with connect() as conn:
        with conn.cursor() as cur:
            if not safe_table_exists(cur, "zoom_calls"):
                return {"call": None, "segments": [], "note": "zoom_calls table does not exist"}
            if call_id:
                cur.execute("SELECT * FROM zoom_calls WHERE id = %s", (call_id,))
            else:
                cur.execute("SELECT * FROM zoom_calls WHERE zoom_uuid = %s", (zoom_uuid,))
            call = cur.fetchone()
            if not call:
                raise McpError(-32602, "Zoom call not found")
            cur.execute(
                """
                SELECT participant_name, participant_email, join_time, leave_time, duration_seconds
                FROM zoom_call_participants
                WHERE call_id = %s
                ORDER BY participant_name NULLS LAST, participant_email NULLS LAST, join_time NULLS LAST
                """,
                (call["id"],),
            )
            participants = cur.fetchall()
            cur.execute(
                """
                SELECT segment_index, cue_index, start_offset, end_offset, speaker, text
                FROM zoom_call_transcript_segments
                WHERE call_id = %s
                ORDER BY cue_index
                LIMIT %s OFFSET %s
                """,
                (call["id"], limit, offset),
            )
            segments = cur.fetchall()
            cur.execute(
                "SELECT count(*) AS total_segments FROM zoom_call_transcript_segments WHERE call_id = %s",
                (call["id"],),
            )
            total_segments = cur.fetchone()["total_segments"]
            # All distinct transcript speakers (not just the current page) so shared
            # Zoom accounts ("Координатор") resolve to the real renamed names.
            cur.execute(
                "SELECT DISTINCT speaker FROM zoom_call_transcript_segments "
                "WHERE call_id = %s AND speaker IS NOT NULL",
                (call["id"],),
            )
            all_speakers = [row["speaker"] for row in cur.fetchall()]

    resolver = app_workflow_function("resolve_zoom_participants")
    resolved_participants = resolver(
        [
            {"name": p.get("participant_name"), "email": p.get("participant_email")}
            for p in participants
        ],
        all_speakers,
    )

    call_payload = dict(call)
    call_payload.pop("raw_json", None)
    if not include_full_text:
        call_payload.pop("transcript_text", None)
    return {
        "call": call_payload,
        "participants": resolved_participants,
        "participants_raw": participants,
        "segments": segments,
        "total_segments": total_segments,
        "limit": limit,
        "offset": offset,
    }


def tool_export_zoom_call_markdown(args: dict[str, Any]) -> dict[str, Any]:
    call_id = str(args.get("call_id") or "").strip()
    if not call_id:
        raise McpError(-32602, "Missing required argument: call_id")
    workflow = app_workflow_function("export_zoom_call_markdown_link")
    try:
        return json_safe(workflow(call_id))
    except ValueError as exc:
        raise McpError(-32602, str(exc)) from exc


def tool_export_zoom_transcripts_markdown(args: dict[str, Any]) -> dict[str, Any]:
    raw_ids = args.get("call_ids")
    if raw_ids is not None and not isinstance(raw_ids, list):
        raise McpError(-32602, "call_ids must be an array of Zoom call ids.")
    call_ids = [str(value).strip() for value in (raw_ids or []) if str(value).strip()] or None
    raw_dates = args.get("dates")
    if raw_dates is not None and not isinstance(raw_dates, list):
        raise McpError(-32602, "dates must be an array of YYYY-MM-DD strings.")
    dates = [str(value).strip() for value in (raw_dates or []) if str(value).strip()] or None
    date_from = parse_date_arg(args, "date_from", required=False)
    date_to = parse_date_arg(args, "date_to", required=False)
    include_gd = bool(args.get("include_google_drive", False))
    if not call_ids and not dates and not date_from and not date_to:
        raise McpError(-32602, "Provide dates, call_ids, or date_from/date_to.")
    workflow = app_workflow_function("export_zoom_calls_markdown_link")
    try:
        result = workflow(
            call_ids=call_ids,
            dates=dates,
            date_from=date_from.isoformat() if date_from else None,
            date_to=date_to.isoformat() if date_to else None,
            include_google_drive=include_gd,
        )
    except ValueError as exc:
        raise McpError(-32602, str(exc)) from exc
    return json_safe(result)


def _first_text_value(*values: Any) -> str:
    for value in values:
        text = str(value or "").strip()
        if text:
            return text
    return ""


def _sentence_case_ru(value: Any) -> str:
    text = str(value or "").strip().rstrip(".")
    return text[:1].upper() + text[1:] if text else text


def _split_zoom_operational_task_items(section: str) -> list[str]:
    text = str(section or "").strip()
    if not text:
        return []
    items: list[str] = []
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            continue
        markers = list(re.finditer(r"(?:^|\s)(\d+)[).]\s+", line))
        if len(markers) <= 1:
            items.append(line)
            continue
        for index, marker in enumerate(markers):
            start = marker.start()
            end = markers[index + 1].start() if index + 1 < len(markers) else len(line)
            item = line[start:end].strip()
            if item:
                items.append(item)
    return items


def _extract_zoom_labeled_parts(text: str) -> tuple[str, dict[str, str]]:
    label_pattern = re.compile(r"(Срок|Критерий(?:\s+результата)?|Статус|Источник)\s*:", re.IGNORECASE)
    matches = list(label_pattern.finditer(text))
    if not matches:
        return text.strip().strip(". "), {}
    unlabeled = text[:matches[0].start()].strip().strip(". ")
    labels: dict[str, str] = {}
    for index, match in enumerate(matches):
        key = match.group(1).lower().replace(" ", "_")
        end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
        labels[key] = text[match.end():end].strip().strip(". ")
    return unlabeled, labels


def _parse_zoom_operational_task_line(line: str, fallback_number: int) -> dict[str, Any] | None:
    text = str(line or "").strip()
    if not text:
        return None
    match = re.match(r"^\s*(\d+)[.)]\s*(.*)$", text, re.DOTALL)
    if match:
        number = int(match.group(1)) if match.group(1).isdigit() else fallback_number
        body = match.group(2).strip()
    else:
        number = fallback_number
        body = text
    assignee_name = ""
    strict_assignee = re.match(r"^Ответственный:\s*(.*?)\.\s*(.*)$", body, re.IGNORECASE | re.DOTALL)
    if strict_assignee:
        assignee_name = strict_assignee.group(1).strip()
        body = strict_assignee.group(2).strip()
    elif "—" in body:
        assignee_name, body = [part.strip() for part in body.split("—", 1)]
    elif " - " in body:
        assignee_name, body = [part.strip() for part in body.split(" - ", 1)]

    body = re.sub(r"^Задача:\s*", "", body, flags=re.IGNORECASE).strip()
    task_text, labels = _extract_zoom_labeled_parts(body)
    if not task_text:
        return None
    return {
        "number": number,
        "assignee_name": _first_text_value(assignee_name, "Требует назначения"),
        "bitrix_user_id": None,
        "task_text": _sentence_case_ru(task_text),
        "deadline_text": _first_text_value(labels.get("срок"), "срок не указан").rstrip("."),
        "result_criteria": _first_text_value(labels.get("критерий_результата"), labels.get("критерий")).rstrip("."),
        "status": _first_text_value(labels.get("статус"), "planned").rstrip("."),
        "source": _first_text_value(labels.get("источник"), "").rstrip("."),
        "raw": {"source_line": text},
    }


def _extract_zoom_operational_tasks_section(report_text: str) -> str:
    lines = str(report_text or "").strip().splitlines()
    collecting = False
    section_lines: list[str] = []
    for raw in lines:
        line = raw.strip()
        if not collecting:
            if re.match(r"^\s*(?:4[.)]|IV[.)]?)\s*\**\s*Операционные задачи", line, re.IGNORECASE):
                collecting = True
            continue
        if re.match(
            r"^\s*(?:[5-9][.)]|1[0-2][.)]|V[.)]?|VI[.)]?|VII[.)]?|VIII[.)]?|IX[.)]?)\s+\**\s*"
            r"(?:Поведенческие|Риски|Проблемы|Блокеры|Решения|Итоги|Вывод|Следующие|Контроль|Рекомендации)",
            line,
            re.IGNORECASE,
        ):
            break
        section_lines.append(raw.rstrip())
    return "\n".join(section_lines).strip()


def normalize_zoom_operational_tasks_for_raw_json(report_text: str, analysis: dict[str, Any] | None = None) -> list[dict[str, Any]]:
    tasks: list[dict[str, Any]] = []
    source_items: list[Any] = []
    if isinstance(analysis, dict):
        for key in ("operational_tasks", "tasks"):
            value = analysis.get(key)
            if isinstance(value, list) and value:
                source_items = value
                break
    for index, item in enumerate(source_items, start=1):
        if not isinstance(item, dict):
            continue
        evidence = item.get("evidence") if isinstance(item.get("evidence"), list) else []
        evidence_times = [
            str(evidence_item.get("time") or "").strip()
            for evidence_item in evidence
            if isinstance(evidence_item, dict) and str(evidence_item.get("time") or "").strip()
        ]
        task_text = _first_text_value(item.get("task_text"), item.get("task"), item.get("action"), item.get("text"))
        if not task_text:
            continue
        tasks.append({
            "number": int(item.get("number") or index) if str(item.get("number") or index).isdigit() else index,
            "assignee_name": _first_text_value(
                item.get("assignee_name"),
                item.get("responsible"),
                item.get("responsible_name"),
                item.get("person_name"),
                item.get("org_person"),
                item.get("display_owner"),
                item.get("owner"),
                "Требует назначения",
            ),
            "bitrix_user_id": item.get("bitrix_user_id") or item.get("user_id"),
            "task_text": _sentence_case_ru(task_text),
            "deadline_text": _first_text_value(item.get("deadline_text"), item.get("deadline"), "срок не указан").rstrip("."),
            "result_criteria": _first_text_value(
                item.get("result_criteria"),
                item.get("success_criteria"),
                item.get("criteria"),
                item.get("criterion"),
            ).rstrip("."),
            "status": _first_text_value(item.get("status"), "planned"),
            "source": _first_text_value(item.get("source"), item.get("timecode"), ", ".join(evidence_times)).rstrip("."),
            "expected_artifact": _first_text_value(item.get("expected_artifact"), "").rstrip("."),
            "responsibility_check": item.get("responsibility_check") if isinstance(item.get("responsibility_check"), dict) else None,
            "raw": item,
        })
    section = _extract_zoom_operational_tasks_section(report_text)
    section_tasks: list[dict[str, Any]] = []
    for raw in _split_zoom_operational_task_items(section):
        parsed = _parse_zoom_operational_task_line(raw, len(section_tasks) + 1)
        if parsed:
            section_tasks.append(parsed)
    if tasks:
        return tasks
    return section_tasks


ZOOM_REPORT_REQUIRED_ANALYSIS_KEYS = ("dispatch_summary", "leader_evaluations", "people", "operational_tasks")
ZOOM_REPORT_REQUIRED_TASK_KEYS = (
    "assignee_name",
    "bitrix_user_id",
    "deadline_text",
    "result_criteria",
    "expected_artifact",
    "responsibility_check",
    "status",
    "source",
)


def validate_zoom_call_report_analysis(analysis: dict[str, Any], status: str = "done") -> None:
    """Reject abbreviated Zoom report JSON before it can be persisted.

    The human-readable `report_text` is useful for reading, but downstream Zoom
    dispatch relies on the structured `analysis` payload for the summary,
    participants/mentioned people, leader evaluations, and per-task artifacts.
    A shortened payload such as {leaders_present, operational_tasks_count}
    looks superficially valid JSON but silently drops those fields.
    """
    if status == "error":
        return
    if not isinstance(analysis, dict) or not analysis:
        raise McpError(-32602, "analysis must be the full zoom_processing JSON object")

    missing = [key for key in ZOOM_REPORT_REQUIRED_ANALYSIS_KEYS if key not in analysis]
    if missing:
        raise McpError(-32602, "analysis is incomplete; missing keys: " + ", ".join(missing))

    if not str(analysis.get("dispatch_summary") or "").strip():
        raise McpError(-32602, "analysis.dispatch_summary must be non-empty")
    if not isinstance(analysis.get("leader_evaluations"), list):
        raise McpError(-32602, "analysis.leader_evaluations must be a list")
    if not isinstance(analysis.get("people"), dict):
        raise McpError(-32602, "analysis.people must be an object")

    operational_tasks = analysis.get("operational_tasks")
    if not isinstance(operational_tasks, list):
        raise McpError(-32602, "analysis.operational_tasks must be a list")
    for index, task in enumerate(operational_tasks, start=1):
        if not isinstance(task, dict):
            raise McpError(-32602, f"analysis.operational_tasks[{index}] must be an object")
        task_missing = [key for key in ZOOM_REPORT_REQUIRED_TASK_KEYS if key not in task]
        if task_missing:
            raise McpError(
                -32602,
                f"analysis.operational_tasks[{index}] is incomplete; missing keys: " + ", ".join(task_missing),
            )


def tool_save_zoom_call_report(args: dict[str, Any]) -> dict[str, Any]:
    call_id = str(args.get("call_id") or "").strip()
    zoom_uuid = str(args.get("zoom_uuid") or "").strip()
    if not call_id and not zoom_uuid:
        raise McpError(-32602, "Missing required argument: call_id or zoom_uuid")

    report_text = str(args.get("report_text") or "").strip()
    summary = str(args.get("summary") or "").strip()
    if not report_text and summary:
        report_text = summary
    if not report_text:
        raise McpError(-32602, "Missing required argument: report_text")

    analysis = args.get("analysis") if isinstance(args.get("analysis"), dict) else {}
    raw_input = args.get("raw_input") if isinstance(args.get("raw_input"), dict) else {}
    model = str(args.get("model") or "").strip()
    status = str(args.get("status") or "done").strip() or "done"
    if status not in {"done", "error"}:
        raise McpError(-32602, "status must be done or error")

    validate_zoom_call_report_analysis(analysis, status)
    operational_tasks = normalize_zoom_operational_tasks_for_raw_json(report_text, analysis)
    report_payload = {
        "source": "mcp_save_zoom_call_report",
        "summary": summary,
        "report_text": report_text,
        "operational_tasks": operational_tasks,
        "analysis": analysis,
        "raw_input": raw_input,
        "model": model or None,
        "status": status,
        "saved_at": datetime.now().isoformat(),
    }

    with connect() as conn:
        with conn.transaction():
            with conn.cursor() as cur:
                if not safe_table_exists(cur, "zoom_calls"):
                    raise McpError(-32602, "zoom_calls table does not exist")
                if call_id:
                    cur.execute("SELECT id, zoom_uuid, call_date, topic, technical_topic FROM zoom_calls WHERE id = %s", (call_id,))
                else:
                    cur.execute("SELECT id, zoom_uuid, call_date, topic, technical_topic FROM zoom_calls WHERE zoom_uuid = %s", (zoom_uuid,))
                call = cur.fetchone()
                if not call:
                    raise McpError(-32602, "Zoom call not found")
                cur.execute(
                    """
                    UPDATE zoom_calls
                    SET analytical_note = %s,
                        raw_json = COALESCE(raw_json, '{}'::jsonb) || jsonb_build_object('ai_report', %s::jsonb),
                        updated_at = now()
                    WHERE id = %s
                    RETURNING id, zoom_uuid, call_date, topic, technical_topic, updated_at
                    """,
                    (report_text, jsonb_arg(report_payload), call["id"]),
                )
                updated = cur.fetchone()
    return {"saved": True, "call": updated, "status": status}


def tool_delete_zoom_call_report(args: dict[str, Any]) -> dict[str, Any]:
    call_id = str(args.get("call_id") or "").strip()
    zoom_uuid = str(args.get("zoom_uuid") or "").strip()
    if not call_id and not zoom_uuid:
        raise McpError(-32602, "Missing required argument: call_id or zoom_uuid")

    with connect() as conn:
        with conn.transaction():
            with conn.cursor() as cur:
                if not safe_table_exists(cur, "zoom_calls"):
                    raise McpError(-32602, "zoom_calls table does not exist")
                if call_id:
                    cur.execute("SELECT id FROM zoom_calls WHERE id = %s", (call_id,))
                else:
                    cur.execute("SELECT id FROM zoom_calls WHERE zoom_uuid = %s", (zoom_uuid,))
                call = cur.fetchone()
                if not call:
                    raise McpError(-32602, "Zoom call not found")
                cur.execute(
                    """
                    UPDATE zoom_calls
                    SET analytical_note = '',
                        raw_json = COALESCE(raw_json, '{}'::jsonb) - 'ai_report',
                        updated_at = now()
                    WHERE id = %s
                    RETURNING id, zoom_uuid, call_date, topic, technical_topic, updated_at
                    """,
                    (call["id"],),
                )
                updated = cur.fetchone()
    return {"deleted": True, "call": updated}


def tool_search_zoom_transcripts(args: dict[str, Any]) -> dict[str, Any]:
    query = str(args.get("query") or "").strip()
    if not query:
        raise McpError(-32602, "Missing required argument: query")
    date_from = parse_date_arg(args, "date_from", required=False)
    date_to = parse_date_arg(args, "date_to", required=False)
    limit = parse_limit(args)
    offset = parse_offset(args)

    filters = ["zcts.text ILIKE %s"]
    params: list[Any] = [f"%{query}%"]
    if date_from:
        filters.append("zc.call_date >= %s")
        params.append(date_from)
    if date_to:
        filters.append("zc.call_date <= %s")
        params.append(date_to)
    params.extend([limit, offset])

    with connect() as conn:
        with conn.cursor() as cur:
            if not safe_table_exists(cur, "zoom_call_transcript_segments"):
                return {"items": [], "limit": limit, "offset": offset, "note": "zoom transcript tables do not exist"}
            cur.execute(
                f"""
                SELECT
                    zc.id AS call_id,
                    zc.zoom_uuid,
                    zc.call_date,
                    zc.start_time_msk,
                    zc.topic,
                    zc.technical_topic,
                    zcts.segment_index,
                    zcts.cue_index,
                    zcts.start_offset,
                    zcts.end_offset,
                    zcts.speaker,
                    zcts.text
                FROM zoom_call_transcript_segments zcts
                JOIN zoom_calls zc ON zc.id = zcts.call_id
                WHERE {' AND '.join(filters)}
                ORDER BY zc.call_date DESC, zc.start_time_msk DESC, zcts.cue_index ASC
                LIMIT %s OFFSET %s
                """,
                params,
            )
            rows = cur.fetchall()
    return {"items": rows, "limit": limit, "offset": offset}


CHAT_REPORT_ITEM_KEYS = {
    "previous_day_tasks": "previous_day_task",
    "goals": "goal",
    "commitments": "commitment",
    "actions": "action",
    "results": "result",
    "questions": "question",
    "unanswered_questions": "unanswered_question",
    "decisions": "decision",
    "risks": "risk",
    "blockers": "blocker",
    "bitrix_references": "bitrix_reference",
    "next_steps": "next_step",
    "recommendation_feedback": "recommendation_feedback",
    "strange_feedback": "strange_feedback",
    "notes": "note",
}


def jsonb_arg(value: Any) -> Jsonb:
    return Jsonb(json_safe(value if value is not None else {}))


def analysis_items(analysis: dict[str, Any]) -> list[tuple[str, dict[str, Any]]]:
    items: list[tuple[str, dict[str, Any]]] = []
    for key, item_type in CHAT_REPORT_ITEM_KEYS.items():
        raw_items = analysis.get(key) if isinstance(analysis, dict) else None
        if not isinstance(raw_items, list):
            continue
        for item in raw_items:
            if isinstance(item, dict):
                items.append((item_type, item))
            elif item not in (None, ""):
                items.append((item_type, {"text": str(item)}))
    return items


def count_analysis_items(analysis: dict[str, Any], keys: tuple[str, ...]) -> int:
    total = 0
    for key in keys:
        value = analysis.get(key)
        if isinstance(value, list):
            total += len([item for item in value if item])
    return total


def item_text(item: dict[str, Any]) -> str:
    for key in ("text", "title", "summary", "question", "decision", "reason"):
        value = str(item.get(key) or "").strip()
        if value:
            return value
    return json.dumps(json_safe(item), ensure_ascii=False)[:1000]


def evidence_message_ids(item: dict[str, Any]) -> list[int]:
    values = item.get("evidence_message_ids") or []
    if not isinstance(values, list):
        return []
    result: list[int] = []
    for value in values:
        try:
            result.append(int(value))
        except (TypeError, ValueError):
            continue
    return result


def to_int(value: Any) -> int | None:
    try:
        if value in (None, ""):
            return None
        return int(value)
    except (TypeError, ValueError):
        return None


def silence_days_label(days: int) -> str:
    if days <= 0:
        days = 1
    last_two = days % 100
    last = days % 10
    if 11 <= last_two <= 14:
        unit = "дней"
    elif last == 1:
        unit = "день"
    elif 2 <= last <= 4:
        unit = "дня"
    else:
        unit = "дней"
    return f"{days} {unit}"


def chat_tail_current_status(item: dict[str, Any]) -> str:
    for key in ("current_status", "status", "answer_status"):
        value = str(item.get(key) or "").strip()
        if value:
            return value
    return ""


def normalize_tail_text(value: Any) -> str:
    text = str(value or "").strip().lower()
    text = re.sub(r"^\s*\[[^\]]+\]\s*", "", text)
    text = re.sub(r"^\s*[^—:-]{1,80}\s*[—:-]\s*", "", text)
    text = re.sub(r"\s+срок\s+\d{1,2}\.\d{1,2}(?:\.\d{4})?\.?", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\s+", " ", text).strip(" .;:")
    return text


def tail_key(item: dict[str, Any]) -> tuple[str, str]:
    person = str(item.get("person_name") or item.get("addressed_to") or item.get("asked_by") or "").strip().lower()
    return (person, normalize_tail_text(item_text(item)))


# --- restored from HEAD (audit cleanup collateral): owner reports + recommendations ---
def tool_get_owner_reports(args: dict[str, Any]) -> dict[str, Any]:
    report_kind = str(args.get("report_kind") or "daily").strip().lower()
    limit = parse_limit(args, 20)
    if report_kind not in {"daily", "weekly"}:
        raise McpError(-32602, "report_kind must be 'daily' or 'weekly'")

    with connect() as conn:
        with conn.cursor() as cur:
            if report_kind == "daily":
                cur.execute(
                    """
                    SELECT id, report_date, version, is_current, generated_at,
                           summary, dynamics_summary, risks_summary, recommendations, report_text, raw_json
                    FROM owner_daily_reports
                    WHERE is_current = TRUE
                    ORDER BY report_date DESC, version DESC
                    LIMIT %s
                    """,
                    (limit,),
                )
            else:
                cur.execute(
                    """
                    SELECT id, period_start, period_end, version, is_current, generated_at,
                           summary, dynamics_summary, risks_summary, recommendations, report_text, raw_json
                    FROM owner_weekly_reports
                    WHERE is_current = TRUE
                    ORDER BY period_start DESC, period_end DESC, version DESC
                    LIMIT %s
                    """,
                    (limit,),
                )
            reports = cur.fetchall()
    return {
        "report_kind": report_kind,
        "reports": reports,
        "rule": "Use these reports as continuity context before making recommendations or judging what is done/open/repeated.",
    }


RECOMMENDATION_STATUSES = {
    "new",
    "draft",
    "queued",
    "sent",
    "seen",
    "acked",
    "accepted",
    "in_progress",
    "needs_clarification",
    "disagreed",
    "delegated",
    "done",
    "rejected",
    "no_response",
    "overdue",
    "requires_manager_review",
    "cancelled",
    "error",
}

OPEN_RECOMMENDATION_STATUSES = {
    "new",
    "draft",
    "queued",
    "sent",
    "seen",
    "acked",
    "accepted",
    "in_progress",
    "needs_clarification",
    "disagreed",
    "delegated",
    "no_response",
    "overdue",
    "requires_manager_review",
    "error",
}


def tool_list_recommendations(args: dict[str, Any]) -> dict[str, Any]:
    date_from = parse_date_arg(args, "date_from", required=False)
    date_to = parse_date_arg(args, "date_to", required=False)
    status = str(args.get("status") or "").strip()
    dialog_id = str(args.get("dialog_id") or "").strip()
    include_events = bool(args.get("include_events", True))
    manager_bitrix_user_id = args.get("manager_bitrix_user_id")
    employee_bitrix_user_id = args.get("employee_bitrix_user_id")
    limit = parse_limit(args, 100)
    offset = parse_offset(args)

    filters = ["1=1"]
    params: list[Any] = []
    if date_from and date_to:
        filters.append("COALESCE(r.report_date, r.period_start, r.created_at::date) BETWEEN %s AND %s")
        params.extend([date_from, date_to])
    elif date_from or date_to:
        raise McpError(-32602, "date_from and date_to must be provided together")
    if status:
        if status == "open":
            filters.append("r.status = ANY(%s::text[])")
            params.append(sorted(OPEN_RECOMMENDATION_STATUSES))
        else:
            if status not in RECOMMENDATION_STATUSES:
                raise McpError(-32602, f"Unknown recommendation status: {status}")
            filters.append("r.status = %s")
            params.append(status)
    if dialog_id:
        filters.append("(sc.dialog_id = %s OR fc.dialog_id = %s OR r.feedback_dialog_id = %s)")
        params.extend([dialog_id, dialog_id, dialog_id])
    if manager_bitrix_user_id not in (None, ""):
        filters.append("r.manager_bitrix_user_id = %s")
        params.append(int(manager_bitrix_user_id))
    if employee_bitrix_user_id not in (None, ""):
        filters.append("r.employee_bitrix_user_id = %s")
        params.append(int(employee_bitrix_user_id))

    params.extend([limit, offset])
    with connect() as conn:
        with conn.cursor() as cur:
            if not safe_table_exists(cur, "owner_manager_recommendations"):
                return {"items": [], "events": [], "message": "Recommendation table does not exist."}
            cur.execute(
                f"""
                SELECT
                    r.id, r.source_scope, r.report_date, r.period_start, r.period_end,
                    r.status, r.priority, r.recommendation_type, r.subject,
                    r.recommendation_text, r.expected_action, r.due_date,
                    r.response_due_at, r.execution_due_at, r.sent_at, r.last_response_at,
                    r.manager_review_required, r.current_interpretation,
                    mu.full_name AS manager_name,
                    r.manager_bitrix_user_id,
                    eu.full_name AS employee_name,
                    r.employee_bitrix_user_id,
                    sc.dialog_id AS source_dialog_id,
                    sc.chat_title AS source_chat_title,
                    fc.dialog_id AS feedback_dialog_id,
                    fc.chat_title AS feedback_chat_title,
                    r.source_payload, r.created_at, r.updated_at
                FROM owner_manager_recommendations r
                LEFT JOIN users mu ON mu.id = r.manager_user_id
                LEFT JOIN users eu ON eu.id = r.employee_user_id
                LEFT JOIN chats sc ON sc.id = r.source_chat_id
                LEFT JOIN chats fc ON fc.id = r.feedback_chat_id
                WHERE {' AND '.join(filters)}
                ORDER BY r.manager_review_required DESC, r.updated_at DESC, r.created_at DESC
                LIMIT %s OFFSET %s
                """,
                params,
            )
            items = cur.fetchall()
            events: list[dict[str, Any]] = []
            if include_events and items and safe_table_exists(cur, "owner_recommendation_events"):
                ids = [str(item["id"]) for item in items]
                cur.execute(
                    """
                    SELECT
                        e.recommendation_id, e.event_type, e.author_type,
                        u.full_name AS author_name, e.author_bitrix_user_id,
                        e.dialog_id, c.chat_title, e.bitrix_message_id,
                        e.chat_message_day, e.old_status, e.new_status,
                        e.event_text, e.interpretation, e.source_payload, e.event_at
                    FROM owner_recommendation_events e
                    LEFT JOIN users u ON u.id = e.author_user_id
                    LEFT JOIN chats c ON c.id = e.chat_id
                    WHERE e.recommendation_id = ANY(%s::uuid[])
                    ORDER BY e.event_at DESC
                    LIMIT 500
                    """,
                    (ids,),
                )
                events = cur.fetchall()
    return {
        "items": items,
        "events": events,
        "limit": limit,
        "offset": offset,
        "rule": "Use recommendation events as the lifecycle log; do not infer final status from chat text without recording an event.",
    }


def tool_get_recommendation_feedback_context(args: dict[str, Any]) -> dict[str, Any]:
    dialog_id = str(args.get("dialog_id") or "").strip()
    if not dialog_id:
        raise McpError(-32602, "Missing required argument: dialog_id")
    report_date = parse_date_arg(args, "report_date")
    previous_date = report_date - timedelta(days=1)
    with connect() as conn:
        with conn.cursor() as cur:
            cur.execute("SELECT id, dialog_id, chat_title FROM chats WHERE dialog_id = %s", (dialog_id,))
            chat = cur.fetchone()
            if not chat:
                raise McpError(-32602, f"Unknown dialog_id: {dialog_id}")
            cur.execute(
                "SELECT user_id FROM chat_members WHERE chat_id = %s AND is_active = TRUE",
                (chat["id"],),
            )
            member_ids = [str(row["user_id"]) for row in cur.fetchall()]
            member_clause = "FALSE"
            member_params: list[Any] = []
            if member_ids:
                member_clause = "(r.manager_user_id = ANY(%s::uuid[]) OR r.employee_user_id = ANY(%s::uuid[]))"
                member_params = [member_ids, member_ids]
            cur.execute(
                f"""
                SELECT
                    r.id, r.status, r.priority, r.recommendation_type, r.subject,
                    r.recommendation_text, r.expected_action, r.due_date,
                    r.response_due_at, r.execution_due_at, r.sent_at, r.last_response_at,
                    r.manager_review_required, r.current_interpretation,
                    mu.full_name AS manager_name,
                    eu.full_name AS employee_name,
                    sc.dialog_id AS source_dialog_id,
                    fc.dialog_id AS feedback_dialog_id
                FROM owner_manager_recommendations r
                LEFT JOIN users mu ON mu.id = r.manager_user_id
                LEFT JOIN users eu ON eu.id = r.employee_user_id
                LEFT JOIN chats sc ON sc.id = r.source_chat_id
                LEFT JOIN chats fc ON fc.id = r.feedback_chat_id
                WHERE (
                    r.source_chat_id = %s
                    OR r.feedback_chat_id = %s
                    OR COALESCE(r.feedback_dialog_id, '') = %s
                    OR {member_clause}
                )
                  AND (
                    r.status = ANY(%s::text[])
                    OR r.created_at::date BETWEEN %s AND %s
                    OR r.sent_at::date BETWEEN %s AND %s
                    OR r.last_response_at::date BETWEEN %s AND %s
                  )
                ORDER BY r.manager_review_required DESC, r.updated_at DESC
                LIMIT 100
                """,
                [
                    chat["id"],
                    chat["id"],
                    dialog_id,
                    *member_params,
                    sorted(OPEN_RECOMMENDATION_STATUSES),
                    previous_date,
                    report_date,
                    previous_date,
                    report_date,
                    previous_date,
                    report_date,
                ],
            )
            recommendations = cur.fetchall()
            events: list[dict[str, Any]] = []
            if recommendations and safe_table_exists(cur, "owner_recommendation_events"):
                ids = [str(item["id"]) for item in recommendations]
                cur.execute(
                    """
                    SELECT recommendation_id, event_type, author_type, dialog_id,
                           bitrix_message_id, chat_message_day, old_status, new_status,
                           event_text, interpretation, source_payload, event_at
                    FROM owner_recommendation_events
                    WHERE recommendation_id = ANY(%s::uuid[])
                    ORDER BY event_at DESC
                    LIMIT 300
                    """,
                    (ids,),
                )
                events = cur.fetchall()
    return {
        "dialog_id": dialog_id,
        "report_date": report_date,
        "period_to_check": {"date_from": previous_date, "date_to": report_date},
        "recommendations": recommendations,
        "events": events,
        "rule": (
            "For daily chat reports, compare previous-day and current-day messages against these recommendations. "
            "Unclear, evasive, contradictory, or off-topic replies must be marked as 'Странная обратная связь' "
            "and saved with status requires_manager_review or needs_clarification."
        ),
    }


def tool_save_recommendation_event(args: dict[str, Any]) -> dict[str, Any]:
    recommendation_id = str(args.get("recommendation_id") or "").strip()
    if not recommendation_id:
        raise McpError(-32602, "Missing required argument: recommendation_id")
    try:
        recommendation_uuid = UUID(recommendation_id)
    except ValueError as exc:
        raise McpError(-32602, "recommendation_id must be a UUID") from exc
    event_type = str(args.get("event_type") or "ai_interpreted").strip()
    if event_type not in {"created", "sent", "delivered", "seen", "employee_replied", "ai_interpreted", "status_changed", "manager_reviewed", "task_created", "closed", "source_found"}:
        raise McpError(-32602, f"Unknown event_type: {event_type}")
    author_type = str(args.get("author_type") or "ai").strip()
    if author_type not in {"system", "ai", "manager", "employee"}:
        raise McpError(-32602, f"Unknown author_type: {author_type}")
    new_status = str(args.get("new_status") or "").strip() or None
    if new_status and new_status not in RECOMMENDATION_STATUSES:
        raise McpError(-32602, f"Unknown recommendation status: {new_status}")
    event_text = str(args.get("event_text") or "").strip() or None
    interpretation = args.get("interpretation") if isinstance(args.get("interpretation"), dict) else {}
    source_payload = args.get("source_payload") if isinstance(args.get("source_payload"), dict) else {}
    dialog_id = str(args.get("dialog_id") or "").strip() or None
    bitrix_message_id = args.get("bitrix_message_id")
    chat_message_day = parse_date_arg(args, "chat_message_day", required=False)
    author_bitrix_user_id = args.get("author_bitrix_user_id")

    with connect() as conn:
        with conn.transaction():
            with conn.cursor() as cur:
                cur.execute("SELECT id, status FROM owner_manager_recommendations WHERE id = %s", (recommendation_uuid,))
                recommendation = cur.fetchone()
                if not recommendation:
                    raise McpError(-32602, f"Unknown recommendation_id: {recommendation_id}")
                chat_id = None
                if dialog_id:
                    cur.execute("SELECT id FROM chats WHERE dialog_id = %s", (dialog_id,))
                    chat = cur.fetchone()
                    chat_id = chat["id"] if chat else None
                old_status = str(args.get("old_status") or recommendation.get("status") or "new")
                cur.execute(
                    """
                    INSERT INTO owner_recommendation_events (
                        recommendation_id, event_type, author_type, author_bitrix_user_id,
                        chat_id, dialog_id, bitrix_message_id, chat_message_day,
                        old_status, new_status, event_text, interpretation, source_payload
                    )
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                    RETURNING id
                    """,
                    (
                        recommendation_uuid,
                        event_type,
                        author_type,
                        int(author_bitrix_user_id) if author_bitrix_user_id not in (None, "") else None,
                        chat_id,
                        dialog_id,
                        int(bitrix_message_id) if bitrix_message_id not in (None, "") else None,
                        chat_message_day,
                        old_status,
                        new_status,
                        event_text,
                        jsonb_arg(interpretation),
                        jsonb_arg(source_payload),
                    ),
                )
                event_id = cur.fetchone()["id"]
                if new_status:
                    review_required = bool(args.get("manager_review_required")) or bool(interpretation.get("requires_manager_review")) or new_status == "requires_manager_review"
                    cur.execute(
                        """
                        UPDATE owner_manager_recommendations
                        SET status = %s,
                            feedback_chat_id = COALESCE(feedback_chat_id, %s),
                            feedback_dialog_id = COALESCE(feedback_dialog_id, %s),
                            current_interpretation = CASE WHEN %s::jsonb = '{}'::jsonb THEN current_interpretation ELSE %s::jsonb END,
                            manager_review_required = manager_review_required OR %s,
                            last_response_at = CASE
                                WHEN %s IN ('accepted','in_progress','needs_clarification','disagreed','delegated','done','requires_manager_review')
                                THEN now()
                                ELSE last_response_at
                            END,
                            closed_at = CASE WHEN %s IN ('done','rejected','cancelled') THEN COALESCE(closed_at, now()) ELSE closed_at END,
                            updated_at = now()
                        WHERE id = %s
                        """,
                        (
                            new_status,
                            chat_id,
                            dialog_id,
                            jsonb_arg(interpretation),
                            jsonb_arg(interpretation),
                            review_required,
                            new_status,
                            new_status,
                            recommendation_uuid,
                        ),
                    )
    return {"event_id": str(event_id), "recommendation_id": recommendation_id, "new_status": new_status}


def tool_get_previous_owner_daily_context(args: dict[str, Any]) -> dict[str, Any]:
    report_date = parse_date_arg(args, "report_date")
    if report_date is None:
        raise McpError(-32602, "Missing required argument: report_date")
    previous_date = report_date - timedelta(days=1)

    with connect() as conn:
        with conn.cursor() as cur:
            cur.execute(
                """
                SELECT id, report_date, version, is_current, generated_at,
                       summary, dynamics_summary, risks_summary, recommendations,
                       report_text, raw_json
                FROM owner_daily_reports
                WHERE report_date = %s AND is_current = TRUE
                ORDER BY version DESC
                LIMIT 1
                """,
                (previous_date,),
            )
            report = cur.fetchone()

    if not report:
        return {
            "report_date": report_date,
            "previous_report_date": previous_date,
            "previous_owner_daily": None,
            "message": "Previous owner daily report was not found for the previous calendar day.",
            "rule": "Do not invent the previous owner report. Use other verified sources and state that previous owner continuity is unavailable.",
        }

    raw_json = report.get("raw_json") if isinstance(report, dict) else None
    analysis = raw_json.get("analysis") if isinstance(raw_json, dict) and isinstance(raw_json.get("analysis"), dict) else None
    return {
        "report_date": report_date,
        "previous_report_date": previous_date,
        "previous_owner_daily": {
            "id": report.get("id"),
            "report_date": report.get("report_date"),
            "version": report.get("version"),
            "generated_at": report.get("generated_at"),
            "summary": report.get("summary"),
            "dynamics_summary": report.get("dynamics_summary"),
            "risks_summary": report.get("risks_summary"),
            "recommendations": report.get("recommendations"),
            "report_text": report.get("report_text"),
            "analysis": analysis,
        },
        "rule": "Use this previous-day owner daily context only for continuity: open questions, repeated risks, completed items, and recommendations. Do not treat it as evidence for the new day without current sources.",
    }


def owner_report_raw_json(args: dict[str, Any], report_text: str, model: str, status: str) -> dict[str, Any]:
    raw_json = args.get("raw_json") if isinstance(args.get("raw_json"), dict) else {}
    analysis = args.get("analysis") if isinstance(args.get("analysis"), dict) else {}
    raw = dict(raw_json or {})
    raw.setdefault("source", "mcp_save_owner_report")
    raw["model"] = model
    raw["status"] = status
    owner_payload_keys = (
        "manager_recommendations",
        "manager_messages",
        "open_tasks",
        "overdue_tasks",
        "no_response",
        "goal_dynamics",
        "weekly_dynamics",
        "daily_owner_reports_manifest",
    )
    owner_payload = {key: args.get(key) for key in owner_payload_keys if key in args}
    if owner_payload:
        raw["owner_payload"] = owner_payload
    if analysis:
        raw["analysis"] = analysis
    if args.get("raw_input") and isinstance(args.get("raw_input"), dict):
        raw["raw_input"] = args.get("raw_input")
    if report_text:
        raw["report_text"] = report_text
    return raw


def tool_save_owner_daily_report(args: dict[str, Any]) -> dict[str, Any]:
    report_date = parse_date_arg(args, "report_date")
    report_text = str(args.get("report_text") or "").strip()
    summary = str(args.get("summary") or "").strip() or report_text
    if not report_text and not summary:
        raise McpError(-32602, "Missing required argument: report_text or summary")
    status = str(args.get("status") or "done").strip()
    if status not in {"done", "no_data", "error"}:
        raise McpError(-32602, "status must be one of: done, no_data, error")
    model = str(args.get("model") or "mcp-manual").strip()
    raw_json = owner_report_raw_json(args, report_text, model, status)
    dynamics_summary = str(args.get("dynamics_summary") or "").strip() or None
    risks_summary = str(args.get("risks_summary") or "").strip() or None
    recommendations = str(args.get("recommendations") or "").strip() or None

    with connect() as conn:
        with conn.transaction():
            with conn.cursor() as cur:
                cur.execute(
                    "SELECT COALESCE(MAX(version), 0) + 1 AS version FROM owner_daily_reports WHERE report_date = %s",
                    (report_date,),
                )
                version = int(cur.fetchone()["version"])
                cur.execute(
                    "UPDATE owner_daily_reports SET is_current = FALSE WHERE report_date = %s AND is_current = TRUE",
                    (report_date,),
                )
                cur.execute(
                    """
                    INSERT INTO owner_daily_reports (
                        report_date, version, is_current, ai_request_id, prompt_id, generated_at,
                        summary, dynamics_summary, risks_summary, recommendations, report_text, raw_json
                    )
                    VALUES (%s, %s, TRUE, NULL, NULL, now(), %s, %s, %s, %s, %s, %s)
                    RETURNING id
                    """,
                    (
                        report_date,
                        version,
                        summary,
                        dynamics_summary,
                        risks_summary,
                        recommendations,
                        report_text,
                        jsonb_arg(raw_json),
                    ),
                )
                report_id = cur.fetchone()["id"]
                analysis_for_messages = {
                    "manager_messages": args.get("manager_messages"),
                    "manager_recommendations": args.get("manager_recommendations"),
                }
                if analysis_for_messages["manager_messages"] or analysis_for_messages["manager_recommendations"]:
                    app_workflow_function("save_owner_daily_manager_messages")(
                        cur, report_id, report_date, analysis_for_messages
                    )
    return {
        "report_id": str(report_id),
        "report_date": report_date.isoformat(),
        "version": version,
        "status": status,
        "is_current": True,
    }


def tool_save_owner_weekly_report(args: dict[str, Any]) -> dict[str, Any]:
    period_start = parse_date_arg(args, "period_start")
    period_end = parse_date_arg(args, "period_end")
    if period_end < period_start:
        raise McpError(-32602, "period_end must be greater than or equal to period_start")
    report_text = str(args.get("report_text") or "").strip()
    summary = str(args.get("summary") or "").strip() or report_text
    if not report_text and not summary:
        raise McpError(-32602, "Missing required argument: report_text or summary")
    status = str(args.get("status") or "done").strip()
    if status not in {"done", "no_data", "error"}:
        raise McpError(-32602, "status must be one of: done, no_data, error")
    model = str(args.get("model") or "mcp-manual").strip()
    raw_json = owner_report_raw_json(args, report_text, model, status)
    dynamics_summary = str(args.get("dynamics_summary") or "").strip() or None
    risks_summary = str(args.get("risks_summary") or "").strip() or None
    recommendations = str(args.get("recommendations") or "").strip() or None

    with connect() as conn:
        with conn.transaction():
            with conn.cursor() as cur:
                cur.execute(
                    """
                    SELECT COALESCE(MAX(version), 0) + 1 AS version
                    FROM owner_weekly_reports
                    WHERE period_start = %s AND period_end = %s
                    """,
                    (period_start, period_end),
                )
                version = int(cur.fetchone()["version"])
                cur.execute(
                    """
                    UPDATE owner_weekly_reports
                    SET is_current = FALSE
                    WHERE period_start = %s AND period_end = %s AND is_current = TRUE
                    """,
                    (period_start, period_end),
                )
                cur.execute(
                    """
                    INSERT INTO owner_weekly_reports (
                        period_start, period_end, version, is_current, ai_request_id, prompt_id, generated_at,
                        summary, dynamics_summary, risks_summary, recommendations, report_text, raw_json
                    )
                    VALUES (%s, %s, %s, TRUE, NULL, NULL, now(), %s, %s, %s, %s, %s, %s)
                    RETURNING id
                    """,
                    (
                        period_start,
                        period_end,
                        version,
                        summary,
                        dynamics_summary,
                        risks_summary,
                        recommendations,
                        report_text,
                        jsonb_arg(raw_json),
                    ),
                )
                report_id = cur.fetchone()["id"]
    return {
        "report_id": str(report_id),
        "period_start": period_start.isoformat(),
        "period_end": period_end.isoformat(),
        "version": version,
        "status": status,
        "is_current": True,
    }


def tool_list_pending_zoom_operational_dispatches(args: dict[str, Any]) -> dict[str, Any]:
    date_from = parse_date_arg(args, "date_from", required=False)
    date_to = parse_date_arg(args, "date_to", required=False)
    workflow = app_workflow_function("list_pending_zoom_operational_dispatches")
    rows = workflow(date_from, date_to)
    return {"pending": json_safe(rows), "pending_count": len(rows)}


def tool_preview_zoom_operational_tasks(args: dict[str, Any]) -> dict[str, Any]:
    call_id = str(args.get("call_id") or "").strip()
    if not call_id:
        raise McpError(-32602, "Missing required argument: call_id")
    workflow = app_workflow_function("preview_zoom_operational_tasks")
    return json_safe(workflow(call_id))


def tool_dispatch_zoom_operational_tasks(args: dict[str, Any]) -> dict[str, Any]:
    if args.get("confirm") is not True:
        raise McpError(
            -32602,
            "Sending requires confirm=true. The owner must have explicitly approved creating aggregated 'Итоги созвона' "
            "Bitrix tasks for this zoom call (one task per responsible person, exactly like the Albery UI 'Отправка задач'). "
            "Only then call dispatch_zoom_operational_tasks with confirm=true.",
        )
    call_id = str(args.get("call_id") or "").strip()
    if not call_id:
        raise McpError(-32602, "Missing required argument: call_id")
    workflow = app_workflow_function("dispatch_zoom_operational_tasks")
    result = workflow(call_id)
    return json_safe(result)


def tool_preview_zoom_participant_reports(args: dict[str, Any]) -> dict[str, Any]:
    call_id = str(args.get("call_id") or "").strip()
    if not call_id:
        raise McpError(-32602, "Missing required argument: call_id")
    workflow = app_workflow_function("preview_zoom_participant_reports")
    return json_safe(workflow(call_id))


def tool_dispatch_zoom_participant_reports(args: dict[str, Any]) -> dict[str, Any]:
    if args.get("confirm") is not True:
        raise McpError(-32602, "Sending requires confirm=true. The owner/user must explicitly approve sending personal participant reports as Bitrix tasks.")
    call_id = str(args.get("call_id") or "").strip()
    if not call_id:
        raise McpError(-32602, "Missing required argument: call_id")
    workflow = app_workflow_function("dispatch_zoom_participant_reports")
    return json_safe(workflow(call_id))


def tool_list_leader_evaluations(args: dict[str, Any]) -> dict[str, Any]:
    date_from = parse_date_arg(args, "date_from", required=False)
    date_to = parse_date_arg(args, "date_to", required=False)
    workflow = app_workflow_function("list_leader_evaluations")
    return json_safe(workflow(date_from, date_to))


def tool_dispatch_owner_weekly_report_task(args: dict[str, Any]) -> dict[str, Any]:
    if args.get("confirm") is not True:
        raise McpError(
            -32602,
            "Sending requires confirm=true. Creates the weekly owner report as a Bitrix task for "
            "Евгений Палей with the report PDF attached and deadline next Monday 10:00 МСК. In the "
            "automated Friday cron this is called automatically after the report is generated and "
            "saved; interactive callers must confirm.",
        )
    report_id = str(args.get("report_id") or "").strip()
    if not report_id:
        raise McpError(-32602, "Missing required argument: report_id")
    workflow = app_workflow_function("dispatch_owner_weekly_report_task")
    return json_safe(workflow(report_id))


def tool_dispatch_leader_evaluations_digest(args: dict[str, Any]) -> dict[str, Any]:
    if args.get("confirm") is not True:
        raise McpError(
            -32602,
            "Sending requires confirm=true. The owner must have explicitly approved the leader-evaluation digest "
            "in Telegram before it is created as a Bitrix task for Евгений Палей.",
        )
    digest_text = str(args.get("digest_text") or "").strip()
    if not digest_text:
        raise McpError(-32602, "Missing required argument: digest_text")
    date_from = parse_date_arg(args, "date_from", required=False)
    date_to = parse_date_arg(args, "date_to", required=False)
    workflow = app_workflow_function("dispatch_leader_evaluations_digest")
    return json_safe(workflow(digest_text, date_from, date_to))


def _resolve_current_owner_daily_report(cur: Any, report_date: date) -> dict[str, Any]:
    cur.execute(
        """
        SELECT id, version, summary, dynamics_summary, risks_summary, recommendations, report_text
        FROM owner_daily_reports
        WHERE report_date = %s AND is_current = TRUE
        LIMIT 1
        """,
        (report_date,),
    )
    row = cur.fetchone()
    if not row:
        raise McpError(-32602, f"Нет текущего ежедневного owner-отчёта за {report_date.isoformat()}.")
    return dict(row)


def tool_list_pending_owner_recommendations(args: dict[str, Any]) -> dict[str, Any]:
    report_date = parse_date_arg(args, "report_date")
    with connect() as conn:
        with conn.cursor() as cur:
            report = _resolve_current_owner_daily_report(cur, report_date)
            cur.execute(
                """
                SELECT r.id, r.manager_user_id, r.manager_bitrix_user_id,
                       u.full_name AS manager_full_name,
                       r.recommendation_text, r.subject, r.priority, r.due_date,
                       r.recommendation_type, r.status, r.created_at
                FROM owner_manager_recommendations r
                LEFT JOIN users u ON u.id = r.manager_user_id
                WHERE r.owner_daily_report_id = %s
                  AND r.manager_bitrix_user_id IS NOT NULL
                  AND r.status NOT IN ('sent','cancelled','done','rejected')
                ORDER BY
                  CASE r.priority
                    WHEN 'critical' THEN 1
                    WHEN 'high' THEN 2
                    WHEN 'medium' THEN 3
                    ELSE 4
                  END,
                  r.created_at
                """,
                (report["id"],),
            )
            rows = [json_safe(dict(row)) for row in cur.fetchall()]
    return {
        "report_id": str(report["id"]),
        "report_date": report_date.isoformat(),
        "report_summary": report.get("summary"),
        "report_dynamics_summary": report.get("dynamics_summary"),
        "report_risks_summary": report.get("risks_summary"),
        "report_text": report.get("report_text"),
        "recommendations": rows,
        "recommendations_count": len(rows),
    }


def tool_send_owner_recommendations_to_bitrix(args: dict[str, Any]) -> dict[str, Any]:
    if args.get("confirm") is not True:
        raise McpError(
            -32602,
            "Creating recommendation tasks requires confirm=true. First show the owner the exact recommendation list "
            "per recipient and get explicit approval. Only then call send_owner_recommendations_to_bitrix with confirm=true.",
        )
    report_date = parse_date_arg(args, "report_date")
    recipient_recommendations = args.get("recipient_recommendations")
    if not isinstance(recipient_recommendations, dict) or not recipient_recommendations:
        raise McpError(
            -32602,
            "recipient_recommendations must be a non-empty object {bitrix_user_id: message_text}.",
        )
    normalized: dict[str, str] = {}
    for key, value in recipient_recommendations.items():
        user_id = to_int(key)
        text = str(value or "").strip()
        if user_id is None or not text:
            continue
        normalized[str(user_id)] = text
    if not normalized:
        raise McpError(-32602, "recipient_recommendations contains no valid (bitrix_user_id, text) entries.")
    recipient_ids = [int(k) for k in normalized.keys()]
    with connect() as conn:
        with conn.cursor() as cur:
            report = _resolve_current_owner_daily_report(cur, report_date)
    workflow = app_workflow_function("send_owner_report_recommendations_to_bitrix")
    result = workflow(str(report["id"]), "daily", recipient_ids, normalized)
    return {
        "report_id": str(report["id"]),
        "report_date": report_date.isoformat(),
        "sent": result.get("sent", 0),
        "failed": result.get("failed", 0),
        "results": json_safe(result.get("results") or []),
        "errors": json_safe(result.get("errors") or []),
    }


def tool_send_owner_weekly_report_pdf(args: dict[str, Any]) -> dict[str, Any]:
    if args.get("confirm") is not True:
        raise McpError(
            -32602,
            "Sending requires confirm=true. First save the weekly report via save_owner_weekly_report, show the owner "
            "the report period and that the PDF goes to Evgeniy, get explicit approval. Only then call "
            "send_owner_weekly_report_pdf with confirm=true.",
        )
    period_start = parse_date_arg(args, "period_start")
    period_end = parse_date_arg(args, "period_end")
    recipient_ids = args.get("recipient_bitrix_user_ids")
    if recipient_ids in (None, "", []):
        recipient_ids = [1]  # Evgeniy Palei by default
    if not isinstance(recipient_ids, list):
        recipient_ids = [recipient_ids]
    normalized_ids: list[int] = []
    for value in recipient_ids:
        user_id = to_int(value)
        if user_id is not None:
            normalized_ids.append(user_id)
    if not normalized_ids:
        raise McpError(-32602, "recipient_bitrix_user_ids contains no valid integer id.")
    loader = app_workflow_function("load_owner_weekly_report")
    report = loader(period_start, period_end)
    # load_owner_weekly_report returns the dict from owner_weekly_report_to_dict,
    # which exposes the primary key as "report_id" (not "id"). Accept both so the
    # lookup never falsely reports "not found" and re-triggers a save/retry loop.
    report_pk = (report.get("report_id") or report.get("id")) if report else None
    if not report_pk:
        raise McpError(
            -32004,
            f"Текущий недельный отчёт за {period_start.isoformat()}–{period_end.isoformat()} не найден. "
            "Сначала сохрани его через save_owner_weekly_report.",
        )
    report_id = str(report_pk)
    workflow = app_workflow_function("send_owner_report_pdf_to_bitrix")
    result = workflow(report_id, "weekly", normalized_ids)
    return {
        "report_id": report_id,
        "period_start": period_start.isoformat(),
        "period_end": period_end.isoformat(),
        "recipient_bitrix_user_ids": normalized_ids,
        "filename": result.get("filename"),
        "pdf_size": result.get("pdf_size"),
        "sent": result.get("sent", 0),
        "failed": result.get("failed", 0),
        "results": json_safe(result.get("results") or []),
        "errors": json_safe(result.get("errors") or []),
    }


def _resolve_message_recipient(
    recipient_bitrix_user_id: Any = None,
    recipient_name: Any = None,
) -> dict[str, Any]:
    user_id = None
    if recipient_bitrix_user_id not in (None, ""):
        try:
            user_id = int(recipient_bitrix_user_id)
        except (TypeError, ValueError) as exc:
            raise McpError(-32602, "recipient_bitrix_user_id must be an integer.") from exc

    name_clean = str(recipient_name or "").strip()
    with connect() as conn:
        with conn.cursor() as cur:
            if user_id is not None:
                cur.execute(
                    """
                    SELECT bitrix_user_id, full_name, email, work_position, is_active
                    FROM users
                    WHERE bitrix_user_id = %s AND is_active = TRUE
                    LIMIT 1
                    """,
                    (user_id,),
                )
                row = cur.fetchone()
                if not row:
                    raise McpError(
                        -32602,
                        f"Не найден активный сотрудник Bitrix с id {user_id}. Не перебирай id — "
                        "уточни у пользователя правильного сотрудника.",
                    )
                return dict(row)

            if not name_clean:
                raise McpError(
                    -32602,
                    "Нужно указать получателя: recipient_name или recipient_bitrix_user_id.",
                )

            cur.execute(
                """
                SELECT bitrix_user_id, full_name, email, work_position, is_active
                FROM users
                WHERE is_active = TRUE AND bitrix_user_id IS NOT NULL
                ORDER BY full_name
                """
            )
            rows = [dict(row) for row in cur.fetchall()]

    exact = [row for row in rows if str(row.get("full_name") or "").strip().lower() == name_clean.lower()]
    matches = exact or [row for row in rows if _person_names_match(row.get("full_name"), name_clean)]
    if not matches:
        raise McpError(-32602, f"Не удалось найти получателя в оргструктуре: {name_clean}.")
    if len(matches) > 1:
        candidates = [
            {
                "bitrix_user_id": row.get("bitrix_user_id"),
                "full_name": row.get("full_name"),
                "work_position": row.get("work_position"),
            }
            for row in matches[:10]
        ]
        raise McpError(
            -32602,
            "Получатель найден неоднозначно. Уточните recipient_bitrix_user_id. Кандидаты: "
            + json.dumps(candidates, ensure_ascii=False),
        )
    return matches[0]


def tool_send_bitrix_message(args: dict[str, Any]) -> dict[str, Any]:
    if args.get("confirm") is not True:
        raise McpError(
            -32602,
            "Sending a Bitrix message requires confirm=true. First show the user the exact final message_text "
            "and the resolved recipient (full_name + work_position + bitrix_user_id) and get explicit approval. "
            "Only then call send_bitrix_message with confirm=true.",
        )
    message_text = str(args.get("message_text") or "").strip()
    if not message_text:
        raise McpError(-32602, "message_text must be a non-empty string.")
    if len(message_text) > 20000:
        raise McpError(-32602, "message_text is too long (max 20000 characters).")
    recipient = _resolve_message_recipient(
        args.get("recipient_bitrix_user_id"),
        args.get("recipient_name"),
    )
    workflow = app_workflow_function("send_bitrix_personal_message")
    try:
        result = workflow(int(recipient["bitrix_user_id"]), message_text)
    except ValueError as exc:
        raise McpError(-32602, str(exc)) from exc
    except Exception as exc:  # noqa: BLE001
        raise McpError(-32010, f"Bitrix send failed: {exc}") from exc
    return {
        "sent": True,
        "recipient": {
            "bitrix_user_id": recipient.get("bitrix_user_id"),
            "full_name": recipient.get("full_name"),
            "work_position": recipient.get("work_position"),
        },
        "channel": result.get("channel"),
        "message_text": message_text,
        "bitrix_response": json_safe(result.get("response")),
        "im_message_error": result.get("im_message_error"),
    }


def tool_write_company_sheet(args: dict[str, Any]) -> dict[str, Any]:
    if args.get("confirm") is not True:
        raise McpError(
            -32602,
            "Запись в Google-таблицу требует confirm=true. Сначала покажи пользователю, что именно впишешь "
            "(таблица, лист, строки или диапазон), получи явное согласие, и только потом вызови с confirm=true.",
        )
    spreadsheet_id = str(args.get("spreadsheet_id") or "").strip()
    match = re.search(r"/spreadsheets/d/([a-zA-Z0-9_-]+)", spreadsheet_id)
    if match:
        spreadsheet_id = match.group(1)
    if not spreadsheet_id:
        raise McpError(-32602, "spreadsheet_id (или ссылка на таблицу) обязателен.")
    mode = str(args.get("mode") or "append").strip().lower()
    workflow = app_workflow_function("write_company_google_sheet")
    try:
        if mode == "update":
            result = workflow(spreadsheet_id, args.get("sheet"), "update", None, args.get("range"), args.get("values"))
        else:
            rows = args.get("rows") or []
            if not rows:
                raise McpError(-32602, "Для mode=append нужен непустой rows (список строк, каждая — список ячеек).")
            result = workflow(spreadsheet_id, args.get("sheet"), "append", rows, None, None)
    except McpError:
        raise
    except ValueError as exc:
        raise McpError(-32602, str(exc)) from exc
    except Exception as exc:  # noqa: BLE001
        raise McpError(-32010, f"Sheet write failed: {exc}") from exc
    return result


def tool_create_google_sheet(args: dict[str, Any]) -> dict[str, Any]:
    if args.get("confirm") is not True:
        raise McpError(
            -32602,
            "Создание Google-таблицы требует confirm=true. Сначала покажи пользователю, что создашь "
            "(название, какие данные впишешь, как будет оформлена таблица, и что доступ будет «по ссылке — редактор»), "
            "получи согласие, и только потом вызови с confirm=true.",
        )
    title = str(args.get("title") or "").strip()
    if not title:
        raise McpError(-32602, "title (название таблицы) обязателен.")
    rows = args.get("rows")
    if rows is not None and not isinstance(rows, list):
        raise McpError(-32602, "rows должен быть списком строк (каждая — список ячеек).")
    share = args.get("share_anyone_writer", True)
    workflow = app_workflow_function("create_google_sheet")
    try:
        return workflow(title, rows, bool(share))
    except McpError:
        raise
    except Exception as exc:  # noqa: BLE001
        raise McpError(-32010, f"Create sheet failed: {exc}") from exc


def tool_get_google_sheet_meta(args: dict[str, Any]) -> dict[str, Any]:
    sid = str(args.get("spreadsheet_id") or "").strip()
    if not sid:
        raise McpError(-32602, "spreadsheet_id is required.")
    try:
        return app_workflow_function("get_google_sheet_meta")(sid)
    except McpError:
        raise
    except Exception as exc:  # noqa: BLE001
        raise McpError(-32010, f"get_google_sheet_meta failed: {exc}") from exc


def tool_write_google_sheet_values(args: dict[str, Any]) -> dict[str, Any]:
    sid = str(args.get("spreadsheet_id") or "").strip()
    rng = str(args.get("range") or "").strip()
    values = args.get("values")
    if not sid or not rng:
        raise McpError(-32602, "spreadsheet_id and range are required.")
    if not isinstance(values, list):
        raise McpError(-32602, "values must be a list of rows (each row a list).")
    vio = str(args.get("value_input_option") or "USER_ENTERED")
    try:
        return app_workflow_function("write_google_sheet_values")(sid, rng, values, vio)
    except McpError:
        raise
    except Exception as exc:  # noqa: BLE001
        raise McpError(-32010, f"write_google_sheet_values failed: {exc}") from exc


def tool_format_google_sheet(args: dict[str, Any]) -> dict[str, Any]:
    sid = str(args.get("spreadsheet_id") or "").strip()
    requests = args.get("requests")
    if not sid:
        raise McpError(-32602, "spreadsheet_id is required.")
    if not isinstance(requests, list) or not requests:
        raise McpError(-32602, "requests must be a non-empty list of Sheets API request objects.")
    try:
        return app_workflow_function("format_google_sheet")(sid, requests)
    except McpError:
        raise
    except Exception as exc:  # noqa: BLE001
        raise McpError(-32010, f"format_google_sheet failed: {exc}") from exc


def tool_move_drive_file_to_folder(args: dict[str, Any]) -> dict[str, Any]:
    fid = str(args.get("file_id") or args.get("item_id") or "").strip()
    folder = str(args.get("folder") or "").strip()
    if not fid or not folder:
        raise McpError(-32602, "file_id/item_id and folder (id or URL) are required.")
    try:
        return app_workflow_function("move_drive_file_to_folder")(fid, folder)
    except McpError:
        raise
    except Exception as exc:  # noqa: BLE001
        raise McpError(-32010, f"move_drive_file_to_folder failed: {exc}") from exc


def tool_get_webapp_template(args: dict[str, Any]) -> dict[str, Any]:
    title = str(args.get("title") or "").strip() or None
    try:
        return app_workflow_function("webapp_design_template")(title)
    except McpError:
        raise
    except Exception as exc:  # noqa: BLE001
        raise McpError(-32010, f"get_webapp_template failed: {exc}") from exc


# Incremental document drafts — the model assembles a LONG document (contract) in small sections
# so no single tool-call output is huge. The Codex backend drops the stream connection when the
# model has to emit one very large tool argument (a full contract HTML ~ 15k+ output tokens); that
# is the real reason big contracts failed. Small sections (~5-8k chars each) never trigger that.
_DOC_DRAFT_DIR = Path(os.getenv("EXPORT_DRAFT_DIR", "/var/www/albery/.doc_drafts"))
_DOC_HTML_MAX = 400_000


def _doc_draft_path(token: str) -> Path:
    safe = re.sub(r"[^A-Za-z0-9_\-]", "", str(token or ""))
    return _DOC_DRAFT_DIR / f"{safe}.json"


def _render_and_save_doc(title: str, html: str, args: dict[str, Any]) -> str:
    from docformat import html_to_docx
    data = html_to_docx(
        html,
        font_size_pt=float(args.get("font_size_pt") or 12),
        line_spacing=float(args.get("line_spacing") or 1.15),
    )
    from b24bot import _b24_save_export
    return _b24_save_export(data, title, "docx")


def tool_export_document(args: dict[str, Any]) -> dict[str, Any]:
    import json as _json
    import secrets as _secrets
    title = str(args.get("title") or "").strip() or "Документ"
    fmt = str(args.get("format") or "docx").strip().lower()
    if fmt != "docx":
        raise McpError(-32602, "Пока поддерживается только format='docx'.")

    html = str(args.get("html") or "")
    section = str(args.get("section") or "")
    doc_token = str(args.get("doc_token") or "").strip()
    finalize = args.get("finalize") is True or str(args.get("finalize") or "").strip().lower() in {"true", "1", "yes", "да"}

    # --- Incremental mode: doc_token / section / finalize are present -------------------------
    if doc_token or section or finalize:
        try:
            _DOC_DRAFT_DIR.mkdir(parents=True, exist_ok=True)
            if not doc_token:
                # First section — open a new draft.
                doc_token = "doc_" + _secrets.token_urlsafe(8)
                draft = {"title": title, "html": "",
                         "font_size_pt": args.get("font_size_pt"), "line_spacing": args.get("line_spacing")}
            else:
                p = _doc_draft_path(doc_token)
                if not p.is_file():
                    raise McpError(-32602, f"Черновик {doc_token} не найден (истёк или не начат). Начни заново: "
                                           "export_document(title=..., section=...) без doc_token.")
                draft = _json.loads(p.read_text(encoding="utf-8"))
                if title and title != "Документ":
                    draft["title"] = title
            if section:
                draft["html"] += section
            if len(draft["html"]) > _DOC_HTML_MAX:
                raise McpError(-32602, "Документ превысил лимит 400 тыс. символов.")
            # Preserve settings from the opening call.
            for k in ("font_size_pt", "line_spacing"):
                if args.get(k) is not None:
                    draft[k] = args.get(k)

            if finalize:
                if not draft["html"].strip():
                    raise McpError(-32602, "Черновик пуст — добавь секции перед finalize.")
                url = _render_and_save_doc(
                    draft.get("title") or title,
                    draft["html"],
                    {"font_size_pt": draft.get("font_size_pt"), "line_spacing": draft.get("line_spacing")},
                )
                try:
                    _doc_draft_path(doc_token).unlink(missing_ok=True)
                except Exception:  # noqa: BLE001
                    pass
                return {"url": url, "format": "docx", "doc_token": doc_token,
                        "chars_total": len(draft["html"]),
                        "note": "Готово: документ собран из всех секций. Пришли пользователю эту ссылку."}
            _doc_draft_path(doc_token).write_text(_json.dumps(draft, ensure_ascii=False), encoding="utf-8")
            return {
                "doc_token": doc_token,
                "chars_total": len(draft["html"]),
                "finalized": False,
                "note": (f"Секция принята (всего {len(draft['html'])} символов). Добавь следующую часть: "
                         f"export_document(doc_token='{doc_token}', section='<HTML следующей части>'). "
                         f"Когда документ готов — export_document(doc_token='{doc_token}', finalize=true). "
                         "Держи каждую секцию небольшой (примерно до 6000 символов)."),
            }
        except McpError:
            raise
        except Exception as exc:  # noqa: BLE001
            raise McpError(-32010, f"export_document (incremental) failed: {exc}") from exc

    # --- One-shot mode (small documents) ------------------------------------------------------
    if not html.strip():
        raise McpError(-32602, "Передайте html — полное содержимое документа. Для ДЛИННОГО документа "
                               "(договор) собирай его по частям: export_document(title=..., section=...) → "
                               "export_document(doc_token=..., section=...) → export_document(doc_token=..., finalize=true).")
    if len(html) > _DOC_HTML_MAX:
        raise McpError(-32602, "HTML слишком большой (лимит 400 тыс. символов) — сократите документ.")
    try:
        url = _render_and_save_doc(title, html, args)
    except McpError:
        raise
    except Exception as exc:  # noqa: BLE001
        raise McpError(-32010, f"export_document failed: {exc}") from exc
    return {
        "url": url,
        "format": "docx",
        "note": "Файл готов. Пришли пользователю эту ссылку (она token-защищена и временная). "
                "Если пользователь просит поправить оформление — исправь свой HTML и вызови инструмент снова.",
    }


def tool_make_sheet_applet(args: dict[str, Any]) -> dict[str, Any]:
    sid = str(args.get("spreadsheet_id") or args.get("spreadsheet") or args.get("url") or "").strip()
    if not sid:
        raise McpError(-32602, "spreadsheet_id (id or URL) is required.")
    sheet = str(args.get("sheet") or "").strip() or None
    try:
        return app_workflow_function("make_sheet_applet")(sid, sheet)
    except McpError:
        raise
    except Exception as exc:  # noqa: BLE001
        raise McpError(-32010, f"make_sheet_applet failed: {exc}") from exc


def tool_share_drive_item_for_everyone(args: dict[str, Any]) -> dict[str, Any]:
    item = str(args.get("item") or args.get("file_id") or args.get("url") or "").strip()
    if not item:
        raise McpError(-32602, "item (Drive id or URL) is required.")
    role = str(args.get("role") or "writer").strip().lower()
    try:
        return app_workflow_function("share_drive_item_for_everyone")(item, role)
    except McpError:
        raise
    except Exception as exc:  # noqa: BLE001
        raise McpError(-32010, f"share_drive_item_for_everyone failed: {exc}") from exc


def tool_remove_drive_item_from_folder(args: dict[str, Any]) -> dict[str, Any]:
    if args.get("confirm") is not True:
        raise McpError(
            -32602,
            "Нужно confirm=true: сначала покажи пользователю, какой файл/папку убираешь и из какой папки Drive. "
            "Операция убирает элемент только из указанной папки, не удаляя его из Google Drive полностью.",
        )
    item_id = str(args.get("item_id") or args.get("file_id") or "").strip()
    folder = str(args.get("folder") or "").strip()
    if not item_id or not folder:
        raise McpError(-32602, "item_id/file_id and folder (id or URL) are required.")
    try:
        return app_workflow_function("remove_drive_item_from_folder")(item_id, folder)
    except McpError:
        raise
    except Exception as exc:  # noqa: BLE001
        raise McpError(-32010, f"remove_drive_item_from_folder failed: {exc}") from exc


def tool_list_drive_folder_items(args: dict[str, Any]) -> dict[str, Any]:
    folder = str(args.get("folder") or "").strip()
    if not folder:
        raise McpError(-32602, "folder (id or URL) is required.")
    page_size = int(args.get("page_size") or 200)
    try:
        return app_workflow_function("list_drive_folder_items")(folder, page_size)
    except McpError:
        raise
    except Exception as exc:  # noqa: BLE001
        raise McpError(-32010, f"list_drive_folder_items failed: {exc}") from exc


def tool_create_drive_folder(args: dict[str, Any]) -> dict[str, Any]:
    if args.get("confirm") is not True:
        raise McpError(-32602, "Нужно confirm=true: сначала покажи пользователю имя новой папки и где она будет создана.")
    name = str(args.get("name") or "").strip()
    parent = str(args.get("parent_folder") or args.get("folder") or "").strip()
    if not name or not parent:
        raise McpError(-32602, "name and parent_folder/folder (id or URL) are required.")
    try:
        return app_workflow_function("create_drive_folder")(name, parent, bool(args.get("reuse_existing", True)))
    except McpError:
        raise
    except Exception as exc:  # noqa: BLE001
        raise McpError(-32010, f"create_drive_folder failed: {exc}") from exc


def tool_organize_drive_folder(args: dict[str, Any]) -> dict[str, Any]:
    folder = str(args.get("folder") or "").strip()
    if not folder:
        raise McpError(-32602, "folder (id or URL) is required.")
    dry_run = bool(args.get("dry_run", True))
    if not dry_run and args.get("confirm") is not True:
        raise McpError(-32602, "Нужно confirm=true: сначала покажи пользователю план сортировки папки Drive. Для безопасной проверки используй dry_run=true.")
    categories = args.get("categories")
    if categories is not None and not isinstance(categories, list):
        raise McpError(-32602, "categories must be a list of folder names.")
    try:
        return app_workflow_function("organize_drive_folder")(folder, categories, dry_run)
    except McpError:
        raise
    except Exception as exc:  # noqa: BLE001
        raise McpError(-32010, f"organize_drive_folder failed: {exc}") from exc


def tool_manage_apps_script(args: dict[str, Any]) -> dict[str, Any]:
    if args.get("confirm") is not True:
        raise McpError(-32602, "Set confirm=true to run an Apps Script action.")
    action = str(args.get("action") or "").strip().lower()
    if action not in ("create", "get", "update", "deploy", "run", "publish_web_app"):
        raise McpError(-32602, "action must be one of: create, get, update, deploy, run, publish_web_app.")
    try:
        return app_workflow_function("manage_apps_script")(
            action,
            script_id=str(args.get("script_id") or "") or None,
            title=str(args.get("title") or "") or None,
            files=args.get("files"),
            function_name=str(args.get("function_name") or "") or None,
            parameters=args.get("parameters"),
            description=str(args.get("description") or "") or None,
            web_app=bool(args.get("web_app", True)),
            access=str(args.get("access") or "") or None,
            execute_as=str(args.get("execute_as") or "") or None,
            advanced_services=args.get("advanced_services"),
            oauth_scopes=args.get("oauth_scopes"),
            share=bool(args.get("share", True)),
        )
    except McpError:
        raise
    except Exception as exc:  # noqa: BLE001
        raise McpError(-32010, f"manage_apps_script failed: {exc}") from exc


def tool_cancel_owner_recommendation(args: dict[str, Any]) -> dict[str, Any]:
    rec_id_raw = str(args.get("recommendation_id") or "").strip()
    if not rec_id_raw:
        raise McpError(-32602, "Missing required argument: recommendation_id")
    try:
        rec_uuid = UUID(rec_id_raw)
    except (TypeError, ValueError) as exc:
        raise McpError(-32602, "recommendation_id must be a UUID.") from exc
    reason = str(args.get("reason") or "").strip()
    with connect() as conn:
        with conn.transaction():
            with conn.cursor() as cur:
                cur.execute(
                    "SELECT id, status FROM owner_manager_recommendations WHERE id = %s",
                    (rec_uuid,),
                )
                row = cur.fetchone()
                if not row:
                    raise McpError(-32602, f"Recommendation {rec_id_raw} not found.")
                old_status = str(row.get("status") or "new")
                cur.execute(
                    """
                    UPDATE owner_manager_recommendations
                    SET status = 'cancelled', closed_at = COALESCE(closed_at, now()), updated_at = now()
                    WHERE id = %s
                    """,
                    (rec_uuid,),
                )
                cur.execute(
                    """
                    INSERT INTO owner_recommendation_events (
                        recommendation_id, event_type, author_type,
                        old_status, new_status, event_text, source_payload
                    )
                    VALUES (%s, 'cancelled', 'system', %s, 'cancelled', %s, %s)
                    """,
                    (rec_uuid, old_status, reason or "Cancelled by owner.", jsonb_arg({"reason": reason})),
                )
    return {
        "recommendation_id": rec_id_raw,
        "old_status": old_status,
        "new_status": "cancelled",
        "reason": reason or None,
    }


def tool_upsert_ai_instruction(args: dict[str, Any]) -> dict[str, Any]:
    raw_path = str(args.get("path") or "").strip()
    content = str(args.get("content") or "")
    if not raw_path:
        raise McpError(-32602, "Missing required argument: path")
    if len(content) > 500_000:
        raise McpError(-32602, "content is too large; maximum is 500000 characters")
    path_parts = [part.strip() for part in raw_path.replace("\\", "/").split("/") if part.strip()]
    if not path_parts:
        raise McpError(-32602, "path must include at least one folder name")

    with connect() as conn:
        with conn.transaction():
            with conn.cursor() as cur:
                if not safe_table_exists(cur, "ai_instruction_folders"):
                    raise McpError(-32000, "ai_instruction_folders table is missing. Apply database migrations before writing AI instructions.")
                parent_id = None
                current = None
                for index, name in enumerate(path_parts):
                    cur.execute(
                        """
                        SELECT id, name
                        FROM ai_instruction_folders
                        WHERE ((%s::uuid IS NULL AND parent_id IS NULL) OR parent_id = %s::uuid)
                          AND lower(name) = lower(%s)
                        LIMIT 1
                        """,
                        (parent_id, parent_id, name),
                    )
                    current = cur.fetchone()
                    if not current:
                        cur.execute(
                            """
                            SELECT COALESCE(max(sort_order), -1) + 1 AS sort_order
                            FROM ai_instruction_folders
                            WHERE ((%s::uuid IS NULL AND parent_id IS NULL) OR parent_id = %s::uuid)
                            """,
                            (parent_id, parent_id),
                        )
                        sort_order = cur.fetchone()["sort_order"]
                        cur.execute(
                            """
                            INSERT INTO ai_instruction_folders (parent_id, name, content, sort_order)
                            VALUES (%s, %s, '', %s)
                            RETURNING id, name
                            """,
                            (parent_id, name, sort_order),
                        )
                        current = cur.fetchone()
                    parent_id = current["id"]
                    if index == len(path_parts) - 1:
                        cur.execute(
                            """
                            UPDATE ai_instruction_folders
                            SET content = %s, updated_at = now()
                            WHERE id = %s
                            RETURNING id, name, content, updated_at
                            """,
                            (content, current["id"]),
                        )
                        current = cur.fetchone()
    ttl_cache_delete_prefix(("ai_instruction_rows",))
    # Mirror into the git registry so the change actually reaches the agent (reads git).
    try:
        import agent_center  # lazy: avoid import cycle at module load
        agent_center.resync_instructions_to_git()
    except Exception:  # noqa: BLE001
        logger.warning("upsert_ai_instruction: git resync failed", exc_info=True)
    return {"folder": current, "path": " / ".join(path_parts)}


def _compact_owner_report(report: dict[str, Any]) -> dict[str, Any]:
    """Keep only the lightweight summary fields; drop heavy report_text/raw_json.

    The full text stays available through get_owner_reports and
    get_previous_owner_daily_context when the assistant actually needs it.
    """
    return {
        "id": report.get("id"),
        "report_date": report.get("report_date"),
        "period_start": report.get("period_start"),
        "period_end": report.get("period_end"),
        "version": report.get("version"),
        "generated_at": report.get("generated_at"),
        "summary": report.get("summary"),
        "dynamics_summary": report.get("dynamics_summary"),
        "risks_summary": report.get("risks_summary"),
        "recommendations": report.get("recommendations"),
    }


def _compact_company_profile() -> dict[str, Any]:
    """Profile header plus a folder index without the heavy document bodies.

    Full document text stays available through search_company_knowledge and
    get_company_file; bundling every mirrored doc here costs tens of thousands
    of tokens on every export.
    """
    profile = tool_get_company_profile({})
    folders = profile.get("folders") or []
    index = [
        {
            "id": folder.get("id"),
            "parent_id": folder.get("parent_id"),
            "name": folder.get("name"),
            "content_chars": len(folder.get("content") or ""),
            "updated_at": folder.get("updated_at"),
        }
        for folder in folders
    ]
    return {
        "title": profile.get("title"),
        "content": profile.get("content"),
        "updated_at": profile.get("updated_at"),
        "folders_index": index,
        "note": "Folder bodies omitted. Read a document with get_company_file(folder_id) or search with search_company_knowledge.",
    }


def tool_get_compact_export(args: dict[str, Any]) -> dict[str, Any]:
    date_from = parse_date_arg(args, "date_from")
    date_to = parse_date_arg(args, "date_to")
    include_messages = bool(args.get("include_messages", True))
    include_zoom_calls = bool(args.get("include_zoom_calls", True))
    message_limit = parse_limit({"limit": args.get("message_limit", 100)})
    zoom_limit = parse_limit({"limit": args.get("zoom_limit", 50)})

    return {
        "manifest": tool_get_period_index({"date_from": date_from.isoformat(), "date_to": date_to.isoformat()}),
        "company_profile": _compact_company_profile(),
        "org": tool_get_org_structure({"include_inactive": False}),
        "tasks": tool_search_tasks(
            {"date_from": date_from.isoformat(), "date_to": date_to.isoformat(), "limit": args.get("task_limit", 100)}
        )["items"],
        "messages": tool_search_messages(
            {"date_from": date_from.isoformat(), "date_to": date_to.isoformat(), "limit": message_limit}
        )["items"]
        if include_messages
        else [],
        "zoom_calls": tool_list_zoom_calls(
            {
                "date_from": date_from.isoformat(),
                "date_to": date_to.isoformat(),
                "limit": zoom_limit,
            }
        )["items"]
        if include_zoom_calls
        else [],
        "recent_owner_daily_reports": [
            _compact_owner_report(r) for r in tool_get_owner_reports({"report_kind": "daily", "limit": 5})["reports"]
        ],
        "recent_owner_weekly_reports": [
            _compact_owner_report(r) for r in tool_get_owner_reports({"report_kind": "weekly", "limit": 2})["reports"]
        ],
        "notes": [
            "This compact export is read-only and generated on demand from PostgreSQL.",
            "Task descriptions are previews; read one task in full via search_tasks(bitrix_task_id=..., include_full_description=true).",
            "Owner reports here are summaries only; read full report_text via get_owner_reports or get_previous_owner_daily_context when needed.",
            "Zoom calls are available via list_zoom_calls, get_zoom_call_transcript, and search_zoom_transcripts.",
        ],
    }


_GOOGLE_SHEETS_RE = re.compile(r"https?://docs\.google\.com/spreadsheets/d/([a-zA-Z0-9_\-]+)")
_GOOGLE_DOCS_RE = re.compile(r"https?://docs\.google\.com/document/d/([a-zA-Z0-9_\-]+)")
_GOOGLE_GID_RE = re.compile(r"[?&#]gid=(\d+)")


def _rewrite_url_for_export(url: str) -> tuple[str, str]:
    # Non-Google URLs are still fetched directly. Google Docs/Sheets are read through
    # the agent's OAuth account in _fetch_google_url_with_oauth, so private files shared
    # with a9ent.ai@gmail.com work without opening public link access.
    return (url, "raw")


def _google_url_kind(url: str) -> tuple[str, str | None, str | None]:
    m = _GOOGLE_SHEETS_RE.search(url)
    if m:
        gid_match = _GOOGLE_GID_RE.search(url)
        return ("google_sheet_csv", m.group(1), gid_match.group(1) if gid_match else None)
    m = _GOOGLE_DOCS_RE.search(url)
    if m:
        return ("google_doc_text", m.group(1), None)
    return ("raw", None, None)


def _google_creds_for_fetch() -> Any:
    try:
        return app_workflow_function("_google_user_credentials")()
    except Exception as exc:  # noqa: BLE001
        raise McpError(-32010, f"Google OAuth credentials are unavailable: {exc}") from exc


def _csv_text_from_sheet_values(values: list[list[Any]]) -> str:
    out = io.StringIO()
    writer = csv.writer(out)
    for row in values or []:
        writer.writerow(["" if cell is None else cell for cell in row])
    return out.getvalue()


def _extract_google_doc_text(document: dict[str, Any]) -> str:
    parts: list[str] = []

    def walk(elements: Any) -> None:
        if not isinstance(elements, list):
            return
        for el in elements:
            if not isinstance(el, dict):
                continue
            text_run = el.get("textRun")
            if isinstance(text_run, dict):
                parts.append(str(text_run.get("content") or ""))
            paragraph = el.get("paragraph")
            if isinstance(paragraph, dict):
                walk(paragraph.get("elements"))
            table = el.get("table")
            if isinstance(table, dict):
                for row in table.get("tableRows") or []:
                    for cell in (row or {}).get("tableCells") or []:
                        walk((cell or {}).get("content"))
                    parts.append("\n")
            if el.get("sectionBreak") is not None:
                parts.append("\n")

    body = document.get("body") if isinstance(document, dict) else {}
    walk((body or {}).get("content"))
    text = "".join(parts)
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _fetch_google_url_with_oauth(url: str, max_chars: int) -> dict[str, Any] | None:
    kind, file_id, gid = _google_url_kind(url)
    if kind == "raw" or not file_id:
        return None
    try:
        from googleapiclient.discovery import build
        from googleapiclient.errors import HttpError

        creds = _google_creds_for_fetch()
        if kind == "google_sheet_csv":
            sheets = build("sheets", "v4", credentials=creds, cache_discovery=False)
            meta = sheets.spreadsheets().get(
                spreadsheetId=file_id,
                fields="properties(title),sheets(properties(sheetId,title,index))",
            ).execute()
            sheets_props = [s.get("properties") or {} for s in (meta.get("sheets") or [])]
            chosen = None
            if gid is not None:
                for props in sheets_props:
                    if str(props.get("sheetId")) == str(gid):
                        chosen = props
                        break
            if chosen is None and sheets_props:
                chosen = sorted(sheets_props, key=lambda item: item.get("index", 0))[0]
            if not chosen:
                raise McpError(-32010, "Google Sheet has no readable tabs.")
            title = str(chosen.get("title") or "Sheet1")
            safe_title = title.replace("'", "''")
            resp = sheets.spreadsheets().values().get(
                spreadsheetId=file_id,
                range=f"'{safe_title}'",
                valueRenderOption="FORMATTED_VALUE",
            ).execute()
            text = _csv_text_from_sheet_values(resp.get("values") or [])
            return {"status": 200, "content_type": "text/csv; charset=utf-8", "text": text, "final_url": url, "kind": kind}

        drive = build("drive", "v3", credentials=creds, cache_discovery=False)
        exported = drive.files().export(fileId=file_id, mimeType="text/plain").execute()
        if isinstance(exported, bytes):
            text = exported.decode("utf-8", "replace")
        else:
            text = str(exported or "")
        return {"status": 200, "content_type": "text/plain; charset=utf-8", "text": text.strip(), "final_url": url, "kind": kind}
    except McpError:
        raise
    except HttpError as exc:
        status = getattr(getattr(exc, "resp", None), "status", None) or 0
        raise McpError(-32010, f"Google API HTTP {status}: {exc}") from exc
    except Exception as exc:  # noqa: BLE001
        raise McpError(-32010, f"Google API fetch failed: {exc}") from exc


def _strip_html_to_text(html: str) -> str:
    text = re.sub(r"<script\b[^<]*(?:(?!</script>)<[^<]*)*</script>", " ", html, flags=re.IGNORECASE)
    text = re.sub(r"<style\b[^<]*(?:(?!</style>)<[^<]*)*</style>", " ", text, flags=re.IGNORECASE)
    text = re.sub(r"<br\s*/?>", "\n", text, flags=re.IGNORECASE)
    text = re.sub(r"</p\s*>", "\n", text, flags=re.IGNORECASE)
    text = re.sub(r"<[^>]+>", " ", text)
    text = text.replace("&nbsp;", " ").replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">").replace("&quot;", '"').replace("&#39;", "'")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n\s*\n+", "\n\n", text)
    return text.strip()


# --- Deep link reading (parity with the Hermes Brain read-links flow) -------------------------
# Three historical gaps vs the brain's fetch_url.py: a bot-looking UA (anti-bot sites like Dzen
# redirect to auth), no JS rendering, and no binary-document extraction. Fixed below:
# real browser headers; pdf/docx/xlsx by URL are extracted to text locally; and when a public
# page comes back as an auth-wall/JS-shell, we retry through a reader proxy that renders the
# page and returns clean text. Private/internal hosts NEVER go to the external reader.

_FETCH_BROWSER_HEADERS = {
    "User-Agent": ("Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
                   "(KHTML, like Gecko) Chrome/126.0.0.0 Safari/537.36"),
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    "Accept-Language": "ru-RU,ru;q=0.9,en;q=0.5",
}

_BINARY_DOC_EXTS = {"pdf", "docx", "xlsx", "xlsm"}
_BINARY_DOC_CTYPES = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "xlsx",
}
_FETCH_DOC_MAX_BYTES = int(os.getenv("FETCH_URL_DOC_MAX_BYTES", str(12 * 1024 * 1024)) or str(12 * 1024 * 1024))


def _binary_doc_ext(url: str, content_type: str) -> str | None:
    """'pdf'/'docx'/'xlsx' when the URL/Content-Type points at a binary document, else None."""
    ct = (content_type or "").split(";")[0].strip().lower()
    if ct in _BINARY_DOC_CTYPES:
        return _BINARY_DOC_CTYPES[ct]
    path = urlparse(url).path
    from urllib.parse import unquote
    ext = unquote(path).rsplit(".", 1)[-1].lower() if "." in path else ""
    return ext if ext in _BINARY_DOC_EXTS else None


def _extract_binary_document(data: bytes, ext: str) -> str:
    """Extract readable text from pdf/docx/xlsx bytes (same pure-python extractors the chat
    bot uses for inbound attachments). Returns '' when nothing extractable."""
    import io as _io
    try:
        if ext == "pdf":
            from pypdf import PdfReader
            reader = PdfReader(_io.BytesIO(data))
            return "\n".join((pg.extract_text() or "") for pg in reader.pages).strip()
        if ext == "docx":
            from docx import Document
            doc = Document(_io.BytesIO(data))
            parts = [p.text for p in doc.paragraphs]
            for tbl in doc.tables:
                for row in tbl.rows:
                    parts.append(" | ".join(cell.text for cell in row.cells))
            return "\n".join(parts).strip()
        if ext in ("xlsx", "xlsm"):
            from openpyxl import load_workbook
            wb = load_workbook(_io.BytesIO(data), read_only=True, data_only=True)
            out = []
            for ws in wb.worksheets:
                out.append("# Лист: " + str(ws.title))
                for row in ws.iter_rows(values_only=True):
                    cells = ["" if v is None else str(v) for v in row]
                    if any(c.strip() for c in cells):
                        out.append(" | ".join(cells))
            return "\n".join(out).strip()
    except Exception as exc:  # noqa: BLE001
        logging.warning("fetch_url: binary extract failed (%s): %s", ext, repr(exc)[:160])
    return ""


def _reader_allowed_for(url: str) -> bool:
    """Whether the external reader proxy may see this URL. Internal hosts and links that carry
    access tokens (our export links, the Bitrix portal, local addresses) must never leak out."""
    if os.getenv("FETCH_URL_READER", "1").strip().lower() in {"0", "false", "no", "off"}:
        return False
    host = (urlparse(url).netloc or "").lower()
    if not host:
        return False
    private = os.getenv("FETCH_URL_READER_EXCLUDE",
                        "mcp.m4s.ru,m4s.ru,.bitrix24.ru,localhost,127.,10.,192.168.")
    for pat in (p.strip().lower() for p in private.split(",") if p.strip()):
        if host == pat or host.endswith(pat) or host.startswith(pat):
            return False
    return True


def _looks_like_auth_wall(final_url: str, text: str) -> bool:
    """A 200 that is actually a login/consent page or an empty JS shell."""
    host = (urlparse(final_url or "").netloc or "").lower()
    if any(m in host for m in ("passport.", "sso.", "login.", "auth.")):
        return True
    return len((text or "").strip()) < 500


def _fetch_via_reader(url: str, max_chars: int) -> dict[str, Any] | None:
    """Read a JS-heavy / anti-bot page through the reader proxy (renders the page, returns
    markdown text). Best-effort: None on any failure so the caller keeps the direct result."""
    base = os.getenv("FETCH_URL_READER_BASE", "https://r.jina.ai/").strip()
    if not base:
        return None
    try:
        req = urllib.request.Request(base.rstrip("/") + "/" + url, headers=_FETCH_BROWSER_HEADERS)
        with urllib.request.urlopen(req, timeout=60) as resp:
            raw = resp.read(max_chars * 6 + 4096)
        text = raw.decode("utf-8", "replace").strip()
        if len(text) < 200:
            return None
        return {"text": text, "status": 200, "content_type": "text/markdown", "kind": "reader"}
    except Exception as exc:  # noqa: BLE001
        logging.warning("fetch_url: reader fallback failed: %s", repr(exc)[:160])
        return None


def tool_fetch_url(args: dict[str, Any]) -> dict[str, Any]:
    url = str(args.get("url") or "").strip()
    if not url:
        raise McpError(-32602, "url is required.")
    if not re.match(r"^https?://", url):
        raise McpError(-32602, "url must start with http:// or https://.")
    max_chars_raw = args.get("max_chars")
    if max_chars_raw in (None, ""):
        max_chars = 50_000
    else:
        try:
            max_chars = int(max_chars_raw)
        except (TypeError, ValueError) as exc:
            raise McpError(-32602, "max_chars must be a positive integer.") from exc
        if max_chars < 100 or max_chars > 200_000:
            raise McpError(-32602, "max_chars must be between 100 and 200000.")
    strip_html_flag = bool(args.get("strip_html", True))

    fetched_url, kind = _rewrite_url_for_export(url)
    google_result = _fetch_google_url_with_oauth(url, max_chars)
    if google_result is not None:
        status = int(google_result["status"])
        content_type = str(google_result["content_type"])
        text = str(google_result.get("text") or "")
        final_url = str(google_result.get("final_url") or url)
        kind = str(google_result.get("kind") or kind)
    else:
        request = urllib.request.Request(fetched_url, headers=dict(_FETCH_BROWSER_HEADERS), method="GET")
        try:
            with urllib.request.urlopen(request, timeout=30) as response:
                status = response.status
                content_type = response.headers.get("Content-Type", "") or ""
                doc_ext = _binary_doc_ext(fetched_url, content_type)
                # Binary documents need the whole file for extraction; text needs only a slice.
                raw_bytes = response.read(_FETCH_DOC_MAX_BYTES if doc_ext else max_chars * 6 + 4096)
                final_url = response.geturl() or fetched_url
        except urllib.error.HTTPError as exc:
            try:
                body_preview = exc.read().decode("utf-8", "replace")[:500]
            except Exception:  # noqa: BLE001
                body_preview = ""
            # A hard 4xx/5xx on a public page is often anti-bot — the reader proxy still works.
            if _reader_allowed_for(url):
                reader = _fetch_via_reader(url, max_chars)
                if reader:
                    text = reader["text"][:max_chars]
                    return {
                        "ok": True, "original_url": url, "fetched_url": fetched_url,
                        "final_url": url, "kind": "reader", "status": 200,
                        "content_type": reader["content_type"], "char_count": len(text),
                        "truncated": len(reader["text"]) > max_chars, "text": text,
                        "note": f"Прямой запрос вернул HTTP {exc.code}; содержимое получено через reader-прокси.",
                    }
            return {
                "ok": False,
                "original_url": url,
                "fetched_url": fetched_url,
                "kind": kind,
                "status": exc.code,
                "error": f"HTTP {exc.code}",
                "body_preview": body_preview,
                "hint": "",
            }
        except Exception as exc:  # noqa: BLE001
            raise McpError(-32010, f"Fetch failed: {exc}") from exc

        # Binary document (Word/PDF/Excel by URL) -> extract full text locally.
        doc_ext = _binary_doc_ext(final_url or fetched_url, content_type)
        if doc_ext:
            doc_text = _extract_binary_document(raw_bytes, doc_ext)
            if doc_text:
                truncated = len(doc_text) > max_chars
                return {
                    "ok": True, "original_url": url, "fetched_url": fetched_url,
                    "final_url": final_url, "kind": f"document-{doc_ext}", "status": status,
                    "content_type": content_type, "char_count": min(len(doc_text), max_chars),
                    "truncated": truncated, "text": doc_text[:max_chars],
                    "note": ("Это бинарный документ; извлечён его текст."
                             + (" Показано начало — вызови ещё раз с большим max_chars." if truncated else "")),
                }
            return {
                "ok": False, "original_url": url, "fetched_url": fetched_url, "kind": f"document-{doc_ext}",
                "status": status, "error": "binary document without extractable text",
                "hint": "Файл скачан, но текст извлечь не удалось (возможно скан без текстового слоя).",
            }

        charset_match = re.search(r"charset=([a-zA-Z0-9_\-]+)", content_type)
        charset = charset_match.group(1) if charset_match else "utf-8"
        try:
            text = raw_bytes.decode(charset, errors="replace")
        except LookupError:
            text = raw_bytes.decode("utf-8", errors="replace")

    looks_html = ("html" in content_type.lower()) or text.lstrip().lower().startswith(("<!doctype", "<html"))
    if strip_html_flag and looks_html and kind == "raw":
        text = _strip_html_to_text(text)

    # Auth-wall / empty JS shell on a PUBLIC page -> render it through the reader proxy.
    if kind == "raw" and _looks_like_auth_wall(final_url, text) and _reader_allowed_for(url):
        reader = _fetch_via_reader(url, max_chars)
        if reader:
            text = reader["text"]
            kind = "reader"
            content_type = reader["content_type"]
            status = 200
            final_url = url

    truncated = len(text) > max_chars
    if truncated:
        text = text[:max_chars]

    return {
        "ok": True,
        "original_url": url,
        "fetched_url": fetched_url,
        "final_url": final_url,
        "kind": kind,
        "status": status,
        "content_type": content_type,
        "char_count": len(text),
        "truncated": truncated,
        "text": text,
    }


def _b24_bot_user_names(user_ids: Any) -> dict[Any, str]:
    """Resolve bot-portal user ids -> 'First Last' via the bot's own Bitrix webhook (cached).
    The synced ``users`` table holds the main portal; the chat-bot may run on a different
    portal where ids differ, so names there won't match — fall back to user.get."""
    out: dict[Any, str] = {}
    base = (shared_load_env_value("B24_TESTBOT_WEBHOOK_BASE") or "").strip().rstrip("/")
    if not base:
        return out
    import urllib.parse
    for uid in {u for u in user_ids if u not in (None, "")}:
        ck = ("b24_user_name", str(uid))
        cached = ttl_cache_get(ck)
        if cached is not None:
            out[uid] = cached
            continue
        name = ""
        try:
            data = urllib.parse.urlencode({"ID": uid}).encode()
            with urllib.request.urlopen(base + "/user.get.json", data=data, timeout=15) as resp:
                res = json.loads(resp.read().decode()).get("result") or []
            if res:
                u0 = res[0]
                name = " ".join(x for x in (u0.get("NAME"), u0.get("LAST_NAME")) if x).strip() or (u0.get("EMAIL") or "")
        except Exception:  # noqa: BLE001
            name = ""
        ttl_cache_set(ck, name, 3600)
        out[uid] = name
    return out


def _enrich_bot_user_names(items: list[dict[str, Any]]) -> None:
    missing = [it.get("bitrix_user_id") for it in items
               if not it.get("full_name") and it.get("bitrix_user_id") not in (None, "")]
    if not missing:
        return
    names = _b24_bot_user_names(missing)
    for it in items:
        if not it.get("full_name"):
            nm = names.get(it.get("bitrix_user_id"))
            if nm:
                it["full_name"] = nm


def tool_list_bitrix_bot_sessions(args: dict[str, Any]) -> dict[str, Any]:
    """Overview of conversations employees had with the AI assistant inside Bitrix24 chat —
    one row per dialog/user from bitrix_bot_interactions (the bot's own log)."""
    limit = parse_limit(args, 50)
    date_from = parse_date_arg(args, "date_from", required=False)
    date_to = parse_date_arg(args, "date_to", required=False)
    where: list[str] = []
    params: dict[str, Any] = {"limit": limit}
    if date_from:
        where.append("i.created_at >= %(df)s")
        params["df"] = date_from
    if date_to:
        where.append("i.created_at < (%(dt)s::date + 1)")
        params["dt"] = date_to
    where_sql = ("WHERE " + " AND ".join(where)) if where else ""
    with connect() as conn:
        with conn.cursor() as cur:
            if not safe_table_exists(cur, "bitrix_bot_interactions"):
                return {"items": [], "limit": limit, "message": "bitrix_bot_interactions table does not exist yet."}
            cur.execute(
                f"""
                SELECT
                    i.dialog_id,
                    i.bitrix_user_id,
                    u.full_name,
                    u.work_position,
                    count(*) AS messages,
                    count(*) FILTER (WHERE i.status <> 'ok') AS errors,
                    min(i.created_at) AS first_at,
                    max(i.created_at) AS last_at,
                    array_remove(array_agg(DISTINCT i.tier), NULL) AS tiers,
                    round(avg(i.latency_ms))::int AS avg_latency_ms,
                    max(s.epoch) AS epoch,
                    max(s.turns) AS turns
                FROM bitrix_bot_interactions i
                LEFT JOIN users u ON u.bitrix_user_id = i.bitrix_user_id
                LEFT JOIN bitrix_bot_sessions s ON s.dialog_id = i.dialog_id
                {where_sql}
                GROUP BY i.dialog_id, i.bitrix_user_id, u.full_name, u.work_position
                ORDER BY max(i.created_at) DESC
                LIMIT %(limit)s
                """,
                params,
            )
            items = cur.fetchall()
    _enrich_bot_user_names(items)
    return {"items": items, "limit": limit,
            "note": "Conversations employees had with the AI assistant in Bitrix24. "
                    "Use get_bitrix_bot_chat(dialog_id|bitrix_user_id) to read a full transcript."}


def tool_get_bitrix_bot_chat(args: dict[str, Any]) -> dict[str, Any]:
    """Full question→answer transcript of one person's conversation with the AI assistant in
    Bitrix24 (by dialog_id or bitrix_user_id), for quality analysis."""
    limit = parse_limit(args, 100)
    dialog_id = str(args.get("dialog_id") or "").strip()
    raw_user = args.get("bitrix_user_id")
    query = str(args.get("query") or "").strip()
    date_from = parse_date_arg(args, "date_from", required=False)
    date_to = parse_date_arg(args, "date_to", required=False)
    if not dialog_id and raw_user in (None, ""):
        raise McpError(-32602, "Provide dialog_id or bitrix_user_id")
    where: list[str] = []
    params: dict[str, Any] = {"limit": limit}
    if dialog_id:
        where.append("i.dialog_id = %(d)s")
        params["d"] = dialog_id
    if raw_user not in (None, ""):
        try:
            params["u"] = int(raw_user)
        except (TypeError, ValueError) as exc:
            raise McpError(-32602, "bitrix_user_id must be an integer") from exc
        where.append("i.bitrix_user_id = %(u)s")
    if query:
        where.append("(i.question ILIKE %(q)s OR i.answer ILIKE %(q)s)")
        params["q"] = f"%{query}%"
    if date_from:
        where.append("i.created_at >= %(df)s")
        params["df"] = date_from
    if date_to:
        where.append("i.created_at < (%(dt)s::date + 1)")
        params["dt"] = date_to
    where_sql = "WHERE " + " AND ".join(where)
    with connect() as conn:
        with conn.cursor() as cur:
            if not safe_table_exists(cur, "bitrix_bot_interactions"):
                return {"items": [], "limit": limit, "message": "bitrix_bot_interactions table does not exist yet."}
            cur.execute(
                f"""
                SELECT i.id, i.created_at, i.dialog_id, i.bitrix_user_id, u.full_name,
                       i.tier, i.question, i.answer, i.latency_ms, i.status, i.error
                FROM bitrix_bot_interactions i
                LEFT JOIN users u ON u.bitrix_user_id = i.bitrix_user_id
                {where_sql}
                ORDER BY i.id DESC
                LIMIT %(limit)s
                """,
                params,
            )
            rows = cur.fetchall()
    rows.reverse()  # chronological order for reading
    _enrich_bot_user_names(rows)
    return {"items": rows, "limit": limit}


def tool_get_ai_capabilities(args: dict[str, Any]) -> dict[str, Any]:
    """Return what the assistant can do for the caller's current connector/tool set.
    The legacy full/faq key is still used only to select the stored note."""
    connector_id = args.get("_connector_id") or "full"
    tier = "faq" if connector_id == "faq" else "full"
    with connect() as conn:
        with conn.cursor() as cur:
            if not safe_table_exists(cur, "ai_agent_capabilities"):
                return {"tier": tier, "capabilities": "", "message": "ai_agent_capabilities table does not exist yet."}
            cur.execute("SELECT content, updated_at, updated_by FROM ai_agent_capabilities WHERE tier = %s", (tier,))
            row = cur.fetchone()
    can_edit = tier == "full"
    return {
        "tier": tier,
        "capabilities": row["content"] if row else "",
        "updated_at": row["updated_at"] if row else None,
        "note": (
            f"Это твои возможности для текущего коннектора/набора инструментов ({tier}). Отвечай пользователю "
            "СТРОГО в рамках этого списка — не обещай того, чего тут нет. "
            + ("Если узнал/получил новую возможность — дополни список через update_ai_capabilities."
               if can_edit else "Обновлять список из этого коннектора нельзя.")
        ),
    }


def tool_update_ai_capabilities(args: dict[str, Any]) -> dict[str, Any]:
    """Self-update the capabilities note when this tool is enabled for the connector.
    mode 'append' (default) adds a line; 'replace' overwrites. Stored note key:
    legacy 'full' (default) or 'faq'."""
    content = str(args.get("content") or "").strip()
    if not content:
        raise McpError(-32602, "content is required")
    tier = str(args.get("tier") or "full").strip().lower()
    if tier not in ("full", "faq"):
        raise McpError(-32602, "tier must be 'full' or 'faq'")
    mode = str(args.get("mode") or "append").strip().lower()
    if mode not in ("append", "replace"):
        raise McpError(-32602, "mode must be 'append' or 'replace'")
    with connect() as conn:
        with conn.cursor() as cur:
            cur.execute("SELECT content FROM ai_agent_capabilities WHERE tier = %s", (tier,))
            row = cur.fetchone()
            existing = (row["content"] if row else "") or ""
            new_content = content if mode == "replace" else ((existing + "\n" + content).strip() if existing else content)
            cur.execute(
                """
                INSERT INTO ai_agent_capabilities (tier, content, updated_at, updated_by)
                VALUES (%s, %s, now(), 'agent')
                ON CONFLICT (tier) DO UPDATE
                  SET content = EXCLUDED.content, updated_at = now(), updated_by = 'agent'
                """,
                (tier, new_content),
            )
    return {"tier": tier, "mode": mode, "ok": True, "content": new_content}


def _b24_webhook_call(method: str, params: dict[str, Any] | None = None) -> dict[str, Any]:
    """Call the bot-portal (b24-0xrp3s) incoming webhook — that is the portal the company is moving
    to and where the absence chart («График отсутствий») + calendar scope live. Returns the parsed
    Bitrix response dict; raises McpError on transport/API error."""
    base = (shared_load_env_value("B24_TESTBOT_WEBHOOK_BASE") or "").strip().rstrip("/")
    if not base:
        raise McpError(-32011, "B24_TESTBOT_WEBHOOK_BASE (bot-portal webhook) is not configured.")
    import urllib.parse
    qs = urllib.parse.urlencode(params or {}, doseq=True)
    url = f"{base}/{method}.json" + (f"?{qs}" if qs else "")
    try:
        with urllib.request.urlopen(url, timeout=30) as resp:
            data = json.loads(resp.read().decode("utf-8"))
    except urllib.error.HTTPError as exc:
        detail = exc.read().decode("utf-8", "replace")[:300]
        raise McpError(-32011, f"{method}: HTTP {exc.code} {detail}")
    except Exception as exc:  # noqa: BLE001
        raise McpError(-32011, f"{method}: {str(exc)[:200]}")
    if isinstance(data, dict) and data.get("error"):
        raise McpError(-32011, f"{method}: {data.get('error')} {data.get('error_description') or ''}".strip())
    return data if isinstance(data, dict) else {"result": data}


def _b24_active_users() -> list[dict[str, Any]]:
    """Active users on the bot portal (b24-0xrp3s), cached 1h. Each: id, full_name, position."""
    cached = ttl_cache_get(("b24_active_users",))
    if cached is not None:
        return cached
    users: list[dict[str, Any]] = []
    start = 0
    for _ in range(20):
        data = _b24_webhook_call("user.get", {"ACTIVE": "true", "start": start})
        rows = data.get("result") or []
        if not isinstance(rows, list) or not rows:
            break
        for u in rows:
            users.append({
                "id": u.get("ID"),
                "full_name": " ".join(x for x in (u.get("NAME"), u.get("LAST_NAME")) if x).strip(),
                "position": u.get("WORK_POSITION") or "",
            })
        nxt = data.get("next")
        if nxt in (None, "", start):
            break
        start = nxt
    ttl_cache_set(("b24_active_users",), users, 3600)
    return users


def _b24_absence_periods(user_ids: list[Any], date_from: str, date_to: str) -> dict[str, list[dict[str, Any]]]:
    """{user_id: [absence periods]} from calendar.accessibility.get. Keeps only entries the absence
    chart marks as 'absent' (vacation/sick/trip); regular busy meetings ('busy') are ignored."""
    if not user_ids:
        return {}
    params = {"from": date_from, "to": date_to, "users[]": [str(u) for u in user_ids]}
    data = _b24_webhook_call("calendar.accessibility.get", params)
    result = data.get("result") or {}
    out: dict[str, list[dict[str, Any]]] = {}
    if isinstance(result, dict):
        for uid, entries in result.items():
            periods = []
            for e in (entries or []):
                if not isinstance(e, dict):
                    continue
                acc = str(e.get("ACCESSIBILITY") or e.get("accessibility") or "").lower()
                name = str(e.get("NAME") or e.get("name") or "")
                is_absence = acc == "absent" or "отпуск" in name.lower() or "vacation" in name.lower()
                if not is_absence:
                    continue
                periods.append({
                    "from": e.get("DATE_FROM") or e.get("dateFrom"),
                    "to": e.get("DATE_TO") or e.get("dateTo"),
                    "name": name or "Отсутствие",
                    "accessibility": acc,
                })
            out[str(uid)] = periods
    return out


def tool_get_agent_monitoring(args: dict[str, Any]) -> dict[str, Any]:
    """Live self-monitoring + usage accounting (same data as the Центр Агента pages):
    health of every integration, speed chart, event feed, and per-employee usage
    (turns, time with the agent, real token spend from the Hermes session store)."""
    period = str(args.get("period") or "7")
    return app_workflow_function("agent_center_report")(period)


def tool_get_employee_absences(args: dict[str, Any]) -> dict[str, Any]:
    """Whether employees are on vacation/absent per the Bitrix «График отсутствий» (bot portal
    b24-0xrp3s). Resolves by bitrix_user_id or employee_name; with neither, returns everyone who is
    absent in the period. Use date_from=date_to to check a specific date (e.g. a task deadline)."""
    df = parse_date_arg(args, "date_from", required=False)
    dt = parse_date_arg(args, "date_to", required=False)
    today = date.today()
    date_from = (df or today).isoformat()
    date_to = (dt or df or today).isoformat()

    bitrix_user_id = args.get("bitrix_user_id")
    employee_name = str(args.get("employee_name") or "").strip()
    actives = _b24_active_users()

    if bitrix_user_id not in (None, ""):
        try:
            uid = int(bitrix_user_id)
        except (TypeError, ValueError) as exc:
            raise McpError(-32602, "bitrix_user_id must be an integer.") from exc
        match = next((u for u in actives if str(u["id"]) == str(uid)), None)
        targets = [match or {"id": uid, "full_name": "", "position": ""}]
    elif employee_name:
        matches = [u for u in actives if _person_names_match(u["full_name"], employee_name)]
        if not matches:
            return {"query": {"employee_name": employee_name}, "matched": 0,
                    "note": "Сотрудник не найден среди активных на портале b24-0xrp3s.",
                    "candidates": [u["full_name"] for u in actives if u["full_name"]][:50]}
        if len(matches) > 1:
            return {"query": {"employee_name": employee_name}, "matched": len(matches), "ambiguous": True,
                    "candidates": [{"bitrix_user_id": u["id"], "full_name": u["full_name"]} for u in matches]}
        targets = matches
    else:
        targets = [u for u in actives if u["full_name"]]

    ids = [u["id"] for u in targets if u.get("id") not in (None, "")]
    absences = _b24_absence_periods(ids, date_from, date_to)
    results = []
    for u in targets:
        periods = absences.get(str(u["id"]), [])
        results.append({
            "bitrix_user_id": u["id"],
            "full_name": u["full_name"],
            "position": u.get("position") or "",
            "on_vacation": bool(periods),
            "absences": periods,
        })
    scanning_all = bitrix_user_id in (None, "") and not employee_name
    if scanning_all:
        results = [r for r in results if r["on_vacation"]]
    return {
        "portal": "b24-0xrp3s",
        "date_from": date_from,
        "date_to": date_to,
        "count": len(results),
        "employees": results,
        "note": ("Источник — Bitrix «График отсутствий» (calendar.accessibility.get, accessibility='absent'). "
                 "Пустой список при проверке конкретного сотрудника = он НЕ в отпуске в этот период."),
    }


# --- CRM: воронки сделок (deal pipelines) -------------------------------------------------------
# The incoming webhooks have no `crm` scope (probed live 2026-07-08), but the bot local-app OAuth
# token does — every funnel/deal tool calls Bitrix through b24bot.b24_app_method_call (auto-refresh).
# CRM lives on the bot portal (b24-0xrp3s): user ids here are bot-portal ids, the same id space
# _b24_active_users() / get_employee_absences use. userfieldconfig.* is NOT allowed for this token —
# custom deal fields go through the classic crm.deal.userfield.* API.

DEAL_ENTITY_TYPE_ID = 2
_CRM_COLOR_RE = re.compile(r"^#[0-9A-Fa-f]{6}$")
_CRM_CODE_SANITIZE_RE = re.compile(r"[^A-Z0-9_]+")
_CRM_DEAL_LIST_SELECT = [
    "ID", "TITLE", "STAGE_ID", "CATEGORY_ID", "OPPORTUNITY", "CURRENCY_ID", "ASSIGNED_BY_ID",
    "COMPANY_ID", "CONTACT_ID", "DATE_CREATE", "BEGINDATE", "CLOSEDATE", "CLOSED", "COMMENTS",
]


def _crm_call(method: str, payload: dict[str, Any] | None = None) -> dict[str, Any]:
    try:
        return app_workflow_function("b24_app_method_call")(method, payload or {})
    except McpError:
        raise
    except Exception as exc:  # noqa: BLE001
        raise McpError(-32012, f"Bitrix CRM {method}: {str(exc)[:300]}") from exc


def _crm_portal_deal_url(deal_id: int) -> str:
    base = (shared_load_env_value("B24_TESTBOT_WEBHOOK_BASE") or "").strip()
    host = urlparse(base).netloc if base else ""
    return f"https://{host}/crm/deal/details/{int(deal_id)}/" if host else ""


def _crm_categories() -> list[dict[str, Any]]:
    res = _crm_call("crm.category.list", {"entityTypeId": DEAL_ENTITY_TYPE_ID}).get("result") or {}
    cats = (res.get("categories") or []) if isinstance(res, dict) else []
    return sorted(cats, key=lambda c: (int(c.get("sort") or 0), int(c.get("id") or 0)))


def _crm_stage_entity(category_id: int) -> str:
    return "DEAL_STAGE" if int(category_id) == 0 else f"DEAL_STAGE_{int(category_id)}"


def _crm_stages(category_id: int) -> list[dict[str, Any]]:
    rows = _crm_call(
        "crm.status.list", {"filter": {"ENTITY_ID": _crm_stage_entity(category_id)}}
    ).get("result") or []
    stages = []
    for r in rows:
        extra = r.get("EXTRA") or {}
        stages.append({
            "id": int(r.get("ID") or 0),
            "stage_id": r.get("STATUS_ID"),
            "name": r.get("NAME"),
            "sort": int(r.get("SORT") or 0),
            "color": extra.get("COLOR") or r.get("COLOR"),
            "semantics": extra.get("SEMANTICS") or r.get("SEMANTICS") or "process",
            "system": str(r.get("SYSTEM") or "") == "Y",
        })
    return sorted(stages, key=lambda s: s["sort"])


def _crm_resolve_category(args: dict[str, Any], required: bool = True) -> dict[str, Any] | None:
    """Resolve a pipeline by category_id (exact) or pipeline_name (case-insensitive, then substring).
    Ambiguity / not found -> readable refusal listing what exists."""
    cats = _crm_categories()
    cid = args.get("category_id")
    if cid not in (None, ""):
        try:
            cid = int(cid)
        except (TypeError, ValueError) as exc:
            raise McpError(-32602, "category_id must be an integer.") from exc
        match = next((c for c in cats if str(c.get("id")) == str(cid)), None)
        if not match:
            raise McpError(-32602, "Воронка id=%s не найдена. Существуют: %s" % (
                cid, "; ".join(f"{c['id']} — {c.get('name')}" for c in cats)))
        return match
    name = str(args.get("pipeline_name") or "").strip()
    if not name:
        if required:
            raise McpError(-32602, "Укажи category_id или pipeline_name (список — list_crm_pipelines).")
        return None
    nl = name.casefold()
    matches = [c for c in cats if str(c.get("name") or "").casefold() == nl]
    if not matches:
        matches = [c for c in cats if nl in str(c.get("name") or "").casefold()]
    if not matches:
        raise McpError(-32602, "Воронка «%s» не найдена. Существуют: %s" % (
            name, "; ".join(f"{c['id']} — {c.get('name')}" for c in cats)))
    if len(matches) > 1:
        raise McpError(-32602, "Название «%s» неоднозначно: %s. Уточни category_id." % (
            name, "; ".join(f"{c['id']} — {c.get('name')}" for c in matches)))
    return matches[0]


def _crm_resolve_stage(category_id: int, ref: str) -> dict[str, Any]:
    """Resolve a stage inside a pipeline by STATUS_ID (full 'C8:NEW' or bare 'NEW') or by name."""
    ref = str(ref or "").strip()
    if not ref:
        raise McpError(-32602, "Стадия не указана.")
    stages = _crm_stages(category_id)
    rl = ref.casefold()
    for s in stages:
        sid = str(s.get("stage_id") or "")
        bare = sid.split(":", 1)[1] if ":" in sid else sid
        if rl in (sid.casefold(), bare.casefold()):
            return s
    matches = [s for s in stages if str(s.get("name") or "").casefold() == rl]
    if not matches:
        matches = [s for s in stages if rl in str(s.get("name") or "").casefold()]
    if not matches:
        raise McpError(-32602, "Стадия «%s» не найдена в воронке %s. Стадии: %s" % (
            ref, category_id, "; ".join(f"{s['stage_id']} — {s['name']}" for s in stages)))
    if len(matches) > 1:
        raise McpError(-32602, "Стадия «%s» неоднозначна: %s" % (
            ref, "; ".join(f"{s['stage_id']} — {s['name']}" for s in matches)))
    return matches[0]


def _crm_resolve_portal_user(args: dict[str, Any], id_key: str, name_key: str) -> int | None:
    """Bot-portal user by id or fuzzy name (same id space as get_employee_absences)."""
    uid = args.get(id_key)
    if uid not in (None, ""):
        try:
            return int(uid)
        except (TypeError, ValueError) as exc:
            raise McpError(-32602, f"{id_key} must be an integer.") from exc
    name = str(args.get(name_key) or "").strip()
    if not name:
        return None
    actives = _b24_active_users()
    matches = [u for u in actives if _person_names_match(u["full_name"], name)]
    if not matches:
        raise McpError(-32602, "Сотрудник «%s» не найден среди активных на портале. Есть: %s" % (
            name, ", ".join(u["full_name"] for u in actives if u["full_name"])[:1500]))
    if len(matches) > 1:
        raise McpError(-32602, "Имя «%s» неоднозначно: %s" % (
            name, "; ".join(f"{u['id']} — {u['full_name']}" for u in matches)))
    return int(matches[0]["id"])


def _crm_custom_fields_arg(args: dict[str, Any]) -> dict[str, Any]:
    custom = args.get("custom_fields")
    if custom in (None, ""):
        return {}
    if not isinstance(custom, dict):
        raise McpError(-32602, "custom_fields должен быть объектом {\"UF_CRM_...\": значение}.")
    bad = [k for k in custom if not str(k).upper().startswith("UF_")]
    if bad:
        raise McpError(-32602, "custom_fields: коды должны начинаться с UF_ (получены: %s). "
                               "Реальные коды — list_crm_deal_fields." % ", ".join(map(str, bad)))
    return {str(k).upper(): v for k, v in custom.items()}


def _crm_deal_brief(row: dict[str, Any], names: dict[str, Any] | None = None) -> dict[str, Any]:
    names = names or {}
    deal_id = int(row.get("ID") or 0)
    brief = {
        "deal_id": deal_id,
        "title": row.get("TITLE"),
        "category_id": int(row.get("CATEGORY_ID") or 0),
        "pipeline_name": names.get("pipeline_name"),
        "stage_id": row.get("STAGE_ID"),
        "stage_name": names.get("stage_name"),
        "amount": row.get("OPPORTUNITY"),
        "currency": row.get("CURRENCY_ID"),
        "assigned_by_id": row.get("ASSIGNED_BY_ID"),
        "assigned_name": names.get("assigned_name"),
        "closed": str(row.get("CLOSED") or "") == "Y",
        "date_create": row.get("DATE_CREATE"),
        "close_date": row.get("CLOSEDATE"),
        "url": _crm_portal_deal_url(deal_id) if deal_id else "",
    }
    return {k: v for k, v in brief.items() if v not in (None, "")}


def _crm_deal_enrich_names(rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Per-row {pipeline_name, stage_name, assigned_name} lookups, каждое — best-effort."""
    cat_names: dict[int, str] = {}
    stage_names: dict[int, dict[str, str]] = {}
    try:
        cat_names = {int(c["id"]): str(c.get("name") or "") for c in _crm_categories()}
    except McpError:
        pass
    users: dict[str, str] = {}
    try:
        users = {str(u["id"]): u["full_name"] for u in _b24_active_users()}
    except Exception:  # noqa: BLE001
        pass
    out = []
    for row in rows:
        cid = int(row.get("CATEGORY_ID") or 0)
        if cid not in stage_names:
            try:
                stage_names[cid] = {s["stage_id"]: s["name"] for s in _crm_stages(cid)}
            except McpError:
                stage_names[cid] = {}
        out.append({
            "pipeline_name": cat_names.get(cid),
            "stage_name": stage_names[cid].get(row.get("STAGE_ID")),
            "assigned_name": users.get(str(row.get("ASSIGNED_BY_ID") or "")),
        })
    return out


def _crm_stage_code(spec_code: Any, name: str) -> str:
    code = _CRM_CODE_SANITIZE_RE.sub("_", str(spec_code or "").strip().upper()).strip("_")
    if not code:
        latin = _CRM_CODE_SANITIZE_RE.sub("_", str(name or "").upper()).strip("_")
        code = latin[:20] if latin else ""
    return code or f"S{int(time.time()) % 10**8}"


def _crm_add_stage(category_id: int, spec: dict[str, Any]) -> dict[str, Any]:
    name = str(spec.get("name") or "").strip()
    if not name:
        raise McpError(-32602, "У каждой стадии обязателен name.")
    stages = _crm_stages(category_id)
    code = _crm_stage_code(spec.get("stage_code"), name)
    status_id = code if int(category_id) == 0 else f"C{int(category_id)}:{code}"
    if any(str(s.get("stage_id") or "").casefold() == status_id.casefold() for s in stages):
        raise McpError(-32602, f"Стадия с кодом {status_id} уже существует — задай другой stage_code.")
    semantics_failure = str(spec.get("semantics") or "").strip().lower() in ("failure", "fail", "провал")
    sort = spec.get("sort")
    if sort in (None, ""):
        final = [s["sort"] for s in stages if s.get("semantics") in ("success", "failure")]
        process = [s["sort"] for s in stages if s.get("semantics") not in ("success", "failure")]
        if semantics_failure:
            # Bitrix requires a losing stage to sort AFTER the final stages.
            sort = (max([s["sort"] for s in stages]) + 10) if stages else 10
        else:
            cand = (max(process) + 10) if process else 10
            sort = cand if not final or cand < min(final) else max(min(final) - 1, 1)
    fields: dict[str, Any] = {
        "ENTITY_ID": _crm_stage_entity(category_id),
        "STATUS_ID": status_id,
        "NAME": name,
        "SORT": int(sort),
    }
    color = str(spec.get("color") or "").strip()
    if color:
        if not _CRM_COLOR_RE.match(color):
            raise McpError(-32602, "color должен быть в формате #RRGGBB.")
        fields["COLOR"] = color
    semantics = str(spec.get("semantics") or "").strip().lower()
    if semantics_failure:
        fields["SEMANTICS"] = "F"  # extra losing stage; success beyond the system WON is not supported
    elif semantics and semantics != "process":
        raise McpError(-32602, "semantics: только 'process' (обычная) или 'failure' (доп. проигрышная).")
    row_id = _crm_call("crm.status.add", {"fields": fields}).get("result")
    return {"id": row_id, "stage_id": status_id, "name": name, "sort": int(sort)}


def tool_list_crm_pipelines(args: dict[str, Any]) -> dict[str, Any]:
    include_stages = args.get("include_stages") is not False
    include_counts = args.get("include_deal_counts") is not False
    pipelines = []
    for c in _crm_categories():
        cid = int(c.get("id") or 0)
        item: dict[str, Any] = {
            "category_id": cid,
            "name": c.get("name"),
            "sort": int(c.get("sort") or 0),
            "is_default": str(c.get("isDefault") or "") == "Y",
        }
        if include_stages:
            item["stages"] = _crm_stages(cid)
        if include_counts:
            r = _crm_call("crm.deal.list", {"filter": {"CATEGORY_ID": cid}, "select": ["ID"]})
            item["deals_total"] = r.get("total", len(r.get("result") or []))
        pipelines.append(item)
    return {"portal": "b24-0xrp3s", "count": len(pipelines), "pipelines": pipelines}


def tool_create_crm_pipeline(args: dict[str, Any]) -> dict[str, Any]:
    name = str(args.get("name") or "").strip()
    if not name:
        raise McpError(-32602, "name (название воронки) обязателен.")
    dup = [c for c in _crm_categories() if str(c.get("name") or "").casefold() == name.casefold()]
    if dup and args.get("allow_duplicate_name") is not True:
        raise McpError(-32602, "Воронка «%s» уже существует (id=%s). Если дубль нужен намеренно — "
                               "передай allow_duplicate_name=true." % (name, dup[0].get("id")))
    fields: dict[str, Any] = {"name": name}
    if args.get("sort") not in (None, ""):
        fields["sort"] = int(args["sort"])
    res = _crm_call("crm.category.add", {"entityTypeId": DEAL_ENTITY_TYPE_ID, "fields": fields})
    cat = ((res.get("result") or {}).get("category") or {}) if isinstance(res.get("result"), dict) else {}
    cid = int(cat.get("id") or 0)
    added, stage_errors = [], []
    for spec in (args.get("stages") or []):
        if isinstance(spec, str):
            spec = {"name": spec}
        try:
            added.append(_crm_add_stage(cid, spec))
        except McpError as exc:
            stage_errors.append(str(exc))
    out = {
        "created": True,
        "category_id": cid,
        "name": cat.get("name"),
        "stages": _crm_stages(cid),
        "note": "Bitrix автоматически создаёт стандартный набор стадий; лишние можно удалить "
                "через manage_crm_pipeline_stage (системные «Сделка успешна»/«Сделка провалена» не удаляются).",
    }
    if added:
        out["stages_added"] = added
    if stage_errors:
        out["stage_errors"] = stage_errors
    return out


def tool_update_crm_pipeline(args: dict[str, Any]) -> dict[str, Any]:
    cat = _crm_resolve_category(args)
    fields: dict[str, Any] = {}
    if str(args.get("new_name") or "").strip():
        fields["name"] = str(args["new_name"]).strip()
    if args.get("sort") not in (None, ""):
        fields["sort"] = int(args["sort"])
    if not fields:
        raise McpError(-32602, "Нечего менять: укажи new_name и/или sort.")
    res = _crm_call("crm.category.update",
                    {"entityTypeId": DEAL_ENTITY_TYPE_ID, "id": int(cat["id"]), "fields": fields})
    updated = ((res.get("result") or {}).get("category") or {}) if isinstance(res.get("result"), dict) else {}
    return {"updated": True, "category_id": int(cat["id"]),
            "name": updated.get("name") or fields.get("name") or cat.get("name")}


def tool_delete_crm_pipeline(args: dict[str, Any]) -> dict[str, Any]:
    if args.get("confirm") is not True:
        raise McpError(-32602, "Удаление воронки требует подтверждения: покажи пользователю точную "
                               "воронку (id, название, сколько в ней сделок) и вызови с confirm=true.")
    cat = _crm_resolve_category(args)
    cid = int(cat["id"])
    if str(cat.get("isDefault") or "") == "Y" or cid == 0:
        raise McpError(-32602, "Основную (default) воронку удалять нельзя.")
    expected = str(args.get("expected_name") or "").strip()
    if expected and expected.casefold() != str(cat.get("name") or "").casefold():
        raise McpError(-32602, "expected_name не совпал: воронка id=%s называется «%s». Удаление отменено."
                       % (cid, cat.get("name")))
    r = _crm_call("crm.deal.list", {"filter": {"CATEGORY_ID": cid}, "select": ["ID"]})
    total = r.get("total", len(r.get("result") or []))
    if total:
        raise McpError(-32602, "В воронке «%s» осталось сделок: %s. Сначала перенеси их в другую "
                               "воронку (update_crm_deal) или удали — потом удаляй воронку." % (cat.get("name"), total))
    _crm_call("crm.category.delete", {"entityTypeId": DEAL_ENTITY_TYPE_ID, "id": cid})
    return {"deleted": True, "category_id": cid, "name": cat.get("name")}


def tool_manage_crm_pipeline_stage(args: dict[str, Any]) -> dict[str, Any]:
    action = str(args.get("action") or "").strip().lower()
    if action not in ("add", "update", "delete"):
        raise McpError(-32602, "action: add | update | delete.")
    if action == "delete" and args.get("confirm") is not True:
        raise McpError(-32602, "Удаление стадии требует подтверждения: покажи пользователю стадию "
                               "и воронку, затем вызови с confirm=true.")
    cat = _crm_resolve_category(args)
    cid = int(cat["id"])

    if action == "add":
        created = _crm_add_stage(cid, {
            "name": args.get("name"), "stage_code": args.get("stage_code"), "sort": args.get("sort"),
            "color": args.get("color"), "semantics": args.get("semantics"),
        })
        return {"added": created, "pipeline": cat.get("name"), "stages": _crm_stages(cid)}

    stage = _crm_resolve_stage(cid, args.get("stage") or args.get("name"))
    if action == "delete":
        if stage.get("system"):
            raise McpError(-32602, "Стадия «%s» системная — Bitrix не даёт её удалить." % stage.get("name"))
        r = _crm_call("crm.deal.list", {"filter": {"STAGE_ID": stage["stage_id"], "CATEGORY_ID": cid},
                                        "select": ["ID"]})
        total = r.get("total", len(r.get("result") or []))
        if total:
            raise McpError(-32602, "На стадии «%s» стоят сделки: %s. Сначала перенеси их "
                                   "(update_crm_deal), потом удаляй стадию." % (stage.get("name"), total))
        _crm_call("crm.status.delete", {"id": int(stage["id"])})
        return {"deleted": True, "stage_id": stage["stage_id"], "name": stage.get("name"),
                "pipeline": cat.get("name")}

    fields: dict[str, Any] = {}
    if str(args.get("new_name") or "").strip():
        fields["NAME"] = str(args["new_name"]).strip()
    if args.get("sort") not in (None, ""):
        fields["SORT"] = int(args["sort"])
    color = str(args.get("color") or "").strip()
    if color:
        if not _CRM_COLOR_RE.match(color):
            raise McpError(-32602, "color должен быть в формате #RRGGBB.")
        fields["COLOR"] = color
    if not fields:
        raise McpError(-32602, "Нечего менять: укажи new_name, sort и/или color.")
    _crm_call("crm.status.update", {"id": int(stage["id"]), "fields": fields})
    return {"updated": True, "stage_id": stage["stage_id"], "pipeline": cat.get("name"),
            "stages": _crm_stages(cid)}


def tool_list_crm_deal_fields(args: dict[str, Any]) -> dict[str, Any]:
    rows = _crm_call("crm.deal.userfield.list", {}).get("result") or []
    custom = []
    for r in rows:
        label = r.get("EDIT_FORM_LABEL")
        if isinstance(label, dict):
            label = label.get("ru") or label.get("en") or next(iter(label.values()), "")
        custom.append({
            "id": int(r.get("ID") or 0),
            "field_code": r.get("FIELD_NAME"),
            "type": r.get("USER_TYPE_ID"),
            "label": label,
            "multiple": str(r.get("MULTIPLE") or "") == "Y",
            "mandatory": str(r.get("MANDATORY") or "") == "Y",
            "show_in_list": str(r.get("SHOW_IN_LIST") or "") == "Y",
            "settings": r.get("SETTINGS") or {},
        })
    out: dict[str, Any] = {"custom_fields": custom, "custom_count": len(custom)}
    if args.get("include_standard") is True:
        std = _crm_call("crm.deal.fields", {}).get("result") or {}
        out["standard_fields"] = {
            code: {"type": meta.get("type"), "title": meta.get("formLabel") or meta.get("title") or code}
            for code, meta in std.items() if isinstance(meta, dict) and not code.startswith("UF_")
        }
    else:
        out["note"] = "Стандартные поля сделки — вызови с include_standard=true."
    return out


_CRM_UF_TYPES = {"string", "integer", "double", "boolean", "date", "datetime", "money", "url",
                 "enumeration", "employee", "file", "address"}


def tool_manage_crm_deal_field(args: dict[str, Any]) -> dict[str, Any]:
    action = str(args.get("action") or "").strip().lower()
    if action not in ("add", "update", "delete"):
        raise McpError(-32602, "action: add | update | delete.")
    if action == "delete" and args.get("confirm") is not True:
        raise McpError(-32602, "Удаление поля СТИРАЕТ его значения во всех сделках. Покажи "
                               "пользователю точное поле и вызови с confirm=true (после подтверждения).")

    def _resolve_existing() -> dict[str, Any]:
        rows = _crm_call("crm.deal.userfield.list", {}).get("result") or []
        fid = args.get("field_id")
        if fid not in (None, ""):
            match = next((r for r in rows if str(r.get("ID")) == str(fid)), None)
            if not match:
                raise McpError(-32602, f"Пользовательское поле id={fid} не найдено.")
            return match
        code = str(args.get("field_code") or "").strip().upper()
        if not code:
            raise McpError(-32602, "Укажи field_code (UF_CRM_...) или field_id.")
        if not code.startswith("UF_"):
            code = "UF_CRM_" + code
        match = next((r for r in rows if str(r.get("FIELD_NAME") or "").upper() == code), None)
        if not match:
            raise McpError(-32602, "Поле %s не найдено. Существуют: %s" % (
                code, ", ".join(str(r.get("FIELD_NAME")) for r in rows) or "(нет пользовательских полей)"))
        return match

    if action == "add":
        label = str(args.get("label") or "").strip()
        if not label:
            raise McpError(-32602, "label (человеческое название поля) обязателен.")
        ftype = str(args.get("type") or "string").strip().lower()
        if ftype not in _CRM_UF_TYPES:
            raise McpError(-32602, "type: %s." % ", ".join(sorted(_CRM_UF_TYPES)))
        code = _CRM_CODE_SANITIZE_RE.sub("_", str(args.get("field_code") or "").strip().upper()).strip("_")
        if code and not code.startswith("UF_"):
            code = "UF_CRM_" + code
        if not code:
            code = f"UF_CRM_F{int(time.time()) % 10**8}"
        fields: dict[str, Any] = {
            "FIELD_NAME": code,
            "USER_TYPE_ID": ftype,
            "EDIT_FORM_LABEL": {"ru": label, "en": label},
            "LIST_COLUMN_LABEL": {"ru": label, "en": label},
            "LIST_FILTER_LABEL": {"ru": label, "en": label},
            "MANDATORY": "Y" if args.get("mandatory") is True else "N",
            "MULTIPLE": "Y" if args.get("multiple") is True else "N",
            "SHOW_IN_LIST": "N" if args.get("show_in_list") is False else "Y",
        }
        items = args.get("list_items")
        if ftype == "enumeration":
            if not (isinstance(items, list) and items):
                raise McpError(-32602, "Для type=enumeration обязателен list_items — список вариантов.")
            fields["LIST"] = [{"VALUE": str(v), "SORT": (i + 1) * 100} for i, v in enumerate(items)]
        elif items:
            raise McpError(-32602, "list_items имеет смысл только при type=enumeration.")
        uf_id = _crm_call("crm.deal.userfield.add", {"fields": fields}).get("result")
        return {"added": True, "field_id": uf_id, "field_code": code, "type": ftype, "label": label,
                "note": "Код поля используй в custom_fields инструментов create_crm_deal/update_crm_deal."}

    existing = _resolve_existing()
    code = existing.get("FIELD_NAME")
    if action == "delete":
        _crm_call("crm.deal.userfield.delete", {"id": int(existing["ID"])})
        return {"deleted": True, "field_code": code}

    fields = {}
    label = str(args.get("label") or "").strip()
    if label:
        fields["EDIT_FORM_LABEL"] = {"ru": label, "en": label}
        fields["LIST_COLUMN_LABEL"] = {"ru": label, "en": label}
        fields["LIST_FILTER_LABEL"] = {"ru": label, "en": label}
    if args.get("mandatory") in (True, False):
        fields["MANDATORY"] = "Y" if args["mandatory"] else "N"
    if args.get("show_in_list") in (True, False):
        fields["SHOW_IN_LIST"] = "Y" if args["show_in_list"] else "N"
    items = args.get("list_items")
    if items:
        if str(existing.get("USER_TYPE_ID")) != "enumeration":
            raise McpError(-32602, "list_items можно менять только у поля типа enumeration.")
        fields["LIST"] = [{"VALUE": str(v), "SORT": (i + 1) * 100} for i, v in enumerate(items)]
    if not fields:
        raise McpError(-32602, "Нечего менять: укажи label, mandatory, show_in_list и/или list_items.")
    _crm_call("crm.deal.userfield.update", {"id": int(existing["ID"]), "fields": fields})
    return {"updated": True, "field_code": code}


def tool_list_crm_deals(args: dict[str, Any]) -> dict[str, Any]:
    limit = min(int(args.get("limit") or 50), 200)
    offset = max(int(args.get("offset") or 0), 0)
    filt: dict[str, Any] = {}
    cat = _crm_resolve_category(args, required=False)
    if cat is not None:
        filt["CATEGORY_ID"] = int(cat["id"])
    stage_ref = str(args.get("stage") or "").strip()
    if stage_ref:
        if cat is not None:
            filt["STAGE_ID"] = _crm_resolve_stage(int(cat["id"]), stage_ref)["stage_id"]
        elif ":" in stage_ref or stage_ref.upper() == stage_ref:
            filt["STAGE_ID"] = stage_ref
        else:
            raise McpError(-32602, "Для поиска по названию стадии укажи и воронку "
                                   "(category_id/pipeline_name), либо передай точный STAGE_ID ('C8:NEW').")
    if args.get("include_closed") is False:
        filt["CLOSED"] = "N"
    assigned = _crm_resolve_portal_user(args, "assigned_bitrix_user_id", "assigned_name")
    if assigned is not None:
        filt["ASSIGNED_BY_ID"] = assigned
    search = str(args.get("search") or "").strip()
    if search:
        filt["%TITLE"] = search
    select = list(_CRM_DEAL_LIST_SELECT) + (["UF_*"] if args.get("include_custom_fields") is True else [])

    rows: list[dict[str, Any]] = []
    start = offset
    total = 0
    while len(rows) < limit:
        r = _crm_call("crm.deal.list", {"order": {"DATE_CREATE": "DESC"}, "filter": filt,
                                        "select": select, "start": start})
        page = r.get("result") or []
        total = r.get("total", len(page))
        rows.extend(page)
        nxt = r.get("next")
        if not page or nxt in (None, "", start):
            break
        start = nxt
    rows = rows[:limit]
    names = _crm_deal_enrich_names(rows)
    deals = []
    for row, nm in zip(rows, names):
        item = _crm_deal_brief(row, nm)
        if args.get("include_custom_fields") is True:
            uf = {k: v for k, v in row.items() if k.startswith("UF_") and v not in (None, "", [])}
            if uf:
                item["custom_fields"] = uf
        deals.append(item)
    return {"total": total, "returned": len(deals), "offset": offset, "deals": deals}


def tool_get_crm_deal(args: dict[str, Any]) -> dict[str, Any]:
    try:
        deal_id = int(args.get("deal_id"))
    except (TypeError, ValueError) as exc:
        raise McpError(-32602, "deal_id must be an integer.") from exc
    row = _crm_call("crm.deal.get", {"id": deal_id}).get("result") or {}
    names = _crm_deal_enrich_names([row])[0]
    out = _crm_deal_brief(row, names)
    out["comments"] = row.get("COMMENTS") or ""
    out["custom_fields"] = {k: v for k, v in row.items() if k.startswith("UF_") and v not in (None, "", [])}
    out["fields"] = {k: v for k, v in row.items()
                     if not k.startswith("UF_") and v not in (None, "", []) and k not in
                     ("TITLE", "STAGE_ID", "CATEGORY_ID", "OPPORTUNITY", "CURRENCY_ID",
                      "ASSIGNED_BY_ID", "DATE_CREATE", "CLOSEDATE", "CLOSED", "COMMENTS", "ID")}
    return out


def _crm_deal_common_fields(args: dict[str, Any]) -> dict[str, Any]:
    fields: dict[str, Any] = {}
    if args.get("amount") not in (None, ""):
        fields["OPPORTUNITY"] = float(args["amount"])
        fields["CURRENCY_ID"] = str(args.get("currency") or "RUB").upper()
    elif args.get("currency"):
        fields["CURRENCY_ID"] = str(args["currency"]).upper()
    assigned = _crm_resolve_portal_user(args, "responsible_bitrix_user_id", "responsible_name")
    if assigned is not None:
        fields["ASSIGNED_BY_ID"] = assigned
    if args.get("comments") not in (None, ""):
        fields["COMMENTS"] = str(args["comments"])
    if args.get("begin_date") not in (None, ""):
        fields["BEGINDATE"] = str(args["begin_date"])
    if args.get("close_date") not in (None, ""):
        fields["CLOSEDATE"] = str(args["close_date"])
    if args.get("contact_id") not in (None, ""):
        fields["CONTACT_ID"] = int(args["contact_id"])
    if args.get("company_id") not in (None, ""):
        fields["COMPANY_ID"] = int(args["company_id"])
    fields.update(_crm_custom_fields_arg(args))
    return fields


def tool_create_crm_deal(args: dict[str, Any]) -> dict[str, Any]:
    title = str(args.get("title") or "").strip()
    if not title:
        raise McpError(-32602, "title (название сделки) обязателен.")
    fields: dict[str, Any] = {"TITLE": title}
    cat = _crm_resolve_category(args, required=False)
    if cat is not None:
        fields["CATEGORY_ID"] = int(cat["id"])
    stage_ref = str(args.get("stage") or "").strip()
    if stage_ref:
        cid = int(cat["id"]) if cat is not None else next(
            (int(c["id"]) for c in _crm_categories() if str(c.get("isDefault") or "") == "Y"), 0)
        fields["STAGE_ID"] = _crm_resolve_stage(cid, stage_ref)["stage_id"]
    fields.update(_crm_deal_common_fields(args))
    deal_id = _crm_call("crm.deal.add", {"fields": fields,
                                         "params": {"REGISTER_SONET_EVENT": "Y"}}).get("result")
    row = _crm_call("crm.deal.get", {"id": int(deal_id)}).get("result") or {}
    return {"created": True, **_crm_deal_brief(row, _crm_deal_enrich_names([row])[0])}


def tool_update_crm_deal(args: dict[str, Any]) -> dict[str, Any]:
    try:
        deal_id = int(args.get("deal_id"))
    except (TypeError, ValueError) as exc:
        raise McpError(-32602, "deal_id must be an integer.") from exc
    current = _crm_call("crm.deal.get", {"id": deal_id}).get("result") or {}
    expected = str(args.get("expected_title") or "").strip()
    if expected and expected.casefold() != str(current.get("TITLE") or "").casefold():
        raise McpError(-32602, "expected_title не совпал: сделка %s называется «%s». Изменение отменено."
                       % (deal_id, current.get("TITLE")))

    target_cat: int | None = None
    cat = _crm_resolve_category(args, required=False)
    if cat is not None and int(cat["id"]) != int(current.get("CATEGORY_ID") or 0):
        target_cat = int(cat["id"])

    fields: dict[str, Any] = {}
    if str(args.get("title") or "").strip():
        fields["TITLE"] = str(args["title"]).strip()
    stage_ref = str(args.get("stage") or "").strip()
    if stage_ref:
        cid = target_cat if target_cat is not None else int(current.get("CATEGORY_ID") or 0)
        fields["STAGE_ID"] = _crm_resolve_stage(cid, stage_ref)["stage_id"]
    fields.update(_crm_deal_common_fields(args))

    if target_cat is None and not fields:
        raise McpError(-32602, "Нечего менять: передай хотя бы одно поле (title/stage/amount/"
                               "responsible_*/comments/custom_fields/воронку и т.д.).")

    moved = False
    if target_cat is not None:
        # Перенос между воронками — только через универсальный crm.item.update
        # (crm.deal.update молча игнорирует CATEGORY_ID).
        item_fields: dict[str, Any] = {"categoryId": target_cat}
        if "STAGE_ID" in fields:
            item_fields["stageId"] = fields.pop("STAGE_ID")
        _crm_call("crm.item.update", {"entityTypeId": DEAL_ENTITY_TYPE_ID, "id": deal_id,
                                      "fields": item_fields})
        moved = True
    if fields:
        _crm_call("crm.deal.update", {"id": deal_id, "fields": fields})
    row = _crm_call("crm.deal.get", {"id": deal_id}).get("result") or {}
    out = {"updated": True, **_crm_deal_brief(row, _crm_deal_enrich_names([row])[0])}
    if moved:
        out["moved_to_pipeline"] = cat.get("name")
    return out


def tool_delete_crm_deal(args: dict[str, Any]) -> dict[str, Any]:
    if args.get("confirm") is not True:
        raise McpError(-32602, "Удаление сделки требует подтверждения: покажи пользователю точную "
                               "сделку (id, название, воронка, сумма) и вызови с confirm=true.")
    try:
        deal_id = int(args.get("deal_id"))
    except (TypeError, ValueError) as exc:
        raise McpError(-32602, "deal_id must be an integer.") from exc
    row = _crm_call("crm.deal.get", {"id": deal_id}).get("result") or {}
    expected = str(args.get("expected_title") or "").strip()
    if expected and expected.casefold() != str(row.get("TITLE") or "").casefold():
        raise McpError(-32602, "expected_title не совпал: сделка %s называется «%s». Удаление отменено."
                       % (deal_id, row.get("TITLE")))
    _crm_call("crm.deal.delete", {"id": deal_id})
    return {"deleted": True, "deal_id": deal_id, "title": row.get("TITLE")}


# --- Telegram news watchlist (отраслевые каналы WB/маркетплейсов) ------------------------------
# The TG bot's watchlist (.tg_agent_state.json, managed via /add_channel) is read here so a
# Bitrix agent («Новостной агент») can see fresh posts of ALL watched channels in one fast call.
# Fetching goes through tg_digest's public-preview scraper in parallel; per-(channel,days)
# results are cached 30 minutes, so repeated tool calls within one turn/session are instant.

def _tg_watchlist() -> list[str]:
    try:
        state = json.loads((ROOT / ".tg_agent_state.json").read_text(encoding="utf-8"))
        return [str(c) for c in (state.get("channels") or [])]
    except (OSError, ValueError):
        return []


def tool_get_tg_news(args: dict[str, Any]) -> dict[str, Any]:
    import tg_digest  # lazy: standalone module, no Flask/циркулярок

    days = max(1, min(int(args.get("days") or 7), 30))
    only = [str(c).strip().lstrip("@") for c in (args.get("channels") or []) if str(c).strip()]
    names = only or _tg_watchlist()
    if not names:
        raise McpError(-32000, "Список отслеживаемых каналов пуст (владелец наполняет его через "
                               "TG-бота командой /add_channel).")
    max_posts = max(1, min(int(args.get("max_posts_per_channel") or 12), 50))
    post_chars = max(200, min(int(args.get("post_chars") or 700), 2000))
    since = datetime.now(timezone.utc) - timedelta(days=days)

    def _one(name: str) -> tuple[str, list[str], str | None]:
        cached = ttl_cache_get(("tg_news", name, days))
        if cached is not None:
            return name, cached[0], cached[1]
        posts, err = tg_digest.fetch_channel_posts(name, since)
        ttl_cache_set(("tg_news", name, days), (posts, err), 1800)
        return name, posts, err

    from concurrent.futures import ThreadPoolExecutor
    with ThreadPoolExecutor(max_workers=12) as pool:
        results = list(pool.map(_one, names))

    channels_out, empty, problems = [], [], []
    total_posts, budget = 0, 90000
    for name, posts, err in results:
        if err:
            problems.append(f"t.me/{name} — {err}")
            continue
        if not posts:
            empty.append(name)
            continue
        fresh = [p[:post_chars] for p in posts[-max_posts:]]
        piece_len = sum(len(p) for p in fresh)
        if budget - piece_len < 0:
            problems.append(f"t.me/{name} — не влез в лимит ответа (запроси его отдельно через channels=[...])")
            continue
        budget -= piece_len
        total_posts += len(fresh)
        channels_out.append({"channel": name, "url": f"https://t.me/{name}",
                             "posts_count": len(fresh), "posts": fresh})
    return {
        "period_days": days,
        "channels_with_posts": len(channels_out),
        "total_posts": total_posts,
        "channels": channels_out,
        "empty_channels": empty,
        "problems": problems,
        "note": ("Посты за период по каждому каналу (старые выше). Нужен полный текст постов "
                 "конкретного канала — вызови ещё раз с channels=['имя'] и большим post_chars."),
    }


# --- Оргструктура Bitrix (отделы + принадлежность людей) ----------------------------------------
# Правит ТОЛЬКО вебхук: scope `department` есть у входящего вебхука и НЕТ у app-токена (разведано
# 2026-07-13). Менять оргструктуру разрешено строго ограниченному кругу (ORG_STRUCTURE_ADMINS,
# по умолчанию 14=Евгений Палей, 22=ИИ Агент) — агент обязан передать id того, кто попросил, и
# получить явное подтверждение (confirm=true). Любое изменение пишется в журнал.

# Разведано вживую 2026-07-13: Bitrix САМ добавляет назначенного руководителя в этот отдел
# (UF_DEPARTMENT пополняется). Значит «пустых» отделов с руководителем не бывает — предупреждаем
# человека, иначе он удивится, что руководитель «переехал», и не сможет удалить отдел.
_ORG_HEAD_AUTOJOIN_NOTE = ("Bitrix автоматически включил руководителя в состав этого отдела. "
                           "Сообщи это человеку. Чтобы удалить отдел, сначала переведи из него "
                           "ВСЕХ, включая руководителя.")


def _org_admin_ids() -> set[int]:
    raw = os.getenv("ORG_STRUCTURE_ADMINS", "14,22")
    return {int(x) for x in re.findall(r"\d+", raw)}


def _org_assert_allowed(args: dict[str, Any], action: str) -> int:
    """Кто просит? Разрешено только Евгению (14) и ИИ Агенту (22)."""
    uid = _int_or_none(args.get("requested_by_bitrix_user_id"))
    if uid is None:
        raise McpError(-32602, "Укажи requested_by_bitrix_user_id — id того, КТО просит изменить "
                               "оргструктуру (менять её могут только Евгений Палей и ИИ Агент).")
    allowed = _org_admin_ids()
    if uid not in allowed:
        who = _resolve_active_bitrix_user(uid, None) if uid else None
        name = (who or {}).get("full_name") or f"id {uid}"
        raise McpError(-32602, f"У {name} нет прав менять оргструктуру. Это могут делать только "
                               f"Евгений Палей и ИИ Агент. Скажи это человеку прямо и не выполняй "
                               f"изменение.")
    if args.get("confirm") is not True:
        raise McpError(-32602, f"Оргструктура — чувствительные данные. Сначала покажи человеку "
                               f"ТОЧНЫЙ план ({action}: что и с какими id изменится), дождись явного "
                               f"«да» и вызови повторно с confirm=true.")
    return uid


def _org_webhook(method: str, payload: dict[str, Any] | None = None) -> Any:
    data = _webhook_raw(method, payload or {})
    return data.get("result") if isinstance(data, dict) else data


def _org_resync_team() -> str:
    """После правки оргструктуры пересинхронизировать наш справочник (users), иначе
    get_org_structure будет отдавать старое."""
    try:
        base = os.getenv("BITRIX_WEBHOOK_BASE", "").strip()
        if base:
            res = app_workflow_function("sync_bitrix_team")(base)
            return f"справочник обновлён (сотрудников: {res.get('saved', '?')})"
    except Exception as exc:  # noqa: BLE001
        logging.warning("org: team resync failed: %s", repr(exc)[:150])
    return "справочник обновится ближайшей синхронизацией"


def tool_get_bitrix_departments(args: dict[str, Any]) -> dict[str, Any]:
    """Живая оргструктура портала: отделы (id, название, родитель, руководитель) + кто в каждом."""
    deps = _org_webhook("department.get") or []
    users = _org_webhook("user.get", {"ACTIVE": "true"}) or []
    by_dep: dict[str, list[dict[str, Any]]] = {}
    for u in users:
        name = " ".join(x for x in (u.get("NAME"), u.get("LAST_NAME")) if x).strip()
        for dep_id in (u.get("UF_DEPARTMENT") or []):
            by_dep.setdefault(str(dep_id), []).append({
                "bitrix_user_id": _int_or_none(u.get("ID")),
                "full_name": name or f"id {u.get('ID')}",
                "position": u.get("WORK_POSITION") or "",
            })
    heads = {str(d.get("ID")): _int_or_none(d.get("UF_HEAD")) for d in deps}
    names = {str(_int_or_none(u.get("ID"))): " ".join(x for x in (u.get("NAME"), u.get("LAST_NAME")) if x).strip()
             for u in users}
    out = []
    for d in deps:
        did = str(d.get("ID"))
        head_id = heads.get(did)
        out.append({
            "department_id": _int_or_none(d.get("ID")),
            "name": d.get("NAME"),
            "parent_id": _int_or_none(d.get("PARENT")),
            "head_bitrix_user_id": head_id,
            "head_name": names.get(str(head_id)) if head_id else None,
            "employees": sorted(by_dep.get(did, []), key=lambda e: e["full_name"]),
        })
    return {
        "portal": "b24-0xrp3s",
        "departments": sorted(out, key=lambda d: (d["parent_id"] or 0, d["department_id"] or 0)),
        "count": len(out),
        "note": ("Это ЖИВАЯ структура портала и ЕДИНСТВЕННЫЙ источник правды об отделах: схемы из "
                 "документов базы знаний могут описывать старую/задуманную структуру — не выдавай "
                 "их за факт. Менять структуру могут только Евгений Палей и ИИ Агент "
                 "(manage_bitrix_department / assign_employee_department) — после показа плана и "
                 "явного подтверждения."),
    }


def tool_manage_bitrix_department(args: dict[str, Any]) -> dict[str, Any]:
    """Создать / переименовать / переподчинить / назначить руководителя / удалить отдел."""
    action = str(args.get("action") or "").strip().lower()
    if action not in ("create", "update", "delete"):
        raise McpError(-32602, "action: create | update | delete.")
    requester = _org_assert_allowed(args, f"отдел: {action}")

    if action == "create":
        name = str(args.get("name") or "").strip()
        if not name:
            raise McpError(-32602, "Для create нужен name (название отдела).")
        fields: dict[str, Any] = {"NAME": name, "PARENT": _int_or_none(args.get("parent_id")) or 1}
        head = _int_or_none(args.get("head_bitrix_user_id"))
        if head:
            fields["UF_HEAD"] = head
        dep_id = _org_webhook("department.add", fields)
        logging.info("org: user %s created department %s «%s»", requester, dep_id, name)
        out = {"created": True, "department_id": _int_or_none(dep_id), "name": name,
               "parent_id": fields["PARENT"], "head_bitrix_user_id": head,
               "sync": _org_resync_team()}
        if head:
            out["note"] = _ORG_HEAD_AUTOJOIN_NOTE
        return out

    dep_id = _int_or_none(args.get("department_id"))
    if not dep_id:
        raise McpError(-32602, "Укажи department_id (см. get_bitrix_departments).")
    current = (_org_webhook("department.get", {"ID": dep_id}) or [{}])[0]
    if not current:
        raise McpError(-32602, f"Отдел id={dep_id} не найден.")
    if dep_id == 1:
        raise McpError(-32602, "Корневой отдел портала менять/удалять нельзя.")

    if action == "delete":
        members = [u for u in (_org_webhook("user.get", {"ACTIVE": "true"}) or [])
                   if dep_id in [_int_or_none(x) for x in (u.get("UF_DEPARTMENT") or [])]]
        if members:
            names = ", ".join(" ".join(x for x in (u.get("NAME"), u.get("LAST_NAME")) if x) for u in members[:10])
            raise McpError(-32602, f"В отделе «{current.get('NAME')}» ещё есть сотрудники: {names}. "
                                   f"Сначала переведи их (assign_employee_department), потом удаляй отдел.")
        _org_webhook("department.delete", {"ID": dep_id})
        logging.info("org: user %s deleted department %s «%s»", requester, dep_id, current.get("NAME"))
        return {"deleted": True, "department_id": dep_id, "name": current.get("NAME"),
                "sync": _org_resync_team()}

    fields = {"ID": dep_id}
    if str(args.get("name") or "").strip():
        fields["NAME"] = str(args["name"]).strip()
    if args.get("parent_id") not in (None, ""):
        fields["PARENT"] = _int_or_none(args["parent_id"])
    if args.get("head_bitrix_user_id") not in (None, ""):
        head = _int_or_none(args["head_bitrix_user_id"])
        if head and not _resolve_active_bitrix_user(head, None):
            raise McpError(-32602, f"Сотрудник id={head} не найден среди активных — руководителя не назначил.")
        fields["UF_HEAD"] = head
    if len(fields) == 1:
        raise McpError(-32602, "Нечего менять: передай name, parent_id и/или head_bitrix_user_id.")
    _org_webhook("department.update", fields)
    after = (_org_webhook("department.get", {"ID": dep_id}) or [{}])[0]
    logging.info("org: user %s updated department %s: %s", requester, dep_id, fields)
    out = {"updated": True, "department_id": dep_id,
           "was": {"name": current.get("NAME"), "parent_id": _int_or_none(current.get("PARENT")),
                   "head_bitrix_user_id": _int_or_none(current.get("UF_HEAD"))},
           "now": {"name": after.get("NAME"), "parent_id": _int_or_none(after.get("PARENT")),
                   "head_bitrix_user_id": _int_or_none(after.get("UF_HEAD"))},
           "sync": _org_resync_team()}
    if "UF_HEAD" in fields:
        out["note"] = _ORG_HEAD_AUTOJOIN_NOTE
    return out


def tool_assign_employee_department(args: dict[str, Any]) -> dict[str, Any]:
    """Перевести сотрудника(ов) в отдел и/или задать должность."""
    requester = _org_assert_allowed(args, "перевод сотрудника")
    dep_ids = [_int_or_none(d) for d in (args.get("department_ids") or [])]
    dep_ids = [d for d in dep_ids if d]
    position = args.get("position")
    if not dep_ids and position is None:
        raise McpError(-32602, "Нечего менять: передай department_ids и/или position.")

    known = {_int_or_none(d.get("ID")) for d in (_org_webhook("department.get") or [])}
    bad = [d for d in dep_ids if d not in known]
    if bad:
        raise McpError(-32602, f"Отделов с id {bad} нет. Актуальные — get_bitrix_departments.")

    targets: list[dict[str, Any]] = []
    for item in (args.get("employees") or []):
        user = _resolve_active_bitrix_user(item if str(item).isdigit() else None,
                                           None if str(item).isdigit() else str(item))
        if not user:
            raise McpError(-32602, f"Сотрудник «{item}» не найден среди активных. Сверься с "
                                   f"get_org_structure и передай точный id.")
        targets.append(user)
    if not targets:
        raise McpError(-32602, "Укажи employees — список id или ФИО сотрудников.")

    results = []
    for user in targets:
        uid = int(user["bitrix_user_id"])
        payload: dict[str, Any] = {"ID": uid}
        if dep_ids:
            payload["UF_DEPARTMENT"] = dep_ids
        if position is not None:
            payload["WORK_POSITION"] = str(position)
        _org_webhook("user.update", payload)
        results.append({"bitrix_user_id": uid, "full_name": user.get("full_name"),
                        "department_ids": dep_ids or "не менял",
                        "position": position if position is not None else "не менял"})
        logging.info("org: user %s moved employee %s -> deps=%s pos=%s", requester, uid, dep_ids, position)
    return {"updated": True, "employees": results, "sync": _org_resync_team()}


def tool_save_news_digest(args: dict[str, Any]) -> dict[str, Any]:
    """Store a weekly news digest so ad-hoc questions reuse it instead of rebuilding."""
    summary = str(args.get("summary") or "").strip()
    if not summary:
        raise McpError(-32602, "summary (текст сводки) обязателен.")
    period_days = max(1, min(int(args.get("period_days") or 7), 30))
    with connect() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "INSERT INTO tg_news_digests (period_days, summary, meta) VALUES (%s, %s, %s) "
                "RETURNING id, created_at",
                (period_days, summary[:20000], Jsonb(args.get("meta") or {})))
            row = cur.fetchone()
    return {"saved": True, "id": row["id"], "created_at": _to_msk(row["created_at"]).isoformat()}


def tool_get_latest_news_digest(args: dict[str, Any]) -> dict[str, Any]:
    """The most recent stored digest + its age. Use it to answer questions about the news
    WITHOUT rebuilding — rebuild (get_tg_news) only if it is missing or older than max_age_days."""
    max_age = int(args.get("max_age_days") or 7)
    with connect() as conn:
        with conn.cursor() as cur:
            cur.execute("SELECT id, created_at, period_days, summary FROM tg_news_digests "
                        "ORDER BY created_at DESC LIMIT 1")
            row = cur.fetchone()
    if not row:
        return {"found": False, "note": "Сохранённых сводок ещё нет — собери свежую через get_tg_news."}
    created = _to_msk(row["created_at"])
    age_days = (datetime.now(_MSK_TZ) - created).total_seconds() / 86400
    return {
        "found": True,
        "id": row["id"],
        "created_at": created.isoformat(),
        "age_days": round(age_days, 1),
        "is_fresh": age_days <= max_age,
        "period_days": row["period_days"],
        "summary": row["summary"],
        "note": ("Свежая — отвечай на её основе, НЕ пересобирай." if age_days <= max_age
                 else f"Устарела ({round(age_days,1)} дн) — можно пересобрать через get_tg_news."),
    }


TOOLS: dict[str, dict[str, Any]] = {
    "get_agent_monitoring": {
        "description": (
            "Мониторинг и учёт использования самого агента (живые данные страниц «Центра Агента»): "
            "здоровье всех систем (БД, MCP, мозг, Bitrix REST, Zoom, Google Drive, память сервера) "
            "с полем problems (что не ок — подсвети владельцу), скорость каждого хода за 24 часа, "
            "лента событий (ошибки, медленные ходы, жалобы, деплои) и расход по сотрудникам за период: "
            "ходы, время в работе с агентом, время работы агента и реальные токены из сессий Hermes. "
            "Используй для анализа «кто сколько потребляет», «почему агент тормозил», «всё ли работает»."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "period": {
                    "type": "string",
                    "description": "Период учёта использования: today, либо число дней (7, 30, 90). По умолчанию 7.",
                },
            },
        },
        "handler": tool_get_agent_monitoring,
    },
    "get_employee_absences": {
        "description": (
            "Узнать, в отпуске ли / отсутствует ли сотрудник (Bitrix «График отсутствий», портал b24-0xrp3s). "
            "Передай bitrix_user_id ИЛИ employee_name (нечёткий поиск среди активных; при неоднозначности "
            "возвращается список кандидатов). Без обоих — вернёт всех, кто отсутствует в период (по умолчанию "
            "сегодня). Для проверки на конкретную дату задай date_from=date_to (например дедлайн задачи). "
            "Возвращает on_vacation и периоды отсутствия. ПРАВИЛО постановки задач: если исполнитель on_vacation "
            "на дедлайн — задачу не ставить, сообщить владельцу. Использовать только если этот инструмент "
            "включён в текущем наборе агента/коннектора."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "employee_name": {"type": "string", "description": "ФИО или его отличительная часть. Игнорируется, если задан bitrix_user_id."},
                "bitrix_user_id": {"type": "integer", "description": "Точный id сотрудника на портале b24-0xrp3s (приоритетнее имени)."},
                "date_from": {"type": "string", "description": "Начало периода YYYY-MM-DD (по умолчанию сегодня)."},
                "date_to": {"type": "string", "description": "Конец периода YYYY-MM-DD (по умолчанию = date_from)."},
            },
            "additionalProperties": False,
        },
        "handler": tool_get_employee_absences,
    },
    "start_here_always_read_ai_instructions": {
        "description": "MANDATORY FIRST TOOL. Always call this before any company analysis, report, recommendation, or answer. It reads live rules from Настройки -> Инструкции для ИИ and tells the assistant exactly how to work.",
        "inputSchema": {"type": "object", "properties": {}, "additionalProperties": False},
        "handler": tool_start_here_always_read_ai_instructions,
    },
    "health": {
        "description": "Check PostgreSQL connectivity and MCP server status.",
        "inputSchema": {"type": "object", "properties": {}, "additionalProperties": False},
        "handler": tool_health,
    },
    "get_runtime_status": {
        "description": "Inspect MCP-first/PostgreSQL-only runtime mode, database target, cache TTL, and whether legacy HTTP API compatibility is enabled.",
        "inputSchema": {"type": "object", "properties": {}, "additionalProperties": False},
        "handler": tool_get_runtime_status,
    },
    "get_context_guide": {
        "description": "Read navigation rules after start_here_always_read_ai_instructions: where to search first, which tools map to which business sources, and how to avoid chaotic database exploration. Pass intent to get only the workflow and sources for the current task instead of the whole guide.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "intent": {
                    "type": "string",
                    "description": "Optional task route. One of: company_rule_question, employee_period_question, chat_event_question, bitrix_task_creation, recommendation_answer, owner_daily_report_creation, owner_weekly_report_creation. Returns only that workflow and its sources. Omit for the full guide.",
                },
            },
            "additionalProperties": False,
        },
        "handler": tool_get_context_guide,
    },
    "get_ai_instructions": {
        "description": "Read live editable AI behavior and answer-format instructions from Настройки -> Инструкции для ИИ. start_here_always_read_ai_instructions already returns the full text, so call this only to re-read one folder by path. Use get_context_guide for the index of available paths.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": "Optional folder path prefix, e.g. 'Формирование отчетов / Ежедневный отчет по компании'. Returns only matching folders. Omit to read the full tree.",
                },
            },
            "additionalProperties": False,
        },
        "handler": tool_get_ai_instructions,
    },
    "get_report_contract": {
        "description": "Read the active report-generation contract for a configured report category. Use this before creating daily, weekly, owner, or Zoom reports.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "category_key": {"type": "string", "description": "For daily chat reports use chat_analysis."},
            },
            "required": ["category_key"],
            "additionalProperties": False,
        },
        "handler": tool_get_report_contract,
    },
    "list_available_sources": {
        "description": "Show which known context tables exist and how many rows each has.",
        "inputSchema": {"type": "object", "properties": {}, "additionalProperties": False},
        "handler": tool_list_available_sources,
    },
    "get_company_profile": {
        "description": "Read the editable company profile text from PostgreSQL for business context.",
        "inputSchema": {"type": "object", "properties": {}, "additionalProperties": False},
        "handler": tool_get_company_profile,
    },
    "list_company_files": {
        "description": "List all files/folders available in the company knowledge section, including Google Drive mirrored documents. Use before reading specific company files.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "include_empty": {"type": "boolean", "description": "Include folders/files that have no text content."},
                "limit": {"type": "integer", "minimum": 1, "maximum": MAX_LIMIT},
                "offset": {"type": "integer", "minimum": 0},
            },
            "additionalProperties": False,
        },
        "handler": tool_list_company_files,
    },
    "get_company_file": {
        "description": "Read the full text and source metadata for one company knowledge file by folder_id or google_file_id.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "folder_id": {"type": "string", "description": "company_folders.id from list_company_files/search_company_knowledge."},
                "google_file_id": {"type": "string", "description": "Google Drive file id from list_company_files/search_company_knowledge."},
            },
            "additionalProperties": False,
        },
        "handler": tool_get_company_file,
    },
    "search_company_knowledge": {
        "description": "Search the persistent 'О компании' knowledge base, including Google Drive mirrored docs/sheets. Use for rules, regulations, processes, and company facts before searching chats. Returns focused passages (chunks); read a full document with get_company_file(folder_id). If a search returns no results, RETRY with synonyms/rephrasings (e.g. 'созвоны'→'встречи/планёрки/Zoom', 'график'→'расписание/периодичность/ритм') before concluding nothing exists.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "query": {"type": "string", "description": "Search text. Empty returns the latest folders/documents."},
                "limit": {"type": "integer", "minimum": 1, "maximum": MAX_LIMIT},
                "offset": {"type": "integer", "minimum": 0},
            },
            "additionalProperties": False,
        },
        "handler": tool_search_company_knowledge,
    },
    "list_periods": {
        "description": "List recent dates available in chat messages and Zoom/owner sources.",
        "inputSchema": {
            "type": "object",
            "properties": {"limit": {"type": "integer", "minimum": 1, "maximum": MAX_LIMIT}},
            "additionalProperties": False,
        },
        "handler": tool_list_periods,
    },
    "get_period_index": {
        "description": "Return counts and top chats for a date period.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "date_from": {"type": "string", "description": "YYYY-MM-DD"},
                "date_to": {"type": "string", "description": "YYYY-MM-DD"},
            },
            "required": ["date_from", "date_to"],
            "additionalProperties": False,
        },
        "handler": tool_get_period_index,
    },
    "get_report_readiness": {
        "description": (
            "Report-building readiness for a date range in one call: per day, which active chats have "
            "messages and which already have a current daily report (missing_daily_reports), which Zoom "
            "calls already have an analytical_note (missing_zoom_reports), and whether the current and "
            "previous owner daily reports exist. Call this before daily/weekly/owner reports instead of "
            "probing each chat and Zoom call separately."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "date_from": {"type": "string", "description": "YYYY-MM-DD"},
                "date_to": {"type": "string", "description": "YYYY-MM-DD; defaults to date_from for a single day"},
            },
            "required": ["date_from"],
            "additionalProperties": False,
        },
        "handler": tool_get_report_readiness,
    },
    "get_org_structure": {
        "description": "Return departments and users with managers and department memberships.",
        "inputSchema": {
            "type": "object",
            "properties": {"include_inactive": {"type": "boolean"}},
            "additionalProperties": False,
        },
        "handler": tool_get_org_structure,
    },
    "search_tasks": {
        "description": (
            "Search Bitrix tasks by id, period, text, or responsible user. When the request mentions a task "
            "number (e.g. 318241), pass it as bitrix_task_id for an instant single-task lookup — do NOT put the "
            "number in query (query matches title/description text only). Each row includes comments_total_count "
            "and comments_human_count; when a task has human comments, read them with get_task_comments(bitrix_task_id). "
            "Descriptions are truncated by default to keep results small; set include_full_description=true (best with "
            "bitrix_task_id) to read one task's full description."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "bitrix_task_id": {"type": "integer", "description": "Exact Bitrix task number. Fastest path; ignores other filters."},
                "date_from": {"type": "string"},
                "date_to": {"type": "string"},
                "query": {"type": "string", "description": "Text match on title/description. Not for task numbers — use bitrix_task_id instead."},
                "responsible_bitrix_user_id": {"type": "integer"},
                "include_full_description": {"type": "boolean", "description": "Return full descriptions instead of a 500-char preview. Use only with bitrix_task_id or a small result set."},
                "limit": {"type": "integer", "minimum": 1, "maximum": MAX_LIMIT},
                "offset": {"type": "integer", "minimum": 0},
            },
            "additionalProperties": False,
        },
        "handler": tool_search_tasks,
    },
    "get_task_comments": {
        "description": (
            "Read the discussion/comments of one Bitrix task by bitrix_task_id (from search_tasks). "
            "Returns human comments with author name, timestamp, and BB-code-cleaned text. "
            "ВЛОЖЕНИЯ КОММЕНТАРИЕВ ЧИТАЮТСЯ: скрины распознаются (OCR), документы извлекаются — поле "
            "attachments у комментария, полный текст через get_attachment_text(attachment_id). Комментарий "
            "без текста, но со скрином — НЕ пустой. Ссылки из комментариев открывай fetch_url. "
            "Auto-generated notifications (overdue reminders, status cards, completion notices) are "
            "excluded by default; set include_service=true to include them. Use this to find what "
            "people asked, decided, committed, or blocked on a task — search_tasks alone does not show it."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "bitrix_task_id": {"type": "integer", "description": "Bitrix task id from search_tasks."},
                "include_service": {
                    "type": "boolean",
                    "description": "Include system notifications and status-card messages. Default false.",
                },
                "order": {
                    "type": "string",
                    "enum": ["asc", "desc"],
                    "description": "Chronological order by date. Default asc (oldest first).",
                },
                "limit": {"type": "integer", "minimum": 1, "maximum": MAX_LIMIT},
                "offset": {"type": "integer", "minimum": 0},
            },
            "required": ["bitrix_task_id"],
            "additionalProperties": False,
        },
        "handler": tool_get_task_comments,
    },
    "create_bitrix_task": {
        "description": (
            "Create one Bitrix task through the configured Bitrix webhook. Supports one-off and recurring tasks, "
            "and optional observers (наблюдатели) resolved against the org structure. STRICT RULE: do not call "
            "this tool unless the user provided a task title, exactly one responsible person, and a deadline. "
            "If title, responsible_name/responsible_bitrix_user_id, or deadline is missing, ask the user for it "
            "first. The tool resolves responsible_name and auditor_names through the active org structure and "
            "refuses ambiguous matches. To make the task recurring, pass periodic={type, ...}; otherwise the "
            "task is one-off. Before creating, confirm with the user: title, responsible, deadline, full list "
            "of observers (if any), and periodic schedule (if any). Every task is created with "
            "«результат обязателен» — it cannot be completed without a result. The "
            "постановщик (creator_bitrix_user_id/creator_name) defaults to the current chat user; "
            "set another person only when explicitly asked. If the deadline is already in the past, the "
            "tool refuses unless confirm_past_deadline=true — ask the user whether to keep it as-is (then "
            "re-call with confirm_past_deadline=true) or give a new future deadline. Every task MUST have "
            "result_criteria (what counts as done + how it is proven); if the user did not specify it, ASK — "
            "never invent it. The tool refuses to create a task without result_criteria. "
            "Optional task fields (set only when asked): соисполнители (accomplice_names/ids), "
            "родительская задача/подзадача (parent_task_id), проект (group_id), планирование сроков "
            "(start_plan/end_plan/time_estimate_hours), элементы CRM (crm_elements), пользовательские поля "
            "(custom_fields, коды из list_task_userfields), файлы (attachment_ids), чек-лист (checklist)."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "title": {"type": "string", "description": "Required task title."},
                "description": {"type": "string", "description": "Task description. If omitted, title is used."},
                "responsible_name": {"type": "string", "description": "Responsible employee name from org structure (fuzzy-matched against active users)."},
                "responsible_bitrix_user_id": {"type": "integer", "description": "Exact Bitrix user id of the responsible employee. Preferred over responsible_name when known."},
                "deadline": {"type": "string", "description": "Required deadline: YYYY-MM-DD, DD.MM.YYYY, or ISO datetime. For recurring tasks this is the first instance deadline."},
                "result_criteria": {"type": "string", "description": "ОБЯЗАТЕЛЬНО. Что должно быть результатом задачи: критерий выполнения (по чему поймём, что сделано) + чем подтверждается (скрин/ссылка/файл/артефакт). Если пользователь не указал — СПРОСИ, не выдумывай. Без него инструмент откажет."},
                "priority": {"type": "string", "enum": ["normal", "high", "critical"]},
                "creator_bitrix_user_id": {"type": "integer", "description": "Bitrix user id of the task CREATOR (постановщик). Default: the CURRENT chat user (whoever asked to create the task) — pass their id. Use a different id ONLY if the user explicitly asks to make someone else the постановщик."},
                "creator_name": {"type": "string", "description": "Full name of the task CREATOR (постановщик) from the org structure, used when the id is unknown. Same default rule as creator_bitrix_user_id."},
                "confirm_past_deadline": {"type": "boolean", "description": "Set to true ONLY after the user explicitly confirmed creating a task whose deadline is already in the past, as-is. Normally omit. If the deadline is in the past and this is not set, the tool refuses — ask the user first, then either re-call with confirm_past_deadline=true (keep as-is) or with a new future deadline."},
                "tags": {"type": "array", "items": {"type": "string"}},
                "auditor_names": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "Optional list of observer (наблюдатель) full names from the active org structure. Each is fuzzy-matched against users; ambiguity is refused.",
                },
                "auditor_bitrix_user_ids": {
                    "type": "array",
                    "items": {"type": "integer"},
                    "description": "Optional list of observer Bitrix user ids. Preferred over auditor_names when known. Merged with names; duplicates are dropped.",
                },
                "accomplice_names": {
                    "type": "array", "items": {"type": "string"},
                    "description": "Соисполнители (co-executors) — полные имена из оргструктуры (fuzzy-matched; ambiguity refused).",
                },
                "accomplice_bitrix_user_ids": {
                    "type": "array", "items": {"type": "integer"},
                    "description": "Соисполнители по Bitrix id (приоритетнее имён).",
                },
                "parent_task_id": {"type": "integer", "description": "Родительская задача: id задачи-родителя (эта задача станет подзадачей)."},
                "group_id": {"type": "integer", "description": "Проект/рабочая группа: id группы Bitrix, к которой привязать задачу."},
                "start_plan": {"type": "string", "description": "Планирование сроков — плановое начало (YYYY-MM-DD[ HH:MM] / DD.MM.YYYY[ HH:MM] / ISO)."},
                "end_plan": {"type": "string", "description": "Планирование сроков — плановое завершение (тот же формат)."},
                "time_estimate_hours": {"type": "number", "description": "Оценка трудозатрат в часах (планирование)."},
                "crm_elements": {
                    "type": "array", "items": {"type": "string"},
                    "description": "Элементы CRM: привязки вида 'D_123' (сделка), 'L_45' (лид), 'C_7' (контакт), 'CO_9' (компания).",
                },
                "custom_fields": {
                    "type": "object",
                    "description": "Пользовательские поля задачи: {\"UF_...\": значение}. Коды полей — list_task_userfields.",
                    "additionalProperties": True,
                },
                "attachment_ids": {
                    "type": "array", "items": {"type": "string"},
                    "description": "Файлы к задаче: токены вложений (att_…), присланные пользователем боту; прикрепляются к файлам задачи.",
                },
                "checklist": {
                    "type": "array",
                    "items": {"type": ["string", "object"]},
                    "description": "Чек-лист: список пунктов — строки или {title, complete}. Каждый добавляется в чек-лист задачи.",
                },
                "periodic": {
                    "type": "object",
                    "description": "Optional recurrence schedule. Omit for a one-off task. When present, the task is created as IS_REGULAR=Y with the corresponding REGULAR_PARAMETERS in Bitrix.",
                    "properties": {
                        "type": {"type": "string", "enum": ["daily", "weekly", "monthly"], "description": "Recurrence type."},
                        "interval": {"type": "integer", "minimum": 1, "description": "Every N units (default 1). E.g. interval=2 with type=weekly = every other week."},
                        "weekdays": {
                            "type": "array",
                            "items": {"type": "string", "enum": ["MO", "TU", "WE", "TH", "FR", "SA", "SU"]},
                            "description": "Required for type=weekly. Two-letter day codes (MO/TU/WE/TH/FR/SA/SU).",
                        },
                        "day_of_month": {"type": "integer", "minimum": 1, "maximum": 31, "description": "Required for type=monthly. Day of the month (1-31)."},
                        "daily_mode": {"type": "string", "enum": ["all", "workdays"], "description": "Optional for type=daily (default 'all'). 'workdays' = skip weekends."},
                        "until": {"type": "string", "description": "Optional end date YYYY-MM-DD. After this date, no new instances are created."},
                    },
                    "required": ["type"],
                    "additionalProperties": False,
                },
            },
            "required": ["title", "deadline"],
            "additionalProperties": False,
        },
        "handler": tool_create_bitrix_task,
    },
    "delete_bitrix_task": {
        "description": (
            "Delete one Bitrix task through the configured Bitrix webhook. STRICT CONFIRMATION RULE: do not call "
            "this tool until the user has confirmed deletion of one exact bitrix_task_id after seeing the task "
            "title, status, responsible person, and deadline. Never delete by title/search text/name or ambiguous "
            "reference. First use search_tasks(bitrix_task_id=...) to show the exact task, then ask for confirmation. "
            "Only after the user confirms, call delete_bitrix_task(bitrix_task_id=..., confirm=true)."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "bitrix_task_id": {"type": "integer", "description": "Exact Bitrix task number to delete."},
                "confirm": {
                    "type": "boolean",
                    "description": "Must be true. Means the user explicitly confirmed deleting this exact task after seeing its details.",
                },
                "expected_title": {
                    "type": "string",
                    "description": "Optional safety check: if provided, must exactly match the locally indexed task title.",
                },
            },
            "required": ["bitrix_task_id", "confirm"],
            "additionalProperties": False,
        },
        "handler": tool_delete_bitrix_task,
    },
    "add_bitrix_task_comment": {
        "description": (
            "Add one comment to an existing Bitrix task discussion. Use only with an exact bitrix_task_id; "
            "if the user points to a task by title/person/text, first resolve it with search_tasks and avoid "
            "ambiguous matches. This is the normal tool for ведение переписки внутри задачи. BY DEFAULT the "
            "comment is posted ОТ ЛИЦА the current chat user — pass author_bitrix_user_id = the id of the person "
            "who asked (the prompt gives it); the comment then shows as authored by them, not the bot. You can "
            "attach the user's screenshots/documents by passing their attachment_ids. Set as_result=true to mark "
            "the comment as the task RESULT (posts a «✅ РЕЗУЛЬТАТ» comment and pins the attached file(s) to the "
            "task — this is how you «прикрепить скрин как результат»)."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "bitrix_task_id": {"type": "integer", "description": "Exact Bitrix task number."},
                "comment_text": {"type": "string", "description": "Comment text to add to the task discussion."},
                "author_bitrix_user_id": {"type": "integer", "description": "Bitrix user id of the comment AUTHOR. Default: the CURRENT chat user (whoever asked to add the comment) — pass their id so the comment is «от лица» them. Use another id only if explicitly asked to comment on someone else's behalf."},
                "author_name": {"type": "string", "description": "Full name of the comment author from the org structure, used when the id is unknown. Same default rule as author_bitrix_user_id."},
                "attachment_ids": {"type": "array", "items": {"type": "string"}, "description": "Optional list of attachment tokens (att_…) the user sent the bot, to attach to this comment. The tokens are given in the prompt when the user sends a file."},
                "as_result": {"type": "boolean", "description": "If true, post this comment as the task RESULT (labelled «✅ РЕЗУЛЬТАТ» + attached files pinned to the task). Use for «отметить комментарий/скрин как результат»."},
                "expected_title": {
                    "type": "string",
                    "description": "Optional safety check: if the task is locally indexed, the title must match before posting.",
                },
            },
            "required": ["bitrix_task_id"],
            "additionalProperties": False,
        },
        "handler": tool_add_bitrix_task_comment,
    },
    "complete_bitrix_task": {
        "description": (
            "Complete (close/«завершить») one Bitrix task. BY DEFAULT the closure is attributed to the current "
            "chat user (on_behalf_bitrix_user_id → the status change is recorded as done by that person). If the "
            "task requires a result, pass result_text and/or attachment_ids — a «✅ РЕЗУЛЬТАТ» comment is posted "
            "and the file(s) pinned to the task before closing, so the proof is attached. Resolve the exact task "
            "with search_tasks first; confirm with the user before closing someone else's task."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "bitrix_task_id": {"type": "integer", "description": "Exact Bitrix task number to complete."},
                "on_behalf_bitrix_user_id": {"type": "integer", "description": "Bitrix user id the completion is on behalf of. Default: the CURRENT chat user (whoever asked to close it)."},
                "on_behalf_name": {"type": "string", "description": "Full name of the person the completion is on behalf of, used when the id is unknown."},
                "result_text": {"type": "string", "description": "Optional result text; posted as a «✅ РЕЗУЛЬТАТ» comment before closing (needed for result-required tasks)."},
                "attachment_ids": {"type": "array", "items": {"type": "string"}, "description": "Optional attachment tokens (att_…) to attach as the result before closing."},
                "expected_title": {"type": "string", "description": "Optional safety check on the task title."},
            },
            "required": ["bitrix_task_id"],
            "additionalProperties": False,
        },
        "handler": tool_complete_bitrix_task,
    },
    "attach_files_to_task": {
        "description": (
            "Attach one or more files the user sent the bot (screenshots, documents — referenced by their "
            "attachment tokens att_…) to a Bitrix task. Default: pin them to the task's files. Pass as_result=true "
            "to deliver them as the task RESULT (labelled comment + pinned files), or as_comment=true with a note "
            "to attach them inside a discussion comment. This is how the agent forwards the user's attachments "
            "into a task."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "bitrix_task_id": {"type": "integer", "description": "Exact Bitrix task number."},
                "attachment_ids": {"type": "array", "items": {"type": "string"}, "description": "Attachment tokens (att_…) from the prompt to attach."},
                "as_result": {"type": "boolean", "description": "Deliver as the task RESULT (labelled comment + pinned files)."},
                "as_comment": {"type": "boolean", "description": "Deliver inside a discussion comment (with note)."},
                "note": {"type": "string", "description": "Optional text to accompany the files when as_comment/as_result."},
                "author_bitrix_user_id": {"type": "integer", "description": "When posting a comment/result: the author (default current chat user)."},
                "author_name": {"type": "string", "description": "Author full name when the id is unknown."},
                "expected_title": {"type": "string", "description": "Optional safety check on the task title."},
            },
            "required": ["bitrix_task_id", "attachment_ids"],
            "additionalProperties": False,
        },
        "handler": tool_attach_files_to_task,
    },
    "get_attachment_text": {
        "description": (
            "Read the FULL text of a file the user sent the bot (contract, document, screenshot OCR), by its "
            "attachment token att_… The prompt shows a preview of long documents; call this to read the WHOLE "
            "thing — nothing is truncated. For a long contract, read it in chunks: call with offset=0, then "
            "offset=next_offset until has_more is false. Use before drafting/analysing a document the user attached."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "attachment_id": {"type": "string", "description": "Attachment token att_… from the prompt."},
                "offset": {"type": "integer", "minimum": 0, "description": "Character offset to start from (default 0)."},
                "max_chars": {"type": "integer", "description": "Max characters to return (default 40000, cap 120000)."},
            },
            "required": ["attachment_id"],
            "additionalProperties": False,
        },
        "handler": tool_get_attachment_text,
    },
    "reopen_bitrix_task": {
        "description": (
            "Reopen/renew one completed Bitrix task and write a comment explaining why. Use after checking "
            "search_tasks plus task result/comments and deciding the result is unsatisfactory. Requires exact "
            "bitrix_task_id, reason, and confirm=true after explicit user confirmation or a standing review instruction. "
            "Pass new_deadline to «возобновить с новым сроком». The explanatory comment is posted on behalf of the "
            "current chat user (on_behalf_bitrix_user_id)."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "bitrix_task_id": {"type": "integer", "description": "Exact Bitrix task number to reopen."},
                "reason": {"type": "string", "description": "Why the result is unsatisfactory / what must be fixed."},
                "confirm": {
                    "type": "boolean",
                    "description": "Must be true. Means reopening was explicitly approved or follows a standing task-review instruction.",
                },
                "new_deadline": {"type": "string", "description": "Optional new deadline (YYYY-MM-DD, DD.MM.YYYY, or with time) to set when reopening — «возобновить с новым сроком»."},
                "confirm_past_deadline": {"type": "boolean", "description": "Set true only if the user explicitly wants a new_deadline that is already in the past."},
                "on_behalf_bitrix_user_id": {"type": "integer", "description": "Bitrix user id the reopen comment is authored by (default current chat user)."},
                "on_behalf_name": {"type": "string", "description": "Full name of the person the reopen is on behalf of, when the id is unknown."},
                "expected_title": {
                    "type": "string",
                    "description": "Optional safety check: if the task is locally indexed, the title must match before reopening.",
                },
            },
            "required": ["bitrix_task_id", "reason", "confirm"],
            "additionalProperties": False,
        },
        "handler": tool_reopen_bitrix_task,
    },
    "update_bitrix_task": {
        "description": (
            "Изменить существующую задачу Bitrix по точному bitrix_task_id — любой набор полей из "
            "конструктора задачи: соисполнители (accomplice_*), наблюдатели (auditor_*), теги, "
            "родительская задача (parent_task_id), проект (group_id), планирование сроков "
            "(start_plan/end_plan/time_estimate_hours), элементы CRM (crm_elements), пользовательские "
            "поля (custom_fields), срок (deadline), приоритет, ответственный, название/описание. "
            "Списочные поля ЗАМЕНЯЮТСЯ переданным списком (передавай полный набор). Смену ответственного "
            "или родительской задачи сначала подтверди у пользователя. Файлы к задаче — attach_files_to_task; "
            "чек-лист — add_task_checklist; учёт времени — log_task_time; связи — link_tasks."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "bitrix_task_id": {"type": "integer", "description": "Точный номер задачи."},
                "title": {"type": "string", "description": "Новое название."},
                "description": {"type": "string", "description": "Новое описание."},
                "deadline": {"type": "string", "description": "Новый срок (YYYY-MM-DD[ HH:MM] / DD.MM.YYYY[ HH:MM] / ISO)."},
                "confirm_past_deadline": {"type": "boolean", "description": "true, если новый срок в прошлом и пользователь подтвердил."},
                "priority": {"type": "string", "enum": ["normal", "high", "critical"]},
                "responsible_name": {"type": "string", "description": "Новый ответственный по имени (подтверди у пользователя)."},
                "responsible_bitrix_user_id": {"type": "integer", "description": "Новый ответственный по id."},
                "accomplice_names": {"type": "array", "items": {"type": "string"}, "description": "Соисполнители по именам (заменяют текущих)."},
                "accomplice_bitrix_user_ids": {"type": "array", "items": {"type": "integer"}, "description": "Соисполнители по id."},
                "auditor_names": {"type": "array", "items": {"type": "string"}, "description": "Наблюдатели по именам (заменяют текущих)."},
                "auditor_bitrix_user_ids": {"type": "array", "items": {"type": "integer"}, "description": "Наблюдатели по id."},
                "tags": {"type": "array", "items": {"type": "string"}, "description": "Теги (заменяют текущие)."},
                "parent_task_id": {"type": "integer", "description": "Родительская задача (сделать подзадачей)."},
                "group_id": {"type": "integer", "description": "Проект/рабочая группа."},
                "start_plan": {"type": "string", "description": "Плановое начало."},
                "end_plan": {"type": "string", "description": "Плановое завершение."},
                "time_estimate_hours": {"type": "number", "description": "Оценка трудозатрат, часы."},
                "crm_elements": {"type": "array", "items": {"type": "string"}, "description": "Элементы CRM ('D_123','L_45','C_7','CO_9')."},
                "custom_fields": {"type": "object", "description": "Пользовательские поля {\"UF_...\": значение}.", "additionalProperties": True},
                "expected_title": {"type": "string", "description": "Опц. проверка: должно совпасть с названием индексированной задачи."},
            },
            "required": ["bitrix_task_id"],
            "additionalProperties": False,
        },
        "handler": tool_update_bitrix_task,
    },
    "add_task_checklist": {
        "description": (
            "Добавить пункты ЧЕК-ЛИСТА к задаче Bitrix. items — список строк или объектов {title, "
            "complete}. complete=true отмечает пункт выполненным."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "bitrix_task_id": {"type": "integer", "description": "Точный номер задачи."},
                "items": {
                    "type": "array",
                    "items": {"type": ["string", "object"]},
                    "description": "Пункты чек-листа: строки или {title, complete}.",
                },
                "expected_title": {"type": "string", "description": "Опц. проверка названия задачи."},
            },
            "required": ["bitrix_task_id", "items"],
            "additionalProperties": False,
        },
        "handler": tool_add_task_checklist,
    },
    "log_task_time": {
        "description": (
            "Записать УЧЁТ ВРЕМЕНИ по задаче (затраченное время). Время из hours и/или minutes (или "
            "seconds). Опц. комментарий и от чьего лица (on_behalf_*)."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "bitrix_task_id": {"type": "integer", "description": "Точный номер задачи."},
                "hours": {"type": "number", "description": "Часы (можно вместе с minutes)."},
                "minutes": {"type": "number", "description": "Минуты."},
                "seconds": {"type": "number", "description": "Секунды (если нужно точно)."},
                "comment": {"type": "string", "description": "Комментарий к записи учёта времени."},
                "on_behalf_bitrix_user_id": {"type": "integer", "description": "От чьего лица записать время (id)."},
                "on_behalf_name": {"type": "string", "description": "От чьего лица записать время (имя)."},
                "expected_title": {"type": "string", "description": "Опц. проверка названия задачи."},
            },
            "required": ["bitrix_task_id"],
            "additionalProperties": False,
        },
        "handler": tool_log_task_time,
    },
    "link_tasks": {
        "description": (
            "Связать две задачи (СВЯЗАННЫЕ ЗАДАЧИ / зависимость для Ганта). task_id_from зависит от "
            "task_id_to. link_type: finish_start (по умолчанию) / start_start / start_finish / "
            "finish_finish, либо число 0..3."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "task_id_from": {"type": "integer", "description": "Задача-источник (зависит от целевой)."},
                "task_id_to": {"type": "integer", "description": "Целевая задача."},
                "link_type": {"type": "string", "description": "finish_start | start_start | start_finish | finish_finish (или 0..3)."},
            },
            "required": ["task_id_from", "task_id_to"],
            "additionalProperties": False,
        },
        "handler": tool_link_tasks,
    },
    "add_task_reminder": {
        "description": (
            "Добавить НАПОМИНАНИЕ по задаче на конкретное время. remind_at — когда напомнить. Кому: "
            "по умолчанию ответственный задачи, либо user_name/user_bitrix_user_id. Best-effort: если "
            "портал не поддерживает напоминания через REST — вернёт внятную ошибку."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "bitrix_task_id": {"type": "integer", "description": "Точный номер задачи."},
                "remind_at": {"type": "string", "description": "Когда напомнить (YYYY-MM-DD HH:MM / DD.MM.YYYY HH:MM / ISO)."},
                "user_name": {"type": "string", "description": "Кому напомнить (имя). По умолчанию — ответственный."},
                "user_bitrix_user_id": {"type": "integer", "description": "Кому напомнить (id)."},
            },
            "required": ["bitrix_task_id", "remind_at"],
            "additionalProperties": False,
        },
        "handler": tool_add_task_reminder,
    },
    "list_task_userfields": {
        "description": (
            "Показать ПОЛЬЗОВАТЕЛЬСКИЕ ПОЛЯ задач (UF_*), заведённые на портале, с кодами и подписями — "
            "чтобы использовать реальные коды в custom_fields инструментов create_bitrix_task / "
            "update_bitrix_task, а не угадывать. Только чтение."
        ),
        "inputSchema": {"type": "object", "properties": {}, "additionalProperties": False},
        "handler": tool_list_task_userfields,
    },
    "create_recurring_task": {
        "description": (
            "Создать ПОВТОРЯЮЩУЮСЯ (регулярную) задачу — она будет создаваться автоматически по "
            "расписанию (например, каждую пятницу в 10:00 с дедлайном 19:00 того же дня). Расписание "
            "ведёт СОБСТВЕННЫЙ планировщик агента (в Bitrix нет подписки на регулярные задачи, поэтому "
            "НИКОГДА не создавай шаблон регулярной задачи в Bitrix — он не будет спавнить задачи) — "
            "приложение само создаёт обычную разовую задачу в срок БЕЗ хода LLM. Запись сразу видна "
            "владельцу в Центре Агента → Агенты → вкладка «Автоматизации» (чип «регулярная задача»). "
            "Это правильный инструмент для «ставь задачу каждый день/неделю/месяц»; schedule_my_automation "
            "для этого НЕ использовать (та — для регулярных ОТЧЁТОВ/действий и жжёт полноценный ход). "
            "Укажи period (weekly/daily/monthly), для weekly — weekdays "
            "(MO/TU/WE/TH/FR/SA/SU), create_time (во сколько создавать, ЧЧ:ММ) и срок каждой задачи: "
            "deadline_time (ЧЧ:ММ того же дня; если раньше create_time — считается следующий день) ИЛИ "
            "deadline_after_hours. Постановщик (creator) по умолчанию = текущий собеседник. Как и обычная "
            "задача, требует result_criteria. Можно задать соисполнителей/наблюдателей/теги/проект/CRM/"
            "пользовательские поля/чек-лист — они попадут в каждую созданную задачу. Просмотр — "
            "list_recurring_tasks, остановка — delete_recurring_task."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "title": {"type": "string", "description": "Название задачи."},
                "responsible_name": {"type": "string", "description": "Имя ответственного из оргструктуры."},
                "responsible_bitrix_user_id": {"type": "integer", "description": "Точный Bitrix id ответственного (приоритетнее имени)."},
                "period": {"type": "string", "enum": ["daily", "weekly", "monthly"], "description": "Тип повтора."},
                "weekdays": {"type": "array", "items": {"type": "string", "enum": ["MO", "TU", "WE", "TH", "FR", "SA", "SU"]}, "description": "Дни недели для weekly (например ['FR'] = каждую пятницу)."},
                "day_of_month": {"type": "integer", "minimum": 1, "maximum": 31, "description": "День месяца для monthly."},
                "interval": {"type": "integer", "minimum": 1, "description": "Каждые N единиц (по умолчанию 1). Напр. interval=2 + weekly = раз в две недели."},
                "create_time": {"type": "string", "description": "Время создания каждой задачи, ЧЧ:ММ (по умолчанию 10:00)."},
                "deadline_time": {"type": "string", "description": "Срок каждой задачи, ЧЧ:ММ того же дня (напр. 19:00). Если раньше create_time — срок на следующий день."},
                "deadline_after_hours": {"type": "number", "description": "Альтернатива deadline_time: срок через N часов после создания."},
                "result_criteria": {"type": "string", "description": "ОБЯЗАТЕЛЬНО: по чему поймём, что задача выполнена."},
                "description": {"type": "string", "description": "Описание задачи (необязательно)."},
                "creator_name": {"type": "string", "description": "Постановщик по имени (по умолчанию — текущий собеседник)."},
                "creator_bitrix_user_id": {"type": "integer", "description": "Постановщик по id (по умолчанию — текущий собеседник)."},
                "until": {"type": "string", "description": "Дата окончания повторов YYYY-MM-DD или DD.MM.YYYY (по умолчанию бессрочно)."},
                "priority": {"type": "string", "enum": ["normal", "high", "critical"], "description": "Приоритет каждой задачи."},
                "accomplice_names": {"type": "array", "items": {"type": "string"}, "description": "Соисполнители по именам."},
                "accomplice_bitrix_user_ids": {"type": "array", "items": {"type": "integer"}, "description": "Соисполнители по id."},
                "auditor_names": {"type": "array", "items": {"type": "string"}, "description": "Наблюдатели по именам."},
                "auditor_bitrix_user_ids": {"type": "array", "items": {"type": "integer"}, "description": "Наблюдатели по id."},
                "tags": {"type": "array", "items": {"type": "string"}, "description": "Теги задачи."},
                "group_id": {"type": "integer", "description": "Проект/рабочая группа."},
                "crm_elements": {"type": "array", "items": {"type": "string"}, "description": "Элементы CRM ('D_123','L_45','C_7','CO_9')."},
                "custom_fields": {"type": "object", "description": "Пользовательские поля {\"UF_...\": значение}.", "additionalProperties": True},
                "checklist": {"type": "array", "items": {"type": ["string", "object"]}, "description": "Чек-лист: строки или {title, complete}."},
            },
            "required": ["title", "period", "result_criteria"],
            "additionalProperties": False,
        },
        "handler": tool_create_recurring_task,
    },
    "list_recurring_tasks": {
        "description": (
            "Показать ПОВТОРЯЮЩИЕСЯ (регулярные) задачи — все или по одному человеку. Для каждой отдаёт "
            "расписание (человекочитаемо), ближайшее авто-создание (next_execution), ответственного и "
            "id шаблона. Используй, когда просят «какие повторяющиеся задачи у <человека>» или «покажи все "
            "регулярные задачи». Bitrix не отдаёт список шаблонов через REST, поэтому список ведётся в "
            "реестре агента (задачи, созданные через create_recurring_task). Эти же записи владелец "
            "видит в Центре Агента → Агенты → «Автоматизации»."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "responsible_name": {"type": "string", "description": "Показать задачи этого человека (по имени)."},
                "responsible_bitrix_user_id": {"type": "integer", "description": "Показать задачи этого человека (по id)."},
                "include_inactive": {"type": "boolean", "description": "Включить неактивные/удалённые из Bitrix (по умолчанию нет)."},
            },
            "additionalProperties": False,
        },
        "handler": tool_list_recurring_tasks,
    },
    "update_recurring_task": {
        "description": (
            "Изменить ПОВТОРЯЮЩУЮСЯ (регулярную) задачу: дни недели, время создания, дедлайн, "
            "название/описание/чек-лист/критерий результата. recurring_id бери из list_recurring_tasks. "
            "Дни недели: weekdays=['MO','TU','WE','TH','FR'] = по будням (просьба «не присылай в "
            "выходные» решается именно так); все 7 дней = ежедневно. Время: create_time (ЧЧ:ММ МСК); "
            "если дедлайн был «до 18:00», он сохранится автоматически. Расписание пересчитывается "
            "сразу (next_run в ответе); запись видна во вкладке «Автоматизации»."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "recurring_id": {"type": "integer", "description": "id повторяющейся задачи из list_recurring_tasks."},
                "weekdays": {"type": "array", "items": {"type": ["string", "integer"]}, "description": "Новые дни: MO/TU/WE/TH/FR/SA/SU или 1-7 (Пн=1). Все 7 = ежедневно."},
                "day_of_month": {"type": "integer", "minimum": 1, "maximum": 31, "description": "Перевести на ежемесячный повтор в этот день месяца."},
                "create_time": {"type": "string", "description": "Новое время создания, ЧЧ:ММ МСК."},
                "deadline_time": {"type": "string", "description": "Новый срок каждой задачи, ЧЧ:ММ того же дня."},
                "deadline_after_hours": {"type": "number", "description": "Альтернатива deadline_time: срок через N часов после создания."},
                "until": {"type": "string", "description": "Дата окончания повторов YYYY-MM-DD или DD.MM.YYYY."},
                "title": {"type": "string", "description": "Новое название задачи."},
                "description": {"type": "string", "description": "Новое описание задачи."},
                "checklist": {"type": "array", "items": {"type": ["string", "object"]}, "description": "Новый чек-лист (полностью заменяет старый)."},
                "result_criteria": {"type": "string", "description": "Новый критерий результата."},
                "priority": {"type": "string", "enum": ["normal", "high", "critical"], "description": "Приоритет создаваемых задач."},
            },
            "required": ["recurring_id"],
            "additionalProperties": False,
        },
        "handler": tool_update_recurring_task,
    },
    "get_employee_dossier": {
        "description": (
            "ДОСЬЕ по сотрудникам (внутреннее, для владельца/админа — рядовым не показывать): кто "
            "реально работает с агентом (ходы за 30 дней, из них в задачах), реакция на предложения "
            "помощи (offers_made/engaged/declined), какие задачи человека агент может ускорить "
            "(automatable), заметки. Обновляется ежедневным обходом задач. Без аргументов — все; "
            "name/bitrix_user_id — один человек."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "name": {"type": "string", "description": "Сотрудник по имени."},
                "bitrix_user_id": {"type": "integer", "description": "Сотрудник по id."},
            },
            "additionalProperties": False,
        },
        "handler": tool_get_employee_dossier,
    },
    "update_employee_dossier": {
        "description": (
            "Записать наблюдение в ДОСЬЕ сотрудника (внутреннее): паттерны его задач, что удалось/не "
            "удалось автоматизировать, как человек взаимодействует с агентом. По умолчанию note "
            "ДОПИСЫВАЕТСЯ с датой; replace=true — заменить заметки целиком."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "name": {"type": "string", "description": "Сотрудник по имени."},
                "bitrix_user_id": {"type": "integer", "description": "Сотрудник по id."},
                "note": {"type": "string", "description": "Наблюдение (до 1000 символов)."},
                "replace": {"type": "boolean", "description": "true = заменить все заметки этой."},
            },
            "required": ["note"],
            "additionalProperties": False,
        },
        "handler": tool_update_employee_dossier,
    },
    "delete_recurring_task": {
        "description": (
            "Остановить ПОВТОРЯЮЩУЮСЯ задачу — планировщик перестанет её создавать. Определи по "
            "recurring_id из list_recurring_tasks. Уже созданные задачи не трогаются. Требует "
            "confirm=true после показа задачи пользователю."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "recurring_id": {"type": "integer", "description": "id повторяющейся задачи из list_recurring_tasks."},
                "confirm": {"type": "boolean", "description": "Должно быть true — пользователь подтвердил остановку."},
            },
            "required": ["recurring_id", "confirm"],
            "additionalProperties": False,
        },
        "handler": tool_delete_recurring_task,
    },
    "list_chats": {
        "description": "List active non-excluded chats, optionally with message counts for a period.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "date_from": {"type": "string"},
                "date_to": {"type": "string"},
                "query": {"type": "string"},
                "limit": {"type": "integer", "minimum": 1, "maximum": MAX_LIMIT},
                "offset": {"type": "integer", "minimum": 0},
            },
            "additionalProperties": False,
        },
        "handler": tool_list_chats,
    },
    "search_messages": {
        "description": "Search raw chat messages and OCR text from attached images for a period. Returns original message text plus file OCR transcripts.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "date_from": {"type": "string"},
                "date_to": {"type": "string"},
                "query": {"type": "string"},
                "dialog_id": {"type": "string"},
                "include_ocr": {"type": "boolean"},
                "limit": {"type": "integer", "minimum": 1, "maximum": MAX_LIMIT},
                "offset": {"type": "integer", "minimum": 0},
            },
            "required": ["date_from", "date_to"],
            "additionalProperties": False,
        },
        "handler": tool_search_messages,
    },
    "list_bitrix_bot_sessions": {
        "description": "List conversations/sessions employees had with the AI assistant (Гермес-ассистент) INSIDE Bitrix24 chat — one row per dialog/user with message count, first/last activity, access tier, error count, current session epoch. Use this when asked about chats/sessions/interactions people had with the AI agent/bot in Bitrix (this is the bot's OWN log; human-to-human chats are list_chats/search_messages). Then read a full transcript with get_bitrix_bot_chat.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "date_from": {"type": "string", "description": "YYYY-MM-DD inclusive (optional)."},
                "date_to": {"type": "string", "description": "YYYY-MM-DD inclusive (optional)."},
                "limit": {"type": "integer", "minimum": 1, "maximum": MAX_LIMIT},
            },
            "additionalProperties": False,
        },
        "handler": tool_list_bitrix_bot_sessions,
    },
    "get_bitrix_bot_chat": {
        "description": "Read the full question→answer transcript of one person's conversation with the AI assistant in Bitrix24 (by dialog_id or bitrix_user_id), to analyze and improve answer quality. Optional `query` filters by substring in the question/answer.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "dialog_id": {"type": "string", "description": "Bitrix dialog id (from list_bitrix_bot_sessions)."},
                "bitrix_user_id": {"type": "integer", "description": "Bitrix user id (alternative to dialog_id)."},
                "query": {"type": "string", "description": "Optional substring filter over question/answer."},
                "date_from": {"type": "string", "description": "YYYY-MM-DD inclusive (optional)."},
                "date_to": {"type": "string", "description": "YYYY-MM-DD inclusive (optional)."},
                "limit": {"type": "integer", "minimum": 1, "maximum": MAX_LIMIT},
            },
            "additionalProperties": False,
        },
        "handler": tool_get_bitrix_bot_chat,
    },
    "get_ai_capabilities": {
        "description": "Return what YOU (this AI assistant) can do for the current connector/tool set. Call this whenever the user asks 'что ты умеешь', 'твои возможности', 'what can you do', or before promising an action — and answer strictly within the returned list. Your exact tool set is in start_here.available_tools; this is the human-readable, owner-maintained description.",
        "inputSchema": {"type": "object", "properties": {}, "additionalProperties": False},
        "handler": tool_get_ai_capabilities,
    },
    "update_ai_capabilities": {
        "description": "Record/update the assistant's capabilities note when this tool is enabled. Use when you gain or are told about a new capability so future sessions know it. mode 'append' (default) adds, 'replace' overwrites; stored note key 'full' (default) or 'faq'. Keep it concise and truthful.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "content": {"type": "string", "description": "Text to append or the full replacement."},
                "mode": {"type": "string", "enum": ["append", "replace"], "description": "Default 'append'."},
                "tier": {"type": "string", "enum": ["full", "faq"], "description": "Legacy stored note key. Default 'full'."},
            },
            "required": ["content"],
            "additionalProperties": False,
        },
        "handler": tool_update_ai_capabilities,
    },
    "get_chat_transcript": {
        "description": "Get raw chat transcript messages by dialog_id and period, including OCR transcripts for attached images and PDFs.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "dialog_id": {"type": "string"},
                "date_from": {"type": "string"},
                "date_to": {"type": "string"},
                "include_ocr": {"type": "boolean"},
                "limit": {"type": "integer", "minimum": 1, "maximum": MAX_LIMIT},
                "offset": {"type": "integer", "minimum": 0},
            },
            "required": ["dialog_id", "date_from", "date_to"],
            "additionalProperties": False,
        },
        "handler": tool_get_chat_transcript,
    },
    "get_chat_ocr_status": {
        "description": "Check whether image/PDF attachments for one chat/date already have OCR text before generating a daily report.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "dialog_id": {"type": "string"},
                "report_date": {"type": "string", "description": "YYYY-MM-DD"},
            },
            "required": ["dialog_id", "report_date"],
            "additionalProperties": False,
        },
        "handler": tool_get_chat_ocr_status,
    },
    "process_chat_ocr": {
        "description": "Run OCR processing for image/PDF chat attachments through the local app workflow. Use when get_chat_ocr_status shows missing/pending OCR before a daily chat report.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "date_from": {"type": "string", "description": "YYYY-MM-DD"},
                "date_to": {"type": "string", "description": "YYYY-MM-DD"},
                "dialog_id": {"type": "string"},
                "force": {"type": "boolean"},
            },
            "required": ["date_from"],
            "additionalProperties": False,
        },
        "handler": tool_process_chat_ocr,
    },
    "list_zoom_calls": {
        "description": "List Zoom cloud recordings/calls with dates, technical topics, participants, and transcript segment counts.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "date_from": {"type": "string", "description": "YYYY-MM-DD"},
                "date_to": {"type": "string", "description": "YYYY-MM-DD"},
                "query": {"type": "string"},
                "limit": {"type": "integer", "minimum": 1, "maximum": MAX_LIMIT},
                "offset": {"type": "integer", "minimum": 0},
            },
            "additionalProperties": False,
        },
        "handler": tool_list_zoom_calls,
    },
    "get_zoom_call_transcript": {
        "description": "Get one Zoom call with factual Zoom participants and raw transcript segments. For Zoom reports, combine with get_org_structure to map participants and mentioned people to roles/departments.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "call_id": {"type": "string"},
                "zoom_uuid": {"type": "string"},
                "include_full_text": {"type": "boolean"},
                "limit": {"type": "integer", "minimum": 1, "maximum": ZOOM_TRANSCRIPT_MAX_LIMIT},
                "offset": {"type": "integer", "minimum": 0},
            },
            "additionalProperties": False,
        },
        "handler": tool_get_zoom_call_transcript,
    },
    "export_zoom_call_markdown": {
        "description": (
            "Export ONE Zoom call as a Markdown document: header with topic, date, time (МСК), duration "
            "and participants, then the FULL transcript line by line (speaker + timecode). The file is "
            "saved server-side and the tool returns {download_url, filename, call_id, chars, bytes, preview} "
            "— the FULL transcript is NOT inlined (clients truncate it). Use when the owner asks to get/send "
            "a call's transcript 'в md'/'markdown'/'файлом': give them the `download_url` (public, login-free, "
            "unguessable link that EXPIRES in ~30 min — send it promptly); `preview` is only the first lines for context. "
            "SCOPE: this is for exactly ONE call the owner pointed at. If the owner asks for a DATE/DAY or "
            "'все встречи/созвоны за <дата>', do NOT call this per-call and do NOT export just one — a day "
            "usually has several meetings; use export_zoom_transcripts_markdown with date_from=date_to=<day> "
            "to get ALL of them in one file."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "call_id": {"type": "string", "description": "Zoom call id from list_zoom_calls/get_zoom_call_transcript."},
            },
            "required": ["call_id"],
            "additionalProperties": False,
        },
        "handler": tool_export_zoom_call_markdown,
    },
    "export_zoom_transcripts_markdown": {
        "description": (
            "Export MANY Zoom calls into ONE Markdown document with a table of contents and clear '---' "
            "boundaries between meetings; each meeting has metadata (topic, date, МСК time, duration, "
            "participants) plus its FULL transcript. THIS IS THE DEFAULT TOOL for any 'выгрузи "
            "созвоны/встречи/транскрипты за <даты/период>' request, and it ALWAYS returns EVERY meeting on "
            "the chosen days (a day often has 2–3 meetings). Choose ONE selector: "
            "(1) `dates` — an array of specific YYYY-MM-DD days (USE THIS for a list of non-contiguous days, "
            "e.g. 7,8,14,15 мая → dates=['2026-05-07','2026-05-08','2026-05-14','2026-05-15']); "
            "(2) `date_from`+`date_to` — a continuous range (for one day pass the same date twice); "
            "(3) `call_ids` — only if you already have exact ids. "
            "DO NOT call list_zoom_calls and paginate it to collect ids — just pass `dates`; this tool "
            "resolves all meetings itself in ONE call. Google Drive transcript imports (noisy duplicates) "
            "are excluded unless include_google_drive=true. The file is saved server-side and the tool "
            "returns {download_url, filename, calls, chars, bytes, preview} — the FULL document is NOT "
            "inlined (it would be truncated by the client). Give the owner the `download_url` (a public, "
            "login-free, unguessable link that EXPIRES in ~30 min — send it promptly); `preview` is just the first lines for context."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "dates": {"type": "array", "items": {"type": "string"}, "description": "Specific days as YYYY-MM-DD. Best for a list of non-contiguous days; returns ALL meetings on each."},
                "call_ids": {"type": "array", "items": {"type": "string"}, "description": "Explicit Zoom call ids; overrides dates and the range."},
                "date_from": {"type": "string", "description": "YYYY-MM-DD (continuous range start)"},
                "date_to": {"type": "string", "description": "YYYY-MM-DD (continuous range end)"},
                "include_google_drive": {"type": "boolean", "description": "Keep Google Drive transcript imports (default false)."},
            },
            "additionalProperties": False,
        },
        "handler": tool_export_zoom_transcripts_markdown,
    },
    "search_zoom_transcripts": {
        "description": "Search Zoom transcript segments by text and optional date range. For chat context, search keywords derived from chat OCR, tasks, risks, project names, and owner names.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "query": {"type": "string"},
                "date_from": {"type": "string", "description": "YYYY-MM-DD"},
                "date_to": {"type": "string", "description": "YYYY-MM-DD"},
                "limit": {"type": "integer", "minimum": 1, "maximum": MAX_LIMIT},
                "offset": {"type": "integer", "minimum": 0},
            },
            "required": ["query"],
            "additionalProperties": False,
        },
        "handler": tool_search_zoom_transcripts,
    },
    "save_zoom_call_report": {
        "description": "Save a generated AI report for one Zoom call directly to zoom_calls.analytical_note in PostgreSQL. Standalone Zoom reports must include factual participants, mentioned people mapped to org structure, strict task ownership/deadline gaps, and behavioral factors.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "call_id": {"type": "string", "description": "Zoom call id from list_zoom_calls/get_zoom_call_transcript."},
                "zoom_uuid": {"type": "string", "description": "Zoom UUID, alternative to call_id."},
                "summary": {"type": "string"},
                "report_text": {"type": "string", "description": "Human-readable Zoom report to store in zoom_calls.analytical_note."},
                "analysis": {"type": "object", "description": "ОБЯЗАТЕЛЬНО передай ВЕСЬ JSON-объект из контракта zoom_processing ЦЕЛИКОМ и без сокращений: dispatch_summary, leader_evaluations, people, и operational_tasks где у каждой задачи assignee_name, bitrix_user_id, deadline_text, result_criteria, expected_artifact, responsibility_check, status, source. НЕ передавай урезанный/сводный analysis (например {leaders_present, operational_tasks_count}) — иначе рассылка теряет сводку, оценку руководителя и артефакты."},
                "model": {"type": "string"},
                "status": {"type": "string", "enum": ["done", "error"]},
                "raw_input": {"type": "object", "description": "Source manifest: transcript ids/segments, prompt id, checked context."},
            },
            "additionalProperties": False,
        },
        "handler": tool_save_zoom_call_report,
    },
    "delete_zoom_call_report": {
        "description": "Delete only the AI report for one Zoom call from zoom_calls.analytical_note and raw_json.ai_report. Does not delete the Zoom call, participants, or transcript.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "call_id": {"type": "string", "description": "Zoom call id from list_zoom_calls/get_zoom_call_transcript."},
                "zoom_uuid": {"type": "string", "description": "Zoom UUID, alternative to call_id."},
            },
            "additionalProperties": False,
        },
        "handler": tool_delete_zoom_call_report,
    },
    "get_owner_reports": {
        "description": "Read recent current owner daily or weekly reports. Use before recommendations and management answers to understand prior context, done/open items, and repeated issues.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "report_kind": {"type": "string", "enum": ["daily", "weekly"]},
                "limit": {"type": "integer", "minimum": 1, "maximum": 500},
            },
            "additionalProperties": False,
        },
        "handler": tool_get_owner_reports,
    },
    "list_recommendations": {
        "description": "List addressable recommendations with lifecycle status and optional event history. Use this before checking recommendation feedback or repeated recommendations.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "date_from": {"type": "string", "description": "YYYY-MM-DD"},
                "date_to": {"type": "string", "description": "YYYY-MM-DD"},
                "status": {"type": "string", "description": "Use 'open' for all non-final statuses, or a concrete recommendation status."},
                "dialog_id": {"type": "string"},
                "manager_bitrix_user_id": {"type": "integer"},
                "employee_bitrix_user_id": {"type": "integer"},
                "include_events": {"type": "boolean"},
                "limit": {"type": "integer", "minimum": 1, "maximum": MAX_LIMIT},
                "offset": {"type": "integer", "minimum": 0},
            },
            "additionalProperties": False,
        },
        "handler": tool_list_recommendations,
    },
    "get_recommendation_feedback_context": {
        "description": "Read active recommendations and event history relevant to one chat/date. Use with raw chat transcripts to analyze previous-day and current-day replies.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "dialog_id": {"type": "string"},
                "report_date": {"type": "string", "description": "YYYY-MM-DD"},
            },
            "required": ["dialog_id", "report_date"],
            "additionalProperties": False,
        },
        "handler": tool_get_recommendation_feedback_context,
    },
    "save_recommendation_event": {
        "description": "Append one lifecycle event for an addressable recommendation and optionally update its status. Use after interpreting a human reply.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "recommendation_id": {"type": "string"},
                "event_type": {"type": "string", "enum": ["created", "sent", "delivered", "seen", "employee_replied", "ai_interpreted", "status_changed", "manager_reviewed", "task_created", "closed", "source_found"]},
                "author_type": {"type": "string", "enum": ["system", "ai", "manager", "employee"]},
                "author_bitrix_user_id": {"type": "integer"},
                "dialog_id": {"type": "string"},
                "bitrix_message_id": {"type": "integer"},
                "chat_message_day": {"type": "string", "description": "YYYY-MM-DD"},
                "old_status": {"type": "string"},
                "new_status": {"type": "string"},
                "event_text": {"type": "string"},
                "interpretation": {"type": "object"},
                "source_payload": {"type": "object"},
                "manager_review_required": {"type": "boolean"},
            },
            "required": ["recommendation_id"],
            "additionalProperties": False,
        },
        "handler": tool_save_recommendation_event,
    },
    "get_previous_owner_daily_context": {
        "description": "Read only the previous calendar day's current owner daily report as continuity context for creating or checking a new owner daily report.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "report_date": {"type": "string", "description": "Target owner daily report date in YYYY-MM-DD. The tool reads report_date minus one day."},
            },
            "required": ["report_date"],
            "additionalProperties": False,
        },
        "handler": tool_get_previous_owner_daily_context,
    },
    "save_owner_daily_report": {
        "description": "Save a generated daily owner report directly to owner_daily_reports. Use for the general daily owner report, not for per-chat reports.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "report_date": {"type": "string", "description": "YYYY-MM-DD"},
                "summary": {"type": "string"},
                "dynamics_summary": {"type": "string"},
                "risks_summary": {"type": "string"},
                "recommendations": {"type": "string"},
                "manager_recommendations": {"type": ["array", "object", "string"]},
                "manager_messages": {"type": ["array", "object", "string"]},
                "open_tasks": {"type": ["array", "object", "string"]},
                "overdue_tasks": {"type": ["array", "object", "string"]},
                "no_response": {"type": ["array", "object", "string"]},
                "goal_dynamics": {"type": ["array", "object", "string"]},
                "report_text": {"type": "string"},
                "raw_json": {"type": "object"},
                "analysis": {"type": "object"},
                "raw_input": {"type": "object"},
                "model": {"type": "string"},
                "status": {"type": "string", "enum": ["done", "no_data", "error"]},
            },
            "required": ["report_date"],
            "additionalProperties": False,
        },
        "handler": tool_save_owner_daily_report,
    },
    "save_owner_weekly_report": {
        "description": "Save a generated weekly owner report directly to owner_weekly_reports. Use for the general weekly owner report, not for per-chat reports.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "period_start": {"type": "string", "description": "YYYY-MM-DD"},
                "period_end": {"type": "string", "description": "YYYY-MM-DD"},
                "summary": {"type": "string"},
                "dynamics_summary": {"type": "string"},
                "risks_summary": {"type": "string"},
                "recommendations": {"type": "string"},
                "manager_recommendations": {"type": ["array", "object", "string"]},
                "manager_messages": {"type": ["array", "object", "string"]},
                "open_tasks": {"type": ["array", "object", "string"]},
                "overdue_tasks": {"type": ["array", "object", "string"]},
                "no_response": {"type": ["array", "object", "string"]},
                "goal_dynamics": {"type": ["array", "object", "string"]},
                "weekly_dynamics": {"type": ["array", "object", "string"]},
                "daily_owner_reports_manifest": {"type": ["array", "object", "string"]},
                "report_text": {"type": "string"},
                "raw_json": {"type": "object"},
                "analysis": {"type": "object"},
                "raw_input": {"type": "object"},
                "model": {"type": "string"},
                "status": {"type": "string", "enum": ["done", "no_data", "error"]},
            },
            "required": ["period_start", "period_end"],
            "additionalProperties": False,
        },
        "handler": tool_save_owner_weekly_report,
    },
    "send_owner_weekly_report_pdf": {
        "description": (
            "Send the current weekly owner report as a PDF into Bitrix personal messages (default recipient: "
            "Evgeniy Palei, bitrix_user_id 1). Builds the PDF from the saved owner_weekly_report for the given period "
            "and attaches it via im.disk.file.commit. STRICT RULE: confirm=true is mandatory — first build and save "
            "the weekly report via save_owner_weekly_report, show the owner the report period and that the PDF goes to "
            "Evgeniy, get explicit approval, and only then call this tool with confirm=true. The weekly report for that "
            "period must already be saved (is_current)."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "period_start": {"type": "string", "description": "YYYY-MM-DD (Monday of the reported week)."},
                "period_end": {"type": "string", "description": "YYYY-MM-DD (Friday/Sunday of the reported week)."},
                "recipient_bitrix_user_ids": {
                    "type": "array",
                    "items": {"type": "integer"},
                    "description": "Bitrix user ids to send the PDF to. Default [1] (Evgeniy Palei).",
                },
                "confirm": {
                    "type": "boolean",
                    "description": "Must be true. Means the owner explicitly approved sending the weekly PDF to Evgeniy.",
                },
            },
            "required": ["period_start", "period_end", "confirm"],
            "additionalProperties": False,
        },
        "handler": tool_send_owner_weekly_report_pdf,
    },
    "list_pending_zoom_operational_dispatches": {
        "description": (
            "List Zoom calls that have a saved analytical report but were NOT yet dispatched as aggregated "
            "'Итоги созвона' tasks to Bitrix. Each pending entry has call_id, call_date, topic, time, and duration. "
            "Use this AS THE FIRST STEP when the owner answers 'ставь' / 'создавай' / 'да' after a zoom-to-tasks "
            "Telegram summary: it tells you exactly which call_ids still need dispatching. "
            "DEFAULT PERIOD: today only (Europe/Moscow) — covers exactly the calls the owner just approved. "
            "Pass date_from explicitly if you really need to dispatch older calls (rare; usually the owner only "
            "wants the call they just saw in the summary)."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "date_from": {"type": "string", "description": "YYYY-MM-DD (default: today minus 2 days)"},
                "date_to": {"type": "string", "description": "YYYY-MM-DD (default: today)"},
            },
            "additionalProperties": False,
        },
        "handler": tool_list_pending_zoom_operational_dispatches,
    },
    "preview_zoom_operational_tasks": {
        "description": (
            "Preview the aggregated 'Итоги созвона' Bitrix tasks that would be created for one Zoom call WITHOUT sending. "
            "Returns title (e.g. 'Итоги созвона 09:28'), per-recipient cards (one card = one aggregated task per "
            "responsible person, grouping all of their operational_tasks from the call), deadline (18:00 МСК сегодня; "
            "если до 18:00 меньше 3 часов или выходной — следующий рабочий день 11:00), "
            "and the standard description '"
            "Ознакомьтесь со списком выделенных из созвона задач и поставьте себе самые важные в Битрикс...'. "
            "Do not call this if you only need to send — just call dispatch_zoom_operational_tasks directly."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "call_id": {"type": "string", "description": "Zoom call UUID (from list_pending_zoom_operational_dispatches or list_zoom_calls)."},
            },
            "required": ["call_id"],
            "additionalProperties": False,
        },
        "handler": tool_preview_zoom_operational_tasks,
    },
    "dispatch_zoom_operational_tasks": {
        "description": (
            "Create aggregated 'Итоги созвона <ЧЧ:ММ>' Bitrix tasks for ONE Zoom call: one task per responsible person, "
            "deadline = 18:00 МСК того же дня (если до 18:00 меньше 3 часов или выходной — следующий рабочий день "
            "11:00), description = standard 'Ознакомьтесь со списком...' header plus the list of "
            "this person's operational_tasks from the call. Behaves EXACTLY like Albery UI button 'Отправка задач'. "
            "STRICT CONFIRMATION RULE: do not call unless the owner has just explicitly approved sending (replied "
            "'ставь' / 'создавай' / 'да' to the zoom-to-tasks Telegram summary). confirm=true is mandatory. "
            "DO NOT use create_bitrix_task to recreate individual tasks here — this is the right tool for zoom "
            "operational tasks; it groups them and produces the exact card format the owner expects. "
            "After success the call is marked with raw_json.ai_report.bitrix_dispatch.dispatched_at so "
            "list_pending_zoom_operational_dispatches will not return it again. "
            "PARTIAL SUCCESS: the result returns 'sent' (number of tasks created) and 'skipped_assignees' "
            "(responsible people with no matching Bitrix user, e.g. not in the team sync). This is a SUCCESS, "
            "not an error — do NOT retry. Report it to the owner like 'Поставил N задач; не нашёл в Битриксе: "
            "<names> — добавьте их в Битрикс или поставьте задачу вручную.'"
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "call_id": {"type": "string", "description": "Zoom call UUID."},
                "confirm": {
                    "type": "boolean",
                    "description": "Must be true. Means the owner explicitly approved creating aggregated 'Итоги созвона' tasks for this call.",
                },
            },
            "required": ["call_id", "confirm"],
            "additionalProperties": False,
        },
        "handler": tool_dispatch_zoom_operational_tasks,
    },
    "preview_zoom_participant_reports": {
        "description": (
            "Preview personal participant report Bitrix tasks for one Zoom call WITHOUT sending: one supportive task per matched participant, "
            "with shared call outcomes plus personal soft evaluation. Uses raw_json.ai_report.analysis dispatch_summary, people, "
            "leader_evaluations and person_summaries when present. Separate from operational task dispatch."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {"call_id": {"type": "string", "description": "Zoom call UUID."}},
            "required": ["call_id"],
            "additionalProperties": False,
        },
        "handler": tool_preview_zoom_participant_reports,
    },
    "dispatch_zoom_participant_reports": {
        "description": (
            "Create personal participant report Bitrix tasks for one Zoom call: one task per matched participant, containing shared summary "
            "and supportive personal feedback. Requires confirm=true. Marks raw_json.ai_report.participant_reports_dispatched_at and "
            "participant_report_task_ids; does NOT affect operational bitrix_dispatch status."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "call_id": {"type": "string", "description": "Zoom call UUID."},
                "confirm": {"type": "boolean", "description": "Must be true after explicit approval."},
            },
            "required": ["call_id", "confirm"],
            "additionalProperties": False,
        },
        "handler": tool_dispatch_zoom_participant_reports,
    },
    "list_leader_evaluations": {
        "description": (
            "Read aggregated leader evaluations (how Артур, Наталья, Евгений, Сергей run their calls) across saved "
            "zoom reports in a date range. Returns by_leader: per leader a list of call evaluations "
            "{call_date, topic, role (host/co_leader), verdict (good/minor_issue/issue), result_for_owner} plus "
            "verdict counts, and calls_count. DEFAULT PERIOD: previous Friday .. today (the Wednesday digest window — "
            "Fri + Mon + Tue + Wed). Use this on the Wednesday owner-digest cron to compose a per-leader summary with "
            "dynamics, show it to the owner in Telegram for approval, then send via dispatch_leader_evaluations_digest."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "date_from": {"type": "string", "description": "YYYY-MM-DD (default: today minus 5 days = previous Friday)"},
                "date_to": {"type": "string", "description": "YYYY-MM-DD (default: today)"},
            },
            "additionalProperties": False,
        },
        "handler": tool_list_leader_evaluations,
    },
    "dispatch_leader_evaluations_digest": {
        "description": (
            "Create the weekly leader-evaluation digest as ONE Bitrix task for the owner Евгений Палей: title "
            "'Ознакомиться с оценкой руководителей за период <даты>', description = the approved digest_text, deadline = "
            "next calendar day 10:00 МСК (sent Wednesday evening, reviewed by Thursday 10:00). STRICT CONFIRMATION RULE: "
            "call only after the owner explicitly approved the digest in Telegram ('отправляй'); confirm=true is mandatory. "
            "digest_text must be the final approved per-leader summary text."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "digest_text": {"type": "string", "description": "Final approved per-leader digest text for the task description."},
                "date_from": {"type": "string", "description": "YYYY-MM-DD (default: previous Friday)"},
                "date_to": {"type": "string", "description": "YYYY-MM-DD (default: today)"},
                "confirm": {"type": "boolean", "description": "Must be true. Owner explicitly approved sending the digest."},
            },
            "required": ["digest_text", "confirm"],
            "additionalProperties": False,
        },
        "handler": tool_dispatch_leader_evaluations_digest,
    },
    "list_pending_owner_recommendations": {
        "description": (
            "List addressed manager recommendations from the current owner_daily_report for a given date. "
            "Returns rows with id, manager_full_name, manager_bitrix_user_id, recommendation_text, subject, priority, due_date, status. "
            "Also returns the report summary, dynamics_summary, risks_summary, and report_text so the agent can build "
            "the per-recipient recommendation list that will become each person's Bitrix task. "
            "Excludes recommendations already sent, cancelled, done, or rejected. "
            "Use this after save_owner_daily_report and before send_owner_recommendations_to_bitrix."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "report_date": {"type": "string", "description": "YYYY-MM-DD"},
            },
            "required": ["report_date"],
            "additionalProperties": False,
        },
        "handler": tool_list_pending_owner_recommendations,
    },
    "send_owner_recommendations_to_bitrix": {
        "description": (
            "Create owner_daily_report recommendation TASKS in Bitrix from the configured BITRIX_WEBHOOK_BASE account "
            "(owner) — one task per recipient, NOT personal messages. Each task is titled 'Рекомендации DD.MM', the "
            "recipient is the responsible person, the description asks them to react (неактуально / взял в работу / "
            "уже в работе / уже есть результат / уже есть рабочий файл — «ссылка») and then lists their recommendations. "
            "The deadline is fixed at 12:00 of the next WORKING day (Europe/Moscow; a Friday report is due "
            "Monday 12:00) and the assignee cannot move it "
            "(ALLOW_CHANGE_DEADLINE=N). STRICT RULE: do not call this tool unless the owner has just approved the exact "
            "recommendation texts. Confirm=true is mandatory. Each entry in recipient_recommendations is the final "
            "recommendation body for that Bitrix user id (a clean numbered list, no greeting) and becomes the task "
            "description body verbatim. Uses tasks.task.add. Logs the outcome to owner_recommendation_dispatches and "
            "updates owner_manager_recommendations.status to 'sent' for matching rows."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "report_date": {"type": "string", "description": "YYYY-MM-DD — the owner_daily_report date."},
                "recipient_recommendations": {
                    "type": "object",
                    "description": "Map of bitrix_user_id -> final recommendation body (clean numbered list, no greeting). Becomes the task description body verbatim.",
                    "additionalProperties": {"type": "string"},
                },
                "confirm": {
                    "type": "boolean",
                    "description": "Must be true. Means the owner has explicitly approved creating these recommendation tasks.",
                },
            },
            "required": ["report_date", "recipient_recommendations", "confirm"],
            "additionalProperties": False,
        },
        "handler": tool_send_owner_recommendations_to_bitrix,
    },
    "dispatch_owner_weekly_report_task": {
        "description": (
            "Create the weekly owner report as a Bitrix TASK for Евгений Палей with the report PDF attached "
            "(UF_TASK_WEBDAV_FILES). Title 'Ознакомиться с недельным отчётом за <период>', deadline next Monday "
            "10:00 МСК, cannot be completed without a result (SE_PARAMETER code 3). Uses build_owner_report_pdf + "
            "upload_pdf_to_bitrix_disk + tasks.task.add. confirm=true is mandatory (the Friday cron passes it "
            "automatically after generating the report)."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "report_id": {"type": "string", "description": "owner_weekly_reports.id of the report to attach."},
                "confirm": {"type": "boolean", "description": "Must be true."},
            },
            "required": ["report_id", "confirm"],
            "additionalProperties": False,
        },
        "handler": tool_dispatch_owner_weekly_report_task,
    },
    "send_bitrix_message": {
        "description": (
            "Send one personal Bitrix message to a single employee via the configured BITRIX_WEBHOOK_BASE account "
            "(your own user — there is no separate bot user). Uses im.message.add with fallback to "
            "im.notify.personal.add when the private chat is blocked. STRICT RULES: "
            "(1) confirm=true is mandatory — first resolve the recipient and show the user the exact final "
            "message_text together with the recipient's full_name + work_position + bitrix_user_id, then ask "
            "for explicit approval; only then call this tool with confirm=true. "
            "(2) Pass either recipient_bitrix_user_id (preferred — exact integer) or recipient_name (fuzzy lookup "
            "against active users). If recipient_name is ambiguous, the tool returns the candidates and refuses "
            "to send — re-ask the user and call again with recipient_bitrix_user_id. "
            "(3) message_text is sent verbatim — do not wrap, edit, or prepend signatures."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "recipient_bitrix_user_id": {
                    "type": "integer",
                    "description": "Exact Bitrix user id of the recipient. Preferred over recipient_name.",
                },
                "recipient_name": {
                    "type": "string",
                    "description": "Full name (or distinctive part) of the recipient. Used only when recipient_bitrix_user_id is not provided. Fuzzy-matched against active users; ambiguity is refused.",
                },
                "message_text": {
                    "type": "string",
                    "description": "Final message to send, verbatim. Up to 20000 characters.",
                },
                "confirm": {
                    "type": "boolean",
                    "description": "Must be true. Means the user has explicitly approved sending this exact text to this exact recipient.",
                },
            },
            "required": ["message_text", "confirm"],
            "additionalProperties": False,
        },
        "handler": tool_send_bitrix_message,
    },
    "write_company_sheet": {
        "description": (
            "Записать данные в Google-таблицу компании (выполняется от имени владельца через Apps Script, "
            "поэтому может редактировать таблицы владельца). mode='append' добавляет строки в конец листа "
            "(rows — список списков ячеек); mode='update' пишет values в A1-диапазон range (напр. 'A2:D11'). "
            "Принимает spreadsheet_id ИЛИ полную ссылку на таблицу. Перед записью ОБЯЗАТЕЛЬНО покажи пользователю, "
            "что и куда впишешь, получи согласие, затем вызови с confirm=true. Инструмент только для полного "
            "доступа (в FAQ-коннекторе недоступен)."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "spreadsheet_id": {"type": "string", "description": "ID таблицы или полная ссылка на Google Sheet"},
                "sheet": {"type": "string", "description": "Имя листа (опц.; по умолчанию первый лист)"},
                "mode": {"type": "string", "enum": ["append", "update"], "description": "append — строки в конец; update — диапазон"},
                "rows": {"type": "array", "items": {"type": "array"}, "description": "Для append: список строк (каждая — список ячеек)"},
                "range": {"type": "string", "description": "Для update: A1-диапазон, напр. 'A2:D11'"},
                "values": {"type": "array", "items": {"type": "array"}, "description": "Для update: значения для диапазона (список списков)"},
                "confirm": {"type": "boolean", "description": "Должно быть true после явного согласия пользователя"},
            },
            "required": ["spreadsheet_id", "confirm"],
            "additionalProperties": False,
        },
        "handler": tool_write_company_sheet,
    },
    "create_google_sheet": {
        "description": (
            "Создать НОВУЮ Google-таблицу (от имени Google-аккаунта агента). По умолчанию выдаёт доступ "
            "«по ссылке — редактор» (anyone with link = editor) и возвращает ссылку. Можно сразу вписать "
            "данные через rows (список строк, каждая — список ячеек, с A1). Если rows переданы, сервер "
            "автоматически применяет спокойное читабельное оформление: перенос текста, автоширины под русский "
            "текст/числа/₽, аккуратные границы, лёгкая зебра и без лишней пёстрой заливки. Перед созданием "
            "ОБЯЗАТЕЛЬНО покажи пользователю название и что впишешь, получи согласие, затем вызови с confirm=true. "
            "Инструмент только для полного/операционного доступа (в FAQ-коннекторе недоступен)."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "title": {"type": "string", "description": "Название новой таблицы"},
                "rows": {"type": "array", "items": {"type": "array"}, "description": "Опц.: начальные данные с A1 (список строк)"},
                "share_anyone_writer": {"type": "boolean", "description": "Доступ «по ссылке — редактор» (по умолчанию true)"},
                "confirm": {"type": "boolean", "description": "Должно быть true после явного согласия пользователя"},
            },
            "required": ["title", "confirm"],
            "additionalProperties": False,
        },
        "handler": tool_create_google_sheet,
    },
    "get_google_sheet_meta": {
        "description": (
            "Read a Google Sheet's structure: its tabs (sheetId / title / grid size). Call this "
            "before format_google_sheet so you know each tab's sheetId for batchUpdate requests. Read-only."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {"spreadsheet_id": {"type": "string", "description": "Google Sheet id"}},
            "required": ["spreadsheet_id"],
            "additionalProperties": False,
        },
        "handler": tool_get_google_sheet_meta,
    },
    "write_google_sheet_values": {
        "description": (
            "Write a 2D array of values/formulas into an A1 range of a Google Sheet (USER_ENTERED, "
            "so formulas work). The server normalizes formula argument separators for the target Sheet locale "
            "(ru_RU uses semicolons, not commas) and validates the updated range; if formulas still produce "
            "#ERROR, the tool fails instead of letting the assistant claim success. Use for sheets the agent "
            "created via create_google_sheet."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "spreadsheet_id": {"type": "string"},
                "range": {"type": "string", "description": "A1 range, e.g. 'G3' or 'Sheet1!A1'"},
                "values": {"type": "array", "items": {"type": "array"}, "description": "2D array of rows"},
                "value_input_option": {"type": "string", "enum": ["USER_ENTERED", "RAW"]},
            },
            "required": ["spreadsheet_id", "range", "values"],
            "additionalProperties": False,
        },
        "handler": tool_write_google_sheet_values,
    },
    "format_google_sheet": {
        "description": (
            "Style a sheet and build dashboards: apply a list of Sheets API batchUpdate request objects. "
            "A Google Sheet or task/list table is NOT finished if it is a raw grid or a cramped rainbow dashboard. "
            "Every delivered table must be beautiful and readable: clear title/header, frozen header row, readable "
            "column widths, wrapping, enough row height, light borders, number/currency/date formats, and charts only "
            "when they clarify the data. Use a calm executive-dashboard palette: mostly white/light neutral background, "
            "one main accent color, muted section headers, and status colors only for small status cells; do NOT flood "
            "whole tables with red/green/yellow or multiple bright colors. Before giving a link to the user, make sure "
            "texts and ₽/number cells are not clipped. The server automatically runs a readability polish after your "
            "custom requests: wrap text, auto-fit rows/columns, then enforce sane widths from actual cell contents. "
            "create_google_sheet auto-applies a default readable style when rows are provided; use this tool for any "
            "extra formatting, dashboards, charts, task/status colors, merged title blocks or polished layout. "
            "Before saying a dashboard is ready, write/check formulas with write_google_sheet_values so formula "
            "errors are caught; for ru_RU sheets use semicolons in formulas. "
            "Get each tab's sheetId from get_google_sheet_meta. requests = standard Sheets API request objects "
            "(repeatCell, mergeCells, updateSheetProperties, addChart, addConditionalFormatRule, "
            "updateDimensionProperties, ...)."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "spreadsheet_id": {"type": "string"},
                "requests": {"type": "array", "items": {"type": "object"}, "description": "List of Sheets API batchUpdate request objects"},
            },
            "required": ["spreadsheet_id", "requests"],
            "additionalProperties": False,
        },
        "handler": tool_format_google_sheet,
    },
    "move_drive_file_to_folder": {
        "description": (
            "Move a Google Drive item — file, spreadsheet, document OR folder — into another Drive folder "
            "(folder id or URL). This replaces the item's previous parent folders with the target folder, so use "
            "it for moving files inside folders and moving nested folders. IMPORTANT: the target folder must be "
            "shared with the agent's account (a9ent.ai@gmail.com) as Editor, otherwise it returns 404 - in that "
            "case ask the owner to give a9ent.ai@gmail.com edit access to the folder first."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "file_id": {"type": "string", "description": "Drive file/folder id or URL"},
                "item_id": {"type": "string", "description": "Alias for file_id; Drive file/folder id or URL"},
                "folder": {"type": "string", "description": "Target Drive folder id or folder URL"},
            },
            "required": ["file_id", "folder"],
            "additionalProperties": False,
        },
        "handler": tool_move_drive_file_to_folder,
    },
    "get_webapp_template": {
        "description": (
            "Get the Albery-branded HTML/CSS web-app template (matches the prod React site: light bg, "
            "white rounded cards, primary purple #5440F6, Inter font, soft shadows, styled inputs/buttons/"
            "tables/badges). Returns html_skeleton with {{TITLE}}/{{CONTENT}}/{{APPLET}} placeholders, css, "
            "a working content_example and how_to. ALWAYS use this as the base for any web app so it looks "
            "consistent and beautiful. Combine with make_sheet_applet for data."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {"title": {"type": "string", "description": "App title to inline (optional)"}},
            "additionalProperties": False,
        },
        "handler": tool_get_webapp_template,
    },
    "export_document": {
        "description": (
            "Собери НАСТОЯЩИЙ файл документа Word (.docx) из СВОЕГО HTML — ты сам полностью управляешь "
            "оформлением (это главный инструмент для договоров, актов, официальных документов). "
            "Верни пользователю ссылку url из ответа. ПОДДЕРЖИВАЕМЫЙ HTML: блоки h1..h4, p, div, "
            "ul/ol/li, table/tr/th/td, br, hr; инлайн b/strong, i/em, u, s; стили text-align "
            "(left|center|right|justify), font-size:14pt, text-indent:1.25cm, line-height:1.5; "
            "НОВАЯ СТРАНИЦА: style=\"page-break-before:always\" на блоке (каждое приложение к договору "
            "начинай с него); таблица БЕЗ рамок: <table border=\"0\"> (реквизиты сторон делай "
            "двухколоночной таблицей без рамок). Шрифт всегда Times New Roman, поля ГОСТ (лево 3 см), "
            "А4. Эмодзи вырезаются автоматически — НЕ используй их и BB-коды [b] в html. Нумерацию "
            "разделов/пунктов (1., 1.1., 5.2.1.) пиши ЯВНО в тексте — так надёжнее списков. "
            "В html клади ТОЛЬКО сам документ: без комментариев, рекомендаций и пояснений — их пиши "
            "отдельно в чате.\n\n"
            "⚠️ ДЛИННЫЙ ДОКУМЕНТ (договор, соглашение с приложениями) СОБИРАЙ ПО ЧАСТЯМ — это ОБЯЗАТЕЛЬНО, "
            "иначе один огромный вызов обрывает связь с ИИ и документ НЕ создаётся. Схема: "
            "(1) первый вызов export_document(title=..., section='<HTML шапки + преамбула + раздел 1>') — "
            "вернётся doc_token; (2) следующие вызовы export_document(doc_token='<тот же>', section='<HTML "
            "следующих 1-2 разделов>') — по очереди; (3) финал export_document(doc_token='<тот же>', "
            "finalize=true) — вернётся ссылка url. Держи КАЖДУЮ section небольшой (примерно до 6000 "
            "символов). Одним вызовом html=... шли ТОЛЬКО короткие документы (справка, простое письмо)."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "title": {"type": "string", "description": "Имя файла (без расширения), например «Договор поставки ткани». Указывай в ПЕРВОМ вызове."},
                "html": {"type": "string", "description": "Полный HTML КОРОТКОГО документа за один вызов. Для длинного документа НЕ используй — собирай через section/doc_token/finalize."},
                "section": {"type": "string", "description": "Очередная ЧАСТЬ HTML длинного документа (до ~6000 символов). Первый вызов с section и без doc_token открывает черновик и возвращает doc_token."},
                "doc_token": {"type": "string", "description": "Идентификатор черновика (возвращается первым вызовом). Передавай во все последующие section-вызовы и в finalize."},
                "finalize": {"type": "boolean", "description": "true в ПОСЛЕДНЕМ вызове — собрать docx из всех накопленных секций и вернуть ссылку url."},
                "format": {"type": "string", "enum": ["docx"], "description": "Формат файла (пока docx)."},
                "font_size_pt": {"type": "number", "description": "Базовый размер шрифта, pt (по умолчанию 12). Указывай в первом вызове."},
                "line_spacing": {"type": "number", "description": "Межстрочный интервал (по умолчанию 1.15; для договоров обычно 1.5). Указывай в первом вызове."},
            },
            "required": ["title"],
            "additionalProperties": False,
        },
        "handler": tool_export_document,
    },
    "make_sheet_applet": {
        "description": (
            "Make a Google Sheet usable by an ANONYMOUS Apps Script web app WITHOUT any Google login or "
            "authorization. Returns a public token-protected applet_url (served by Albery with a9ent.ai's "
            "token): GET it -> {values:[[...]]} (reads rows), POST {values:[...]} -> appends a row, plus a "
            "ready html_snippet (appletRows()/appletAdd()). HOW TO BUILD A DATA WEB APP: the Apps Script "
            "doGet must return ONLY HTML/JS (NEVER call SpreadsheetApp/DriveApp in the code — that makes "
            "Google demand login and return 403 to anonymous users); read/write the sheet from the page's "
            "JS via this applet. Pass the spreadsheet id or URL."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "spreadsheet_id": {"type": "string", "description": "Spreadsheet id or URL"},
                "sheet": {"type": "string", "description": "Optional sheet/tab name (default: first sheet)"},
            },
            "required": ["spreadsheet_id"],
            "additionalProperties": False,
        },
        "handler": tool_make_sheet_applet,
    },
    "share_drive_item_for_everyone": {
        "description": (
            "Open ANY Google Drive item — spreadsheet, document, folder, file, or an Apps Script project — for "
            "ANYONE WITH THE LINK (editor by default). Accepts a Drive/Docs/Sheets id or URL. ALWAYS call this "
            "after creating a Google object, or before sending its link to a person, so the link is never "
            "'Нет доступа'. For a sheet your Apps Script created at runtime, pass its id/url here too."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "item": {"type": "string", "description": "Drive/Docs/Sheets id or URL"},
                "file_id": {"type": "string", "description": "Alias for item"},
                "role": {"type": "string", "description": "writer (default = editor) or reader (viewer)"},
            },
            "required": ["item"],
            "additionalProperties": False,
        },
        "handler": tool_share_drive_item_for_everyone,
    },
    "remove_drive_item_from_folder": {
        "description": (
            "Remove a Google Drive item — file, spreadsheet, document OR folder — from one specified parent "
            "folder without deleting the item from Drive completely. Use this when the owner says: 'удали эту "
            "таблицу/файл/папку из этой папки', 'убери из папки'. Strict rule: ask for confirmation first and "
            "call with confirm=true only after the user approved the exact item and folder. If the item has no "
            "other parent, Google Drive may place it in the owner's My Drive/root or keep it accessible by link; "
            "this is still not permanent deletion."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "item_id": {"type": "string", "description": "Drive file/folder id or URL"},
                "file_id": {"type": "string", "description": "Alias for item_id; Drive file/folder id or URL"},
                "folder": {"type": "string", "description": "Parent Drive folder id or folder URL to remove from"},
                "confirm": {"type": "boolean"},
            },
            "required": ["item_id", "folder", "confirm"],
            "additionalProperties": False,
        },
        "handler": tool_remove_drive_item_from_folder,
    },

    "list_drive_folder_items": {
        "description": (
            "List direct contents of a Google Drive folder: files AND subfolders, with ids, names, mime types and links. "
            "Use this before creating/moving/removing Drive items, and before proposing a folder-sorting plan."
        ),
        "inputSchema": {"type": "object", "properties": {"folder": {"type": "string", "description": "Drive folder id or URL"}, "page_size": {"type": "integer", "default": 200}}, "required": ["folder"], "additionalProperties": False},
        "handler": tool_list_drive_folder_items,
    },
    "create_drive_folder": {
        "description": "Create a Google Drive subfolder inside a specified parent folder, or reuse an existing exact-name subfolder. Ask for confirmation first and call with confirm=true.",
        "inputSchema": {"type": "object", "properties": {"name": {"type": "string", "description": "New subfolder name"}, "parent_folder": {"type": "string", "description": "Parent Drive folder id or URL"}, "folder": {"type": "string", "description": "Alias for parent_folder"}, "reuse_existing": {"type": "boolean", "default": True}, "confirm": {"type": "boolean"}}, "required": ["name", "confirm"], "additionalProperties": False},
        "handler": tool_create_drive_folder,
    },
    "organize_drive_folder": {
        "description": "Smartly organize a Google Drive folder: create/reuse category subfolders and move files AND folders into categories. Use dry_run=true first; after approval use dry_run=false and confirm=true.",
        "inputSchema": {"type": "object", "properties": {"folder": {"type": "string", "description": "Drive folder id or URL to organize"}, "categories": {"type": "array", "items": {"type": "string"}}, "dry_run": {"type": "boolean", "default": True}, "confirm": {"type": "boolean"}}, "required": ["folder"], "additionalProperties": False},
        "handler": tool_organize_drive_folder,
    },
    "manage_apps_script": {
        "description": (
            "Google Apps Script via the Apps Script API. To MAKE A WORKING WEB APP in one call use "
            "action=publish_web_app with files=[{name,type:HTML|SERVER_JS,source}] (the code needs a "
            "doGet/doPost) -> returns a ready web_app_url (https://script.google.com/macros/s/.../exec) "
            "open by link to everyone (access ANYONE_ANONYMOUS, executeAs USER_DEPLOYING). Other actions: "
            "create (new project) | get (files) | update (overwrite files=[{name,type:SERVER_JS|JSON|HTML,"
            "source}], manifest preserved) | deploy (version+deploy, web app by default -> web_app_url) | "
            "run (run function_name). advanced_services=['drive','sheets','calendar','gmail','docs',...] "
            "enables Apps Script advanced services in the manifest; oauth_scopes adds runtime scopes. "
            "share=true (default) makes the project editable by link. ALWAYS give the owner the web_app_url, "
            "not the editor_url. Requires confirm=true. If the API is disabled, ask the owner to enable the "
            "Apps Script API in Google Cloud for a9ent.ai (and for an advanced service, its matching API)."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "action": {"type": "string", "enum": ["create", "get", "update", "deploy", "run", "publish_web_app"]},
                "script_id": {"type": "string"},
                "title": {"type": "string"},
                "files": {"type": "array", "items": {"type": "object"}, "description": "[{name,type:SERVER_JS|HTML|JSON,source}]"},
                "function_name": {"type": "string"},
                "parameters": {"type": "array"},
                "description": {"type": "string"},
                "web_app": {"type": "boolean", "description": "deploy as a web app (default true)"},
                "access": {"type": "string", "description": "ANYONE_ANONYMOUS (default, open to all) | ANYONE | DOMAIN | MYSELF"},
                "execute_as": {"type": "string", "description": "USER_DEPLOYING (default) | USER_ACCESSING"},
                "advanced_services": {"type": "array", "items": {"type": "string"}, "description": "e.g. ['drive','sheets','calendar']"},
                "oauth_scopes": {"type": "array", "items": {"type": "string"}},
                "share": {"type": "boolean", "description": "share project editable by link (default true)"},
                "confirm": {"type": "boolean", "description": "must be true"},
            },
            "required": ["action", "confirm"],
            "additionalProperties": False,
        },
        "handler": tool_manage_apps_script,
    },
    "cancel_owner_recommendation": {
        "description": (
            "Mark one owner_manager_recommendations row as cancelled (e.g. the owner decided not to send it). "
            "Writes an owner_recommendation_events 'cancelled' entry with the reason. Use when the owner explicitly "
            "rejects a draft for a specific person."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "recommendation_id": {"type": "string", "description": "UUID of the recommendation row."},
                "reason": {"type": "string", "description": "Short owner-provided reason (optional)."},
            },
            "required": ["recommendation_id"],
            "additionalProperties": False,
        },
        "handler": tool_cancel_owner_recommendation,
    },
    "fetch_url": {
        "description": (
            "Fetch the contents of a web URL the user sent in chat and return it as readable text. THE tool for "
            "«вот ссылка — о чём страница / прочитай / вытащи данные». Reads: articles and normal pages (HTML "
            "stripped to text); JS-heavy or anti-bot pages (Дзен, новостные сайты, SPA) — automatically re-read "
            "through a rendering reader proxy, so a login-redirect or an empty shell still yields the article text; "
            "Word/PDF/Excel files by URL — the document text is extracted (use for «прочитай договор по ссылке»); "
            "Google Sheets and Google Docs are read with the agent's authorized Google account (private files shared "
            "with that account work without public access). Size is hard-capped (default 50000 chars, max 200000). "
            "If the result has kind='reader', the text is a rendered markdown of the page. On 403/404 from Google "
            "docs, the file is not accessible to the agent Google account. Do NOT use this for company knowledge "
            "that already lives in Albery — prefer search_company_knowledge, list_company_files, get_company_file, "
            "search_messages, get_zoom_call_transcript for that."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "url": {
                    "type": "string",
                    "description": "Full http(s) URL to fetch. Google Sheets/Docs links are read through the agent authorized Google account when possible.",
                },
                "max_chars": {
                    "type": "integer",
                    "minimum": 100,
                    "maximum": 200000,
                    "description": "Hard cap on returned text length. Default 50000. Bump only if a larger document is genuinely needed.",
                },
                "strip_html": {
                    "type": "boolean",
                    "description": "Strip HTML tags to plain text when the response looks like HTML. Default true.",
                },
            },
            "required": ["url"],
            "additionalProperties": False,
        },
        "handler": tool_fetch_url,
    },
    "upsert_ai_instruction": {
        "description": "Create or update one editable AI instruction folder by path in Настройки -> Инструкции для ИИ.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "path": {"type": "string", "description": "Folder path separated by /, for example: Формирование отчетов/Ежедневный отчет по компании"},
                "content": {"type": "string"},
            },
            "required": ["path", "content"],
            "additionalProperties": False,
        },
        "handler": tool_upsert_ai_instruction,
    },
    "get_compact_export": {
        "description": "Generate a compact export bundle for a period from live PostgreSQL data.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "date_from": {"type": "string"},
                "date_to": {"type": "string"},
                "include_messages": {"type": "boolean"},
                "include_zoom_calls": {"type": "boolean"},
                "message_limit": {"type": "integer", "minimum": 1, "maximum": MAX_LIMIT},
                "task_limit": {"type": "integer", "minimum": 1, "maximum": MAX_LIMIT},
                "zoom_limit": {"type": "integer", "minimum": 1, "maximum": MAX_LIMIT},
            },
            "required": ["date_from", "date_to"],
            "additionalProperties": False,
        },
        "handler": tool_get_compact_export,
    },
    "get_tg_news": {
        "description": (
            "СВЕЖИЕ ПОСТЫ отраслевых Telegram-каналов (список ведёт владелец: WB/маркетплейсы, "
            "оргпрактики, ИИ). Один вызов = все отслеживаемые каналы за period_days (по умолчанию 7), "
            "результат кэшируется 30 минут. channels=['имя'] — только выбранные и подробнее. "
            "Источник — публичные веб-превью t.me/s/."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "days": {"type": "integer", "minimum": 1, "maximum": 30, "description": "За сколько дней собрать посты (по умолчанию 7)."},
                "channels": {"type": "array", "items": {"type": "string"}, "description": "Только эти каналы (@имя или имя). Пусто = весь список владельца."},
                "max_posts_per_channel": {"type": "integer", "minimum": 1, "maximum": 50, "description": "Максимум постов на канал (по умолчанию 12, самые свежие)."},
                "post_chars": {"type": "integer", "minimum": 200, "maximum": 2000, "description": "Обрезка текста поста (по умолчанию 700 символов)."},
            },
            "additionalProperties": False,
        },
        "handler": tool_get_tg_news,
    },
    "get_bitrix_departments": {
        "description": (
            "ЖИВАЯ оргструктура портала: отделы (id, название, родитель, руководитель) и кто в "
            "каждом отделе (с id и должностями). Всегда начинай работу с оргструктурой отсюда — "
            "здесь точные id, которые нужны остальным инструментам. Только чтение."
        ),
        "inputSchema": {"type": "object", "properties": {}, "additionalProperties": False},
        "handler": tool_get_bitrix_departments,
    },
    "manage_bitrix_department": {
        "description": (
            "ОТДЕЛЫ оргструктуры: create (name, опц. parent_id, head_bitrix_user_id), update "
            "(department_id + name/parent_id/head_bitrix_user_id — переименовать, переподчинить, "
            "назначить руководителя), delete (только пустой отдел). МЕНЯТЬ ОРГСТРУКТУРУ МОГУТ "
            "ТОЛЬКО Евгений Палей и ИИ Агент — обязателен requested_by_bitrix_user_id (id того, кто "
            "просит); остальным инструмент откажет. Порядок: get_bitrix_departments → показать "
            "человеку ТОЧНЫЙ план с id → дождаться «да» → вызвать с confirm=true."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "action": {"type": "string", "enum": ["create", "update", "delete"]},
                "department_id": {"type": "integer", "description": "id отдела для update/delete."},
                "name": {"type": "string", "description": "Название отдела (create / переименование)."},
                "parent_id": {"type": "integer", "description": "Родительский отдел (по умолчанию корневой 1)."},
                "head_bitrix_user_id": {"type": "integer", "description": "id руководителя отдела."},
                "requested_by_bitrix_user_id": {"type": "integer", "description": "id того, КТО просит изменение (только Евгений Палей / ИИ Агент)."},
                "confirm": {"type": "boolean", "description": "true — после показа плана и явного подтверждения человека."},
            },
            "required": ["action", "requested_by_bitrix_user_id"],
            "additionalProperties": False,
        },
        "handler": tool_manage_bitrix_department,
    },
    "assign_employee_department": {
        "description": (
            "Перевести сотрудника(ов) в отдел и/или задать должность. employees — список id или ФИО "
            "(резолвится по активным сотрудникам), department_ids — id отделов (человек может быть "
            "в нескольких), position — должность. МЕНЯТЬ МОГУТ ТОЛЬКО Евгений Палей и ИИ Агент "
            "(requested_by_bitrix_user_id). Сначала get_bitrix_departments + показать план с id, "
            "затем confirm=true."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "employees": {"type": "array", "items": {"type": ["integer", "string"]}, "description": "id или ФИО сотрудников."},
                "department_ids": {"type": "array", "items": {"type": "integer"}, "description": "id отделов, куда перевести."},
                "position": {"type": "string", "description": "Должность (WORK_POSITION)."},
                "requested_by_bitrix_user_id": {"type": "integer", "description": "id того, КТО просит (только Евгений Палей / ИИ Агент)."},
                "confirm": {"type": "boolean", "description": "true — после показа плана и явного подтверждения."},
            },
            "required": ["employees", "requested_by_bitrix_user_id"],
            "additionalProperties": False,
        },
        "handler": tool_assign_employee_department,
    },
    "save_news_digest": {
        "description": (
            "Сохранить недельную новостную сводку, чтобы на повторные вопросы отвечать из неё, "
            "а не пересобирать. Вызывай в конце еженедельной автоматизации, summary = финальный "
            "текст сводки."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "summary": {"type": "string", "description": "Финальный текст сводки."},
                "period_days": {"type": "integer", "minimum": 1, "maximum": 30, "description": "За сколько дней (по умолчанию 7)."},
            },
            "required": ["summary"],
            "additionalProperties": False,
        },
        "handler": tool_save_news_digest,
    },
    "get_latest_news_digest": {
        "description": (
            "Последняя сохранённая новостная сводка + её возраст. Отвечая на вопрос о новостях, "
            "СНАЧАЛА вызови это: если is_fresh=true — отвечай на её основе, НЕ пересобирай через "
            "get_tg_news (экономия). Пересобирай только если сводки нет или устарела."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "max_age_days": {"type": "integer", "description": "Свежей считать сводку не старше стольких дней (по умолчанию 7)."},
            },
            "additionalProperties": False,
        },
        "handler": tool_get_latest_news_digest,
    },
    "list_crm_pipelines": {
        "description": (
            "ВОРОНКИ CRM (сделки Bitrix): показать все воронки с их стадиями и количеством сделок. "
            "Первый шаг любой работы с воронками/сделками — отсюда берутся category_id и коды стадий."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "include_stages": {"type": "boolean", "description": "Включить стадии каждой воронки (по умолчанию true)."},
                "include_deal_counts": {"type": "boolean", "description": "Посчитать сделки в каждой воронке (по умолчанию true)."},
            },
            "additionalProperties": False,
        },
        "handler": tool_list_crm_pipelines,
    },
    "create_crm_pipeline": {
        "description": (
            "Создать НОВУЮ ВОРОНКУ CRM (направление сделок). Bitrix сам создаст стандартные стадии; "
            "дополнительные можно передать сразу (stages) или добавить потом через "
            "manage_crm_pipeline_stage. Перед созданием покажи пользователю название и план стадий."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "name": {"type": "string", "description": "Название воронки."},
                "sort": {"type": "integer", "description": "Порядок сортировки среди воронок."},
                "stages": {
                    "type": "array",
                    "items": {"type": ["string", "object"]},
                    "description": "Доп. стадии: строки-названия или {name, stage_code?, sort?, color?, semantics?}.",
                },
                "allow_duplicate_name": {"type": "boolean", "description": "true — создать, даже если воронка с таким названием уже есть."},
            },
            "required": ["name"],
            "additionalProperties": False,
        },
        "handler": tool_create_crm_pipeline,
    },
    "update_crm_pipeline": {
        "description": "Переименовать воронку CRM или поменять её порядок (sort). Воронка — по category_id или pipeline_name.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "category_id": {"type": "integer", "description": "id воронки (из list_crm_pipelines)."},
                "pipeline_name": {"type": "string", "description": "Название воронки (если id неизвестен)."},
                "new_name": {"type": "string", "description": "Новое название."},
                "sort": {"type": "integer", "description": "Новый порядок сортировки."},
            },
            "additionalProperties": False,
        },
        "handler": tool_update_crm_pipeline,
    },
    "delete_crm_pipeline": {
        "description": (
            "УДАЛИТЬ воронку CRM. Жёсткое правило: сначала list_crm_pipelines, показать пользователю "
            "точную воронку и число сделок в ней, дождаться явного подтверждения — и только затем "
            "вызвать с confirm=true. Воронка с сделками не удаляется (сначала перенести их)."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "category_id": {"type": "integer", "description": "id воронки."},
                "pipeline_name": {"type": "string", "description": "Название воронки (если id неизвестен)."},
                "expected_name": {"type": "string", "description": "Safety-check: должно совпасть с названием воронки."},
                "confirm": {"type": "boolean", "description": "Обязательно true после явного подтверждения пользователя."},
            },
            "additionalProperties": False,
        },
        "handler": tool_delete_crm_pipeline,
    },
    "manage_crm_pipeline_stage": {
        "description": (
            "СТАДИИ воронки CRM: add — добавить стадию (name; опц. stage_code/sort/color/"
            "semantics='failure' для доп. проигрышной), update — переименовать/пересортировать/"
            "перекрасить (stage + new_name/sort/color), delete — удалить пустую несистемную стадию "
            "(confirm=true после подтверждения). Стадия задаётся кодом ('C8:NEW') или названием."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "action": {"type": "string", "enum": ["add", "update", "delete"]},
                "category_id": {"type": "integer", "description": "id воронки."},
                "pipeline_name": {"type": "string", "description": "Название воронки (если id неизвестен)."},
                "stage": {"type": "string", "description": "Для update/delete: код стадии ('C8:NEW') или её название."},
                "name": {"type": "string", "description": "Для add: название новой стадии."},
                "new_name": {"type": "string", "description": "Для update: новое название."},
                "stage_code": {"type": "string", "description": "Для add: латинский код стадии (A-Z, 0-9, _). По умолчанию генерируется."},
                "sort": {"type": "integer", "description": "Порядок стадии (меньше — левее на канбане)."},
                "color": {"type": "string", "description": "Цвет #RRGGBB."},
                "semantics": {"type": "string", "enum": ["process", "failure"], "description": "Для add: failure — дополнительная проигрышная стадия."},
                "confirm": {"type": "boolean", "description": "Для delete: обязательно true после подтверждения."},
            },
            "required": ["action"],
            "additionalProperties": False,
        },
        "handler": tool_manage_crm_pipeline_stage,
    },
    "list_crm_deal_fields": {
        "description": (
            "ПОЛЯ СДЕЛОК CRM: пользовательские поля (UF_CRM_*) с кодами/типами/подписями — реальные "
            "коды для custom_fields в create_crm_deal/update_crm_deal. include_standard=true добавит "
            "стандартные поля сделки."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "include_standard": {"type": "boolean", "description": "Включить стандартные поля сделки (по умолчанию false)."},
            },
            "additionalProperties": False,
        },
        "handler": tool_list_crm_deal_fields,
    },
    "manage_crm_deal_field": {
        "description": (
            "СОБСТВЕННЫЕ ПОЛЯ СДЕЛОК (UF_CRM_*): add — создать поле (label; type: string/integer/"
            "double/boolean/date/datetime/money/url/enumeration/employee/file/address; для "
            "enumeration обязателен list_items), update — поменять подпись/обязательность/варианты "
            "списка, delete — удалить поле (СТИРАЕТ значения во всех сделках — только с confirm=true "
            "после явного подтверждения). Поле задаётся field_code или field_id."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "action": {"type": "string", "enum": ["add", "update", "delete"]},
                "label": {"type": "string", "description": "Человеческое название поля."},
                "field_code": {"type": "string", "description": "Код поля UF_CRM_... (для add — опционально, сгенерируется)."},
                "field_id": {"type": "integer", "description": "id поля (альтернатива field_code для update/delete)."},
                "type": {"type": "string", "description": "Тип поля для add (по умолчанию string)."},
                "list_items": {"type": "array", "items": {"type": "string"}, "description": "Варианты для type=enumeration (при update — полный новый список ДОБАВЛЯЕМЫХ вариантов)."},
                "mandatory": {"type": "boolean", "description": "Обязательное поле."},
                "multiple": {"type": "boolean", "description": "Множественное значение (только при add)."},
                "show_in_list": {"type": "boolean", "description": "Показывать в списке сделок (по умолчанию true)."},
                "confirm": {"type": "boolean", "description": "Для delete: обязательно true после подтверждения."},
            },
            "required": ["action"],
            "additionalProperties": False,
        },
        "handler": tool_manage_crm_deal_field,
    },
    "list_crm_deals": {
        "description": (
            "СДЕЛКИ CRM: список с фильтрами по воронке (category_id/pipeline_name), стадии (stage — "
            "код или название), ответственному (assigned_name/assigned_bitrix_user_id), тексту в "
            "названии (search). include_closed=false скроет закрытые; include_custom_fields=true "
            "добавит UF_CRM_* поля. Сортировка — новые сверху."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "category_id": {"type": "integer", "description": "id воронки."},
                "pipeline_name": {"type": "string", "description": "Название воронки."},
                "stage": {"type": "string", "description": "Стадия: код ('C8:NEW') или название (тогда нужна и воронка)."},
                "assigned_name": {"type": "string", "description": "Ответственный по имени."},
                "assigned_bitrix_user_id": {"type": "integer", "description": "Ответственный по id."},
                "search": {"type": "string", "description": "Подстрока в названии сделки."},
                "include_closed": {"type": "boolean", "description": "false — только открытые сделки (по умолчанию true, показываются все)."},
                "include_custom_fields": {"type": "boolean", "description": "Вернуть и пользовательские поля UF_CRM_* (по умолчанию false)."},
                "limit": {"type": "integer", "minimum": 1, "maximum": 200, "description": "Максимум сделок (по умолчанию 50)."},
                "offset": {"type": "integer", "minimum": 0, "description": "Смещение для пагинации."},
            },
            "additionalProperties": False,
        },
        "handler": tool_list_crm_deals,
    },
    "get_crm_deal": {
        "description": "Одна СДЕЛКА CRM целиком: все заполненные поля, воронка/стадия по-человечески, пользовательские поля, ссылка на портал.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "deal_id": {"type": "integer", "description": "Точный id сделки."},
            },
            "required": ["deal_id"],
            "additionalProperties": False,
        },
        "handler": tool_get_crm_deal,
    },
    "create_crm_deal": {
        "description": (
            "Создать СДЕЛКУ CRM. Обязателен title; воронка — category_id/pipeline_name (без неё — "
            "основная), стадия — stage (без неё — первая), сумма amount (+currency, по умолчанию RUB), "
            "ответственный responsible_name/responsible_bitrix_user_id, комментарий comments, "
            "пользовательские поля custom_fields {UF_CRM_...: значение} (коды — list_crm_deal_fields). "
            "Перед созданием покажи пользователю, что именно будет создано."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "title": {"type": "string", "description": "Название сделки."},
                "category_id": {"type": "integer", "description": "id воронки."},
                "pipeline_name": {"type": "string", "description": "Название воронки."},
                "stage": {"type": "string", "description": "Стадия: код или название."},
                "amount": {"type": "number", "description": "Сумма сделки."},
                "currency": {"type": "string", "description": "Валюта (по умолчанию RUB)."},
                "responsible_name": {"type": "string", "description": "Ответственный по имени."},
                "responsible_bitrix_user_id": {"type": "integer", "description": "Ответственный по id."},
                "comments": {"type": "string", "description": "Комментарий к сделке."},
                "begin_date": {"type": "string", "description": "Дата начала (YYYY-MM-DD)."},
                "close_date": {"type": "string", "description": "Плановая дата закрытия (YYYY-MM-DD)."},
                "contact_id": {"type": "integer", "description": "id контакта CRM."},
                "company_id": {"type": "integer", "description": "id компании CRM."},
                "custom_fields": {"type": "object", "description": "Пользовательские поля {\"UF_CRM_...\": значение}.", "additionalProperties": True},
            },
            "required": ["title"],
            "additionalProperties": False,
        },
        "handler": tool_create_crm_deal,
    },
    "update_crm_deal": {
        "description": (
            "Изменить СДЕЛКУ CRM: название, стадию (stage — движение по воронке), сумму, "
            "ответственного, комментарий, пользовательские поля; перенос в ДРУГУЮ воронку — "
            "category_id/pipeline_name (+опц. stage целевой воронки). Указывай только то, что меняется."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "deal_id": {"type": "integer", "description": "Точный id сделки."},
                "title": {"type": "string", "description": "Новое название."},
                "stage": {"type": "string", "description": "Новая стадия: код или название."},
                "category_id": {"type": "integer", "description": "Перенести в воронку с этим id."},
                "pipeline_name": {"type": "string", "description": "Перенести в воронку с этим названием."},
                "amount": {"type": "number", "description": "Новая сумма."},
                "currency": {"type": "string", "description": "Валюта."},
                "responsible_name": {"type": "string", "description": "Новый ответственный по имени."},
                "responsible_bitrix_user_id": {"type": "integer", "description": "Новый ответственный по id."},
                "comments": {"type": "string", "description": "Новый комментарий (заменяет прежний)."},
                "begin_date": {"type": "string", "description": "Дата начала."},
                "close_date": {"type": "string", "description": "Плановая дата закрытия."},
                "contact_id": {"type": "integer", "description": "id контакта CRM."},
                "company_id": {"type": "integer", "description": "id компании CRM."},
                "custom_fields": {"type": "object", "description": "Пользовательские поля {\"UF_CRM_...\": значение}.", "additionalProperties": True},
                "expected_title": {"type": "string", "description": "Safety-check: должно совпасть с текущим названием сделки."},
            },
            "required": ["deal_id"],
            "additionalProperties": False,
        },
        "handler": tool_update_crm_deal,
    },
    "delete_crm_deal": {
        "description": (
            "УДАЛИТЬ сделку CRM. Жёсткое правило: сначала get_crm_deal, показать пользователю точную "
            "сделку (id, название, воронка, сумма), дождаться явного подтверждения — и только затем "
            "вызвать с confirm=true. expected_title — защита от удаления не той сделки."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "deal_id": {"type": "integer", "description": "Точный id сделки."},
                "expected_title": {"type": "string", "description": "Safety-check: должно совпасть с названием сделки."},
                "confirm": {"type": "boolean", "description": "Обязательно true после явного подтверждения пользователя."},
            },
            "required": ["deal_id"],
            "additionalProperties": False,
        },
        "handler": tool_delete_crm_deal,
    },
}


# Google Sheets writing runs through the Apps Script web-app deployed under xotizwf@gmail.com,
# so every edit is attributed to that account. Owner asked (2026-06-16) to stop that account
# appearing anywhere until a replacement account is connected. Disable the write tool by default;
# re-enable with ALBERY_ALLOW_SHEET_WRITE=1 once a new Apps Script account is in place.
if os.getenv("ALBERY_ALLOW_SHEET_WRITE", "").strip().lower() not in {"1", "true", "yes", "on"}:
    TOOLS.pop("write_company_sheet", None)


FAQ_TOOL_NAMES: set[str] = {
    "start_here_always_read_ai_instructions",
    "health",
    "get_runtime_status",
    "get_context_guide",
    "get_ai_instructions",
    "list_available_sources",
    "get_company_profile",
    "list_company_files",
    "get_company_file",
    "search_company_knowledge",
    "get_org_structure",
    "list_zoom_calls",
    "get_zoom_call_transcript",
    "search_zoom_transcripts",
    "get_ai_capabilities",
    # Read-only: an agent reads the full text of a file the user sent it (token-gated, unguessable).
    "get_attachment_text",
}


# Tools that change global configuration or destroy data. Public scoped connectors (/mcp-ops,
# /mcp-faq, core variants) never receive these. Per-agent connectors are different: the agent's
# own whitelist is the capability boundary, so an owner/destructive tool can be exposed there
# only when it is explicitly enabled for that particular agent.
OWNER_ONLY_TOOL_NAMES: set[str] = {
    "upsert_ai_instruction",
    "update_ai_capabilities",
    "delete_bitrix_task",
    "delete_zoom_call_report",
    # Per-employee usage/monitoring is management data — admin connector only.
    "get_agent_monitoring",
    # Employee dossiers are internal management data — never on the public scoped connectors.
    "get_employee_dossier",
    "update_employee_dossier",
    # CRM: deleting a whole funnel or a deal destroys business data — admin / explicitly-enabled only.
    "delete_crm_pipeline",
    "delete_crm_deal",
}

# Оргструктура: читать может любой тир кроме FAQ; ПРАВИТЬ — только Евгений(14)/ИИ Агент(22),
# и это проверяется В САМИХ инструментах (requested_by_bitrix_user_id + confirm), потому что
# тиры Bitrix-доступа не совпадают с этим списком (у Александра admin, но менять оргструктуру
# ему не разрешено).

# Operational-full connector: every registered tool EXCEPT the admin-only ones above.
OPS_TOOL_NAMES: set[str] = set(TOOLS) - OWNER_ONLY_TOOL_NAMES

# --- Core toolset: two-stage tool loading for the chat bot (/mcp-core, /mcp-ops-core) --------
# The chat bot registers only this curated core (picked from real usage stats in the Hermes
# session DB: ~82% of all historical calls, plus every tool the bot prompt names explicitly)
# and two meta-tools. Everything else is discovered via find_tool and invoked via call_tool.
# Cron agents keep the full /mcp and /mcp-ops connectors, so their scripted tool names are
# unaffected by this list.
CORE_TOOL_NAMES: set[str] = {
    # entry / self-knowledge
    "start_here_always_read_ai_instructions",
    "get_ai_instructions",
    "get_ai_capabilities",
    "get_context_guide",
    # company knowledge
    "search_company_knowledge",
    "list_company_files",
    "get_company_file",
    "get_org_structure",
    "get_employee_absences",
    # tasks
    "search_tasks",
    "get_task_comments",
    "add_bitrix_task_comment",
    "create_bitrix_task",
    "update_bitrix_task",
    "add_task_checklist",
    "log_task_time",
    "link_tasks",
    "create_recurring_task",
    "list_recurring_tasks",
    "update_recurring_task",
    "delete_recurring_task",
    "complete_bitrix_task",
    "reopen_bitrix_task",
    "delete_bitrix_task",
    "attach_files_to_task",
    "get_attachment_text",
    # zoom
    "list_zoom_calls",
    "get_zoom_call_transcript",
    "search_zoom_transcripts",
    # dialog memory
    "get_bitrix_bot_chat",
    "list_bitrix_bot_sessions",
    # crm funnels & deals
    "list_crm_pipelines",
    "list_crm_deals",
    "get_crm_deal",
    "create_crm_deal",
    "update_crm_deal",
    # messaging / web
    "send_bitrix_message",
    "fetch_url",
    # google workflow the bot prompt teaches
    "create_google_sheet",
    "get_google_sheet_meta",
    "write_google_sheet_values",
    "share_drive_item_for_everyone",
    "get_webapp_template",
    "make_sheet_applet",
    "manage_apps_script",
}


# --- Agent-system management tools (owner-only): let a top-access agent run the agent center ---
# These wrap the same /api/agent-center/* logic the UI uses (via a nested request context, so
# validation/enforcement is identical) and let a trusted agent create/edit/delete agents, toggle
# their tools and manage teams from chat. They are OWNER_ONLY (admin class) and only reach an
# agent whose owner explicitly enabled them — keep such an agent's team to trusted people.

def _mgmt_endpoint(method: str, path: str, view, *view_args, json_body=None) -> dict[str, Any]:
    from app import app
    with app.test_request_context(path, method=method, json=(json_body or {})):
        resp = view(*view_args)
    if isinstance(resp, tuple):
        data, status = resp[0].get_json(), resp[1]
    else:
        data, status = resp.get_json(), 200
    if status >= 400 or (isinstance(data, dict) and data.get("error")):
        raise McpError(-32000, (data or {}).get("error") if isinstance(data, dict) else f"HTTP {status}")
    return data or {}


def _mgmt_resolve(items) -> tuple[list[int], list[str]]:
    import agent_center
    ids, unresolved = [], []
    for it in (items or []):
        uid, _ = agent_center.resolve_bitrix_user(it)
        (ids.append(uid) if uid is not None else unresolved.append(str(it)))
    return ids, unresolved


def tool_list_agents(args: dict[str, Any]) -> dict[str, Any]:
    import agent_center
    return agent_center.mgmt_list_agents()


def tool_create_agent(args: dict[str, Any]) -> dict[str, Any]:
    import agent_center
    name = str(args.get("name") or "").strip()
    if not name:
        raise McpError(-32602, "Укажите name нового агента.")
    member_ids, unresolved = _mgmt_resolve(args.get("members"))
    body = {"name": name, "tier": "ops", "role_prompt": str(args.get("role_prompt") or ""),
            "position": str(args.get("position") or ""), "members": member_ids}
    data = _mgmt_endpoint("POST", "/api/agent-center/agents", agent_center.agent_center_create_agent, json_body=body)
    if unresolved:
        data["unresolved_members"] = unresolved
    return data


def tool_update_agent(args: dict[str, Any]) -> dict[str, Any]:
    import agent_center
    slug = str(args.get("slug") or "").strip()
    if not slug:
        raise McpError(-32602, "Укажите slug агента (см. list_agents).")
    body = {k: args[k] for k in ("name", "position", "role_prompt", "is_active") if k in args}
    if not body:
        raise McpError(-32602, "Нечего менять: передайте name/position/role_prompt/is_active.")
    return _mgmt_endpoint("PATCH", f"/api/agent-center/agents/{slug}", agent_center.agent_center_agent_update, slug, json_body=body)


def tool_delete_agent(args: dict[str, Any]) -> dict[str, Any]:
    import agent_center
    slug = str(args.get("slug") or "").strip()
    if not slug:
        raise McpError(-32602, "Укажите slug агента.")
    return _mgmt_endpoint("DELETE", f"/api/agent-center/agents/{slug}", agent_center.agent_center_agent_delete, slug)


def tool_set_agent_tools(args: dict[str, Any]) -> dict[str, Any]:
    import agent_center
    slug = str(args.get("slug") or "").strip()
    if not slug:
        raise McpError(-32602, "Укажите slug агента.")
    cfg = _mgmt_endpoint("GET", f"/api/agent-center/agents/{slug}/config", agent_center.agent_center_agent_config, slug)
    valid = {t["name"] for t in cfg["tools"]}
    enabled = {t["name"] for t in cfg["tools"] if t["enabled"]}
    enable = {str(t) for t in (args.get("enable") or [])}
    bad = enable - valid
    if bad:
        raise McpError(-32602, f"Неизвестные инструменты: {', '.join(sorted(bad))}. Точные имена — в list_agents/config.")
    new_enabled = (enabled | (enable & valid)) - {str(t) for t in (args.get("disable") or [])}
    body = {"tools": sorted(new_enabled),
            "instructions": [i["id"] for i in cfg["instructions"] if i["selected"]],
            "skills": [s["id"] for s in cfg["skills"] if s["selected"]]}
    return _mgmt_endpoint("PUT", f"/api/agent-center/agents/{slug}/config", agent_center.agent_center_agent_config_save, slug, json_body=body)


def tool_set_agent_team(args: dict[str, Any]) -> dict[str, Any]:
    import agent_center
    slug = str(args.get("slug") or "").strip()
    if not slug:
        raise McpError(-32602, "Укажите slug агента ('main' для универсального).")
    detail = _mgmt_endpoint("GET", f"/api/agent-center/agents/{slug}", agent_center.agent_center_agent_detail, slug)
    current = {int(m["id"]) for m in detail.get("members", [])}
    add_ids, un1 = _mgmt_resolve(args.get("add"))
    rem_ids, un2 = _mgmt_resolve(args.get("remove"))
    new_members = sorted((current | set(add_ids)) - set(rem_ids))
    data = _mgmt_endpoint("PATCH", f"/api/agent-center/agents/{slug}", agent_center.agent_center_agent_update, slug,
                          json_body={"members": new_members})
    data["team_size"] = len(new_members)
    if un1 or un2:
        data["unresolved"] = un1 + un2
    return data


AGENT_MGMT_TOOL_SPECS: dict[str, dict[str, Any]] = {
    "list_agents": {
        "description": "Список всех агентов системы (универсальный + субагенты) с их настройками: имя, должность, вкл/выкл, сколько инструментов включено, команда. Начинай управление отсюда — здесь точные slug'и.",
        "inputSchema": {"type": "object", "properties": {}, "additionalProperties": False},
        "handler": tool_list_agents,
    },
    "create_agent": {
        "description": "Создать нового субагента (Bitrix-бот зарегистрируется автоматически). Стартует с широкого набора «все функции» — дальше настрой инструменты через set_agent_tools.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "name": {"type": "string", "description": "Имя агента (как в Bitrix)."},
                "position": {"type": "string", "description": "Должность (WORK_POSITION в Bitrix)."},
                "role_prompt": {"type": "string", "description": "Роль/системный промпт агента."},
                "members": {"type": "array", "items": {"type": "string"}, "description": "Кому доступен: имена или id сотрудников (пусто = всем с доступом)."},
            },
            "required": ["name"],
            "additionalProperties": False,
        },
        "handler": tool_create_agent,
    },
    "update_agent": {
        "description": "Изменить агента: имя/должность (синхронизируются с Bitrix), роль-промпт, вкл/выкл. slug бери из list_agents ('main' — универсальный).",
        "inputSchema": {
            "type": "object",
            "properties": {
                "slug": {"type": "string"},
                "name": {"type": "string"},
                "position": {"type": "string"},
                "role_prompt": {"type": "string"},
                "is_active": {"type": "boolean"},
            },
            "required": ["slug"],
            "additionalProperties": False,
        },
        "handler": tool_update_agent,
    },
    "delete_agent": {
        "description": "Удалить субагента (разрегистрирует Bitrix-бота, чистит коннектор и данные). Универсального ('main') удалить нельзя.",
        "inputSchema": {
            "type": "object",
            "properties": {"slug": {"type": "string"}},
            "required": ["slug"],
            "additionalProperties": False,
        },
        "handler": tool_delete_agent,
    },
    "set_agent_tools": {
        "description": "Включить/выключить MCP-инструменты у агента. enable/disable — точные имена инструментов (см. list_agents → возможности). Базовые инструменты остаются всегда.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "slug": {"type": "string"},
                "enable": {"type": "array", "items": {"type": "string"}},
                "disable": {"type": "array", "items": {"type": "string"}},
            },
            "required": ["slug"],
            "additionalProperties": False,
        },
        "handler": tool_set_agent_tools,
    },
    "set_agent_team": {
        "description": "Добавить/убрать людей из команды агента (кому он доступен). add/remove — имена или id. Для универсального (slug='main') это выдаёт/снимает доступ к основному боту.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "slug": {"type": "string"},
                "add": {"type": "array", "items": {"type": "string"}},
                "remove": {"type": "array", "items": {"type": "string"}},
            },
            "required": ["slug"],
            "additionalProperties": False,
        },
        "handler": tool_set_agent_team,
    },
}

TOOLS.update(AGENT_MGMT_TOOL_SPECS)
# Owner-only/admin class: they manage the agent system itself. Kept out of the ops/faq
# connectors (OPS_TOOL_NAMES was already materialised above without them) and shown with the
# 'admin' chip + confirm in the UI; only reach an agent whose owner explicitly enabled them.
OWNER_ONLY_TOOL_NAMES.update(AGENT_MGMT_TOOL_SPECS.keys())


META_TOOL_SPECS: dict[str, dict[str, Any]] = {
    "find_tool": {
        "description": (
            "Найди инструмент по задаче. Твой список инструментов — ЯДРО самых частых; у "
            "коннектора есть и другие. Если нужного действия нет в списке — НЕ отвечай «не "
            "умею»: вызови этот поиск, получи имя/описание/схему аргументов и выполни действие "
            "через call_tool. Query — короткие английские ключевые слова по смыслу действия "
            "(например 'delete task', 'zoom report', 'drive folder', 'owner recommendations')."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "query": {
                    "type": "string",
                    "description": "Английские ключевые слова: что нужно сделать.",
                },
                "limit": {
                    "type": "integer",
                    "description": "Сколько кандидатов вернуть (по умолчанию 5).",
                },
            },
            "required": ["query"],
        },
    },
    "call_tool": {
        "description": (
            "Вызови любой инструмент коннектора по точному имени — в том числе не входящий в "
            "ядро. Сначала найди его через find_tool и заполни arguments по его inputSchema."
        ),
        "inputSchema": {
            "type": "object",
            "properties": {
                "name": {"type": "string", "description": "Точное имя инструмента (из find_tool)."},
                "arguments": {"type": "object", "description": "Аргументы по схеме инструмента."},
            },
            "required": ["name"],
        },
    },
}


# Live activity feed for the chat bot's status message: every tools/call on the CORE
# connectors (the bot's own) is recorded here, so the bridge can show what the agent is doing
# right now. Cron agents use the full connectors (core=False) and never pollute this feed.
import collections as _collections

_RECENT_CORE_TOOL_CALLS: "_collections.deque" = _collections.deque(maxlen=64)


def record_core_tool_call(name: str) -> None:
    _RECENT_CORE_TOOL_CALLS.append((time.time(), str(name)))


def recent_core_tool_calls(since_ts: float) -> list:
    return [(ts, name) for ts, name in list(_RECENT_CORE_TOOL_CALLS) if ts >= since_ts]


def _find_tool_matches(query: Any, tool_names: set[str] | None, limit: int) -> list[dict[str, Any]]:
    tokens = [t for t in re.split(r"[^0-9a-zA-Zа-яА-ЯёЁ_]+", str(query or "").lower()) if len(t) >= 3]
    if not tokens:
        raise McpError(-32602, "Нужен запрос: find_tool(query='что нужно сделать', английскими словами).")
    scored: list[tuple[int, str, dict[str, Any]]] = []
    for name, spec in _allowed_tools(tool_names).items():
        hay_name = name.lower()
        hay_desc = str(spec.get("description") or "").lower()
        score = 0
        for tok in tokens:
            if tok in hay_name:
                score += 30
            score += 3 * hay_desc.count(tok)
        if score:
            scored.append((score, name, spec))
    scored.sort(key=lambda item: (-item[0], item[1]))
    return [
        {
            "name": name,
            "description": spec["description"],
            "inputSchema": spec["inputSchema"],
            "how_to_call": "call_tool(name='" + name + "', arguments={...})",
        }
        for _score, name, spec in scored[: max(1, limit)]
    ]


def _allowed_tools(tool_names: set[str] | None = None) -> dict[str, dict[str, Any]]:
    if tool_names is None:
        return TOOLS
    return {name: TOOLS[name] for name in tool_names if name in TOOLS}


def list_tools(tool_names: set[str] | None = None, core: bool = False) -> list[dict[str, Any]]:
    registry = _allowed_tools(tool_names)
    if core:
        registry = {name: registry[name] for name in sorted(CORE_TOOL_NAMES) if name in registry}
    items = [
        {
            "name": name,
            "description": (
                spec["description"]
                if name == "start_here_always_read_ai_instructions"
                else TOOL_USAGE_CONTRACT + spec["description"]
            ),
            "inputSchema": spec["inputSchema"],
        }
        for name, spec in registry.items()
    ]
    if core:
        items.extend(
            {"name": name, "description": spec["description"], "inputSchema": spec["inputSchema"]}
            for name, spec in META_TOOL_SPECS.items()
        )
    return items


def handle_request(request: dict[str, Any], tool_names: set[str] | None = None,
                   core: bool = False,
                   allow_owner_tools: bool = False,
                   instruction_scope: set[str] | list[str] | None = None,
                   agent_slug: str | None = None) -> dict[str, Any] | None:
    request_id = request.get("id")
    method = request.get("method")
    available_tools = _allowed_tools(tool_names)

    # JSON-RPC notifications (initialized, cancelled, progress, ...) are one-way:
    # answering them with a response object (even an error) violates the protocol.
    if isinstance(method, str) and method.startswith("notifications/"):
        return None

    try:
        if method == "initialize":
            result = {
                "protocolVersion": PROTOCOL_VERSION,
                "capabilities": {"tools": {}},
                "serverInfo": {"name": SERVER_NAME, "version": SERVER_VERSION},
            }
        elif method == "ping":
            # MCP liveness probe. hermes >=0.17 pings every connector each keepalive
            # interval (180s) and treats ANY failure as a dead connection with a
            # finite reconnect budget — an unanswered ping killed every connector
            # within ~15 minutes (incident 2026-07-06, agents lost all tools).
            result = {}
        elif method == "tools/list":
            result = {"tools": list_tools(tool_names, core=core)}
        elif method == "tools/call":
            params = request.get("params") or {}
            name = params.get("name")
            args = params.get("arguments") or {}
            if core and name == "find_tool":
                record_core_tool_call("find_tool")
                matches = _find_tool_matches(args.get("query"), tool_names, int(args.get("limit") or 5))
                return {"jsonrpc": "2.0", "id": request_id, "result": text_response({
                    "matches": matches,
                    "note": "Вызывай выбранный инструмент через call_tool(name=..., arguments={...}) по его inputSchema.",
                })}
            if core and name == "call_tool":
                inner_args = args.get("arguments")
                name = str(args.get("name") or "").strip()
                args = inner_args if isinstance(inner_args, dict) else {}
                logger.info("mcp_call_tool_proxy name=%s", name)
            if name not in available_tools:
                if core:
                    raise McpError(-32601, f"Unknown or unavailable tool: {name}. Найди точное имя через find_tool.")
                raise McpError(-32601, f"Unknown or unavailable tool: {name}")
            # Public scoped connectors (/mcp-ops, /mcp-faq, core variants) must never expose
            # owner/destructive tools. Per-agent connectors pass allow_owner_tools=True: their
            # security boundary is the exact agent whitelist from _agent_tool_names().
            if name in OWNER_ONLY_TOOL_NAMES and tool_names is not None and not allow_owner_tools:
                raise McpError(-32601, f"Unknown or unavailable tool: {name}")
            if core:
                record_core_tool_call(name)
            if name in ("start_here_always_read_ai_instructions", "get_ai_capabilities"):
                connector_id = "faq" if tool_names == FAQ_TOOL_NAMES else "full"
                args = {
                    **args,
                    "_connector_tools": sorted(available_tools.keys()),
                    "_connector_id": connector_id,
                }
                if core:
                    args["_connector_tools"] = sorted(
                        (set(available_tools.keys()) & CORE_TOOL_NAMES) | set(META_TOOL_SPECS)
                    )
                    args["_connector_hidden_tools"] = sorted(
                        set(available_tools.keys()) - CORE_TOOL_NAMES
                    )
            # Per-agent instruction scope: when this connector is bound to an agent,
            # start_here / get_ai_instructions return ONLY the agent's allowed paths
            # (universal + connected). None = no scoping (legacy public/admin connectors).
            if instruction_scope is not None and name in (
                "start_here_always_read_ai_instructions", "get_ai_instructions"
            ):
                args = {**args, "_allowed_instruction_paths": list(instruction_scope)}
            # Recurring tasks are attributed to the agent whose connector created them,
            # so the row lands in THAT agent's «Автоматизации» tab.
            if agent_slug and name == "create_recurring_task":
                args = {**args, "_agent_slug": agent_slug}
            started = time.perf_counter()
            result_payload = available_tools[name]["handler"](args)
            duration_ms = round((time.perf_counter() - started) * 1000, 1)
            if duration_ms >= TOOL_LATENCY_LOG_MS:
                result_size = len(json.dumps(json_safe(result_payload), ensure_ascii=False, default=json_default))
                logger.info("mcp_tool_call name=%s duration_ms=%s result_bytes=%s", name, duration_ms, result_size)
            result = text_response(result_payload)
        else:
            raise McpError(-32601, f"Unknown method: {method}")
        return {"jsonrpc": "2.0", "id": request_id, "result": result}
    except McpError as exc:
        tool_name = ""
        if method == "tools/call" and isinstance(request.get("params"), dict):
            tool_name = " tool=" + str(request["params"].get("name") or "")
        logger.warning(
            "MCP request rejected: method=%s%s id=%s code=%s message=%s",
            method, tool_name, request_id, exc.code, exc.message,
        )
        return {"jsonrpc": "2.0", "id": request_id, "error": {"code": exc.code, "message": exc.message}}
    except Exception:
        logger.exception("Unhandled MCP request error: method=%s id=%s", method, request_id)
        return {"jsonrpc": "2.0", "id": request_id, "error": {"code": -32000, "message": "Internal MCP error."}}


def main() -> None:
    try:
        sys.stdout.reconfigure(encoding="utf-8")
        sys.stdin.reconfigure(encoding="utf-8")
    except Exception:
        pass
    for line in sys.stdin:
        line = line.strip().lstrip("\ufeff")
        if not line:
            continue
        try:
            request = json.loads(line)
            response = handle_request(request)
        except Exception:
            logger.exception("Failed to parse or handle MCP stdin request")
            response = {"jsonrpc": "2.0", "id": None, "error": {"code": -32700, "message": "Invalid MCP request."}}
        if response is not None:
            sys.stdout.write(json.dumps(response, ensure_ascii=False, default=json_default) + "\n")
            sys.stdout.flush()


if __name__ == "__main__":
    main()
