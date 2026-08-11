"""Channel-neutral recognition for employee-sent images, audio and documents.

Groq is deliberately a media preprocessor only.  It extracts pixels/speech into text; the
logical Albery agent still reasons and calls tools through its Hermes/Codex profile.
"""
from __future__ import annotations

import base64
import io
import json
import logging
import os
import re
import urllib.request
from pathlib import Path

import requests


log = logging.getLogger("albery.media")
IMAGE_EXTENSIONS = {"png", "jpg", "jpeg", "gif", "webp", "bmp", "heic"}
AUDIO_EXTENSIONS = {
    "mp3", "m4a", "aac", "ogg", "oga", "opus", "wav", "webm", "flac", "mp4", "mpga",
    "mpeg", "amr",
}
DOCUMENT_EXTENSIONS = {
    "pdf", "docx", "doc", "xlsx", "xlsm", "xls", "md", "markdown", "txt", "csv", "tsv",
    "json", "rtf", "htm", "html", "log", "yaml", "yml", "pptx", "ppt", "mht", "mhtml",
    "odt", "ods", "odp", "zip", "xml",
}
_CONTROL_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")
_THINK_RE = re.compile(r"<think>.*?</think>\s*", re.S | re.I)
_VISION_PROMPT = (
    "Это изображение или скриншот, который сотрудник прислал ИИ-агенту компании. Извлеки весь "
    "видимый текст дословно, затем в 1–2 предложениях опиши содержимое и возможную проблему. "
    "Ответь по-русски, без вступлений и рассуждений."
)


def groq_api_key() -> str:
    value = os.getenv("GROQ_API_KEY", "").strip()
    if value:
        return value
    try:
        for line in Path("/root/.hermes/secure/hermes-gateway.env").read_text(
            encoding="utf-8"
        ).splitlines():
            if line.strip().startswith("GROQ_API_KEY="):
                return line.split("=", 1)[1].strip().strip('"').strip("'")
    except OSError:
        pass
    return ""


def recognize_image(data: bytes, name: str = "image.png") -> str:
    key = groq_api_key()
    if not key or not data:
        return ""
    ext = name.rsplit(".", 1)[-1].lower() if "." in name else "png"
    mime = {"jpg": "jpeg", "jpeg": "jpeg", "png": "png", "gif": "gif", "webp": "webp"}.get(
        ext, "png"
    )
    models = [item.strip() for item in os.getenv(
        "GROQ_VISION_MODELS",
        os.getenv("B24_VISION_MODELS", "qwen/qwen3.6-27b,meta-llama/llama-4-scout-17b-16e-instruct"),
    ).split(",") if item.strip()]
    encoded = base64.b64encode(data).decode("ascii")
    for model in models:
        payload = {
            "model": model,
            "max_tokens": 800,
            "temperature": 0,
            "messages": [{"role": "user", "content": [
                {"type": "text", "text": _VISION_PROMPT},
                {"type": "image_url", "image_url": {
                    "url": f"data:image/{mime};base64,{encoded}"
                }},
            ]}],
        }
        try:
            request = urllib.request.Request(
                "https://api.groq.com/openai/v1/chat/completions",
                data=json.dumps(payload).encode("utf-8"),
                headers={
                    "Authorization": "Bearer " + key,
                    "Content-Type": "application/json",
                    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)",
                },
            )
            with urllib.request.urlopen(request, timeout=70) as response:
                body = json.loads(response.read().decode("utf-8", "ignore"))
            text = ((body.get("choices") or [{}])[0].get("message") or {}).get("content", "")
            text = _THINK_RE.sub("", str(text or "")).strip()
            if text:
                return text
        except Exception as exc:  # noqa: BLE001
            log.warning("Groq vision model %s failed: %s", model, type(exc).__name__)
    return ""


def transcribe_audio(data: bytes, name: str = "voice.ogg") -> str:
    key = groq_api_key()
    if not key or not data:
        return ""
    try:
        response = requests.post(
            "https://api.groq.com/openai/v1/audio/transcriptions",
            headers={
                "Authorization": "Bearer " + key,
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)",
            },
            files={"file": ((name or "voice.ogg").rsplit("/", 1)[-1], data)},
            data={
                "model": os.getenv("GROQ_STT_MODEL", os.getenv("B24_STT_MODEL", "whisper-large-v3")),
                "response_format": "json",
                "temperature": "0",
            },
            timeout=180,
        )
        if response.status_code != 200:
            log.warning("Groq transcription rejected %s with HTTP %s", name, response.status_code)
            return ""
        return _CONTROL_RE.sub("", str((response.json() or {}).get("text") or "")).strip()
    except Exception as exc:  # noqa: BLE001
        log.warning("Groq transcription failed for %s: %s", name, type(exc).__name__)
        return ""


def extract_document(data: bytes, name: str) -> str:
    ext = name.rsplit(".", 1)[-1].lower() if "." in name else ""
    try:
        if ext in {"md", "markdown", "txt", "csv", "tsv", "json", "log", "yaml", "yml"}:
            text = data.decode("utf-8", "ignore")
        elif ext == "pdf":
            from pypdf import PdfReader
            text = "\n".join((page.extract_text() or "") for page in PdfReader(io.BytesIO(data)).pages)
        elif ext == "docx":
            from docx import Document
            document = Document(io.BytesIO(data))
            parts = [paragraph.text for paragraph in document.paragraphs]
            parts.extend(" | ".join(cell.text for cell in row.cells)
                         for table in document.tables for row in table.rows)
            text = "\n".join(parts)
        elif ext in {"xlsx", "xlsm"}:
            import webread
            text = webread.extract_xlsx(data)
        else:
            import docextract
            text = docextract.extract(data, name, inner=extract_document) if ext in docextract.EXTS else ""
        return _CONTROL_RE.sub("", str(text or "")).strip()
    except Exception as exc:  # noqa: BLE001
        log.warning("Document extraction failed for %s: %s", name, type(exc).__name__)
        return ""
